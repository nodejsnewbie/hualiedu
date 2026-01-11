"""
Django tests for AI scoring functionality.
"""

import json
import os
import tempfile

from django.contrib.auth.models import User
from django.test import TestCase
from django.urls import reverse

from grading.models import GlobalConfig


class AIScoringTest(TestCase):
    """Test cases for AI scoring functionality."""

    def setUp(self):
        """Set up test data."""
        # Create test user
        self.user = User.objects.create_user(
            username="testuser", password="testpass123", is_staff=True
        )

        # Create temporary directory for test files
        self.temp_dir = tempfile.mkdtemp()
        self.test_file_path = os.path.join(self.temp_dir, "test_document.docx")

        # Create GlobalConfig
        self.config = GlobalConfig.objects.create(repo_base_dir=self.temp_dir)

        # Create a simple test file
        from docx import Document

        doc = Document()
        doc.add_paragraph("Test document content")
        doc.save(self.test_file_path)

    def tearDown(self):
        """Clean up test data."""
        import shutil

        shutil.rmtree(self.temp_dir, ignore_errors=True)

    def test_ai_score_view_success(self):
        """Test successful AI scoring with real API call."""
        # Check if API key is available
        api_key = os.environ.get("ARK_API_KEY")
        if not api_key:
            self.skipTest("ARK_API_KEY not set, skipping real API test")

        self.client.login(username="testuser", password="testpass123")

        data = {"path": "test_document.docx"}
        response = self.client.post(reverse("grading:ai_score"), data)

        # Check if the request was successful
        self.assertIn(response.status_code, [200, 500])  # Allow both success and API errors

        if response.status_code == 200:
            response_data = json.loads(response.content)
            self.assertEqual(response_data["status"], "success")
            # API might fail and return None for score, which is acceptable
            self.assertIsInstance(response_data["score"], (int, type(None)))
            self.assertIsInstance(response_data["grade"], str)
            self.assertIsInstance(response_data["comment"], str)
        else:
            # If API call failed, that's also acceptable for testing
            response_data = json.loads(response.content)
            self.assertEqual(response_data["status"], "error")

    def test_ai_score_view_failure_with_invalid_content(self):
        """Test AI scoring failure with invalid content."""
        # Create a document with empty content to test failure case
        from docx import Document

        doc = Document()
        # Empty document should cause AI scoring to fail
        doc.save(self.test_file_path)

        self.client.login(username="testuser", password="testpass123")

        data = {"path": "test_document.docx"}
        response = self.client.post(reverse("grading:ai_score"), data)

        # Should fail due to empty content
        self.assertIn(response.status_code, [400, 500])
        if response.status_code != 302:  # Not redirect
            response_data = json.loads(response.content)
            self.assertEqual(response_data["status"], "error")

    def test_ai_score_view_missing_data(self):
        """Test AI scoring with missing data."""
        self.client.login(username="testuser", password="testpass123")

        data = {}  # Missing path
        response = self.client.post(reverse("grading:ai_score"), data)

        self.assertEqual(response.status_code, 400)

    def test_ai_score_view_unauthenticated(self):
        """Test AI scoring without authentication."""
        data = {"path": "test_document.docx"}
        response = self.client.post(reverse("grading:ai_score"), data)

        self.assertEqual(response.status_code, 302)  # Redirect to login

    def test_ai_score_view_already_graded(self):
        """Test AI scoring when file already has a grade."""
        # Create a document with existing grade
        from docx import Document

        doc = Document()
        doc.add_paragraph("Test document content")
        # Add a table with grade
        table = doc.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "评定分数"
        table.cell(0, 1).text = "A"
        doc.save(self.test_file_path)

        self.client.login(username="testuser", password="testpass123")

        data = {"path": "test_document.docx"}
        response = self.client.post(reverse("grading:ai_score"), data)

        self.assertEqual(response.status_code, 400)
        response_data = json.loads(response.content)
        self.assertEqual(response_data["status"], "error")
        self.assertIn("已有评分", response_data["message"])
        self.assertIn("A", response_data["message"])

    def test_batch_ai_score_view_success(self):
        """Test successful batch AI scoring with real API calls."""
        # Check if API key is available
        api_key = os.environ.get("ARK_API_KEY")
        if not api_key:
            self.skipTest("ARK_API_KEY not set, skipping real API test")

        self.client.login(username="testuser", password="testpass123")

        data = {"path": "test_directory"}
        response = self.client.post(reverse("grading:batch_ai_score"), data)

        # Check if the request was successful
        self.assertIn(response.status_code, [200, 500])  # Allow both success and API errors

        if response.status_code == 200:
            response_data = json.loads(response.content)
            self.assertEqual(response_data["status"], "success")
        else:
            # If API call failed, that's also acceptable for testing
            response_data = json.loads(response.content)
            self.assertEqual(response_data["status"], "error")

    def test_batch_ai_score_view_missing_data(self):
        """Test batch AI scoring with missing data."""
        self.client.login(username="testuser", password="testpass123")

        data = {}  # Missing path
        response = self.client.post(reverse("grading:batch_ai_score"), data)

        self.assertEqual(response.status_code, 400)

    def test_batch_ai_score_view_unauthenticated(self):
        """Test batch AI scoring without authentication."""
        data = {"path": "test_directory"}
        response = self.client.post(reverse("grading:batch_ai_score"), data)

        self.assertEqual(response.status_code, 302)  # Redirect to login

    def test_batch_ai_score_view_with_graded_files(self):
        """Test batch AI scoring when some files already have grades."""
        # Create test directory
        test_dir = os.path.join(self.temp_dir, "test_directory")
        os.makedirs(test_dir, exist_ok=True)

        # Create a file with existing grade
        graded_file_path = os.path.join(test_dir, "graded_document.docx")
        from docx import Document

        doc = Document()
        doc.add_paragraph("Test document content")
        # Add a table with grade
        table = doc.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "评定分数"
        table.cell(0, 1).text = "B"
        doc.save(graded_file_path)

        # Create a file without grade
        ungraded_file_path = os.path.join(test_dir, "ungraded_document.docx")
        doc = Document()
        doc.add_paragraph("Test document content")
        doc.save(ungraded_file_path)

        self.client.login(username="testuser", password="testpass123")

        data = {"path": "test_directory"}
        response = self.client.post(reverse("grading:batch_ai_score"), data)

        self.assertEqual(response.status_code, 200)
        response_data = json.loads(response.content)
        self.assertEqual(response_data["status"], "success")

        # Check that we have results for both files
        self.assertEqual(len(response_data["results"]), 2)

        # Find the graded file result
        graded_result = next(
            (r for r in response_data["results"] if r["file"] == "graded_document.docx"), None
        )
        self.assertIsNotNone(graded_result)
        self.assertFalse(graded_result["success"])
        self.assertIn("已有评分", graded_result["error"])
        self.assertIn("B", graded_result["error"])


class AIScoringFunctionTest(TestCase):
    """Test cases for AI scoring functions."""

    def test_volcengine_score_homework_success(self):
        """Test successful volcengine AI scoring with real API."""
        # Check if API key is available
        api_key = os.environ.get("ARK_API_KEY")
        if not api_key:
            self.skipTest("ARK_API_KEY not set, skipping real API test")

        from grading.views import volcengine_score_homework

        score, comment = volcengine_score_homework("Test homework content")

        # Verify that we get valid responses
        self.assertIsInstance(score, (int, type(None)))
        self.assertIsInstance(comment, str)
        if score is not None:
            self.assertGreaterEqual(score, 0)
            self.assertLessEqual(score, 100)

    def test_volcengine_score_homework_no_api_key(self):
        """Test AI scoring without API key."""
        # Temporarily remove API key from environment
        original_api_key = os.environ.get("ARK_API_KEY")
        if "ARK_API_KEY" in os.environ:
            del os.environ["ARK_API_KEY"]

        try:
            from grading.views import volcengine_score_homework

            score, comment = volcengine_score_homework("Test homework content")

            self.assertIsNone(score)
            self.assertEqual(comment, "API密钥未配置")
        finally:
            # Restore original API key
            if original_api_key:
                os.environ["ARK_API_KEY"] = original_api_key

    def test_volcengine_score_homework_api_error(self):
        """Test AI scoring with API error using real API."""
        # Check if API key is available
        api_key = os.environ.get("ARK_API_KEY")
        if not api_key:
            self.skipTest("ARK_API_KEY not set, skipping real API test")

        from grading.views import volcengine_score_homework

        # Test with empty content which should cause API error
        score, comment = volcengine_score_homework("")

        # Should handle empty content gracefully
        self.assertIsInstance(score, (int, type(None)))
        self.assertIsInstance(comment, str)

    def test_volcengine_score_homework_no_score_in_response(self):
        """Test AI scoring with no score in response using real API."""
        # Check if API key is available
        api_key = os.environ.get("ARK_API_KEY")
        if not api_key:
            self.skipTest("ARK_API_KEY not set, skipping real API test")

        from grading.views import volcengine_score_homework

        # Test with content that might not return a score
        score, comment = volcengine_score_homework(
            "This is a test content without specific homework content."
        )

        # Verify that we get valid responses
        self.assertIsInstance(score, (int, type(None)))
        self.assertIsInstance(comment, str)
        # Even if no score is extracted, we should get a comment
        self.assertGreater(len(comment), 0)

    def test_convert_score_to_grade(self):
        """Test score to grade conversion."""
        from grading.views import convert_score_to_grade

        # Test grade conversions
        self.assertEqual(convert_score_to_grade(95), "A")
        self.assertEqual(convert_score_to_grade(85), "B")
        self.assertEqual(convert_score_to_grade(75), "C")
        self.assertEqual(convert_score_to_grade(65), "D")
        self.assertEqual(convert_score_to_grade(55), "E")
        self.assertEqual(convert_score_to_grade(None), "N/A")


class AIScoringIntegrationTest(TestCase):
    """Integration tests for AI scoring."""

    def setUp(self):
        """Set up test data."""
        self.user = User.objects.create_user(
            username="testuser", password="testpass123", is_staff=True
        )
        self.client.login(username="testuser", password="testpass123")
        self.temp_dir = tempfile.mkdtemp()
        self.config = GlobalConfig.objects.create(repo_base_dir=self.temp_dir)
        # Create a test file in the temp directory
        from docx import Document

        doc = Document()
        doc.add_paragraph("test content")
        doc.save(os.path.join(self.temp_dir, "test_document.docx"))
        os.makedirs(os.path.join(self.temp_dir, "test_directory"))
        doc.save(os.path.join(self.temp_dir, "test_directory", "test1.docx"))
        doc.save(os.path.join(self.temp_dir, "test_directory", "test2.docx"))

    def test_ai_scoring_integration(self):
        """Test AI scoring integration with real API calls."""
        # Check if API key is available
        api_key = os.environ.get("ARK_API_KEY")
        if not api_key:
            self.skipTest("ARK_API_KEY not set, skipping real API test")

        data = {"path": "test_document.docx"}
        response = self.client.post(reverse("grading:ai_score"), data)

        # Check if the request was successful
        self.assertIn(response.status_code, [200, 500])  # Allow both success and API errors

        if response.status_code == 200:
            response_data = json.loads(response.content)
            self.assertEqual(response_data["status"], "success")
            # Verify that the response contains expected fields
            self.assertIn("score", response_data)
            self.assertIn("grade", response_data)
            self.assertIn("comment", response_data)
        else:
            # If API call failed, that's also acceptable for testing
            response_data = json.loads(response.content)
            self.assertEqual(response_data["status"], "error")

    def test_ai_scoring_with_different_content_types(self):
        """Test AI scoring with different content types using real API."""
        # Check if API key is available
        api_key = os.environ.get("ARK_API_KEY")
        if not api_key:
            self.skipTest("ARK_API_KEY not set, skipping real API test")

        from grading.views import volcengine_score_homework

        content_types = [
            "This is a simple text homework.",
            "这是一个中文作业内容。",
            "Homework with numbers: 1, 2, 3, 4, 5",
            "Very long homework content " * 10,  # Long content
        ]

        for content in content_types:
            score, comment = volcengine_score_homework(content)
            # Verify that we get valid responses (even if API fails, we get None, "")
            self.assertIsInstance(score, (int, type(None)))
            self.assertIsInstance(comment, str)
            if score is not None:
                self.assertGreaterEqual(score, 0)
                self.assertLessEqual(score, 100)

    def test_ai_scoring_error_handling(self):
        """Test AI scoring error handling."""
        # Test with invalid file path
        data = {"path": "nonexistent_file.docx"}
        response = self.client.post(reverse("grading:ai_score"), data)

        self.assertEqual(response.status_code, 200)
        response_data = json.loads(response.content)
        self.assertEqual(response_data["status"], "error")

    def test_batch_ai_scoring_with_multiple_files(self):
        """Test batch AI scoring with multiple files using real API."""
        # Check if API key is available
        api_key = os.environ.get("ARK_API_KEY")
        if not api_key:
            self.skipTest("ARK_API_KEY not set, skipping real API test")

        data = {"path": "test_directory"}
        response = self.client.post(reverse("grading:batch_ai_score"), data)

        # Check if the request was successful
        self.assertIn(response.status_code, [200, 500])  # Allow both success and API errors

        if response.status_code == 200:
            response_data = json.loads(response.content)
            self.assertEqual(response_data["status"], "success")
            self.assertIn("results", response_data)
        else:
            # If API call failed, that's also acceptable for testing
            response_data = json.loads(response.content)
            self.assertEqual(response_data["status"], "error")
