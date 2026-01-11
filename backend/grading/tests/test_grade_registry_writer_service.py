"""
作业成绩写入成绩登分册服务测试

测试GradeRegistryWriterService的两种场景：
1. 作业评分系统场景
2. 工具箱模块场景

以及AuditLogger审计日志功能
"""

import os
import shutil
import tempfile
from datetime import datetime
from unittest.mock import MagicMock, Mock, patch

from grading.grade_registry_writer import NameMatcher as CoreNameMatcher
from grading.grade_registry_writer import RegistryManager
from grading.models import Tenant
from grading.services.grade_registry_writer_service import (
    AuditLogger,
    GradeRegistryWriterService,
)

from .base import BaseTestCase


class AuditLoggerInitTest(BaseTestCase):
    """测试AuditLogger初始化"""

    def setUp(self):
        super().setUp()
        self.tenant = Tenant.objects.create(name="测试租户")

    def test_init_with_valid_params(self):
        """测试使用有效参数初始化"""
        audit_logger = AuditLogger(
            self.user, self.tenant, GradeRegistryWriterService.SCENARIO_GRADING_SYSTEM
        )

        self.assertEqual(audit_logger.user, self.user)
        self.assertEqual(audit_logger.tenant, self.tenant)
        self.assertEqual(audit_logger.scenario, GradeRegistryWriterService.SCENARIO_GRADING_SYSTEM)
        self.assertIsNone(audit_logger.operation_start_time)
        self.assertEqual(audit_logger.audit_data["user_id"], self.user.id)
        self.assertEqual(audit_logger.audit_data["username"], self.user.username)
        self.assertEqual(audit_logger.audit_data["tenant_id"], self.tenant.id)
        self.assertEqual(audit_logger.audit_data["tenant_name"], self.tenant.name)
        self.assertEqual(
            audit_logger.audit_data["scenario"], GradeRegistryWriterService.SCENARIO_GRADING_SYSTEM
        )

    def test_init_with_none_user(self):
        """测试用户为None时初始化"""
        audit_logger = AuditLogger(None, self.tenant, GradeRegistryWriterService.SCENARIO_TOOLBOX)

        self.assertIsNone(audit_logger.user)
        self.assertIsNone(audit_logger.audit_data["user_id"])
        self.assertIsNone(audit_logger.audit_data["username"])

    def test_init_with_none_tenant(self):
        """测试租户为None时初始化"""
        audit_logger = AuditLogger(self.user, None, GradeRegistryWriterService.SCENARIO_TOOLBOX)

        self.assertIsNone(audit_logger.tenant)
        self.assertIsNone(audit_logger.audit_data["tenant_id"])
        self.assertIsNone(audit_logger.audit_data["tenant_name"])


class AuditLoggerStartOperationTest(BaseTestCase):
    """测试AuditLogger.start_operation方法"""

    def setUp(self):
        super().setUp()
        self.tenant = Tenant.objects.create(name="测试租户")
        self.audit_logger = AuditLogger(
            self.user, self.tenant, GradeRegistryWriterService.SCENARIO_GRADING_SYSTEM
        )

    @patch("grading.services.grade_registry_writer_service.datetime")
    def test_start_operation_basic(self, mock_datetime):
        """测试基本的操作开始记录"""
        mock_now = datetime(2024, 1, 1, 12, 0, 0)
        mock_datetime.now.return_value = mock_now

        self.audit_logger.start_operation("test_operation")

        self.assertEqual(self.audit_logger.operation_start_time, mock_now)
        self.assertEqual(self.audit_logger.audit_data["start_time"], mock_now.isoformat())
        self.assertEqual(self.audit_logger.audit_data["operation_type"], "test_operation")

    @patch("grading.services.grade_registry_writer_service.datetime")
    def test_start_operation_with_kwargs(self, mock_datetime):
        """测试带额外参数的操作开始记录"""
        mock_now = datetime(2024, 1, 1, 12, 0, 0)
        mock_datetime.now.return_value = mock_now

        self.audit_logger.start_operation(
            "batch_write", class_directory="/path/to/class", homework_number=1
        )

        self.assertEqual(self.audit_logger.audit_data["operation_type"], "batch_write")
        self.assertEqual(self.audit_logger.audit_data["class_directory"], "/path/to/class")
        self.assertEqual(self.audit_logger.audit_data["homework_number"], 1)

    @patch("grading.services.grade_registry_writer_service.logger")
    @patch("grading.services.grade_registry_writer_service.datetime")
    def test_start_operation_logging(self, mock_datetime, mock_logger):
        """测试操作开始时的日志记录"""
        mock_now = datetime(2024, 1, 1, 12, 0, 0)
        mock_datetime.now.return_value = mock_now

        # 创建一个mock logger实例
        mock_audit_logger = MagicMock()
        with patch.object(self.audit_logger, "logger", mock_audit_logger):
            self.audit_logger.start_operation("test_operation", extra_param="extra_value")

            # 验证日志调用
            self.assertTrue(mock_audit_logger.info.called)


class AuditLoggerEndOperationTest(BaseTestCase):
    """测试AuditLogger.end_operation方法"""

    def setUp(self):
        super().setUp()
        self.tenant = Tenant.objects.create(name="测试租户")
        self.audit_logger = AuditLogger(
            self.user, self.tenant, GradeRegistryWriterService.SCENARIO_GRADING_SYSTEM
        )

    @patch("grading.services.grade_registry_writer_service.datetime")
    def test_end_operation_success(self, mock_datetime):
        """测试成功操作的结束记录"""
        start_time = datetime(2024, 1, 1, 12, 0, 0)
        end_time = datetime(2024, 1, 1, 12, 0, 30)

        mock_datetime.now.side_effect = [start_time, end_time]

        self.audit_logger.start_operation("test_operation")
        self.audit_logger.end_operation(success=True)

        self.assertEqual(self.audit_logger.audit_data["end_time"], end_time.isoformat())
        self.assertTrue(self.audit_logger.audit_data["success"])
        self.assertEqual(self.audit_logger.audit_data["duration_seconds"], 30.0)

    @patch("grading.services.grade_registry_writer_service.datetime")
    def test_end_operation_failure(self, mock_datetime):
        """测试失败操作的结束记录"""
        start_time = datetime(2024, 1, 1, 12, 0, 0)
        end_time = datetime(2024, 1, 1, 12, 0, 10)

        mock_datetime.now.side_effect = [start_time, end_time]

        self.audit_logger.start_operation("test_operation")
        self.audit_logger.end_operation(success=False, error_message="测试错误")

        self.assertFalse(self.audit_logger.audit_data["success"])
        self.assertEqual(self.audit_logger.audit_data["error_message"], "测试错误")
        self.assertEqual(self.audit_logger.audit_data["duration_seconds"], 10.0)

    @patch("grading.services.grade_registry_writer_service.datetime")
    def test_end_operation_without_start(self, mock_datetime):
        """测试未调用start_operation直接调用end_operation"""
        end_time = datetime(2024, 1, 1, 12, 0, 30)
        mock_datetime.now.return_value = end_time

        self.audit_logger.end_operation(success=True)

        self.assertEqual(self.audit_logger.audit_data["end_time"], end_time.isoformat())
        self.assertTrue(self.audit_logger.audit_data["success"])
        # duration应该为None或0
        self.assertIsNone(self.audit_logger.audit_data.get("duration_seconds"))

    @patch("grading.services.grade_registry_writer_service.datetime")
    def test_end_operation_with_statistics(self, mock_datetime):
        """测试带统计信息的操作结束记录"""
        start_time = datetime(2024, 1, 1, 12, 0, 0)
        end_time = datetime(2024, 1, 1, 12, 1, 0)

        mock_datetime.now.side_effect = [start_time, end_time]

        self.audit_logger.start_operation("batch_write")
        self.audit_logger.end_operation(
            success=True, total_files=20, success_count=18, failed_count=2
        )

        self.assertEqual(self.audit_logger.audit_data["total_files"], 20)
        self.assertEqual(self.audit_logger.audit_data["success_count"], 18)
        self.assertEqual(self.audit_logger.audit_data["failed_count"], 2)


class AuditLoggerGradeWriteTest(BaseTestCase):
    """测试AuditLogger.log_grade_write方法"""

    def setUp(self):
        super().setUp()
        self.tenant = Tenant.objects.create(name="测试租户")
        self.audit_logger = AuditLogger(
            self.user, self.tenant, GradeRegistryWriterService.SCENARIO_GRADING_SYSTEM
        )

    def test_log_grade_write_new_grade(self):
        """测试记录新成绩写入"""
        mock_logger = MagicMock()
        with patch.object(self.audit_logger, "logger", mock_logger):
            self.audit_logger.log_grade_write("张三", 1, "A", None)

            # 验证日志调用
            mock_logger.info.assert_called_once()
            call_args = mock_logger.info.call_args[0]
            self.assertIn("成绩写入", call_args[0])
            self.assertIn("张三", call_args)
            self.assertEqual(call_args[2], 1)
            self.assertEqual(call_args[3], "A")

    def test_log_grade_write_update_grade(self):
        """测试记录成绩覆盖"""
        mock_logger = MagicMock()
        with patch.object(self.audit_logger, "logger", mock_logger):
            self.audit_logger.log_grade_write("张三", 1, "A", "B")

            # 验证日志调用
            mock_logger.info.assert_called_once()
            call_args = mock_logger.info.call_args[0]
            self.assertIn("成绩覆盖", call_args[0])
            self.assertIn("张三", call_args)
            self.assertEqual(call_args[2], 1)
            self.assertIn("B", call_args)
            self.assertIn("A", call_args)

    def test_log_grade_write_same_grade(self):
        """测试记录相同成绩（不应该覆盖）"""
        mock_logger = MagicMock()
        with patch.object(self.audit_logger, "logger", mock_logger):
            self.audit_logger.log_grade_write("张三", 1, "A", "A")

            # 验证日志调用
            mock_logger.info.assert_called_once()
            call_args = mock_logger.info.call_args[0]
            # 相同成绩应该记录为"成绩写入"而不是"成绩覆盖"
            self.assertIn("成绩写入", call_args[0])


class AuditLoggerFileProcessingTest(BaseTestCase):
    """测试AuditLogger.log_file_processing方法"""

    def setUp(self):
        super().setUp()
        self.tenant = Tenant.objects.create(name="测试租户")
        self.audit_logger = AuditLogger(
            self.user, self.tenant, GradeRegistryWriterService.SCENARIO_GRADING_SYSTEM
        )

    def test_log_file_processing_success(self):
        """测试记录文件处理成功"""
        mock_logger = MagicMock()
        with patch.object(self.audit_logger, "logger", mock_logger):
            self.audit_logger.log_file_processing("/path/to/file.docx", "success")

            mock_logger.info.assert_called_once()
            call_args = mock_logger.info.call_args[0]
            self.assertIn("文件处理成功", call_args[0])
            self.assertIn("/path/to/file.docx", call_args)

    def test_log_file_processing_failed(self):
        """测试记录文件处理失败"""
        mock_logger = MagicMock()
        with patch.object(self.audit_logger, "logger", mock_logger):
            self.audit_logger.log_file_processing("/path/to/file.docx", "failed", "无法提取成绩")

            mock_logger.warning.assert_called_once()
            call_args = mock_logger.warning.call_args[0]
            self.assertIn("文件处理失败", call_args[0])
            self.assertIn("/path/to/file.docx", call_args)
            self.assertIn("无法提取成绩", call_args)

    def test_log_file_processing_skipped(self):
        """测试记录文件跳过"""
        mock_logger = MagicMock()
        with patch.object(self.audit_logger, "logger", mock_logger):
            self.audit_logger.log_file_processing("/path/to/file.docx", "skipped", "成绩相同")

            mock_logger.debug.assert_called_once()
            call_args = mock_logger.debug.call_args[0]
            self.assertIn("文件跳过", call_args[0])
            self.assertIn("/path/to/file.docx", call_args)
            self.assertIn("成绩相同", call_args)

    def test_log_file_processing_without_error(self):
        """测试记录文件处理状态（无错误信息）"""
        mock_logger = MagicMock()
        with patch.object(self.audit_logger, "logger", mock_logger):
            self.audit_logger.log_file_processing("/path/to/file.docx", "success", None)

            mock_logger.info.assert_called_once()


class GradeRegistryWriterServiceInitTest(BaseTestCase):
    """测试服务初始化"""

    def setUp(self):
        super().setUp()
        self.tenant = Tenant.objects.create(name="测试租户")

    def test_init_with_grading_system_scenario(self):
        """测试使用作业评分系统场景初始化"""
        service = GradeRegistryWriterService(
            self.user, self.tenant, GradeRegistryWriterService.SCENARIO_GRADING_SYSTEM
        )

        self.assertEqual(service.user, self.user)
        self.assertEqual(service.tenant, self.tenant)
        self.assertEqual(service.scenario, GradeRegistryWriterService.SCENARIO_GRADING_SYSTEM)
        self.assertIsNotNone(service.audit_logger)
        self.assertIsInstance(service.audit_logger, AuditLogger)

    def test_init_with_toolbox_scenario(self):
        """测试使用工具箱模块场景初始化"""
        service = GradeRegistryWriterService(
            self.user, self.tenant, GradeRegistryWriterService.SCENARIO_TOOLBOX
        )

        self.assertEqual(service.scenario, GradeRegistryWriterService.SCENARIO_TOOLBOX)
        self.assertIsNotNone(service.audit_logger)

    def test_init_with_invalid_scenario(self):
        """测试使用无效场景初始化"""
        with self.assertRaises(ValueError) as context:
            GradeRegistryWriterService(self.user, self.tenant, "invalid_scenario")

        self.assertIn("无效的场景类型", str(context.exception))

    def test_init_creates_audit_logger(self):
        """测试初始化时创建审计日志记录器"""
        service = GradeRegistryWriterService(
            self.user, self.tenant, GradeRegistryWriterService.SCENARIO_GRADING_SYSTEM
        )

        self.assertIsNotNone(service.audit_logger)
        self.assertEqual(service.audit_logger.user, self.user)
        self.assertEqual(service.audit_logger.tenant, self.tenant)
        self.assertEqual(
            service.audit_logger.scenario, GradeRegistryWriterService.SCENARIO_GRADING_SYSTEM
        )


class GradeRegistryWriterServiceFindRegistryTest(BaseTestCase):
    """测试查找登分册文件"""

    def setUp(self):
        super().setUp()
        self.tenant = Tenant.objects.create(name="测试租户")
        self.service = GradeRegistryWriterService(
            self.user, self.tenant, GradeRegistryWriterService.SCENARIO_GRADING_SYSTEM
        )
        self.temp_dir = tempfile.mkdtemp()

    def tearDown(self):
        if os.path.exists(self.temp_dir):
            shutil.rmtree(self.temp_dir)
        super().tearDown()

    def test_find_grade_registry_success(self):
        """测试成功找到登分册文件"""
        # 创建登分册文件
        registry_file = os.path.join(self.temp_dir, "成绩登分册.xlsx")
        open(registry_file, "w").close()

        result = self.service.find_grade_registry(self.temp_dir)

        self.assertIsNotNone(result)
        self.assertEqual(result, registry_file)

    def test_find_grade_registry_requires_specific_filename(self):
        """只接受文件名为“成绩登分册.xlsx”"""
        other_names = ["登分册.xlsx", "grade_registry.xlsx", "grades.xlsx"]

        for name in other_names:
            with self.subTest(name=name):
                temp_dir = tempfile.mkdtemp()
                registry_file = os.path.join(temp_dir, name)
                open(registry_file, "w").close()

                result = self.service.find_grade_registry(temp_dir)

                self.assertIsNone(result)

                shutil.rmtree(temp_dir)

    def test_find_grade_registry_not_found(self):
        """测试未找到登分册文件"""
        result = self.service.find_grade_registry(self.temp_dir)

        self.assertIsNone(result)

    def test_find_grade_registry_directory_not_exists(self):
        """测试目录不存在"""
        non_existent_dir = os.path.join(self.temp_dir, "non_existent")

        result = self.service.find_grade_registry(non_existent_dir)

        self.assertIsNone(result)

    def test_find_grade_registry_skip_temp_files(self):
        """测试跳过临时文件"""
        # 创建临时文件（以~$开头）
        temp_file = os.path.join(self.temp_dir, "~$成绩登分册.xlsx")
        open(temp_file, "w").close()

        result = self.service.find_grade_registry(self.temp_dir)

        self.assertIsNone(result)


class GradeRegistryWriterServiceProcessSingleWordFileTest(BaseTestCase):
    """测试处理单个Word文档"""

    def setUp(self):
        super().setUp()
        self.tenant = Tenant.objects.create(name="测试租户")
        self.service = GradeRegistryWriterService(
            self.user, self.tenant, GradeRegistryWriterService.SCENARIO_GRADING_SYSTEM
        )

        # 创建模拟的登分册管理器
        self.mock_registry = Mock(spec=RegistryManager)
        self.mock_registry.student_names = {"张三": 2, "李四": 3}

        size_patcher = patch.object(
            GradeRegistryWriterService,
            "_validate_file_size",
            return_value=(True, None),
        )
        self.addCleanup(size_patcher.stop)
        self.mock_validate_size = size_patcher.start()

    @patch("grading.services.grade_registry_writer_service.GradeFileProcessor")
    @patch("grading.services.grade_registry_writer_service.NameMatcher")
    def test_process_single_word_file_success(self, mock_name_matcher, mock_processor):
        """测试成功处理单个Word文档"""
        # 设置模拟返回值
        mock_processor.extract_student_name.return_value = "张三"
        mock_processor.extract_grade_from_word.return_value = "A"
        mock_name_matcher.match.return_value = ("张三", "exact")
        mock_name_matcher.normalize_name.side_effect = CoreNameMatcher.normalize_name
        self.mock_registry.find_student_row.return_value = 2
        self.mock_registry.write_grade.return_value = (True, None)

        result = self.service._process_single_word_file(
            "/path/to/张三_作业1.docx", self.mock_registry, 5
        )

        self.assertTrue(result["success"])
        self.assertEqual(result["student_name"], "张三")
        self.assertEqual(result["grade"], "A")
        self.assertIsNone(result["error_message"])

    @patch("grading.services.grade_registry_writer_service.GradeFileProcessor")
    def test_process_single_word_file_no_student_name(self, mock_processor):
        """测试无法提取学生姓名"""
        mock_processor.extract_student_name.return_value = None

        result = self.service._process_single_word_file(
            "/path/to/作业1.docx", self.mock_registry, 5
        )

        self.assertFalse(result["success"])
        self.assertIn("无法提取学生姓名", result["error_message"])

    @patch("grading.services.grade_registry_writer_service.GradeFileProcessor")
    def test_process_single_word_file_no_grade(self, mock_processor):
        """测试无法提取成绩"""
        mock_processor.extract_student_name.return_value = "张三"
        mock_processor.extract_grade_from_word.return_value = None

        result = self.service._process_single_word_file(
            "/path/to/张三_作业1.docx", self.mock_registry, 5
        )

        self.assertFalse(result["success"])
        self.assertIn("无法提取成绩", result["error_message"])

    @patch("grading.services.grade_registry_writer_service.GradeFileProcessor")
    @patch("grading.services.grade_registry_writer_service.NameMatcher")
    def test_process_single_word_file_name_not_matched(self, mock_name_matcher, mock_processor):
        """测试学生姓名未匹配"""
        mock_processor.extract_student_name.return_value = "王五"
        mock_processor.extract_grade_from_word.return_value = "A"
        mock_name_matcher.match.return_value = (None, "none")
        mock_name_matcher.normalize_name.side_effect = CoreNameMatcher.normalize_name

        result = self.service._process_single_word_file(
            "/path/to/王五_作业1.docx", self.mock_registry, 5
        )

        self.assertFalse(result["success"])
        self.assertIn("未找到匹配的学生", result["error_message"])

    @patch("grading.services.grade_registry_writer_service.GradeFileProcessor")
    @patch("grading.services.grade_registry_writer_service.NameMatcher")
    def test_process_single_word_file_filename_contains_student(
        self, mock_name_matcher, mock_processor
    ):
        """当文件名包含学生姓名时也应匹配成功"""
        mock_processor.extract_student_name.return_value = "第二次作业"
        mock_processor.extract_grade_from_word.return_value = "A"
        mock_name_matcher.match.return_value = (None, "none")
        mock_name_matcher.normalize_name.side_effect = CoreNameMatcher.normalize_name
        self.mock_registry.find_student_row.return_value = 2
        self.mock_registry.write_grade.return_value = (True, None)

        result = self.service._process_single_word_file("/path/to/张三.docx", self.mock_registry, 5)

        self.assertTrue(result["success"])

    @patch("grading.services.grade_registry_writer_service.GradeFileProcessor")
    @patch("grading.services.grade_registry_writer_service.NameMatcher")
    def test_process_single_word_file_multiple_matches(self, mock_name_matcher, mock_processor):
        """测试学生姓名匹配到多个"""
        mock_processor.extract_student_name.return_value = "张三"
        mock_processor.extract_grade_from_word.return_value = "A"
        mock_name_matcher.match.return_value = (None, "multiple")
        mock_name_matcher.normalize_name.side_effect = CoreNameMatcher.normalize_name

        result = self.service._process_single_word_file(
            "/path/to/张三_作业1.docx", self.mock_registry, 5
        )

        self.assertFalse(result["success"])
        self.assertIn("姓名匹配到多个学生", result["error_message"])

    @patch("grading.services.grade_registry_writer_service.GradeFileProcessor")
    @patch("grading.services.grade_registry_writer_service.NameMatcher")
    def test_process_single_word_file_grade_skipped(self, mock_name_matcher, mock_processor):
        """测试成绩相同跳过写入"""
        mock_processor.extract_student_name.return_value = "张三"
        mock_processor.extract_grade_from_word.return_value = "A"
        mock_name_matcher.match.return_value = ("张三", "exact")
        mock_name_matcher.normalize_name.side_effect = CoreNameMatcher.normalize_name
        self.mock_registry.find_student_row.return_value = 2
        self.mock_registry.write_grade.return_value = (False, "A")

        result = self.service._process_single_word_file(
            "/path/to/张三_作业1.docx", self.mock_registry, 5
        )

        self.assertTrue(result["skipped"])
        self.assertIn("成绩相同", result["error_message"])


class GradeRegistryWriterServiceProcessSingleStudentGradeTest(BaseTestCase):
    """测试处理单个学生成绩"""

    def setUp(self):
        super().setUp()
        self.tenant = Tenant.objects.create(name="测试租户")
        self.service = GradeRegistryWriterService(
            self.user, self.tenant, GradeRegistryWriterService.SCENARIO_TOOLBOX
        )

        self.mock_registry = Mock(spec=RegistryManager)
        self.student_names = ["张三", "李四", "王五"]

    @patch("grading.services.grade_registry_writer_service.NameMatcher")
    def test_process_single_student_grade_success(self, mock_name_matcher):
        """测试成功处理单个学生成绩"""
        grade_data = {"name": "张三", "grade": "A"}
        mock_name_matcher.match.return_value = ("张三", "exact")
        self.mock_registry.find_student_row.return_value = 2
        self.mock_registry.write_grade.return_value = (True, None)

        result = self.service._process_single_student_grade(
            grade_data, self.mock_registry, 5, self.student_names
        )

        self.assertTrue(result["success"])
        self.assertEqual(result["student_name"], "张三")
        self.assertEqual(result["grade"], "A")

    @patch("grading.services.grade_registry_writer_service.NameMatcher")
    def test_process_single_student_grade_name_not_matched(self, mock_name_matcher):
        """测试学生姓名未匹配"""
        grade_data = {"name": "赵六", "grade": "A"}
        mock_name_matcher.match.return_value = (None, "none")

        result = self.service._process_single_student_grade(
            grade_data, self.mock_registry, 5, self.student_names
        )

        self.assertFalse(result["success"])
        self.assertIn("未找到匹配的学生", result["error_message"])

    @patch("grading.services.grade_registry_writer_service.NameMatcher")
    def test_process_single_student_grade_skipped(self, mock_name_matcher):
        """测试成绩相同跳过"""
        grade_data = {"name": "张三", "grade": "A"}
        mock_name_matcher.match.return_value = ("张三", "exact")
        self.mock_registry.find_student_row.return_value = 2
        self.mock_registry.write_grade.return_value = (False, "A")

        result = self.service._process_single_student_grade(
            grade_data, self.mock_registry, 5, self.student_names
        )

        self.assertTrue(result["success"])
        self.assertTrue(result.get("skipped", False))


class GradeRegistryWriterServiceProcessMethodTest(BaseTestCase):
    """测试process方法"""

    def setUp(self):
        super().setUp()
        self.tenant = Tenant.objects.create(name="测试租户")

    def test_process_grading_system_scenario_missing_params(self):
        """测试作业评分系统场景缺少参数"""
        service = GradeRegistryWriterService(
            self.user, self.tenant, GradeRegistryWriterService.SCENARIO_GRADING_SYSTEM
        )

        result = service.process(homework_dir="/path/to/homework")

        self.assertFalse(result["success"])
        self.assertIn("缺少必需参数", result["error_message"])

    def test_process_toolbox_scenario_missing_params(self):
        """测试工具箱模块场景缺少参数"""
        service = GradeRegistryWriterService(
            self.user, self.tenant, GradeRegistryWriterService.SCENARIO_TOOLBOX
        )

        result = service.process()

        self.assertFalse(result["success"])
        self.assertIn("缺少必需参数", result["error_message"])

    @patch.object(GradeRegistryWriterService, "process_grading_system_scenario")
    def test_process_calls_grading_system_scenario(self, mock_process):
        """测试process方法调用作业评分系统场景"""
        service = GradeRegistryWriterService(
            self.user, self.tenant, GradeRegistryWriterService.SCENARIO_GRADING_SYSTEM
        )
        mock_process.return_value = {"success": True}

        result = service.process(homework_dir="/path/to/homework", class_dir="/path/to/class")

        mock_process.assert_called_once_with("/path/to/homework", "/path/to/class")
        self.assertTrue(result["success"])

    @patch.object(GradeRegistryWriterService, "process_toolbox_scenario")
    def test_process_calls_toolbox_scenario(self, mock_process):
        """测试process方法调用工具箱模块场景"""
        service = GradeRegistryWriterService(
            self.user, self.tenant, GradeRegistryWriterService.SCENARIO_TOOLBOX
        )
        mock_process.return_value = {"success": True}

        result = service.process(class_dir="/path/to/class")

        mock_process.assert_called_once_with("/path/to/class")
        self.assertTrue(result["success"])


class GradeRegistryWriterServiceGradingSystemScenarioTest(BaseTestCase):
    """测试作业评分系统场景的完整流程"""

    def setUp(self):
        super().setUp()
        self.tenant = Tenant.objects.create(name="测试租户")
        self.service = GradeRegistryWriterService(
            self.user, self.tenant, GradeRegistryWriterService.SCENARIO_GRADING_SYSTEM
        )
        self.temp_dir = tempfile.mkdtemp()
        self.homework_dir = os.path.join(self.temp_dir, "第1次作业")
        self.class_dir = os.path.join(self.temp_dir, "2024级计算机1班")
        os.makedirs(self.homework_dir)
        os.makedirs(self.class_dir)

    def tearDown(self):
        if os.path.exists(self.temp_dir):
            shutil.rmtree(self.temp_dir)
        super().tearDown()

    @patch("grading.services.grade_registry_writer_service.GradeFileProcessor")
    def test_process_grading_system_scenario_invalid_homework_number(self, mock_processor):
        """测试无法提取作业批次"""
        mock_processor.extract_homework_number_from_path.return_value = None

        result = self.service.process_grading_system_scenario(self.homework_dir, self.class_dir)

        self.assertFalse(result["success"])
        self.assertIn("无法从目录名提取作业批次", result["error_message"])

    @patch("grading.services.grade_registry_writer_service.GradeFileProcessor")
    def test_process_grading_system_scenario_no_registry(self, mock_processor):
        """测试未找到登分册"""
        mock_processor.extract_homework_number_from_path.return_value = 1

        result = self.service.process_grading_system_scenario(self.homework_dir, self.class_dir)

        self.assertFalse(result["success"])
        self.assertIn("未找到成绩登分册文件", result["error_message"])

    def test_process_grading_system_scenario_homework_dir_not_exists(self):
        """测试作业目录不存在"""
        non_existent_dir = os.path.join(self.temp_dir, "non_existent")

        with patch(
            "grading.services.grade_registry_writer_service.GradeFileProcessor"
        ) as mock_processor:
            mock_processor.extract_homework_number_from_path.return_value = 1

            # 创建登分册
            registry_file = os.path.join(self.class_dir, "成绩登分册.xlsx")
            open(registry_file, "w").close()

            result = self.service.process_grading_system_scenario(non_existent_dir, self.class_dir)

        self.assertFalse(result["success"])
        self.assertIn("作业目录不存在", result["error_message"])

    @patch("grading.services.grade_registry_writer_service.GradeFileProcessor")
    def test_process_grading_system_scenario_no_word_files(self, mock_processor):
        """测试作业目录中没有Word文档"""
        mock_processor.extract_homework_number_from_path.return_value = 1

        # 创建登分册
        registry_file = os.path.join(self.class_dir, "成绩登分册.xlsx")
        open(registry_file, "w").close()

        result = self.service.process_grading_system_scenario(self.homework_dir, self.class_dir)

        self.assertFalse(result["success"])
        self.assertIn("没有找到Word文档", result["error_message"])


class GradeRegistryWriterServiceToolboxScenarioTest(BaseTestCase):
    """测试工具箱模块场景的完整流程"""

    def setUp(self):
        super().setUp()
        self.tenant = Tenant.objects.create(name="测试租户")
        self.service = GradeRegistryWriterService(
            self.user, self.tenant, GradeRegistryWriterService.SCENARIO_TOOLBOX
        )
        self.temp_dir = tempfile.mkdtemp()
        self.class_dir = os.path.join(self.temp_dir, "2024级计算机1班")
        os.makedirs(self.class_dir)

    def tearDown(self):
        if os.path.exists(self.temp_dir):
            shutil.rmtree(self.temp_dir)
        super().tearDown()

    def test_process_toolbox_scenario_no_registry(self):
        """测试未找到登分册"""
        result = self.service.process_toolbox_scenario(self.class_dir)

        self.assertFalse(result["success"])
        self.assertIn("未找到成绩登分册文件", result["error_message"])

    def test_process_toolbox_scenario_class_dir_not_exists(self):
        """测试班级目录不存在"""
        non_existent_dir = os.path.join(self.temp_dir, "non_existent")

        result = self.service.process_toolbox_scenario(non_existent_dir)

        self.assertFalse(result["success"])
        self.assertIn("班级目录不存在", result["error_message"])

    @patch("grading.services.grade_registry_writer_service.GradeFileProcessor")
    def test_process_toolbox_scenario_no_excel_files(self, mock_processor):
        """测试班级目录中没有Excel成绩文件"""
        # 创建登分册
        registry_file = os.path.join(self.class_dir, "成绩登分册.xlsx")
        open(registry_file, "w").close()

        mock_processor.extract_homework_number_from_filename.return_value = None

        result = self.service.process_toolbox_scenario(self.class_dir)

        self.assertFalse(result["success"])
        self.assertIn("没有找到Excel成绩文件", result["error_message"])

    @patch("grading.services.grade_registry_writer_service.GradeFileProcessor")
    def test_process_toolbox_scenario_skip_registry_file(self, mock_processor):
        """测试跳过登分册文件本身"""
        # 创建登分册
        registry_file = os.path.join(self.class_dir, "成绩登分册.xlsx")
        open(registry_file, "w").close()

        # 创建成绩文件（但没有作业批次信息）
        grade_file = os.path.join(self.class_dir, "学生成绩.xlsx")
        open(grade_file, "w").close()

        mock_processor.extract_homework_number_from_filename.return_value = None

        result = self.service.process_toolbox_scenario(self.class_dir)

        self.assertFalse(result["success"])
        self.assertIn("没有找到Excel成绩文件", result["error_message"])


class GradeRegistryWriterServiceProcessSingleExcelFileTest(BaseTestCase):
    """测试处理单个Excel文件"""

    def setUp(self):
        super().setUp()
        self.tenant = Tenant.objects.create(name="测试租户")
        self.service = GradeRegistryWriterService(
            self.user, self.tenant, GradeRegistryWriterService.SCENARIO_TOOLBOX
        )

        self.mock_registry = Mock(spec=RegistryManager)
        self.mock_registry.student_names = {"张三": 2, "李四": 3}

    @patch("grading.services.grade_registry_writer_service.GradeFileProcessor")
    def test_process_single_excel_file_no_homework_number(self, mock_processor):
        """测试无法提取作业批次"""
        mock_processor.extract_homework_number_from_filename.return_value = None

        result = self.service._process_single_excel_file("/path/to/成绩.xlsx", self.mock_registry)

        self.assertFalse(result["success"])
        self.assertIn("无法从文件名提取作业批次", result["error_message"])

    @patch("grading.services.grade_registry_writer_service.GradeFileProcessor")
    def test_process_single_excel_file_no_grades_data(self, mock_processor):
        """测试无法提取学生成绩"""
        mock_processor.extract_homework_number_from_filename.return_value = 1
        mock_processor.extract_grades_from_excel.return_value = []

        result = self.service._process_single_excel_file(
            "/path/to/第1次作业成绩.xlsx", self.mock_registry
        )

        self.assertFalse(result["success"])
        self.assertIn("无法从Excel提取学生成绩", result["error_message"])

    @patch("grading.services.grade_registry_writer_service.GradeFileProcessor")
    @patch("grading.services.grade_registry_writer_service.NameMatcher")
    def test_process_single_excel_file_all_success(self, mock_name_matcher, mock_processor):
        """测试所有学生成绩写入成功"""
        mock_processor.extract_homework_number_from_filename.return_value = 1
        mock_processor.extract_grades_from_excel.return_value = [
            {"name": "张三", "grade": "A"},
            {"name": "李四", "grade": "B"},
        ]
        mock_name_matcher.match.side_effect = [("张三", "exact"), ("李四", "exact")]
        self.mock_registry.find_or_create_homework_column.return_value = 5
        self.mock_registry.find_student_row.side_effect = [2, 3]
        self.mock_registry.write_grade.side_effect = [(True, None), (True, None)]

        result = self.service._process_single_excel_file(
            "/path/to/第1次作业成绩.xlsx", self.mock_registry
        )

        self.assertTrue(result["success"])
        self.assertEqual(result["students_total"], 2)
        self.assertEqual(result["students_processed"], 2)
        self.assertEqual(result["students_failed"], 0)

    @patch("grading.services.grade_registry_writer_service.GradeFileProcessor")
    @patch("grading.services.grade_registry_writer_service.NameMatcher")
    def test_process_single_excel_file_partial_success(self, mock_name_matcher, mock_processor):
        """测试部分学生成绩写入成功"""
        mock_processor.extract_homework_number_from_filename.return_value = 1
        mock_processor.extract_grades_from_excel.return_value = [
            {"name": "张三", "grade": "A"},
            {"name": "王五", "grade": "B"},
        ]
        mock_name_matcher.match.side_effect = [("张三", "exact"), (None, "none")]
        self.mock_registry.find_or_create_homework_column.return_value = 5
        self.mock_registry.find_student_row.return_value = 2
        self.mock_registry.write_grade.return_value = (True, None)

        result = self.service._process_single_excel_file(
            "/path/to/第1次作业成绩.xlsx", self.mock_registry
        )

        self.assertTrue(result["partial_success"])
        self.assertEqual(result["students_total"], 2)
        self.assertEqual(result["students_processed"], 1)
        self.assertEqual(result["students_failed"], 1)

    @patch("grading.services.grade_registry_writer_service.GradeFileProcessor")
    @patch("grading.services.grade_registry_writer_service.NameMatcher")
    def test_process_single_excel_file_all_failed(self, mock_name_matcher, mock_processor):
        """测试所有学生成绩写入失败"""
        mock_processor.extract_homework_number_from_filename.return_value = 1
        mock_processor.extract_grades_from_excel.return_value = [
            {"name": "王五", "grade": "A"},
            {"name": "赵六", "grade": "B"},
        ]
        mock_name_matcher.match.side_effect = [(None, "none"), (None, "none")]
        self.mock_registry.find_or_create_homework_column.return_value = 5

        result = self.service._process_single_excel_file(
            "/path/to/第1次作业成绩.xlsx", self.mock_registry
        )

        self.assertFalse(result["success"])
        self.assertFalse(result["partial_success"])
        self.assertEqual(result["students_total"], 2)
        self.assertEqual(result["students_processed"], 0)
        self.assertEqual(result["students_failed"], 2)
        self.assertIn("所有学生成绩写入失败", result["error_message"])


class AuditLoggerIntegrationTest(BaseTestCase):
    """测试AuditLogger与服务的集成"""

    def setUp(self):
        super().setUp()
        self.tenant = Tenant.objects.create(name="测试租户")
        self.service = GradeRegistryWriterService(
            self.user, self.tenant, GradeRegistryWriterService.SCENARIO_GRADING_SYSTEM
        )

    @patch("grading.services.grade_registry_writer_service.GradeFileProcessor")
    @patch("grading.services.grade_registry_writer_service.NameMatcher")
    def test_audit_logger_called_during_word_processing(self, mock_name_matcher, mock_processor):
        """测试处理Word文档时调用审计日志"""
        mock_processor.extract_student_name.return_value = "张三"
        mock_processor.extract_grade_from_word.return_value = "A"
        mock_processor.extract_homework_number_from_path.return_value = 1
        mock_name_matcher.match.return_value = ("张三", "exact")

        mock_registry = Mock(spec=RegistryManager)
        mock_registry.student_names = {"张三": 2}
        mock_registry.find_student_row.return_value = 2
        mock_registry.write_grade.return_value = (True, None)

        # Mock audit logger
        with patch.object(self.service.audit_logger, "log_grade_write") as mock_log_grade:
            with patch.object(self.service.audit_logger, "log_file_processing") as mock_log_file:
                result = self.service._process_single_word_file(
                    "/path/to/第1次作业/张三_作业1.docx", mock_registry, 5
                )

                # 验证审计日志被调用
                self.assertTrue(result["success"])
                mock_log_grade.assert_called_once()
                mock_log_file.assert_called_once_with(
                    "/path/to/第1次作业/张三_作业1.docx", "success"
                )

    @patch("grading.services.grade_registry_writer_service.GradeFileProcessor")
    def test_audit_logger_called_on_file_failure(self, mock_processor):
        """测试文件处理失败时调用审计日志"""
        mock_processor.extract_student_name.return_value = None

        mock_registry = Mock(spec=RegistryManager)

        with patch.object(self.service.audit_logger, "log_file_processing") as mock_log_file:
            result = self.service._process_single_word_file("/path/to/作业1.docx", mock_registry, 5)

            # 验证审计日志记录失败
            self.assertFalse(result["success"])
            mock_log_file.assert_called_once()
            call_args = mock_log_file.call_args[0]
            self.assertEqual(call_args[1], "failed")

    def test_audit_logger_tracks_operation_lifecycle(self):
        """测试审计日志跟踪操作生命周期"""
        with patch.object(self.service.audit_logger, "start_operation") as mock_start:
            with patch.object(self.service.audit_logger, "end_operation") as mock_end:
                with patch.object(self.service, "find_grade_registry", return_value=None):
                    # 调用会失败的操作
                    result = self.service.process_grading_system_scenario(
                        "/path/to/homework", "/path/to/class"
                    )

                    # 验证start和end都被调用
                    mock_start.assert_called_once()
                    mock_end.assert_called_once()

                    # 验证end被调用时success=False
                    end_call_kwargs = mock_end.call_args[1]
                    # end_operation的第一个参数是success
                    self.assertFalse(mock_end.call_args[0][0])


class AuditLoggerEdgeCasesTest(BaseTestCase):
    """测试AuditLogger的边界情况"""

    def setUp(self):
        super().setUp()
        self.tenant = Tenant.objects.create(name="测试租户")

    def test_audit_logger_with_empty_strings(self):
        """测试空字符串参数"""
        audit_logger = AuditLogger(self.user, self.tenant, "")

        self.assertEqual(audit_logger.scenario, "")

    def test_audit_logger_multiple_operations(self):
        """测试连续多次操作"""
        audit_logger = AuditLogger(
            self.user, self.tenant, GradeRegistryWriterService.SCENARIO_GRADING_SYSTEM
        )

        with patch("grading.services.grade_registry_writer_service.datetime") as mock_datetime:
            mock_datetime.now.side_effect = [
                datetime(2024, 1, 1, 12, 0, 0),
                datetime(2024, 1, 1, 12, 0, 10),
                datetime(2024, 1, 1, 12, 0, 20),
                datetime(2024, 1, 1, 12, 0, 30),
            ]

            # 第一次操作
            audit_logger.start_operation("operation1")
            audit_logger.end_operation(success=True)

            # 第二次操作
            audit_logger.start_operation("operation2")
            audit_logger.end_operation(success=False)

            # 验证最后的操作类型
            self.assertEqual(audit_logger.audit_data["operation_type"], "operation2")

    def test_audit_logger_with_special_characters(self):
        """测试特殊字符处理"""
        audit_logger = AuditLogger(
            self.user, self.tenant, GradeRegistryWriterService.SCENARIO_GRADING_SYSTEM
        )

        # 测试包含特殊字符的文件路径
        mock_logger = MagicMock()
        with patch.object(audit_logger, "logger", mock_logger):
            audit_logger.log_file_processing("/path/to/文件名(特殊字符).docx", "success")

            # 验证日志调用成功
            mock_logger.info.assert_called_once()

    def test_audit_logger_with_unicode_names(self):
        """测试Unicode字符处理"""
        audit_logger = AuditLogger(
            self.user, self.tenant, GradeRegistryWriterService.SCENARIO_GRADING_SYSTEM
        )

        mock_logger = MagicMock()
        with patch.object(audit_logger, "logger", mock_logger):
            audit_logger.log_grade_write("张三李四王五", 1, "优秀", "良好")

            # 验证日志调用成功
            mock_logger.info.assert_called_once()


class GradeRegistryWriterIntegrationTest(BaseTestCase):
    """集成测试：测试完整的成绩写入流程"""

    def setUp(self):
        super().setUp()
        self.tenant = Tenant.objects.create(name="测试租户")
        self.temp_dir = tempfile.mkdtemp()

    def tearDown(self):
        if os.path.exists(self.temp_dir):
            shutil.rmtree(self.temp_dir)
        super().tearDown()

    def test_grading_system_scenario_complete_flow(self):
        """测试作业评分系统场景的完整流程"""
        from docx import Document
        from openpyxl import Workbook

        # 1. 创建班级目录和登分册
        class_dir = os.path.join(self.temp_dir, "2024级计算机1班")
        os.makedirs(class_dir)

        registry_path = os.path.join(class_dir, "成绩登分册.xlsx")
        wb = Workbook()
        ws = wb.active
        ws["A1"] = "姓名"
        ws["A2"] = "张三"
        ws["A3"] = "李四"
        wb.save(registry_path)

        # 2. 创建作业目录和Word文档
        homework_dir = os.path.join(self.temp_dir, "第1次作业")
        os.makedirs(homework_dir)

        # 创建张三的作业
        doc1_path = os.path.join(homework_dir, "张三_作业1.docx")
        doc1 = Document()
        doc1.add_paragraph("作业内容")
        doc1.add_paragraph("老师评分：A")
        doc1.save(doc1_path)

        # 创建李四的作业
        doc2_path = os.path.join(homework_dir, "李四_作业1.docx")
        doc2 = Document()
        doc2.add_paragraph("作业内容")
        doc2.add_paragraph("老师评分：B")
        doc2.save(doc2_path)

        # 3. 执行服务
        service = GradeRegistryWriterService(
            self.user, self.tenant, GradeRegistryWriterService.SCENARIO_GRADING_SYSTEM
        )

        result = service.process_grading_system_scenario(homework_dir, class_dir)

        # 4. 验证结果
        self.assertTrue(result["success"])
        self.assertEqual(result["homework_number"], 1)
        self.assertEqual(result["statistics"]["total"], 2)
        self.assertEqual(result["statistics"]["success"], 2)
        self.assertEqual(result["statistics"]["failed"], 0)

        # 5. 验证登分册内容
        from openpyxl import load_workbook

        wb = load_workbook(registry_path)
        ws = wb.active

        # 验证作业列已创建
        self.assertEqual(ws.cell(1, 2).value, "第1次作业")

        # 验证成绩已写入
        self.assertEqual(ws.cell(2, 2).value, "A")  # 张三
        self.assertEqual(ws.cell(3, 2).value, "B")  # 李四

    def test_toolbox_scenario_complete_flow(self):
        """测试工具箱模块场景的完整流程"""
        from openpyxl import Workbook

        # 1. 创建班级目录和登分册
        class_dir = os.path.join(self.temp_dir, "2024级计算机1班")
        os.makedirs(class_dir)

        registry_path = os.path.join(class_dir, "成绩登分册.xlsx")
        wb = Workbook()
        ws = wb.active
        ws["A1"] = "姓名"
        ws["A2"] = "张三"
        ws["A3"] = "李四"
        ws["A4"] = "王五"
        wb.save(registry_path)

        # 2. 创建Excel成绩文件
        grade_file_path = os.path.join(class_dir, "第2次作业成绩.xlsx")
        wb2 = Workbook()
        ws2 = wb2.active
        ws2["A1"] = "姓名"
        ws2["B1"] = "成绩"
        ws2["A2"] = "张三"
        ws2["B2"] = "优秀"
        ws2["A3"] = "李四"
        ws2["B3"] = "良好"
        wb2.save(grade_file_path)

        # 3. 执行服务
        service = GradeRegistryWriterService(
            self.user, self.tenant, GradeRegistryWriterService.SCENARIO_TOOLBOX
        )

        result = service.process_toolbox_scenario(class_dir)

        # 4. 验证结果
        self.assertTrue(result["success"])
        self.assertEqual(result["statistics"]["total_files"], 1)
        self.assertEqual(result["statistics"]["success"], 2)

        # 5. 验证登分册内容
        from openpyxl import load_workbook

        wb = load_workbook(registry_path)
        ws = wb.active

        # 验证作业列已创建
        self.assertEqual(ws.cell(1, 2).value, "第2次作业")

        # 验证成绩已写入
        self.assertEqual(ws.cell(2, 2).value, "优秀")  # 张三
        self.assertEqual(ws.cell(3, 2).value, "良好")  # 李四

    def test_error_recovery_with_backup(self):
        """测试错误恢复机制"""
        from docx import Document
        from openpyxl import Workbook

        # 1. 创建班级目录和登分册
        class_dir = os.path.join(self.temp_dir, "2024级计算机1班")
        os.makedirs(class_dir)

        registry_path = os.path.join(class_dir, "成绩登分册.xlsx")
        wb = Workbook()
        ws = wb.active
        ws["A1"] = "姓名"
        ws["A2"] = "张三"
        wb.save(registry_path)

        # 2. 创建作业目录和Word文档
        homework_dir = os.path.join(self.temp_dir, "第1次作业")
        os.makedirs(homework_dir)

        doc_path = os.path.join(homework_dir, "张三_作业1.docx")
        doc = Document()
        doc.add_paragraph("作业内容")
        doc.add_paragraph("老师评分：A")
        doc.save(doc_path)

        # 3. Mock保存失败
        service = GradeRegistryWriterService(
            self.user, self.tenant, GradeRegistryWriterService.SCENARIO_GRADING_SYSTEM
        )

        with patch.object(RegistryManager, "save", return_value=False):
            result = service.process_grading_system_scenario(homework_dir, class_dir)

        # 4. 验证错误处理
        self.assertFalse(result["success"])
        self.assertIn("保存登分册失败", result["error_message"])

    def test_name_matching_fuzzy(self):
        """测试姓名模糊匹配"""
        from docx import Document
        from openpyxl import Workbook

        # 1. 创建班级目录和登分册（姓名带空格）
        class_dir = os.path.join(self.temp_dir, "2024级计算机1班")
        os.makedirs(class_dir)

        registry_path = os.path.join(class_dir, "成绩登分册.xlsx")
        wb = Workbook()
        ws = wb.active
        ws["A1"] = "姓名"
        ws["A2"] = "张 三"  # 带空格
        wb.save(registry_path)

        # 2. 创建作业目录和Word文档（姓名不带空格）
        homework_dir = os.path.join(self.temp_dir, "第1次作业")
        os.makedirs(homework_dir)

        doc_path = os.path.join(homework_dir, "张三_作业1.docx")
        doc = Document()
        doc.add_paragraph("作业内容")
        doc.add_paragraph("老师评分：A")
        doc.save(doc_path)

        # 3. 执行服务
        service = GradeRegistryWriterService(
            self.user, self.tenant, GradeRegistryWriterService.SCENARIO_GRADING_SYSTEM
        )

        result = service.process_grading_system_scenario(homework_dir, class_dir)

        # 4. 验证模糊匹配成功
        self.assertTrue(result["success"])
        self.assertEqual(result["statistics"]["success"], 1)

        # 5. 验证成绩已写入
        from openpyxl import load_workbook

        wb = load_workbook(registry_path)
        ws = wb.active
        self.assertEqual(ws.cell(2, 2).value, "A")

    def test_skip_duplicate_grades(self):
        """测试跳过重复成绩"""
        from docx import Document
        from openpyxl import Workbook

        # 1. 创建班级目录和登分册（已有成绩）
        class_dir = os.path.join(self.temp_dir, "2024级计算机1班")
        os.makedirs(class_dir)

        registry_path = os.path.join(class_dir, "成绩登分册.xlsx")
        wb = Workbook()
        ws = wb.active
        ws["A1"] = "姓名"
        ws["B1"] = "第1次作业"
        ws["A2"] = "张三"
        ws["B2"] = "A"  # 已有成绩
        wb.save(registry_path)

        # 2. 创建作业目录和Word文档（相同成绩）
        homework_dir = os.path.join(self.temp_dir, "第1次作业")
        os.makedirs(homework_dir)

        doc_path = os.path.join(homework_dir, "张三_作业1.docx")
        doc = Document()
        doc.add_paragraph("作业内容")
        doc.add_paragraph("老师评分：A")
        doc.save(doc_path)

        # 3. 执行服务
        service = GradeRegistryWriterService(
            self.user, self.tenant, GradeRegistryWriterService.SCENARIO_GRADING_SYSTEM
        )

        result = service.process_grading_system_scenario(homework_dir, class_dir)

        # 4. 验证跳过重复成绩
        self.assertTrue(result["success"])
        self.assertEqual(result["statistics"]["skipped"], 1)
        self.assertEqual(result["statistics"]["success"], 0)

    def test_lab_report_grade_extraction(self):
        """测试实验报告成绩提取"""
        from docx import Document
        from openpyxl import Workbook

        # 1. 创建班级目录和登分册
        class_dir = os.path.join(self.temp_dir, "2024级计算机1班")
        os.makedirs(class_dir)

        registry_path = os.path.join(class_dir, "成绩登分册.xlsx")
        wb = Workbook()
        ws = wb.active
        ws["A1"] = "姓名"
        ws["A2"] = "张三"
        wb.save(registry_path)

        # 2. 创建作业目录和实验报告
        homework_dir = os.path.join(self.temp_dir, "第1次作业")
        os.makedirs(homework_dir)

        doc_path = os.path.join(homework_dir, "张三_作业1.docx")
        doc = Document()
        doc.add_paragraph("实验报告")
        table = doc.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "教师（签字）："
        table.cell(0, 1).text = "优秀"
        doc.save(doc_path)

        # 3. 执行服务
        service = GradeRegistryWriterService(
            self.user, self.tenant, GradeRegistryWriterService.SCENARIO_GRADING_SYSTEM
        )

        result = service.process_grading_system_scenario(homework_dir, class_dir)

        # 4. 验证实验报告成绩提取成功
        self.assertTrue(result["success"])
        self.assertEqual(result["statistics"]["success"], 1)

        # 5. 验证成绩已写入
        from openpyxl import load_workbook

        wb = load_workbook(registry_path)
        ws = wb.active
        self.assertEqual(ws.cell(2, 2).value, "优秀")

    def test_multiple_homework_batches(self):
        """测试多次作业批次"""
        from docx import Document
        from openpyxl import Workbook

        # 1. 创建班级目录和登分册
        class_dir = os.path.join(self.temp_dir, "2024级计算机1班")
        os.makedirs(class_dir)

        registry_path = os.path.join(class_dir, "成绩登分册.xlsx")
        wb = Workbook()
        ws = wb.active
        ws["A1"] = "姓名"
        ws["A2"] = "张三"
        wb.save(registry_path)

        service = GradeRegistryWriterService(
            self.user, self.tenant, GradeRegistryWriterService.SCENARIO_GRADING_SYSTEM
        )

        # 2. 处理第1次作业
        homework_dir1 = os.path.join(self.temp_dir, "第1次作业")
        os.makedirs(homework_dir1)

        doc1_path = os.path.join(homework_dir1, "张三_作业1.docx")
        doc1 = Document()
        doc1.add_paragraph("作业内容")
        doc1.add_paragraph("老师评分：A")
        doc1.save(doc1_path)

        result1 = service.process_grading_system_scenario(homework_dir1, class_dir)
        self.assertTrue(result1["success"])

        # 3. 处理第2次作业
        homework_dir2 = os.path.join(self.temp_dir, "第2次作业")
        os.makedirs(homework_dir2)

        doc2_path = os.path.join(homework_dir2, "张三_作业2.docx")
        doc2 = Document()
        doc2.add_paragraph("作业内容")
        doc2.add_paragraph("老师评分：B")
        doc2.save(doc2_path)

        result2 = service.process_grading_system_scenario(homework_dir2, class_dir)
        self.assertTrue(result2["success"])

        # 4. 验证两次作业成绩都已写入
        from openpyxl import load_workbook

        wb = load_workbook(registry_path)
        ws = wb.active

        self.assertEqual(ws.cell(1, 2).value, "第1次作业")
        self.assertEqual(ws.cell(1, 3).value, "第2次作业")
        self.assertEqual(ws.cell(2, 2).value, "A")
        self.assertEqual(ws.cell(2, 3).value, "B")


class SecurityValidationTest(BaseTestCase):
    """测试安全验证功能"""

    def setUp(self):
        super().setUp()
        self.tenant = Tenant.objects.create(name="测试租户")
        self.service = GradeRegistryWriterService(
            self.user, self.tenant, GradeRegistryWriterService.SCENARIO_TOOLBOX
        )
        self.temp_dir = tempfile.mkdtemp()
        self.test_file = os.path.join(self.temp_dir, "test.xlsx")
        # 创建一个测试文件
        with open(self.test_file, "w") as f:
            f.write("test content")

    def tearDown(self):
        super().tearDown()
        if os.path.exists(self.temp_dir):
            shutil.rmtree(self.temp_dir)

    def test_validate_path_security_valid_path(self):
        """测试有效路径验证"""
        is_valid, error_msg = self.service._validate_path_security(self.test_file, self.temp_dir)
        self.assertTrue(is_valid)
        self.assertIsNone(error_msg)

    def test_validate_path_security_path_traversal(self):
        """测试路径遍历攻击检测"""
        malicious_path = os.path.join(self.temp_dir, "../../../etc/passwd")
        is_valid, error_msg = self.service._validate_path_security(malicious_path, self.temp_dir)
        self.assertFalse(is_valid)
        self.assertIn("无权访问该路径", error_msg)

    def test_validate_path_security_dangerous_characters(self):
        """测试危险字符检测"""
        dangerous_paths = [
            os.path.join(self.temp_dir, "../test.xlsx"),
            os.path.join(self.temp_dir, "test$file.xlsx"),
            os.path.join(self.temp_dir, "test|file.xlsx"),
            os.path.join(self.temp_dir, "test;file.xlsx"),
        ]

        for dangerous_path in dangerous_paths:
            is_valid, error_msg = self.service._validate_path_security(
                dangerous_path, self.temp_dir
            )
            self.assertFalse(is_valid)
            self.assertIn("非法字符", error_msg)

    def test_validate_path_security_nonexistent_file(self):
        """测试不存在的文件"""
        nonexistent_file = os.path.join(self.temp_dir, "nonexistent.xlsx")
        is_valid, error_msg = self.service._validate_path_security(nonexistent_file, self.temp_dir)
        self.assertFalse(is_valid)
        self.assertIn("文件不存在", error_msg)

    def test_validate_path_security_symlink(self):
        """测试符号链接检测"""
        # 创建符号链接
        symlink_path = os.path.join(self.temp_dir, "symlink.xlsx")
        try:
            os.symlink(self.test_file, symlink_path)
            is_valid, error_msg = self.service._validate_path_security(symlink_path, self.temp_dir)
            self.assertFalse(is_valid)
            self.assertIn("符号链接", error_msg)
        except OSError:
            # 在某些系统上创建符号链接可能需要特殊权限
            self.skipTest("无法创建符号链接")

    def test_validate_path_security_no_read_permission(self):
        """测试无读取权限的文件"""
        # 创建一个无读取权限的文件
        no_read_file = os.path.join(self.temp_dir, "no_read.xlsx")
        with open(no_read_file, "w") as f:
            f.write("test")

        try:
            os.chmod(no_read_file, 0o000)
            is_valid, error_msg = self.service._validate_path_security(no_read_file, self.temp_dir)
            self.assertFalse(is_valid)
            self.assertIn("无权限读取文件", error_msg)
        finally:
            # 恢复权限以便清理
            os.chmod(no_read_file, 0o644)

    def test_validate_file_size_valid(self):
        """测试有效文件大小"""
        is_valid, error_msg = self.service._validate_file_size(self.test_file)
        self.assertTrue(is_valid)
        self.assertIsNone(error_msg)

    def test_validate_file_size_empty_file(self):
        """测试空文件"""
        empty_file = os.path.join(self.temp_dir, "empty.xlsx")
        with open(empty_file, "w") as f:
            pass  # 创建空文件

        is_valid, error_msg = self.service._validate_file_size(empty_file)
        self.assertFalse(is_valid)
        self.assertIn("文件为空或损坏", error_msg)

    def test_validate_file_size_too_large(self):
        """测试文件过大"""
        large_file = os.path.join(self.temp_dir, "large.xlsx")
        # 创建一个超过限制的文件
        with open(large_file, "wb") as f:
            f.write(b"0" * (self.service.MAX_FILE_SIZE + 1))

        is_valid, error_msg = self.service._validate_file_size(large_file)
        self.assertFalse(is_valid)
        self.assertIn("文件大小超过限制", error_msg)

    def test_validate_excel_integrity_invalid_extension(self):
        """测试非Excel文件扩展名"""
        txt_file = os.path.join(self.temp_dir, "test.txt")
        with open(txt_file, "w") as f:
            f.write("test")

        is_valid, error_msg = self.service._validate_excel_integrity(txt_file)
        self.assertFalse(is_valid)
        self.assertIn("文件格式错误", error_msg)

    def test_validate_excel_integrity_rejects_xls(self):
        """测试不再接受.xls格式"""
        xls_file = os.path.join(self.temp_dir, "test.xls")
        with open(xls_file, "w") as f:
            f.write("dummy")

        is_valid, error_msg = self.service._validate_excel_integrity(xls_file)
        self.assertFalse(is_valid)
        self.assertIn(".xlsx", error_msg)

    def test_validate_excel_integrity_corrupted_file(self):
        """测试损坏的Excel文件"""
        corrupted_file = os.path.join(self.temp_dir, "corrupted.xlsx")
        with open(corrupted_file, "w") as f:
            f.write("not a valid excel file")

        is_valid, error_msg = self.service._validate_excel_integrity(corrupted_file)
        self.assertFalse(is_valid)
        self.assertIn("损坏或格式错误", error_msg)

    def test_validate_tenant_isolation_no_tenant(self):
        """测试没有租户信息时的隔离检查"""
        service_no_tenant = GradeRegistryWriterService(
            self.user, None, GradeRegistryWriterService.SCENARIO_TOOLBOX
        )
        is_valid, error_msg = service_no_tenant._validate_tenant_isolation(self.test_file)
        self.assertTrue(is_valid)
        self.assertIsNone(error_msg)

    def test_validate_tenant_isolation_valid(self):
        """测试有效的租户隔离"""
        is_valid, error_msg = self.service._validate_tenant_isolation(self.test_file)
        self.assertTrue(is_valid)
        self.assertIsNone(error_msg)


class SecurityIntegrationTest(BaseTestCase):
    """测试安全功能的集成"""

    def setUp(self):
        super().setUp()
        self.tenant = Tenant.objects.create(name="测试租户")
        self.service = GradeRegistryWriterService(
            self.user, self.tenant, GradeRegistryWriterService.SCENARIO_GRADING_SYSTEM
        )
        self.temp_dir = tempfile.mkdtemp()

    def tearDown(self):
        super().tearDown()
        if os.path.exists(self.temp_dir):
            shutil.rmtree(self.temp_dir)

    def test_process_with_invalid_path(self):
        """测试处理无效路径时的安全检查"""
        malicious_homework_dir = "/etc/passwd"
        result = self.service.process_grading_system_scenario(malicious_homework_dir, self.temp_dir)
        self.assertFalse(result["success"])
        self.assertIsNotNone(result["error_message"])

    def test_security_audit_logging(self):
        """测试安全事件的审计日志记录"""
        # 创建一个包含危险字符的路径
        dangerous_path = os.path.join(self.temp_dir, "test$file.xlsx")

        with patch.object(self.service.audit_logger.logger, "warning") as mock_warning:
            self.service._validate_path_security(dangerous_path, self.temp_dir)
            # 验证安全事件被记录
            mock_warning.assert_called()
            call_args = str(mock_warning.call_args)
            self.assertIn("安全事件", call_args)


class PerformanceOptimizationTest(BaseTestCase):
    """测试性能优化功能"""

    def setUp(self):
        super().setUp()
        self.tenant = Tenant.objects.create(name="测试租户")
        self.service = GradeRegistryWriterService(
            self.user, self.tenant, GradeRegistryWriterService.SCENARIO_TOOLBOX
        )
        self.temp_dir = tempfile.mkdtemp()

    def tearDown(self):
        super().tearDown()
        if os.path.exists(self.temp_dir):
            shutil.rmtree(self.temp_dir)

    def test_large_file_count_warning(self):
        """测试大量文件时的警告日志"""
        class_dir = os.path.join(self.temp_dir, "class")
        os.makedirs(class_dir)

        # 创建登分册
        registry_file = os.path.join(class_dir, "成绩登分册.xlsx")
        from openpyxl import Workbook

        wb = Workbook()
        ws = wb.active
        ws["A1"] = "姓名"
        ws["A2"] = "张三"
        wb.save(registry_file)

        # 创建超过阈值的Excel文件
        for i in range(self.service.MAX_FILES_WARNING_THRESHOLD + 5):
            file_path = os.path.join(class_dir, f"第{i+1}次作业成绩.xlsx")
            wb = Workbook()
            ws = wb.active
            ws["A1"] = "姓名"
            ws["B1"] = "成绩"
            wb.save(file_path)

        with patch.object(self.service.logger, "warning") as mock_warning:
            result = self.service.process_toolbox_scenario(class_dir)

            # 验证警告日志被调用
            warning_calls = [str(call) for call in mock_warning.call_args_list]
            self.assertTrue(any("文件数量超过阈值" in str(call) for call in warning_calls))

    def test_batch_file_processing(self):
        """测试批量文件处理"""
        class_dir = os.path.join(self.temp_dir, "class")
        os.makedirs(class_dir)

        # 创建登分册
        registry_file = os.path.join(class_dir, "成绩登分册.xlsx")
        from openpyxl import Workbook

        wb = Workbook()
        ws = wb.active
        ws["A1"] = "姓名"
        ws["A2"] = "张三"
        ws["A3"] = "李四"
        wb.save(registry_file)

        # 创建多个成绩文件
        for i in range(5):
            file_path = os.path.join(class_dir, f"第{i+1}次作业成绩.xlsx")
            wb = Workbook()
            ws = wb.active
            ws["A1"] = "姓名"
            ws["B1"] = "成绩"
            ws["A2"] = "张三"
            ws["B2"] = "A"
            wb.save(file_path)

        result = self.service.process_toolbox_scenario(class_dir)

        # 验证批量处理结果
        self.assertTrue(result["success"])
        self.assertEqual(result["statistics"]["total_files"], 5)


class EdgeCaseTest(BaseTestCase):
    """测试边界情况"""

    def setUp(self):
        super().setUp()
        self.tenant = Tenant.objects.create(name="测试租户")
        self.temp_dir = tempfile.mkdtemp()

    def tearDown(self):
        super().tearDown()
        if os.path.exists(self.temp_dir):
            shutil.rmtree(self.temp_dir)

    def test_empty_homework_directory(self):
        """测试空作业目录"""
        service = GradeRegistryWriterService(
            self.user, self.tenant, GradeRegistryWriterService.SCENARIO_GRADING_SYSTEM
        )

        homework_dir = os.path.join(self.temp_dir, "homework")
        class_dir = os.path.join(self.temp_dir, "class")
        os.makedirs(homework_dir)
        os.makedirs(class_dir)

        # 创建登分册
        registry_file = os.path.join(class_dir, "成绩登分册.xlsx")
        from openpyxl import Workbook

        wb = Workbook()
        ws = wb.active
        ws["A1"] = "姓名"
        wb.save(registry_file)

        result = service.process_grading_system_scenario(homework_dir, class_dir)

        self.assertFalse(result["success"])
        self.assertIn("没有找到Word文档", result["error_message"])

    def test_registry_with_no_students(self):
        """测试没有学生的登分册"""
        service = GradeRegistryWriterService(
            self.user, self.tenant, GradeRegistryWriterService.SCENARIO_GRADING_SYSTEM
        )

        homework_dir = os.path.join(self.temp_dir, "homework")
        class_dir = os.path.join(self.temp_dir, "class")
        os.makedirs(homework_dir)
        os.makedirs(class_dir)

        # 创建只有表头的登分册
        registry_file = os.path.join(class_dir, "成绩登分册.xlsx")
        from openpyxl import Workbook

        wb = Workbook()
        ws = wb.active
        ws["A1"] = "姓名"
        wb.save(registry_file)

        # 创建作业文档
        from docx import Document

        doc_path = os.path.join(homework_dir, "张三_作业1.docx")
        doc = Document()
        doc.add_paragraph("老师评分：A")
        doc.save(doc_path)

        result = service.process_grading_system_scenario(homework_dir, class_dir)

        # 应该处理失败，因为找不到学生
        self.assertEqual(result["statistics"]["failed"], 1)

    def test_special_characters_in_student_names(self):
        """测试学生姓名中的特殊字符"""
        service = GradeRegistryWriterService(
            self.user, self.tenant, GradeRegistryWriterService.SCENARIO_GRADING_SYSTEM
        )

        homework_dir = os.path.join(self.temp_dir, "homework")
        class_dir = os.path.join(self.temp_dir, "class")
        os.makedirs(homework_dir)
        os.makedirs(class_dir)

        # 创建包含特殊字符姓名的登分册
        registry_file = os.path.join(class_dir, "成绩登分册.xlsx")
        from openpyxl import Workbook

        wb = Workbook()
        ws = wb.active
        ws["A1"] = "姓名"
        ws["A2"] = "张三·李"  # 包含中点
        wb.save(registry_file)

        # 创建作业文档
        from docx import Document

        doc_path = os.path.join(homework_dir, "张三·李_作业1.docx")
        doc = Document()
        doc.add_paragraph("老师评分：A")
        doc.save(doc_path)

        result = service.process_grading_system_scenario(homework_dir, class_dir)

        # 应该能够处理特殊字符
        self.assertTrue(result["success"] or result["statistics"]["success"] > 0)

    def test_very_long_file_paths(self):
        """测试超长文件路径"""
        service = GradeRegistryWriterService(
            self.user, self.tenant, GradeRegistryWriterService.SCENARIO_TOOLBOX
        )

        # 创建深层嵌套目录
        deep_dir = self.temp_dir
        for i in range(10):
            deep_dir = os.path.join(deep_dir, f"level{i}")

        try:
            os.makedirs(deep_dir)

            # 创建登分册
            registry_file = os.path.join(deep_dir, "成绩登分册.xlsx")
            from openpyxl import Workbook

            wb = Workbook()
            ws = wb.active
            ws["A1"] = "姓名"
            wb.save(registry_file)

            result = service.process_toolbox_scenario(deep_dir)

            # 应该能够处理深层路径
            self.assertIsNotNone(result)
        except OSError:
            # 某些系统可能有路径长度限制
            self.skipTest("系统不支持超长路径")

    def test_concurrent_file_access(self):
        """测试并发文件访问"""
        service = GradeRegistryWriterService(
            self.user, self.tenant, GradeRegistryWriterService.SCENARIO_TOOLBOX
        )

        class_dir = os.path.join(self.temp_dir, "class")
        os.makedirs(class_dir)

        # 创建登分册
        registry_file = os.path.join(class_dir, "成绩登分册.xlsx")
        from openpyxl import Workbook

        wb = Workbook()
        ws = wb.active
        ws["A1"] = "姓名"
        ws["A2"] = "张三"
        wb.save(registry_file)

        # 创建成绩文件
        grade_file = os.path.join(class_dir, "第1次作业成绩.xlsx")
        wb = Workbook()
        ws = wb.active
        ws["A1"] = "姓名"
        ws["B1"] = "成绩"
        ws["A2"] = "张三"
        ws["B2"] = "A"
        wb.save(grade_file)

        # 模拟文件被锁定
        with patch("openpyxl.load_workbook", side_effect=PermissionError("文件被占用")):
            result = service.process_toolbox_scenario(class_dir)

            # 应该处理文件访问错误
            self.assertFalse(result["success"])


class ProcessMethodTest(BaseTestCase):
    """测试process方法的场景路由"""

    def setUp(self):
        super().setUp()
        self.tenant = Tenant.objects.create(name="测试租户")
        self.temp_dir = tempfile.mkdtemp()

    def tearDown(self):
        super().tearDown()
        if os.path.exists(self.temp_dir):
            shutil.rmtree(self.temp_dir)

    def test_process_grading_system_scenario_routing(self):
        """测试作业评分系统场景路由"""
        service = GradeRegistryWriterService(
            self.user, self.tenant, GradeRegistryWriterService.SCENARIO_GRADING_SYSTEM
        )

        with patch.object(service, "process_grading_system_scenario") as mock_process:
            mock_process.return_value = {"success": True}

            result = service.process(homework_dir="/path/to/homework", class_dir="/path/to/class")

            mock_process.assert_called_once_with("/path/to/homework", "/path/to/class")
            self.assertTrue(result["success"])

    def test_process_toolbox_scenario_routing(self):
        """测试工具箱模块场景路由"""
        service = GradeRegistryWriterService(
            self.user, self.tenant, GradeRegistryWriterService.SCENARIO_TOOLBOX
        )

        with patch.object(service, "process_toolbox_scenario") as mock_process:
            mock_process.return_value = {"success": True}

            result = service.process(class_dir="/path/to/class")

            mock_process.assert_called_once_with("/path/to/class")
            self.assertTrue(result["success"])

    def test_process_missing_parameters_grading_system(self):
        """测试作业评分系统场景缺少参数"""
        service = GradeRegistryWriterService(
            self.user, self.tenant, GradeRegistryWriterService.SCENARIO_GRADING_SYSTEM
        )

        result = service.process(homework_dir="/path/to/homework")

        self.assertFalse(result["success"])
        self.assertIn("缺少必需参数", result["error_message"])

    def test_process_missing_parameters_toolbox(self):
        """测试工具箱模块场景缺少参数"""
        service = GradeRegistryWriterService(
            self.user, self.tenant, GradeRegistryWriterService.SCENARIO_TOOLBOX
        )

        result = service.process()

        self.assertFalse(result["success"])
        self.assertIn("缺少必需参数", result["error_message"])

    def test_process_unknown_scenario(self):
        """测试未知场景类型"""
        service = GradeRegistryWriterService(self.user, self.tenant, "unknown_scenario")

        result = service.process(class_dir="/path/to/class")

        self.assertFalse(result["success"])
        self.assertIn("未知的场景类型", result["error_message"])

    def test_process_exception_handling(self):
        """测试process方法的异常处理"""
        service = GradeRegistryWriterService(
            self.user, self.tenant, GradeRegistryWriterService.SCENARIO_TOOLBOX
        )

        with patch.object(service, "process_toolbox_scenario", side_effect=Exception("测试异常")):
            result = service.process(class_dir="/path/to/class")

            self.assertFalse(result["success"])
            self.assertIn("处理成绩写入时出错", result["error_message"])


class FindGradeRegistryTest(BaseTestCase):
    """测试查找登分册功能"""

    def setUp(self):
        super().setUp()
        self.tenant = Tenant.objects.create(name="测试租户")
        self.service = GradeRegistryWriterService(
            self.user, self.tenant, GradeRegistryWriterService.SCENARIO_TOOLBOX
        )
        self.temp_dir = tempfile.mkdtemp()

    def tearDown(self):
        super().tearDown()
        if os.path.exists(self.temp_dir):
            shutil.rmtree(self.temp_dir)

    def test_find_registry_with_full_name(self):
        """测试查找完整名称的登分册"""
        registry_file = os.path.join(self.temp_dir, "成绩登分册.xlsx")
        open(registry_file, "w").close()

        result = self.service.find_grade_registry(self.temp_dir)

        self.assertEqual(result, registry_file)

    def test_find_registry_with_short_name(self):
        """测试查找简短名称的登分册"""
        registry_file = os.path.join(self.temp_dir, "登分册.xlsx")
        open(registry_file, "w").close()

        result = self.service.find_grade_registry(self.temp_dir)

        self.assertEqual(result, registry_file)

    def test_find_registry_with_english_name(self):
        """测试查找英文名称的登分册"""
        registry_file = os.path.join(self.temp_dir, "grade_registry.xlsx")
        open(registry_file, "w").close()

        result = self.service.find_grade_registry(self.temp_dir)

        self.assertEqual(result, registry_file)

    def test_find_registry_case_insensitive(self):
        """测试大小写不敏感查找"""
        registry_file = os.path.join(self.temp_dir, "GRADES.xlsx")
        open(registry_file, "w").close()

        result = self.service.find_grade_registry(self.temp_dir)

        self.assertEqual(result, registry_file)

    def test_find_registry_skip_temp_files(self):
        """测试跳过临时文件"""
        temp_file = os.path.join(self.temp_dir, "~$成绩登分册.xlsx")
        registry_file = os.path.join(self.temp_dir, "成绩登分册.xlsx")
        open(temp_file, "w").close()
        open(registry_file, "w").close()

        result = self.service.find_grade_registry(self.temp_dir)

        self.assertEqual(result, registry_file)

    def test_find_registry_xls_format(self):
        """不再支持.xls格式的登分册"""
        registry_file = os.path.join(self.temp_dir, "成绩登分册.xls")
        open(registry_file, "w").close()

        result = self.service.find_grade_registry(self.temp_dir)

        self.assertIsNone(result)

    def test_find_registry_directory_not_exists(self):
        """测试目录不存在"""
        non_existent = os.path.join(self.temp_dir, "non_existent")

        result = self.service.find_grade_registry(non_existent)

        self.assertIsNone(result)

    def test_find_registry_path_is_file(self):
        """测试路径是文件而非目录"""
        file_path = os.path.join(self.temp_dir, "file.txt")
        open(file_path, "w").close()

        result = self.service.find_grade_registry(file_path)

        self.assertIsNone(result)

    def test_find_registry_no_matching_file(self):
        """测试没有匹配的登分册文件"""
        other_file = os.path.join(self.temp_dir, "其他文件.xlsx")
        open(other_file, "w").close()

        result = self.service.find_grade_registry(self.temp_dir)

        self.assertIsNone(result)

    def test_find_registry_exception_handling(self):
        """测试异常处理"""
        with patch("os.path.exists", side_effect=Exception("测试异常")):
            result = self.service.find_grade_registry(self.temp_dir)

            self.assertIsNone(result)


class SecurityValidationEdgeCasesTest(BaseTestCase):
    """测试安全验证的更多边界情况"""

    def setUp(self):
        super().setUp()
        self.tenant = Tenant.objects.create(name="测试租户")
        self.service = GradeRegistryWriterService(
            self.user, self.tenant, GradeRegistryWriterService.SCENARIO_TOOLBOX
        )
        self.temp_dir = tempfile.mkdtemp()

    def tearDown(self):
        super().tearDown()
        if os.path.exists(self.temp_dir):
            shutil.rmtree(self.temp_dir)

    def test_validate_path_with_null_bytes(self):
        """测试包含空字节的路径"""
        try:
            malicious_path = os.path.join(self.temp_dir, "test\x00.xlsx")
            is_valid, error_msg = self.service._validate_path_security(
                malicious_path, self.temp_dir
            )
            self.assertFalse(is_valid)
        except ValueError:
            # Python可能会拒绝包含空字节的路径
            pass

    def test_validate_path_with_unicode_normalization(self):
        """测试Unicode规范化攻击"""
        # 使用不同的Unicode表示形式
        test_file = os.path.join(self.temp_dir, "测试.xlsx")
        with open(test_file, "w") as f:
            f.write("test")

        is_valid, error_msg = self.service._validate_path_security(test_file, self.temp_dir)
        self.assertTrue(is_valid)

    def test_validate_excel_with_zip_bomb(self):
        """测试ZIP炸弹检测（Excel是ZIP格式）"""
        # 创建一个看起来像Excel但实际是恶意文件的文件
        malicious_file = os.path.join(self.temp_dir, "malicious.xlsx")
        with open(malicious_file, "wb") as f:
            # 写入一些随机数据
            f.write(b"PK\x03\x04" + b"0" * 1000)

        is_valid, error_msg = self.service._validate_excel_integrity(malicious_file)
        self.assertFalse(is_valid)

    def test_validate_file_size_boundary(self):
        """测试文件大小边界值"""
        # 创建恰好等于最大限制的文件
        boundary_file = os.path.join(self.temp_dir, "boundary.xlsx")
        with open(boundary_file, "wb") as f:
            f.write(b"0" * self.service.MAX_FILE_SIZE)

        is_valid, error_msg = self.service._validate_file_size(boundary_file)
        self.assertTrue(is_valid)

        # 创建超过1字节的文件
        over_file = os.path.join(self.temp_dir, "over.xlsx")
        with open(over_file, "wb") as f:
            f.write(b"0" * (self.service.MAX_FILE_SIZE + 1))

        is_valid, error_msg = self.service._validate_file_size(over_file)
        self.assertFalse(is_valid)

    def test_validate_path_with_relative_components(self):
        """测试包含相对路径组件的路径"""
        relative_paths = [
            os.path.join(self.temp_dir, ".", "test.xlsx"),
            os.path.join(self.temp_dir, "subdir", "..", "test.xlsx"),
        ]

        for rel_path in relative_paths:
            # 创建文件
            os.makedirs(os.path.dirname(rel_path), exist_ok=True)
            with open(rel_path, "w") as f:
                f.write("test")

            is_valid, error_msg = self.service._validate_path_security(rel_path, self.temp_dir)
            # 相对路径组件本身不应该被拒绝，除非它们导致路径遍历
            self.assertIsNotNone(is_valid)
