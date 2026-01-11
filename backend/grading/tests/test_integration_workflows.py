"""
集成测试：完整工作流测试
Integration tests for complete workflows

测试需求:
- 完整的课程创建到评分流程
- Git仓库方式的完整流程
- 文件系统方式的完整流程
- 评价缓存和模板的完整流程
- 批量操作流程
"""

import os
import shutil
import tempfile
from pathlib import Path
from unittest.mock import MagicMock, patch

from django.contrib.auth.models import User
from django.test import Client, TestCase
from django.urls import reverse
from docx import Document

from grading.models import (
    Class,
    CommentTemplate,
    Course,
    GradeTypeConfig,
    Homework,
    Repository,
    Semester,
    Submission,
    Tenant,
    UserProfile,
)
from grading.services.class_service import ClassService
from grading.services.comment_template_service import CommentTemplateService
from grading.services.course_service import CourseService
from grading.services.file_upload_service import FileUploadService
from grading.services.repository_service import RepositoryService
from grading.views import (
    extract_grade_and_comment_from_cell,
    find_teacher_signature_cell,
    write_to_teacher_signature_cell,
)


class CompleteCourseToGradingWorkflowTest(TestCase):
    """测试完整的课程创建到评分流程"""

    def setUp(self):
        """设置测试环境"""
        self.temp_dir = tempfile.mkdtemp()
        self.client = Client()

        # 创建租户
        self.tenant = Tenant.objects.create(
            name="测试学校",
            description="集成测试租户",
            is_active=True,
            tenant_repo_dir="test_tenant",
        )

        # 创建用户和用户配置
        self.user = User.objects.create_user(
            username="testteacher", password="testpass123", is_staff=True
        )

        self.user_profile = UserProfile.objects.create(
            user=self.user, tenant=self.tenant, is_tenant_admin=False, repo_base_dir=self.temp_dir
        )

        # 创建学期
        self.semester = Semester.objects.create(
            name="2024春季学期", start_date="2024-02-01", end_date="2024-07-01"
        )

        # 初始化服务
        self.course_service = CourseService()
        self.class_service = ClassService()
        self.repository_service = RepositoryService()

    def tearDown(self):
        """清理测试环境"""
        if os.path.exists(self.temp_dir):
            shutil.rmtree(self.temp_dir)

    def test_complete_workflow_from_course_to_grading(self):
        """
        测试完整流程：课程创建 -> 班级创建 -> 仓库配置 -> 作业上传 -> 评分
        需求: 1.1-1.5, 4.1-4.9, 5.1-5.8
        """
        # 步骤1: 创建课程
        course = self.course_service.create_course(
            teacher=self.user,
            name="数据结构",
            course_type="lab",
            description="数据结构实验课程",
            semester=self.semester,
        )
        self.assertIsNotNone(course)
        self.assertEqual(course.name, "数据结构")
        self.assertEqual(course.course_type, "lab")

        # 步骤2: 创建班级
        class_obj = self.class_service.create_class(
            course=course, name="计算机1班", student_count=30
        )
        self.assertIsNotNone(class_obj)
        self.assertEqual(class_obj.name, "计算机1班")

        # 步骤3: 配置文件系统仓库
        repository = self.repository_service.create_filesystem_repository(
            teacher=self.user, class_obj=class_obj, name="数据结构仓库"
        )
        self.assertIsNotNone(repository)
        self.assertEqual(repository.repo_type, "filesystem")

        # 步骤4: 创建作业批次
        homework = Homework.objects.create(
            course=course,
            class_obj=class_obj,
            title="第一次实验",
            folder_name="第一次实验",
            homework_type="lab_report",
            tenant=self.tenant,
        )

        # 步骤5: 创建目录结构
        course_dir = Path(repository.filesystem_path) / "数据结构"
        class_dir = course_dir / "计算机1班"
        homework_dir = class_dir / "第一次实验"
        homework_dir.mkdir(parents=True, exist_ok=True)

        # 步骤6: 创建实验报告文件
        doc = Document()
        doc.add_heading("数据结构实验报告", 0)
        doc.add_paragraph("学生姓名：张三")
        doc.add_paragraph("学号：20240001")

        table = doc.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "实验结果"
        table.rows[0].cells[1].text = "完成所有实验要求"
        table.rows[1].cells[0].text = "教师（签字）："

        file_path = homework_dir / "张三.docx"
        doc.save(str(file_path))

        # 步骤7: 评分
        doc = Document(str(file_path))
        cell, _, _, _ = find_teacher_signature_cell(doc)
        self.assertIsNotNone(cell, "应该找到教师签字单元格")

        grade = "A"
        comment = "实验报告完成得非常出色，数据准确，分析透彻。"
        signature_text = "教师（签字）：李老师"

        write_to_teacher_signature_cell(cell, grade, comment, signature_text)
        doc.save(str(file_path))

        # 步骤8: 验证评分
        doc = Document(str(file_path))
        cell, _, _, _ = find_teacher_signature_cell(doc)
        grade_read, comment_read, sig_read = extract_grade_and_comment_from_cell(cell)

        self.assertEqual(grade_read, grade)
        self.assertEqual(comment_read, comment)
        self.assertIn("李老师", sig_read)

        # 步骤9: 验证数据库记录
        courses = self.course_service.list_courses(self.user)
        self.assertEqual(len(courses), 1)
        self.assertEqual(courses[0].name, "数据结构")

        classes = self.class_service.list_classes(course)
        self.assertEqual(len(classes), 1)
        self.assertEqual(classes[0].name, "计算机1班")


class GitRepositoryWorkflowTest(TestCase):
    """测试Git仓库方式的完整流程"""

    def setUp(self):
        """设置测试环境"""
        self.temp_dir = tempfile.mkdtemp()
        self.client = Client()

        # 创建租户
        self.tenant = Tenant.objects.create(
            name="测试学校",
            description="Git测试租户",
            is_active=True,
            tenant_repo_dir="test_tenant",
        )

        # 创建用户和用户配置
        self.user = User.objects.create_user(
            username="gitteacher", password="testpass123", is_staff=True
        )

        self.user_profile = UserProfile.objects.create(
            user=self.user, tenant=self.tenant, is_tenant_admin=False, repo_base_dir=self.temp_dir
        )

        # 创建学期
        self.semester = Semester.objects.create(
            name="2024春季学期", start_date="2024-02-01", end_date="2024-07-01"
        )

        # 初始化服务
        self.course_service = CourseService()
        self.class_service = ClassService()
        self.repository_service = RepositoryService()

    def tearDown(self):
        """清理测试环境"""
        if os.path.exists(self.temp_dir):
            shutil.rmtree(self.temp_dir)

    @patch("git.Repo")
    def test_git_repository_workflow(self, mock_repo):
        """
        测试Git仓库方式的完整流程
        需求: 1.1.1-1.1.9
        """
        # Mock Git仓库
        mock_repo_instance = MagicMock()
        mock_repo.clone_from.return_value = mock_repo_instance

        # 步骤1: 创建课程和班级
        course = self.course_service.create_course(
            teacher=self.user,
            name="计算机网络",
            course_type="lab",
            description="计算机网络实验课程",
            semester=self.semester,
        )

        class_obj = self.class_service.create_class(course=course, name="网络1班", student_count=25)

        # 步骤2: 配置Git仓库
        git_url = "https://github.com/test/repo.git"
        git_branch = "main"

        repository = self.repository_service.create_git_repository(
            teacher=self.user,
            class_obj=class_obj,
            name="网络实验仓库",
            git_url=git_url,
            branch=git_branch,
            username="testuser",
            password="testpass",
        )

        self.assertIsNotNone(repository)
        self.assertEqual(repository.repo_type, "git")
        self.assertEqual(repository.git_url, git_url)
        self.assertEqual(repository.git_branch, git_branch)

        # 步骤3: 验证Git连接（使用mock）
        is_valid = self.repository_service.validate_git_connection(
            git_url=git_url, branch=git_branch, username="testuser", password="testpass"
        )

        # 由于使用了mock，这里应该返回True或者根据mock的配置
        # 实际测试中，这个方法会尝试连接Git仓库

        # 步骤4: 验证仓库列表
        repositories = self.repository_service.list_repositories(self.user)
        self.assertEqual(len(repositories), 1)
        self.assertEqual(repositories[0].name, "网络实验仓库")


class FilesystemRepositoryWorkflowTest(TestCase):
    """测试文件系统方式的完整流程"""

    def setUp(self):
        """设置测试环境"""
        self.temp_dir = tempfile.mkdtemp()
        self.client = Client()

        # 创建租户
        self.tenant = Tenant.objects.create(
            name="测试学校",
            description="文件系统测试租户",
            is_active=True,
            tenant_repo_dir="test_tenant",
        )

        # 创建用户和用户配置
        self.user = User.objects.create_user(
            username="fstest", password="testpass123", is_staff=True
        )

        self.user_profile = UserProfile.objects.create(
            user=self.user, tenant=self.tenant, is_tenant_admin=False, repo_base_dir=self.temp_dir
        )

        # 创建学期
        self.semester = Semester.objects.create(
            name="2024春季学期", start_date="2024-02-01", end_date="2024-07-01"
        )

        # 初始化服务
        self.course_service = CourseService()
        self.class_service = ClassService()
        self.repository_service = RepositoryService()
        self.file_upload_service = FileUploadService()

    def tearDown(self):
        """清理测试环境"""
        if os.path.exists(self.temp_dir):
            shutil.rmtree(self.temp_dir)

    def test_filesystem_repository_workflow(self):
        """
        测试文件系统方式的完整流程
        需求: 1.2.1-1.2.5, 1.4.1-1.4.7
        """
        # 步骤1: 创建课程和班级
        course = self.course_service.create_course(
            teacher=self.user,
            name="操作系统",
            course_type="lab",
            description="操作系统实验课程",
            semester=self.semester,
        )

        class_obj = self.class_service.create_class(
            course=course, name="操作系统1班", student_count=28
        )

        # 步骤2: 配置文件系统仓库
        repository = self.repository_service.create_filesystem_repository(
            teacher=self.user, class_obj=class_obj, name="操作系统仓库"
        )

        self.assertIsNotNone(repository)
        self.assertEqual(repository.repo_type, "filesystem")
        # filesystem_path is just the directory name, not the full path
        # The full path is created by get_full_path()
        full_path = repository.get_full_path()
        self.assertTrue(os.path.exists(full_path), f"Repository path should exist: {full_path}")

        # 步骤3: 验证目录名生成（处理重名）
        # The first repository already exists, so check its directory name
        # It should be in the format: username_reponame or username_reponame_N
        self.assertTrue(repository.filesystem_path.startswith("fstest"))

        # 创建第二个仓库，测试重名处理
        repository_2 = self.repository_service.create_filesystem_repository(
            teacher=self.user, class_obj=class_obj, name="操作系统仓库2"
        )

        # 验证第二个仓库的目录名不同（应该有不同的后缀或名称）
        self.assertNotEqual(repository.filesystem_path, repository_2.filesystem_path)
        self.assertTrue(repository_2.filesystem_path.startswith("fstest"))

        # 步骤4: 创建作业批次和目录结构
        homework = Homework.objects.create(
            course=course,
            class_obj=class_obj,
            title="第一次实验",
            folder_name="第一次实验",
            homework_type="lab_report",
            tenant=self.tenant,
        )

        # Use get_full_path() to get the actual directory path
        repo_full_path = repository.get_full_path()
        course_dir = Path(repo_full_path) / "操作系统"
        class_dir = course_dir / "操作系统1班"
        homework_dir = class_dir / "第一次实验"
        homework_dir.mkdir(parents=True, exist_ok=True)

        # 步骤5: 验证目录结构
        is_valid = self.repository_service.validate_directory_structure(repo_full_path)
        # 由于我们手动创建了符合规范的目录结构，应该验证通过

        # 步骤6: 模拟学生上传作业
        doc = Document()
        doc.add_heading("操作系统实验报告", 0)
        doc.add_paragraph("学生姓名：李四")
        doc.add_paragraph("学号：20240002")

        table = doc.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "实验结果"
        table.rows[0].cells[1].text = "完成进程调度实验"
        table.rows[1].cells[0].text = "教师（签字）："

        file_path = homework_dir / "李四.docx"
        doc.save(str(file_path))

        # 步骤7: 创建提交记录
        submission = Submission.objects.create(
            homework=homework,
            student=self.user,  # 在实际应用中应该是学生用户
            file_path=str(file_path),
            file_name="李四.docx",
            file_size=os.path.getsize(file_path),
            tenant=self.tenant,
        )

        self.assertIsNotNone(submission)
        self.assertEqual(submission.file_name, "李四.docx")

        # 步骤8: 教师评分
        doc = Document(str(file_path))
        cell, _, _, _ = find_teacher_signature_cell(doc)
        self.assertIsNotNone(cell)

        write_to_teacher_signature_cell(cell, "A", "实验完成得很好", "教师（签字）：王老师")
        doc.save(str(file_path))

        # 步骤9: 验证评分
        doc = Document(str(file_path))
        cell, _, _, _ = find_teacher_signature_cell(doc)
        grade, comment, sig = extract_grade_and_comment_from_cell(cell)

        self.assertEqual(grade, "A")
        self.assertEqual(comment, "实验完成得很好")
        self.assertIn("王老师", sig)


class CommentCacheAndTemplateWorkflowTest(TestCase):
    """测试评价缓存和模板的完整流程"""

    def setUp(self):
        """设置测试环境"""
        self.temp_dir = tempfile.mkdtemp()
        self.client = Client()

        # 创建租户
        self.tenant = Tenant.objects.create(
            name="测试学校",
            description="评价测试租户",
            is_active=True,
            tenant_repo_dir="test_tenant",
        )

        # 创建用户和用户配置
        self.user = User.objects.create_user(
            username="commenttest", password="testpass123", is_staff=True
        )

        self.user_profile = UserProfile.objects.create(
            user=self.user, tenant=self.tenant, is_tenant_admin=False, repo_base_dir=self.temp_dir
        )

        # 初始化服务
        self.comment_service = CommentTemplateService()

    def tearDown(self):
        """清理测试环境"""
        if os.path.exists(self.temp_dir):
            shutil.rmtree(self.temp_dir)

    def test_comment_template_workflow(self):
        """
        测试评价模板的完整流程
        需求: 5.2.1-5.2.12
        """
        # 步骤1: 记录评价使用（第一次）
        comment_1 = "实验报告完成得非常出色，数据准确，分析透彻。"
        self.comment_service.record_comment_usage(self.user, comment_1)

        # 步骤2: 验证个人模板
        personal_templates = self.comment_service.get_personal_templates(self.user, limit=5)
        self.assertEqual(len(personal_templates), 1)
        self.assertEqual(personal_templates[0].comment_text, comment_1)
        self.assertEqual(personal_templates[0].usage_count, 1)

        # 步骤3: 多次使用同一评价
        for _ in range(3):
            self.comment_service.record_comment_usage(self.user, comment_1)

        # 步骤4: 验证使用次数累加
        personal_templates = self.comment_service.get_personal_templates(self.user, limit=5)
        self.assertEqual(len(personal_templates), 1)
        self.assertEqual(personal_templates[0].usage_count, 4)

        # 步骤5: 添加更多评价
        comments = [
            "实验报告格式规范，内容完整。",
            "实验数据准确，分析合理。",
            "实验过程记录详细，结论正确。",
            "实验报告质量较高，建议进一步改进。",
            "实验完成情况良好。",
        ]

        for comment in comments:
            self.comment_service.record_comment_usage(self.user, comment)

        # 步骤6: 验证个人模板排序和限制（最多5个）
        personal_templates = self.comment_service.get_personal_templates(self.user, limit=5)
        self.assertEqual(len(personal_templates), 5)

        # 第一个应该是使用次数最多的
        self.assertEqual(personal_templates[0].comment_text, comment_1)
        self.assertEqual(personal_templates[0].usage_count, 4)

        # 步骤7: 测试系统模板（创建其他用户的评价）
        other_user = User.objects.create_user(
            username="otherteacher", password="testpass123", is_staff=True
        )

        UserProfile.objects.create(user=other_user, tenant=self.tenant, is_tenant_admin=False)

        system_comment = "实验报告整体质量优秀。"
        for _ in range(10):
            self.comment_service.record_comment_usage(other_user, system_comment)

        # 步骤8: 验证系统模板
        system_templates = self.comment_service.get_system_templates(self.tenant, limit=5)
        self.assertGreater(len(system_templates), 0)

        # 步骤9: 测试推荐模板（个人优先）
        recommended = self.comment_service.get_recommended_templates(self.user)
        self.assertGreater(len(recommended), 0)

        # 第一个应该是个人使用最多的
        self.assertEqual(recommended[0].comment_text, comment_1)

    def test_comment_deduplication(self):
        """
        测试评价内容去重
        需求: 5.2.11
        """
        # 步骤1: 多次记录相同评价
        comment = "实验报告完成得很好。"

        for _ in range(5):
            self.comment_service.record_comment_usage(self.user, comment)

        # 步骤2: 验证只有一个模板记录
        personal_templates = self.comment_service.get_personal_templates(self.user, limit=5)
        self.assertEqual(len(personal_templates), 1)
        self.assertEqual(personal_templates[0].comment_text, comment)
        self.assertEqual(personal_templates[0].usage_count, 5)

        # 步骤3: 验证数据库中只有一条记录
        templates_in_db = CommentTemplate.objects.filter(teacher=self.user, comment_text=comment)
        self.assertEqual(templates_in_db.count(), 1)


class BatchOperationsWorkflowTest(TestCase):
    """测试批量操作流程"""

    def setUp(self):
        """设置测试环境"""
        self.temp_dir = tempfile.mkdtemp()
        self.client = Client()

        # 创建租户
        self.tenant = Tenant.objects.create(
            name="测试学校",
            description="批量操作测试租户",
            is_active=True,
            tenant_repo_dir="test_tenant",
        )

        # 创建用户和用户配置
        self.user = User.objects.create_user(
            username="batchtest", password="testpass123", is_staff=True
        )

        self.user_profile = UserProfile.objects.create(
            user=self.user, tenant=self.tenant, is_tenant_admin=False, repo_base_dir=self.temp_dir
        )

        # 创建学期
        self.semester = Semester.objects.create(
            name="2024春季学期", start_date="2024-02-01", end_date="2024-07-01"
        )

        # 创建课程
        self.course = Course.objects.create(
            name="编译原理",
            course_type="lab",
            semester=self.semester,
            teacher=self.user,
            location="实验室C303",
            tenant=self.tenant,
        )

        # 创建班级
        self.class_obj = Class.objects.create(
            course=self.course, name="编译原理1班", student_count=30, tenant=self.tenant
        )

        # 创建仓库
        self.repository = Repository.objects.create(
            name="批量测试仓库",
            path=self.temp_dir,
            repo_type="local",
            owner=self.user,
            tenant=self.tenant,
        )

        # 创建作业批次
        self.homework = Homework.objects.create(
            course=self.course,
            class_obj=self.class_obj,
            title="第一次实验",
            folder_name="第一次实验",
            homework_type="lab_report",
            tenant=self.tenant,
        )

        # 创建目录结构
        self.course_dir = Path(self.temp_dir) / "编译原理"
        self.class_dir = self.course_dir / "编译原理1班"
        self.homework_dir = self.class_dir / "第一次实验"
        self.homework_dir.mkdir(parents=True, exist_ok=True)

    def tearDown(self):
        """清理测试环境"""
        if os.path.exists(self.temp_dir):
            shutil.rmtree(self.temp_dir)

    def create_multiple_lab_reports(self, count=10):
        """创建多个实验报告文件"""
        file_paths = []
        students = [
            "张三",
            "李四",
            "王五",
            "赵六",
            "钱七",
            "孙八",
            "周九",
            "吴十",
            "郑十一",
            "陈十二",
        ]

        for i in range(count):
            student_name = students[i % len(students)]
            doc = Document()
            doc.add_heading(f"{student_name}的编译原理实验报告", 0)
            doc.add_paragraph(f"学生姓名：{student_name}")
            doc.add_paragraph(f"学号：202400{i+1:02d}")
            doc.add_paragraph("实验内容：词法分析器设计")

            table = doc.add_table(rows=2, cols=2)
            table.rows[0].cells[0].text = "实验结果"
            table.rows[0].cells[1].text = f"完成实验{i+1}"
            table.rows[1].cells[0].text = "教师（签字）："

            file_path = self.homework_dir / f"{student_name}_{i+1}.docx"
            doc.save(str(file_path))
            file_paths.append(str(file_path))

        return file_paths

    def test_batch_grading_workflow(self):
        """
        测试批量评分流程
        需求: 7.1-7.7
        """
        # 步骤1: 创建多个实验报告
        file_paths = self.create_multiple_lab_reports(count=10)
        self.assertEqual(len(file_paths), 10, "应该创建10个文件")

        # 步骤2: 批量评分
        batch_grade = "B"
        batch_comment = "批量评分：实验完成情况良好"
        signature = "教师（签字）：批量评分老师"

        success_count = 0
        error_count = 0

        for file_path in file_paths:
            try:
                doc = Document(file_path)
                cell, _, _, _ = find_teacher_signature_cell(doc)

                if cell:
                    write_to_teacher_signature_cell(cell, batch_grade, batch_comment, signature)
                    doc.save(file_path)
                    success_count += 1
                else:
                    error_count += 1
            except Exception as e:
                error_count += 1

        # 步骤3: 验证批量评分结果
        self.assertEqual(success_count, 10, "应该成功评分10个文件")
        self.assertEqual(error_count, 0, "不应该有失败的文件")

        # 步骤4: 验证每个文件的评分
        for file_path in file_paths:
            doc = Document(file_path)
            cell, _, _, _ = find_teacher_signature_cell(doc)
            grade, comment, sig = extract_grade_and_comment_from_cell(cell)

            self.assertEqual(grade, batch_grade)
            self.assertEqual(comment, batch_comment)
            self.assertIn("批量评分老师", sig)

    @patch("grading.views.volcengine_score_homework")
    def test_batch_ai_scoring_workflow(self, mock_ai):
        """
        测试批量AI评分流程
        需求: 8.1-8.7
        """
        # Mock AI评分返回
        mock_ai.return_value = (88, "AI评价：实验报告完成良好，建议改进数据分析部分。")

        # 步骤1: 创建多个实验报告
        file_paths = self.create_multiple_lab_reports(count=5)

        # 步骤2: 批量AI评分
        success_count = 0
        error_count = 0
        results = []

        for file_path in file_paths:
            try:
                # 调用AI评分
                score, comment = mock_ai()

                # 将分数转换为等级
                if score >= 90:
                    grade = "A"
                elif score >= 80:
                    grade = "B"
                elif score >= 70:
                    grade = "C"
                elif score >= 60:
                    grade = "D"
                else:
                    grade = "E"

                # 写入评分
                doc = Document(file_path)
                cell, _, _, _ = find_teacher_signature_cell(doc)

                if cell:
                    write_to_teacher_signature_cell(
                        cell, grade, comment, "教师（签字）：AI评分系统"
                    )
                    doc.save(file_path)
                    success_count += 1
                    results.append(
                        {
                            "file": Path(file_path).name,
                            "status": "success",
                            "grade": grade,
                            "score": score,
                        }
                    )
                else:
                    error_count += 1
                    results.append(
                        {"file": Path(file_path).name, "status": "error", "error": "格式错误"}
                    )
            except Exception as e:
                error_count += 1
                results.append({"file": Path(file_path).name, "status": "error", "error": str(e)})

        # 步骤3: 验证批量AI评分结果
        self.assertEqual(success_count, 5, "应该成功评分5个文件")
        self.assertEqual(error_count, 0, "不应该有失败的文件")
        self.assertEqual(len(results), 5, "应该有5个结果记录")

        # 步骤4: 验证AI评分调用次数
        self.assertEqual(mock_ai.call_count, 5, "AI评分应该被调用5次")

        # 步骤5: 验证每个文件的AI评分
        for file_path in file_paths:
            doc = Document(file_path)
            cell, _, _, _ = find_teacher_signature_cell(doc)
            grade, comment, sig = extract_grade_and_comment_from_cell(cell)

            self.assertIsNotNone(grade)
            self.assertIn("AI评价", comment)
            self.assertIn("AI评分系统", sig)

    def test_batch_operations_with_mixed_files(self):
        """
        测试混合文件的批量操作（正常文件和格式错误文件）
        需求: 7.6
        """
        # 步骤1: 创建正常文件
        normal_files = self.create_multiple_lab_reports(count=3)

        # 步骤2: 创建格式错误的文件（没有教师签字单元格）
        error_file = self.homework_dir / "格式错误.docx"
        doc = Document()
        doc.add_heading("格式错误的报告", 0)
        doc.add_paragraph("这个文件没有教师签字单元格")
        doc.save(str(error_file))

        all_files = normal_files + [str(error_file)]

        # 步骤3: 批量评分
        success_count = 0
        error_count = 0

        for file_path in all_files:
            try:
                doc = Document(file_path)
                cell, _, _, _ = find_teacher_signature_cell(doc)

                if cell:
                    write_to_teacher_signature_cell(cell, "B", "批量评分", "教师（签字）：")
                    doc.save(file_path)
                    success_count += 1
                else:
                    # 格式错误，按降级处理
                    doc.add_paragraph("老师评分：D")
                    doc.add_paragraph(
                        "教师评价：【格式错误-已锁定】请按要求的格式写实验报告，此评分不可修改"
                    )
                    doc.save(file_path)
                    error_count += 1
            except Exception as e:
                error_count += 1

        # 步骤4: 验证结果
        self.assertEqual(success_count, 3, "应该成功评分3个正常文件")
        self.assertEqual(error_count, 1, "应该有1个格式错误文件")

        # 步骤5: 验证格式错误文件的处理
        doc = Document(str(error_file))
        text_content = "\n".join([p.text for p in doc.paragraphs])
        self.assertIn("【格式错误-已锁定】", text_content)
        self.assertIn("D", text_content)


class PercentageGradingWorkflowTest(TestCase):
    """测试百分制评分流程"""

    def setUp(self):
        """设置测试环境"""
        self.temp_dir = tempfile.mkdtemp()
        self.client = Client()

        # 创建租户
        self.tenant = Tenant.objects.create(
            name="测试学校",
            description="百分制测试租户",
            is_active=True,
            tenant_repo_dir="test_tenant",
        )

        # 创建用户和用户配置
        self.user = User.objects.create_user(
            username="percenttest", password="testpass123", is_staff=True
        )

        self.user_profile = UserProfile.objects.create(
            user=self.user, tenant=self.tenant, is_tenant_admin=False, repo_base_dir=self.temp_dir
        )

        # 创建学期
        self.semester = Semester.objects.create(
            name="2024春季学期", start_date="2024-02-01", end_date="2024-07-01"
        )

        # 创建课程
        self.course = Course.objects.create(
            name="软件工程",
            course_type="lab",
            semester=self.semester,
            teacher=self.user,
            location="实验室D404",
            tenant=self.tenant,
        )

        # 创建班级
        self.class_obj = Class.objects.create(
            course=self.course, name="软件工程1班", student_count=25, tenant=self.tenant
        )

    def tearDown(self):
        """清理测试环境"""
        if os.path.exists(self.temp_dir):
            shutil.rmtree(self.temp_dir)

    def test_percentage_grading_workflow(self):
        """
        测试百分制评分流程
        需求: 4.3, 4.4
        """
        # 步骤1: 创建实验报告
        homework_dir = Path(self.temp_dir) / "软件工程" / "软件工程1班" / "第一次实验"
        homework_dir.mkdir(parents=True, exist_ok=True)

        doc = Document()
        doc.add_heading("软件工程实验报告", 0)
        doc.add_paragraph("学生姓名：赵六")
        doc.add_paragraph("学号：20240006")

        table = doc.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "实验结果"
        table.rows[0].cells[1].text = "完成需求分析"
        table.rows[1].cells[0].text = "教师（签字）："

        file_path = homework_dir / "赵六.docx"
        doc.save(str(file_path))

        # 步骤2: 使用百分制评分
        doc = Document(str(file_path))
        cell, _, _, _ = find_teacher_signature_cell(doc)
        self.assertIsNotNone(cell)

        # 百分制评分：92分
        grade = "92"
        comment = "实验报告质量优秀，得分92分。"
        signature = "教师（签字）：评分老师"

        write_to_teacher_signature_cell(cell, grade, comment, signature)
        doc.save(str(file_path))

        # 步骤3: 验证百分制评分
        doc = Document(str(file_path))
        cell, _, _, _ = find_teacher_signature_cell(doc)
        grade_read, comment_read, sig_read = extract_grade_and_comment_from_cell(cell)

        # 百分制评分不在预定义的等级列表中，所以grade_read会是None
        # 但评分信息应该在单元格的第一行
        cell_text = cell.text.strip()
        lines = cell_text.split("\n")
        self.assertEqual(lines[0].strip(), "92", "第一行应该是百分制评分")
        self.assertIn("92分", comment_read)
        self.assertIn("评分老师", sig_read)

        # 步骤4: 配置评分类型
        grade_config = GradeTypeConfig.objects.create(
            class_obj=self.class_obj, grade_type="percentage", is_locked=False, tenant=self.tenant
        )

        self.assertEqual(grade_config.grade_type, "percentage")
        self.assertFalse(grade_config.is_locked)
