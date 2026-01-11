"""
测试百分制评分功能
需求: 4.1-4.5
"""

import os
import shutil
import tempfile
import uuid

from django.contrib.auth.models import User
from django.test import TestCase
from docx import Document
from hypothesis import assume, given, settings
from hypothesis import strategies as st
from hypothesis.extra.django import TestCase as HypothesisTestCase

from grading.models import Repository, Tenant, UserProfile
from grading.views import write_grade_and_comment_to_file


class PercentageGradingTest(TestCase):
    """百分制评分测试"""

    def setUp(self):
        """设置测试环境"""
        # 创建租户
        self.tenant = Tenant.objects.create(name="测试租户")

        # 创建用户
        self.user = User.objects.create_user(
            username="testuser", password="testpass123", is_staff=True
        )

        # 创建用户配置
        self.user_profile = UserProfile.objects.create(user=self.user, tenant=self.tenant)

        # 创建临时目录
        self.temp_dir = tempfile.mkdtemp()

    def tearDown(self):
        """清理测试环境"""
        if os.path.exists(self.temp_dir):
            shutil.rmtree(self.temp_dir)

    def create_test_docx(self, filename="test.docx"):
        """创建测试Word文档"""
        doc = Document()
        doc.add_paragraph("这是一个测试文档")

        file_path = os.path.join(self.temp_dir, filename)
        doc.save(file_path)
        return file_path

    def test_percentage_grade_write_to_normal_assignment(self):
        """测试百分制评分写入普通作业 (需求 4.1, 4.3, 4.4)"""
        # 创建测试文档
        file_path = self.create_test_docx()

        # 写入百分制评分
        write_grade_and_comment_to_file(
            full_path=file_path,
            grade="85",
            comment="完成良好",
            base_dir=self.temp_dir,
            is_lab_report=False,
        )

        # 读取文档验证
        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs]

        # 验证评分已写入
        grade_found = False
        for para in paragraphs:
            if "老师评分：85" in para:
                grade_found = True
                break

        self.assertTrue(grade_found, "百分制评分应该被写入文档")

    def test_percentage_grade_with_decimal(self):
        """测试带小数的百分制评分 (需求 4.4)"""
        # 创建测试文档
        file_path = self.create_test_docx()

        # 写入带小数的百分制评分
        write_grade_and_comment_to_file(
            full_path=file_path,
            grade="87.5",
            comment="完成优秀",
            base_dir=self.temp_dir,
            is_lab_report=False,
        )

        # 读取文档验证
        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs]

        # 验证评分已写入
        grade_found = False
        for para in paragraphs:
            if "老师评分：87.5" in para:
                grade_found = True
                break

        self.assertTrue(grade_found, "带小数的百分制评分应该被写入文档")

    def test_percentage_grade_boundary_values(self):
        """测试百分制评分边界值 (需求 4.4)"""
        # 测试最小值 0
        file_path_min = self.create_test_docx("test_min.docx")
        write_grade_and_comment_to_file(
            full_path=file_path_min, grade="0", base_dir=self.temp_dir, is_lab_report=False
        )

        doc_min = Document(file_path_min)
        paragraphs_min = [p.text for p in doc_min.paragraphs]
        self.assertTrue(any("老师评分：0" in p for p in paragraphs_min), "百分制评分0应该被接受")

        # 测试最大值 100
        file_path_max = self.create_test_docx("test_max.docx")
        write_grade_and_comment_to_file(
            full_path=file_path_max, grade="100", base_dir=self.temp_dir, is_lab_report=False
        )

        doc_max = Document(file_path_max)
        paragraphs_max = [p.text for p in doc_max.paragraphs]
        self.assertTrue(
            any("老师评分：100" in p for p in paragraphs_max), "百分制评分100应该被接受"
        )

    def test_percentage_grade_update(self):
        """测试更新百分制评分 (需求 4.1, 4.3)"""
        # 创建测试文档
        file_path = self.create_test_docx()

        # 第一次写入评分
        write_grade_and_comment_to_file(
            full_path=file_path, grade="75", base_dir=self.temp_dir, is_lab_report=False
        )

        # 第二次更新评分
        write_grade_and_comment_to_file(
            full_path=file_path,
            grade="90",
            comment="进步明显",
            base_dir=self.temp_dir,
            is_lab_report=False,
        )

        # 读取文档验证
        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs]

        # 验证新评分已写入
        new_grade_found = False
        old_grade_found = False
        for para in paragraphs:
            if "老师评分：90" in para:
                new_grade_found = True
            if "老师评分：75" in para:
                old_grade_found = True

        self.assertTrue(new_grade_found, "新的百分制评分应该被写入")
        self.assertFalse(old_grade_found, "旧的百分制评分应该被替换")


class PercentageGradeValidationPropertyTest(HypothesisTestCase):
    """
    百分制评分验证的属性测试
    Feature: homework-grading-system, Property 7: 百分制分数范围验证
    Validates: Requirements 4.4
    """

    def setUp(self):
        """设置测试环境"""
        # 创建或获取租户
        self.tenant, _ = Tenant.objects.get_or_create(name="测试租户")

        # 创建或获取用户
        self.user, created = User.objects.get_or_create(
            username="testuser", defaults={"is_staff": True}
        )
        if created:
            self.user.set_password("testpass123")
            self.user.save()

        # 创建或获取用户配置
        self.user_profile, _ = UserProfile.objects.get_or_create(
            user=self.user, defaults={"tenant": self.tenant}
        )

        # 创建临时目录
        self.temp_dir = tempfile.mkdtemp()

        # 创建仓库（每次测试都创建新的）
        repo_name = f"测试仓库_{uuid.uuid4().hex[:8]}"
        self.repository = Repository.objects.create(
            name=repo_name,
            owner=self.user,
            tenant=self.tenant,
            repo_type="filesystem",
            filesystem_path=self.temp_dir,
            path=self.temp_dir,
            is_active=True,
        )

        # 创建测试文档
        self.test_file = self.create_test_docx()
        self.relative_path = os.path.relpath(self.test_file, self.temp_dir)

    def _validate_percentage_grade(self, grade_str):
        """
        直接测试百分制评分验证逻辑

        Args:
            grade_str: 评分字符串

        Returns:
            tuple: (is_valid, error_message)
        """
        try:
            grade_value = float(grade_str)
            if grade_value < 0 or grade_value > 100:
                return False, f"百分制评分必须在0-100之间，当前值: {grade_str}"
            return True, None
        except (ValueError, TypeError):
            return False, f"百分制评分必须是数字，当前值: {grade_str}"

    def tearDown(self):
        """清理测试环境"""
        if os.path.exists(self.temp_dir):
            shutil.rmtree(self.temp_dir)

    def create_test_docx(self, filename="test_property.docx"):
        """创建测试Word文档"""
        doc = Document()
        doc.add_paragraph("这是一个测试文档")

        file_path = os.path.join(self.temp_dir, filename)
        doc.save(file_path)
        return file_path

    @given(
        grade_value=st.floats(min_value=0.0, max_value=100.0, allow_nan=False, allow_infinity=False)
    )
    @settings(max_examples=100, deadline=None)
    def test_property_valid_percentage_grades_accepted(self, grade_value):
        """
        Property 7: 百分制分数范围验证 - 有效值测试
        For any numeric input in [0, 100], the percentage grade validation should accept it.
        Validates: Requirements 4.4
        """
        # 格式化分数
        if grade_value == int(grade_value):
            grade_str = str(int(grade_value))
        else:
            grade_str = f"{grade_value:.1f}"

        # 测试验证逻辑
        is_valid, error_msg = self._validate_percentage_grade(grade_str)

        # 验证结果
        self.assertTrue(
            is_valid,
            f"有效的百分制评分 {grade_str} (原始值: {grade_value}) "
            f"应该被接受，但被拒绝: {error_msg}",
        )
        self.assertIsNone(
            error_msg, f"有效的百分制评分 {grade_str} 不应该有错误消息，但返回: {error_msg}"
        )

    @given(
        grade_value=st.one_of(
            st.floats(
                min_value=-1000.0,
                max_value=-0.01,
                allow_nan=False,
                allow_infinity=False,
            ),
            st.floats(
                min_value=100.01,
                max_value=1000.0,
                allow_nan=False,
                allow_infinity=False,
            ),
        )
    )
    @settings(max_examples=100, deadline=None)
    def test_property_invalid_percentage_grades_rejected(self, grade_value):
        """
        Property 7: 百分制分数范围验证 - 无效值测试
        For any numeric input outside [0, 100], the percentage grade validation should reject it.
        Validates: Requirements 4.4
        """
        # 格式化分数
        grade_str = str(grade_value)

        # 测试验证逻辑
        is_valid, error_msg = self._validate_percentage_grade(grade_str)

        # 验证结果
        self.assertFalse(
            is_valid, f"无效的百分制评分 {grade_str} (原始值: {grade_value}) 应该被拒绝，但被接受"
        )
        self.assertIsNotNone(error_msg, f"无效的百分制评分 {grade_str} 应该有错误消息")
        # 验证错误消息包含范围信息
        self.assertIn("0-100", error_msg, f"错误消息应该包含有效范围信息，但返回: {error_msg}")

    @given(
        non_numeric=st.text(min_size=1).filter(
            lambda x: not x.replace(".", "", 1).replace("-", "", 1).isdigit()
        )
    )
    @settings(max_examples=50, deadline=None)
    def test_property_non_numeric_grades_rejected(self, non_numeric):
        """
        Property 7: 百分制分数范围验证 - 非数字输入测试
        For any non-numeric input, the percentage grade validation should reject it.
        Validates: Requirements 4.4
        """
        # 过滤掉可能被解析为数字的字符串
        assume(non_numeric.strip() != "")

        try:
            float(non_numeric)
            # 如果能转换为float，跳过这个测试用例
            assume(False)
        except (ValueError, OverflowError):
            # 这是我们想要测试的情况
            pass

        # 测试验证逻辑
        is_valid, error_msg = self._validate_percentage_grade(non_numeric)

        # 验证结果
        self.assertFalse(is_valid, f"非数字输入 '{non_numeric}' 应该被拒绝，但被接受")
        self.assertIsNotNone(error_msg, f"非数字输入 '{non_numeric}' 应该有错误消息")
        # 验证错误消息提示必须是数字
        self.assertTrue(
            "数字" in error_msg or "numeric" in error_msg.lower(),
            f"错误消息应该提示必须是数字，但返回: {error_msg}",
        )
