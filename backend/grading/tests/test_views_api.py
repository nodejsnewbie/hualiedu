"""
Tests for API endpoints in grading/views.py

This module tests the API endpoints and AJAX views in the grading views,
including file content retrieval, course information, and homework management.
"""

import json
import os
import tempfile
import shutil
from unittest.mock import Mock, patch, MagicMock
from io import BytesIO

import pytest
from django.test import TestCase, RequestFactory, Client
from django.contrib.auth import get_user_model
from django.http import JsonResponse
from docx import Document

from grading.models import (
    Repository, Course, Homework, Semester, GlobalConfig, 
    Tenant, UserProfile
)
from grading.views import (
    get_file_content,
    get_course_info_api,
    get_homework_list_api,
    update_course_type_api,
    get_homework_info_api,
    get_courses_list_view,
    get_directory_tree_view,
    get_file_grade_info_api,
    save_teacher_comment,
    remove_grade,
)

User = get_user_model()


class BaseAPITestCase(TestCase):
    """Base test case for API tests."""
    
    def setUp(self):
        """Set up test data for API tests."""
        self.factory = RequestFactory()
        self.client = Client()
        
        # Create test user
        self.user = User.objects.create_user(
            username='testuser',
            email='test@example.com',
            password='testpass123',
            is_staff=True
        )
        
        # Create tenant and user profile
        self.tenant = Tenant.objects.create(
            name='Test Tenant',
            domain='test.example.com'
        )
        
        self.user_profile = UserProfile.objects.create(
            user=self.user,
            tenant=self.tenant,
            repo_base_dir='~/test_jobs'
        )
        
        # Create test semester
        self.semester = Semester.objects.create(
            name='Test Semester',
            start_date='2024-01-01',
            end_date='2024-06-30',
            is_active=True
        )
        
        # Create test course
        self.course = Course.objects.create(
            name='Test Course',
            course_type='theory',
            semester=self.semester,
            teacher=self.user,
            location='Room 101'
        )
        
        # Create test repository
        self.repository = Repository.objects.create(
            name='test-repo',
            url='https://github.com/test/repo.git',
            owner=self.user,
            tenant=self.tenant,
            repo_type='local',
            local_path='/tmp/test-repo'
        )
        
        # Create temporary directory
        self.temp_dir = tempfile.mkdtemp()
        self.addCleanup(shutil.rmtree, self.temp_dir)
        
        # Set up global config
        GlobalConfig.objects.create(
            key='default_repo_base_dir',
            value=self.temp_dir
        )


class TestGetFileContentAPI(BaseAPITestCase):
    """Test get_file_content API endpoint."""
    
    def test_get_file_content_text_file(self):
        """Test getting content of a text file."""
        # Create test file
        test_file = os.path.join(self.temp_dir, 'test.txt')
        test_content = 'Hello, World!\nThis is a test file.'
        
        with open(test_file, 'w', encoding='utf-8') as f:
            f.write(test_content)
        
        request = self.factory.post('/get_file_content/', {
            'path': 'test.txt'
        })
        request.user = self.user
        
        response = get_file_content(request)
        
        self.assertEqual(response.status_code, 200)
        data = json.loads(response.content)
        self.assertEqual(data['status'], 'success')
        self.assertEqual(data['type'], 'text')
        self.assertEqual(data['content'], test_content)
    
    def test_get_file_content_docx_file(self):
        """Test getting content of a DOCX file."""
        # Create test DOCX file
        test_file = os.path.join(self.temp_dir, 'test.docx')
        
        doc = Document()
        doc.add_paragraph('First paragraph')
        doc.add_paragraph('Second paragraph')
        doc.save(test_file)
        
        request = self.factory.post('/get_file_content/', {
            'path': 'test.docx'
        })
        request.user = self.user
        
        response = get_file_content(request)
        
        self.assertEqual(response.status_code, 200)
        data = json.loads(response.content)
        self.assertEqual(data['status'], 'success')
        self.assertEqual(data['type'], 'docx')
        self.assertIn('First paragraph', data['content'])
        self.assertIn('Second paragraph', data['content'])
    
    def test_get_file_content_missing_path(self):
        """Test get_file_content with missing path parameter."""
        request = self.factory.post('/get_file_content/', {})
        request.user = self.user
        
        response = get_file_content(request)
        
        self.assertEqual(response.status_code, 200)
        data = json.loads(response.content)
        self.assertEqual(data['status'], 'error')
        self.assertEqual(data['message'], '未提供文件路径')
    
    def test_get_file_content_nonexistent_file(self):
        """Test get_file_content with nonexistent file."""
        request = self.factory.post('/get_file_content/', {
            'path': 'nonexistent.txt'
        })
        request.user = self.user
        
        response = get_file_content(request)
        
        self.assertEqual(response.status_code, 200)
        data = json.loads(response.content)
        self.assertEqual(data['status'], 'error')
        self.assertIn('文件不存在', data['message'])
    
    @patch('grading.views._build_git_adapter')
    def test_get_file_content_git_repository(self, mock_build_adapter):
        """Test get_file_content with Git repository."""
        # Set up Git repository
        self.repository.repo_type = 'git'
        self.repository.save()
        
        # Mock Git adapter
        mock_adapter = Mock()
        mock_adapter.read_file.return_value = b'Git file content'
        mock_build_adapter.return_value = mock_adapter
        
        request = self.factory.post('/get_file_content/', {
            'path': 'test.txt',
            'repo_id': str(self.repository.id),
            'course': 'Test Course'
        })
        request.user = self.user
        
        response = get_file_content(request)
        
        self.assertEqual(response.status_code, 200)
        data = json.loads(response.content)
        self.assertEqual(data['status'], 'success')
        self.assertEqual(data['type'], 'text')
        self.assertEqual(data['content'], 'Git file content')


class TestCourseInfoAPI(BaseAPITestCase):
    """Test course information API endpoints."""
    
    def test_get_course_info_api_existing_course(self):
        """Test get_course_info_api with existing course."""
        request = self.factory.get('/api/course_info/', {
            'course_name': 'Test Course'
        })
        request.user = self.user
        
        response = get_course_info_api(request)
        
        self.assertEqual(response.status_code, 200)
        data = json.loads(response.content)
        self.assertTrue(data['success'])
        self.assertEqual(data['course']['name'], 'Test Course')
        self.assertEqual(data['course']['course_type'], 'theory')
        self.assertTrue(data['course']['in_database'])
        self.assertFalse(data['course']['auto_created'])
    
    def test_get_course_info_api_nonexistent_course_auto_create(self):
        """Test get_course_info_api with nonexistent course and auto-create enabled."""
        request = self.factory.get('/api/course_info/', {
            'course_name': 'New Lab Course',
            'auto_create': 'true'
        })
        request.user = self.user
        
        with patch('grading.views.auto_create_or_update_course') as mock_create:
            # Mock successful course creation
            new_course = Course(
                name='New Lab Course',
                course_type='lab',
                semester=self.semester,
                teacher=self.user
            )
            new_course.id = 999
            mock_create.return_value = new_course
            
            response = get_course_info_api(request)
        
        self.assertEqual(response.status_code, 200)
        data = json.loads(response.content)
        self.assertTrue(data['success'])
        self.assertEqual(data['course']['name'], 'New Lab Course')
        self.assertEqual(data['course']['course_type'], 'lab')
        self.assertTrue(data['course']['auto_created'])
    
    def test_get_course_info_api_missing_course_name(self):
        """Test get_course_info_api with missing course name."""
        request = self.factory.get('/api/course_info/', {})
        request.user = self.user
        
        response = get_course_info_api(request)
        
        self.assertEqual(response.status_code, 200)
        data = json.loads(response.content)
        self.assertFalse(data['success'])
        self.assertEqual(data['message'], '未提供课程名称')
    
    def test_update_course_type_api_success(self):
        """Test update_course_type_api with valid parameters."""
        request = self.factory.post('/api/update_course_type/', {
            'course_name': 'Test Course',
            'course_type': 'lab'
        })
        request.user = self.user
        
        response = update_course_type_api(request)
        
        self.assertEqual(response.status_code, 200)
        data = json.loads(response.content)
        self.assertTrue(data['success'])
        self.assertEqual(data['course']['course_type'], 'lab')
        
        # Verify course was updated in database
        self.course.refresh_from_db()
        self.assertEqual(self.course.course_type, 'lab')
    
    def test_update_course_type_api_invalid_type(self):
        """Test update_course_type_api with invalid course type."""
        request = self.factory.post('/api/update_course_type/', {
            'course_name': 'Test Course',
            'course_type': 'invalid_type'
        })
        request.user = self.user
        
        response = update_course_type_api(request)
        
        self.assertEqual(response.status_code, 200)
        data = json.loads(response.content)
        self.assertFalse(data['success'])
        self.assertEqual(data['message'], '无效的课程类型')


class TestHomeworkAPI(BaseAPITestCase):
    """Test homework-related API endpoints."""
    
    def setUp(self):
        super().setUp()
        # Create test homework
        self.homework = Homework.objects.create(
            course=self.course,
            title='Test Homework',
            homework_type='normal',
            folder_name='homework1',
            description='Test homework description',
            tenant=self.tenant
        )
    
    def test_get_homework_list_api_success(self):
        """Test get_homework_list_api with valid course."""
        request = self.factory.get('/api/homework_list/', {
            'course_name': 'Test Course'
        })
        request.user = self.user
        
        response = get_homework_list_api(request)
        
        self.assertEqual(response.status_code, 200)
        data = json.loads(response.content)
        self.assertTrue(data['success'])
        self.assertEqual(data['course_name'], 'Test Course')
        self.assertEqual(len(data['homeworks']), 1)
        self.assertEqual(data['homeworks'][0]['title'], 'Test Homework')
        self.assertEqual(data['homeworks'][0]['homework_type'], 'normal')
    
    def test_get_homework_list_api_nonexistent_course(self):
        """Test get_homework_list_api with nonexistent course."""
        request = self.factory.get('/api/homework_list/', {
            'course_name': 'Nonexistent Course'
        })
        request.user = self.user
        
        response = get_homework_list_api(request)
        
        self.assertEqual(response.status_code, 200)
        data = json.loads(response.content)
        self.assertFalse(data['success'])
        self.assertEqual(data['message'], '课程不存在')
    
    def test_get_homework_info_api_existing_homework(self):
        """Test get_homework_info_api with existing homework."""
        request = self.factory.get('/api/homework_info/', {
            'course_name': 'Test Course',
            'homework_folder': 'homework1'
        })
        request.user = self.user
        
        response = get_homework_info_api(request)
        
        self.assertEqual(response.status_code, 200)
        data = json.loads(response.content)
        self.assertTrue(data['success'])
        self.assertEqual(data['homework']['title'], 'Test Homework')
        self.assertEqual(data['homework']['homework_type'], 'normal')
        self.assertFalse(data['homework']['is_lab_report'])
    
    def test_get_homework_info_api_auto_create(self):
        """Test get_homework_info_api with auto-create for nonexistent homework."""
        request = self.factory.get('/api/homework_info/', {
            'course_name': 'Test Course',
            'homework_folder': 'new_homework',
            'auto_create': 'true'
        })
        request.user = self.user
        
        response = get_homework_info_api(request)
        
        self.assertEqual(response.status_code, 200)
        data = json.loads(response.content)
        self.assertTrue(data['success'])
        self.assertEqual(data['homework']['folder_name'], 'new_homework')
        self.assertEqual(data['homework']['homework_type'], 'normal')
        
        # Verify homework was created in database
        new_homework = Homework.objects.get(folder_name='new_homework')
        self.assertEqual(new_homework.course, self.course)


class TestDirectoryAPI(BaseAPITestCase):
    """Test directory-related API endpoints."""
    
    def test_get_courses_list_view_local_repository(self):
        """Test get_courses_list_view with local repository."""
        # Create test directory structure
        repo_dir = os.path.join(self.temp_dir, 'test-repo')
        os.makedirs(repo_dir)
        
        # Create course directories
        course_dirs = ['Course1', 'Course2', 'Course3']
        for course_dir in course_dirs:
            os.makedirs(os.path.join(repo_dir, course_dir))
        
        # Update repository path
        self.repository.local_path = repo_dir
        self.repository.save()
        
        request = self.factory.get('/api/courses_list/', {
            'repo_id': str(self.repository.id)
        })
        request.user = self.user
        
        response = get_courses_list_view(request)
        
        self.assertEqual(response.status_code, 200)
        data = json.loads(response.content)
        self.assertEqual(data['status'], 'success')
        self.assertEqual(len(data['courses']), 3)
        
        course_names = [course['name'] for course in data['courses']]
        self.assertIn('Course1', course_names)
        self.assertIn('Course2', course_names)
        self.assertIn('Course3', course_names)
    
    @patch('grading.views._build_git_adapter')
    def test_get_courses_list_view_git_repository(self, mock_build_adapter):
        """Test get_courses_list_view with Git repository."""
        # Set up Git repository
        self.repository.repo_type = 'git'
        self.repository.save()
        
        # Mock Git adapter
        mock_adapter = Mock()
        mock_adapter.list_directory.return_value = [
            {'name': 'Course1', 'type': 'dir'},
            {'name': 'Course2', 'type': 'dir'},
            {'name': '.hidden', 'type': 'dir'},  # Should be filtered out
            {'name': 'file.txt', 'type': 'file'}  # Should be filtered out
        ]
        mock_build_adapter.return_value = mock_adapter
        
        request = self.factory.get('/api/courses_list/', {
            'repo_id': str(self.repository.id)
        })
        request.user = self.user
        
        response = get_courses_list_view(request)
        
        self.assertEqual(response.status_code, 200)
        data = json.loads(response.content)
        self.assertEqual(data['status'], 'success')
        self.assertEqual(len(data['courses']), 2)
        
        course_names = [course['name'] for course in data['courses']]
        self.assertIn('Course1', course_names)
        self.assertIn('Course2', course_names)
        self.assertNotIn('.hidden', course_names)
    
    def test_get_directory_tree_view_local_repository(self):
        """Test get_directory_tree_view with local repository."""
        # Create test directory structure
        repo_dir = os.path.join(self.temp_dir, 'test-repo')
        course_dir = os.path.join(repo_dir, 'TestCourse')
        homework_dir = os.path.join(course_dir, 'Homework1')
        os.makedirs(homework_dir)
        
        # Create test files
        with open(os.path.join(homework_dir, 'file1.docx'), 'w') as f:
            f.write('test')
        with open(os.path.join(homework_dir, 'file2.txt'), 'w') as f:
            f.write('test')
        
        # Update repository path
        self.repository.local_path = repo_dir
        self.repository.save()
        
        request = self.factory.get('/api/directory_tree/', {
            'repo_id': str(self.repository.id),
            'course': 'TestCourse',
            'path': ''
        })
        request.user = self.user
        
        with patch('grading.views.get_directory_tree') as mock_get_tree:
            mock_get_tree.return_value = [
                {
                    'id': 'Homework1',
                    'text': 'Homework1',
                    'type': 'folder',
                    'children': [
                        {'id': 'Homework1/file1.docx', 'text': 'file1.docx', 'type': 'file'},
                        {'id': 'Homework1/file2.txt', 'text': 'file2.txt', 'type': 'file'}
                    ]
                }
            ]
            
            response = get_directory_tree_view(request)
        
        self.assertEqual(response.status_code, 200)
        data = json.loads(response.content)
        self.assertIn('children', data)
        self.assertEqual(len(data['children']), 1)
        self.assertEqual(data['children'][0]['text'], 'Homework1')


class TestGradeAPI(BaseAPITestCase):
    """Test grade-related API endpoints."""
    
    def test_get_file_grade_info_api_success(self):
        """Test get_file_grade_info_api with valid file."""
        # Create test file with grade
        test_file = os.path.join(self.temp_dir, 'test.txt')
        with open(test_file, 'w', encoding='utf-8') as f:
            f.write('老师评分：A\n这是测试内容。')
        
        request = self.factory.get('/api/file_grade_info/', {
            'path': 'test.txt'
        })
        request.user = self.user
        
        response = get_file_grade_info_api(request)
        
        self.assertEqual(response.status_code, 200)
        data = json.loads(response.content)
        self.assertTrue(data['has_grade'])
        self.assertEqual(data['grade'], 'A')
        self.assertEqual(data['grade_type'], 'letter')
    
    def test_get_file_grade_info_api_no_grade(self):
        """Test get_file_grade_info_api with file without grade."""
        # Create test file without grade
        test_file = os.path.join(self.temp_dir, 'test.txt')
        with open(test_file, 'w', encoding='utf-8') as f:
            f.write('这是没有评分的测试内容。')
        
        request = self.factory.get('/api/file_grade_info/', {
            'path': 'test.txt'
        })
        request.user = self.user
        
        response = get_file_grade_info_api(request)
        
        self.assertEqual(response.status_code, 200)
        data = json.loads(response.content)
        self.assertFalse(data['has_grade'])
        self.assertEqual(data['grade'], '')
    
    def test_save_teacher_comment_success(self):
        """Test save_teacher_comment with valid parameters."""
        # Create test file
        test_file = os.path.join(self.temp_dir, 'test.txt')
        with open(test_file, 'w', encoding='utf-8') as f:
            f.write('这是测试内容。')
        
        request = self.factory.post('/save_teacher_comment/', {
            'file_path': 'test.txt',
            'comment': '很好的作业！',
            'grade': 'A'
        })
        request.user = self.user
        
        with patch('grading.views.write_grade_and_comment_to_file') as mock_write:
            mock_write.return_value = None  # No warning
            
            response = save_teacher_comment(request)
        
        self.assertEqual(response.status_code, 200)
        data = json.loads(response.content)
        self.assertTrue(data['success'])
        self.assertEqual(data['message'], '教师评价和评分已保存')
    
    def test_save_teacher_comment_missing_parameters(self):
        """Test save_teacher_comment with missing parameters."""
        request = self.factory.post('/save_teacher_comment/', {
            'file_path': 'test.txt'
            # Missing comment and grade
        })
        request.user = self.user
        
        response = save_teacher_comment(request)
        
        self.assertEqual(response.status_code, 200)
        data = json.loads(response.content)
        self.assertFalse(data['success'])
        self.assertEqual(data['message'], '缺少必要参数')
    
    def test_remove_grade_success(self):
        """Test remove_grade with valid file."""
        # Create test file with grade
        test_file = os.path.join(self.temp_dir, 'test.txt')
        with open(test_file, 'w', encoding='utf-8') as f:
            f.write('这是测试内容。\n老师评分：A\n教师评价：很好！')
        
        request = self.factory.post('/remove_grade/', {
            'path': 'test.txt'
        })
        request.user = self.user
        
        response = remove_grade(request)
        
        self.assertEqual(response.status_code, 200)
        data = json.loads(response.content)
        self.assertEqual(data['status'], 'success')
        self.assertIn('已删除', data['message'])
    
    def test_remove_grade_no_grade_found(self):
        """Test remove_grade with file that has no grade."""
        # Create test file without grade
        test_file = os.path.join(self.temp_dir, 'test.txt')
        with open(test_file, 'w', encoding='utf-8') as f:
            f.write('这是没有评分的测试内容。')
        
        request = self.factory.post('/remove_grade/', {
            'path': 'test.txt'
        })
        request.user = self.user
        
        response = remove_grade(request)
        
        self.assertEqual(response.status_code, 200)
        data = json.loads(response.content)
        self.assertEqual(data['status'], 'success')
        self.assertEqual(data['message'], '文件中没有找到评分或评价')


class TestAPIErrorHandling(BaseAPITestCase):
    """Test error handling in API endpoints."""
    
    def test_api_with_unauthenticated_user(self):
        """Test API endpoints with unauthenticated user."""
        request = self.factory.get('/api/course_info/', {
            'course_name': 'Test Course'
        })
        # No user set - unauthenticated
        
        # This should be handled by @login_required decorator
        # The actual behavior depends on Django settings
        pass
    
    def test_api_with_invalid_repository_id(self):
        """Test API endpoints with invalid repository ID."""
        request = self.factory.get('/api/courses_list/', {
            'repo_id': '99999'  # Non-existent ID
        })
        request.user = self.user
        
        response = get_courses_list_view(request)
        
        self.assertEqual(response.status_code, 200)
        data = json.loads(response.content)
        self.assertEqual(data['status'], 'error')
        self.assertEqual(data['message'], '仓库不存在')
    
    def test_api_exception_handling(self):
        """Test API exception handling."""
        request = self.factory.get('/api/course_info/', {
            'course_name': 'Test Course'
        })
        request.user = self.user
        
        with patch('grading.views.Course.objects.filter') as mock_filter:
            # Simulate database error
            mock_filter.side_effect = Exception('Database error')
            
            response = get_course_info_api(request)
        
        self.assertEqual(response.status_code, 200)
        data = json.loads(response.content)
        self.assertFalse(data['success'])
        self.assertIn('Database error', data['message'])


@pytest.mark.django_db
class TestAPIIntegration:
    """Integration tests for API endpoints."""
    
    def test_course_workflow_integration(self, client, django_user_model):
        """Test complete course management workflow."""
        # Create user and login
        user = django_user_model.objects.create_user(
            username='testuser',
            password='testpass123',
            is_staff=True
        )
        client.force_login(user)
        
        # 1. Get course info (should auto-create)
        response = client.get('/api/course_info/', {
            'course_name': 'Integration Test Course',
            'auto_create': 'true'
        })
        
        assert response.status_code == 200
        data = response.json()
        assert data['success']
        course_id = data['course']['id']
        
        # 2. Update course type
        response = client.post('/api/update_course_type/', {
            'course_name': 'Integration Test Course',
            'course_type': 'lab'
        })
        
        assert response.status_code == 200
        data = response.json()
        assert data['success']
        assert data['course']['course_type'] == 'lab'
        
        # 3. Get homework list (should be empty initially)
        response = client.get('/api/homework_list/', {
            'course_name': 'Integration Test Course'
        })
        
        assert response.status_code == 200
        data = response.json()
        assert data['success']
        assert len(data['homeworks']) == 0
        
        # 4. Get homework info (should auto-create)
        response = client.get('/api/homework_info/', {
            'course_name': 'Integration Test Course',
            'homework_folder': 'lab1',
            'auto_create': 'true'
        })
        
        assert response.status_code == 200
        data = response.json()
        assert data['success']
        assert data['homework']['folder_name'] == 'lab1'


if __name__ == '__main__':
    pytest.main([__file__])