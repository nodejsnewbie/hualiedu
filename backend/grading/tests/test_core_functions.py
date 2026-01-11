"""
测试核心统一函数
Tests for core unified functions: find_teacher_signature_cell, extract_grade_and_comment_from_cell, write_to_teacher_signature_cell
"""

import os
import shutil
import tempfile

from django.test import TestCase
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from grading.views import (
    extract_grade_and_comment_from_cell,
    find_teacher_signature_cell,
    write_to_teacher_signature_cell,
)


class FindTeacherSignatureCellTest(TestCase):
    """测试 find_teacher_signature_cell 函数"""

    def setUp(self):
        self.temp_dir = tempfile.mkdtemp()

    def tearDown(self):
        if os.path.exists(self.temp_dir):
            shutil.rmtree(self.temp_dir)

    def create_test_document_with_signature_cell(self, signature_text="教师（签字）"):
        """创建包含教师签字单元格的测试文档"""
        doc = Document()
        table = doc.add_table(rows=3, cols=2)

        # 在第二行第一列添加"教师（签字）"
        cell = table.rows[1].cells[0]
        cell.text = signature_text

        return doc

    def test_find_signature_cell_with_full_width_parentheses(self):
        """测试查找包含全角括号的"教师（签字）"单元格"""
        doc = self.create_test_document_with_signature_cell("教师（签字）")

        cell, table_idx, row_idx, col_idx = find_teacher_signature_cell(doc)

        self.assertIsNotNone(cell)
        self.assertEqual(table_idx, 0)
        self.assertEqual(row_idx, 1)
        self.assertEqual(col_idx, 0)
        self.assertIn("教师", cell.text)

    def test_find_signature_cell_with_half_width_parentheses(self):
        """测试查找包含半角括号的"教师(签字)"单元格"""
        doc = self.create_test_document_with_signature_cell("教师(签字)")

        cell, table_idx, row_idx, col_idx = find_teacher_signature_cell(doc)

        self.assertIsNotNone(cell)
        self.assertEqual(table_idx, 0)
        self.assertEqual(row_idx, 1)
        self.assertEqual(col_idx, 0)

    def test_find_signature_cell_with_additional_text(self):
        """测试查找包含额外文本的"教师（签字）"单元格"""
        doc = self.create_test_document_with_signature_cell("教师（签字）：张老师 时间：2024-01-01")

        cell, table_idx, row_idx, col_idx = find_teacher_signature_cell(doc)

        self.assertIsNotNone(cell)
        self.assertIn("教师（签字）", cell.text)

    def test_find_signature_cell_not_found(self):
        """测试未找到"教师（签字）"单元格的情况"""
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "学生姓名"
        table.rows[1].cells[0].text = "成绩"

        cell, table_idx, row_idx, col_idx = find_teacher_signature_cell(doc)

        self.assertIsNone(cell)
        self.assertIsNone(table_idx)
        self.assertIsNone(row_idx)
        self.assertIsNone(col_idx)

    def test_find_signature_cell_in_multiple_tables(self):
        """测试在多个表格中查找"教师（签字）"单元格"""
        doc = Document()

        # 第一个表格（不包含签字单元格）
        table1 = doc.add_table(rows=2, cols=2)
        table1.rows[0].cells[0].text = "学生信息"

        # 第二个表格（包含签字单元格）
        table2 = doc.add_table(rows=2, cols=2)
        table2.rows[1].cells[1].text = "教师（签字）"

        cell, table_idx, row_idx, col_idx = find_teacher_signature_cell(doc)

        self.assertIsNotNone(cell)
        self.assertEqual(table_idx, 1)  # 第二个表格
        self.assertEqual(row_idx, 1)
        self.assertEqual(col_idx, 1)

    def test_find_signature_cell_empty_document(self):
        """测试空文档"""
        doc = Document()

        cell, table_idx, row_idx, col_idx = find_teacher_signature_cell(doc)

        self.assertIsNone(cell)


class ExtractGradeAndCommentFromCellTest(TestCase):
    """测试 extract_grade_and_comment_from_cell 函数"""

    def create_cell_with_content(self, lines):
        """创建包含指定内容的单元格"""
        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        cell = table.rows[0].cells[0]
        cell.text = "\n".join(lines)
        return cell

    def test_extract_grade_and_comment_complete(self):
        """测试提取完整的评分、评价和签字文本"""
        lines = ["A", "作业完成得非常出色，内容充实。", "教师（签字）：张老师 时间：2024-01-01"]
        cell = self.create_cell_with_content(lines)

        grade, comment, signature_text = extract_grade_and_comment_from_cell(cell)

        self.assertEqual(grade, "A")
        self.assertEqual(comment, "作业完成得非常出色，内容充实。")
        self.assertIn("教师（签字）", signature_text)
        self.assertIn("张老师", signature_text)

    def test_extract_grade_only(self):
        """测试只有评分没有评价"""
        lines = ["B", "教师（签字）：李老师"]
        cell = self.create_cell_with_content(lines)

        grade, comment, signature_text = extract_grade_and_comment_from_cell(cell)

        self.assertEqual(grade, "B")
        self.assertIsNone(comment)
        self.assertIn("教师（签字）", signature_text)

    def test_extract_chinese_grade(self):
        """测试提取中文评分"""
        lines = ["优秀", "实验报告格式规范，数据准确。", "教师（签字）：王老师"]
        cell = self.create_cell_with_content(lines)

        grade, comment, signature_text = extract_grade_and_comment_from_cell(cell)

        self.assertEqual(grade, "优秀")
        self.assertEqual(comment, "实验报告格式规范，数据准确。")

    def test_extract_all_valid_grades(self):
        """测试所有有效的评分等级"""
        valid_grades = ["A", "B", "C", "D", "E", "优秀", "良好", "中等", "及格", "不及格"]

        for valid_grade in valid_grades:
            with self.subTest(grade=valid_grade):
                lines = [valid_grade, "教师（签字）：测试"]
                cell = self.create_cell_with_content(lines)

                grade, comment, signature_text = extract_grade_and_comment_from_cell(cell)

                self.assertEqual(grade, valid_grade)

    def test_extract_invalid_grade(self):
        """测试无效的评分（第一行不是有效评分）"""
        lines = ["这不是评分", "教师（签字）：测试"]
        cell = self.create_cell_with_content(lines)

        grade, comment, signature_text = extract_grade_and_comment_from_cell(cell)

        self.assertIsNone(grade)

    def test_extract_no_signature(self):
        """测试单元格中没有"教师（签字）"文本"""
        lines = ["A", "评价内容"]
        cell = self.create_cell_with_content(lines)

        grade, comment, signature_text = extract_grade_and_comment_from_cell(cell)

        self.assertIsNone(grade)
        self.assertIsNone(comment)
        self.assertEqual(signature_text, "")

    def test_extract_empty_comment(self):
        """测试评价为空的情况"""
        lines = ["C", "", "教师（签字）：测试"]
        cell = self.create_cell_with_content(lines)

        grade, comment, signature_text = extract_grade_and_comment_from_cell(cell)

        self.assertEqual(grade, "C")
        self.assertIsNone(comment)

    def test_extract_with_half_width_parentheses(self):
        """测试半角括号的教师(签字)"""
        lines = ["A", "很好", "教师(签字)：测试"]
        cell = self.create_cell_with_content(lines)

        grade, comment, signature_text = extract_grade_and_comment_from_cell(cell)

        self.assertEqual(grade, "A")
        self.assertEqual(comment, "很好")
        self.assertIn("教师", signature_text)


class WriteToTeacherSignatureCellTest(TestCase):
    """测试 write_to_teacher_signature_cell 函数"""

    def create_empty_cell(self):
        """创建空单元格"""
        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        return table.rows[0].cells[0], doc

    def test_write_grade_and_comment(self):
        """测试写入评分和评价"""
        cell, doc = self.create_empty_cell()
        grade = "A"
        comment = "作业完成得非常出色"
        signature_text = "教师（签字）：张老师 时间：2024-01-01"

        write_to_teacher_signature_cell(cell, grade, comment, signature_text)

        # 验证单元格内容
        cell_text = cell.text
        self.assertIn(grade, cell_text)
        self.assertIn(comment, cell_text)
        self.assertIn(signature_text, cell_text)

        # 验证段落数量（评分、评价、签字 = 3个段落）
        self.assertEqual(len(cell.paragraphs), 3)

    def test_write_grade_only(self):
        """测试只写入评分（无评价）"""
        cell, doc = self.create_empty_cell()
        grade = "B"
        comment = None
        signature_text = "教师（签字）：李老师"

        write_to_teacher_signature_cell(cell, grade, comment, signature_text)

        cell_text = cell.text
        self.assertIn(grade, cell_text)
        self.assertIn(signature_text, cell_text)

        # 验证段落数量（评分、签字 = 2个段落）
        self.assertEqual(len(cell.paragraphs), 2)

    def test_write_formatting(self):
        """测试写入内容的格式"""
        cell, doc = self.create_empty_cell()
        grade = "A"
        comment = "很好"
        signature_text = "教师（签字）：测试"

        write_to_teacher_signature_cell(cell, grade, comment, signature_text)

        # 验证第一段（评分）格式
        grade_paragraph = cell.paragraphs[0]
        self.assertEqual(grade_paragraph.alignment, WD_ALIGN_PARAGRAPH.CENTER)
        self.assertTrue(grade_paragraph.runs[0].bold)
        self.assertEqual(grade_paragraph.runs[0].font.size, Pt(14))

        # 验证第二段（评价）格式
        comment_paragraph = cell.paragraphs[1]
        self.assertEqual(comment_paragraph.alignment, WD_ALIGN_PARAGRAPH.LEFT)
        self.assertEqual(comment_paragraph.runs[0].font.size, Pt(11))

        # 验证第三段（签字）格式
        signature_paragraph = cell.paragraphs[2]
        self.assertEqual(signature_paragraph.alignment, WD_ALIGN_PARAGRAPH.LEFT)
        self.assertEqual(signature_paragraph.runs[0].font.size, Pt(10))

    def test_write_clears_existing_content(self):
        """测试写入前清空现有内容"""
        cell, doc = self.create_empty_cell()

        # 先写入一些内容
        cell.text = "旧内容\n更多旧内容"

        # 写入新内容
        grade = "C"
        comment = "新评价"
        signature_text = "教师（签字）：新签字"

        write_to_teacher_signature_cell(cell, grade, comment, signature_text)

        # 验证旧内容被清除
        cell_text = cell.text
        self.assertNotIn("旧内容", cell_text)
        self.assertIn(grade, cell_text)
        self.assertIn(comment, cell_text)

    def test_write_chinese_grades(self):
        """测试写入中文评分"""
        chinese_grades = ["优秀", "良好", "中等", "及格", "不及格"]

        for grade in chinese_grades:
            with self.subTest(grade=grade):
                cell, doc = self.create_empty_cell()
                comment = "测试评价"
                signature_text = "教师（签字）：测试"

                write_to_teacher_signature_cell(cell, grade, comment, signature_text)

                self.assertIn(grade, cell.text)

    def test_write_empty_signature_text(self):
        """测试签字文本为空的情况"""
        cell, doc = self.create_empty_cell()
        grade = "A"
        comment = "很好"
        signature_text = ""

        write_to_teacher_signature_cell(cell, grade, comment, signature_text)

        # 应该只有评分和评价两个段落
        self.assertEqual(len(cell.paragraphs), 2)
        self.assertIn(grade, cell.text)
        self.assertIn(comment, cell.text)

    def test_write_long_comment(self):
        """测试写入长评价"""
        cell, doc = self.create_empty_cell()
        grade = "A"
        comment = "这是一个很长的评价内容，" * 10  # 重复10次
        signature_text = "教师（签字）：测试"

        write_to_teacher_signature_cell(cell, grade, comment, signature_text)

        self.assertIn(comment, cell.text)
        self.assertEqual(len(cell.paragraphs), 3)


class CoreFunctionsIntegrationTest(TestCase):
    """核心函数集成测试"""

    def setUp(self):
        self.temp_dir = tempfile.mkdtemp()

    def tearDown(self):
        if os.path.exists(self.temp_dir):
            shutil.rmtree(self.temp_dir)

    def test_complete_workflow(self):
        """测试完整的工作流：查找 -> 提取 -> 写入"""
        # 创建测试文档
        doc = Document()
        table = doc.add_table(rows=2, cols=2)

        # 添加初始内容
        cell = table.rows[1].cells[0]
        cell.text = "B\n作业还不错\n教师（签字）：原始签字"

        # 步骤1：查找单元格
        found_cell, table_idx, row_idx, col_idx = find_teacher_signature_cell(doc)
        self.assertIsNotNone(found_cell)

        # 步骤2：提取内容
        grade, comment, signature_text = extract_grade_and_comment_from_cell(found_cell)
        self.assertEqual(grade, "B")
        self.assertEqual(comment, "作业还不错")
        self.assertIn("原始签字", signature_text)

        # 步骤3：写入新内容
        new_grade = "A"
        new_comment = "修改后的评价"
        write_to_teacher_signature_cell(found_cell, new_grade, new_comment, signature_text)

        # 验证新内容
        self.assertIn(new_grade, found_cell.text)
        self.assertIn(new_comment, found_cell.text)
        self.assertIn("原始签字", found_cell.text)

    def test_workflow_with_file_save(self):
        """测试完整工作流并保存文件"""
        # 创建测试文档
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        cell = table.rows[0].cells[0]
        cell.text = "C\n需要改进\n教师（签字）：测试老师"

        # 保存原始文档
        original_path = os.path.join(self.temp_dir, "test_original.docx")
        doc.save(original_path)

        # 重新加载文档
        doc = Document(original_path)

        # 查找并修改
        found_cell, _, _, _ = find_teacher_signature_cell(doc)
        self.assertIsNotNone(found_cell)

        grade, comment, signature_text = extract_grade_and_comment_from_cell(found_cell)

        new_grade = "A"
        new_comment = "大幅改进"
        write_to_teacher_signature_cell(found_cell, new_grade, new_comment, signature_text)

        # 保存修改后的文档
        modified_path = os.path.join(self.temp_dir, "test_modified.docx")
        doc.save(modified_path)

        # 重新加载并验证
        doc_reloaded = Document(modified_path)
        found_cell_reloaded, _, _, _ = find_teacher_signature_cell(doc_reloaded)

        self.assertIsNotNone(found_cell_reloaded)
        self.assertIn(new_grade, found_cell_reloaded.text)
        self.assertIn(new_comment, found_cell_reloaded.text)

    def test_workflow_update_existing_grade(self):
        """测试更新已有评分的工作流"""
        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        cell = table.rows[0].cells[0]

        # 第一次写入
        write_to_teacher_signature_cell(cell, "B", "初次评价", "教师（签字）：张老师")

        # 提取并验证
        grade1, comment1, sig1 = extract_grade_and_comment_from_cell(cell)
        self.assertEqual(grade1, "B")
        self.assertEqual(comment1, "初次评价")

        # 第二次写入（更新）
        write_to_teacher_signature_cell(cell, "A", "修改后的评价", sig1)

        # 再次提取并验证
        grade2, comment2, sig2 = extract_grade_and_comment_from_cell(cell)
        self.assertEqual(grade2, "A")
        self.assertEqual(comment2, "修改后的评价")

        # 签字应该保持不变
        self.assertEqual(sig1, sig2)


class IsLabReportFileTest(TestCase):
    """测试 is_lab_report_file 函数 - 作业类型判断逻辑"""

    def setUp(self):
        """设置测试数据"""
        from django.contrib.auth.models import User

        from grading.models import Course, Homework, Semester

        # 创建测试用户
        self.user = User.objects.create_user(username="testuser", password="testpass")

        # 创建测试学期
        from datetime import date

        self.semester = Semester.objects.create(
            name="2024年春季学期",
            start_date=date(2024, 3, 1),
            end_date=date(2024, 7, 15),
            is_active=True,
        )

        # 创建不同类型的课程
        self.theory_course = Course.objects.create(
            semester=self.semester,
            teacher=self.user,
            name="数据结构",
            course_type="theory",
            location="教学楼101",
        )

        self.lab_course = Course.objects.create(
            semester=self.semester,
            teacher=self.user,
            name="计算机网络实验",
            course_type="lab",
            location="实验室201",
        )

        self.practice_course = Course.objects.create(
            semester=self.semester,
            teacher=self.user,
            name="软件工程实践",
            course_type="practice",
            location="实验室202",
        )

        self.mixed_course = Course.objects.create(
            semester=self.semester,
            teacher=self.user,
            name="操作系统",
            course_type="mixed",
            location="教学楼102",
        )

        # 创建作业批次
        self.normal_homework = Homework.objects.create(
            course=self.theory_course,
            title="第一次作业",
            homework_type="normal",
            folder_name="作业1",
        )

        self.lab_report_homework = Homework.objects.create(
            course=self.lab_course, title="实验一", homework_type="lab_report", folder_name="实验1"
        )

        # 混合课程的普通作业
        self.mixed_normal_homework = Homework.objects.create(
            course=self.mixed_course,
            title="理论作业",
            homework_type="normal",
            folder_name="理论作业1",
        )

        # 混合课程的实验报告
        self.mixed_lab_homework = Homework.objects.create(
            course=self.mixed_course,
            title="实验报告",
            homework_type="lab_report",
            folder_name="实验报告1",
        )

    def test_priority_1_homework_type_normal(self):
        """测试优先级1：作业批次类型为普通作业"""
        from grading.views import is_lab_report_file

        result = is_lab_report_file(course_name="数据结构", homework_folder="作业1")

        self.assertFalse(result)

    def test_priority_1_homework_type_lab_report(self):
        """测试优先级1：作业批次类型为实验报告"""
        from grading.views import is_lab_report_file

        result = is_lab_report_file(course_name="计算机网络实验", homework_folder="实验1")

        self.assertTrue(result)

    def test_priority_1_mixed_course_normal_homework(self):
        """测试优先级1：混合课程的普通作业"""
        from grading.views import is_lab_report_file

        result = is_lab_report_file(course_name="操作系统", homework_folder="理论作业1")

        self.assertFalse(result)

    def test_priority_1_mixed_course_lab_homework(self):
        """测试优先级1：混合课程的实验报告"""
        from grading.views import is_lab_report_file

        result = is_lab_report_file(course_name="操作系统", homework_folder="实验报告1")

        self.assertTrue(result)

    def test_priority_2_course_type_theory(self):
        """测试优先级2：课程类型为理论课（作业批次不存在）"""
        from grading.views import is_lab_report_file

        result = is_lab_report_file(course_name="数据结构", homework_folder="不存在的作业")

        self.assertFalse(result)

    def test_priority_2_course_type_lab(self):
        """测试优先级2：课程类型为实验课（作业批次不存在）"""
        from grading.views import is_lab_report_file

        result = is_lab_report_file(course_name="计算机网络实验", homework_folder="不存在的实验")

        self.assertTrue(result)

    def test_priority_2_course_type_practice(self):
        """测试优先级2：课程类型为实践课（作业批次不存在）"""
        from grading.views import is_lab_report_file

        result = is_lab_report_file(course_name="软件工程实践", homework_folder="不存在的实践")

        self.assertTrue(result)

    def test_priority_2_course_type_mixed(self):
        """测试优先级2：课程类型为混合课（作业批次不存在）"""
        from grading.views import is_lab_report_file

        result = is_lab_report_file(course_name="操作系统", homework_folder="不存在的作业")

        self.assertTrue(result)

    def test_priority_3_path_parsing(self):
        """测试优先级3：从文件路径提取课程和作业信息"""
        import tempfile

        from grading.views import is_lab_report_file

        # 创建临时目录结构
        with tempfile.TemporaryDirectory() as temp_dir:
            # 模拟路径：base_dir/计算机网络实验/实验1/学生作业.docx
            file_path = os.path.join(temp_dir, "计算机网络实验", "实验1", "学生作业.docx")

            result = is_lab_report_file(file_path=file_path, base_dir=temp_dir)

            # 应该能够从路径提取课程名和作业文件夹，然后查询数据库
            self.assertTrue(result)

    def test_priority_3_path_parsing_with_class(self):
        """测试优先级3：路径包含班级信息"""
        import tempfile

        from grading.views import is_lab_report_file

        with tempfile.TemporaryDirectory() as temp_dir:
            # 模拟路径：base_dir/计算机网络实验/计算机1班/实验1/学生作业.docx
            file_path = os.path.join(
                temp_dir, "计算机网络实验", "计算机1班", "实验1", "学生作业.docx"
            )

            result = is_lab_report_file(file_path=file_path, base_dir=temp_dir)

            self.assertTrue(result)

    def test_priority_4_keyword_matching_lab(self):
        """测试优先级4：关键词匹配 - "实验"关键词"""
        from grading.views import is_lab_report_file

        result = is_lab_report_file(course_name="Python实验课程")

        self.assertTrue(result)

    def test_priority_4_keyword_matching_practice(self):
        """测试优先级4：关键词匹配 - "实训"关键词"""
        from grading.views import is_lab_report_file

        result = is_lab_report_file(course_name="Java实训")

        self.assertTrue(result)

    def test_priority_4_keyword_matching_lab_english(self):
        """测试优先级4：关键词匹配 - "lab"关键词"""
        from grading.views import is_lab_report_file

        result = is_lab_report_file(course_name="Computer Lab")

        self.assertTrue(result)

    def test_priority_4_keyword_matching_experiment(self):
        """测试优先级4：关键词匹配 - "experiment"关键词"""
        from grading.views import is_lab_report_file

        result = is_lab_report_file(course_name="Physics Experiment")

        self.assertTrue(result)

    def test_priority_4_keyword_matching_practice_english(self):
        """测试优先级4：关键词匹配 - "practice"关键词"""
        from grading.views import is_lab_report_file

        result = is_lab_report_file(course_name="Software Practice")

        self.assertTrue(result)

    def test_priority_4_no_keyword_match(self):
        """测试优先级4：无关键词匹配"""
        from grading.views import is_lab_report_file

        result = is_lab_report_file(course_name="高等数学")

        self.assertFalse(result)

    def test_default_behavior_no_params(self):
        """测试默认行为：无任何参数"""
        from grading.views import is_lab_report_file

        result = is_lab_report_file()

        self.assertFalse(result)

    def test_default_behavior_empty_course_name(self):
        """测试默认行为：课程名为空"""
        from grading.views import is_lab_report_file

        result = is_lab_report_file(course_name="", homework_folder="作业1")

        self.assertFalse(result)

    def test_default_behavior_none_course_name(self):
        """测试默认行为：课程名为None"""
        from grading.views import is_lab_report_file

        result = is_lab_report_file(course_name=None, homework_folder="作业1")

        self.assertFalse(result)

    def test_database_query_failure_fallback(self):
        """测试数据库查询失败时的降级处理"""
        from unittest.mock import patch

        from grading.views import is_lab_report_file

        # 模拟数据库查询失败
        with patch("grading.views.Homework.objects.filter") as mock_filter:
            mock_filter.side_effect = Exception("Database error")

            # 应该降级到关键词匹配
            result = is_lab_report_file(course_name="Python实验", homework_folder="实验1")

            self.assertTrue(result)

    def test_path_parsing_insufficient_depth(self):
        """测试路径解析：路径层级不足"""
        import tempfile

        from grading.views import is_lab_report_file

        with tempfile.TemporaryDirectory() as temp_dir:
            # 路径太短，无法提取课程和作业信息
            file_path = os.path.join(temp_dir, "学生作业.docx")

            result = is_lab_report_file(file_path=file_path, base_dir=temp_dir)

            # 应该返回默认值False
            self.assertFalse(result)

    def test_path_parsing_with_repository_dirs(self):
        """测试路径解析：跳过仓库目录"""
        import tempfile

        from grading.views import is_lab_report_file

        with tempfile.TemporaryDirectory() as temp_dir:
            # 模拟路径：base_dir/homework/计算机网络实验/实验1/学生作业.docx
            file_path = os.path.join(
                temp_dir, "homework", "计算机网络实验", "实验1", "学生作业.docx"
            )

            result = is_lab_report_file(file_path=file_path, base_dir=temp_dir)

            # 应该能够跳过"homework"目录，正确提取课程名
            self.assertTrue(result)

    def test_case_insensitive_keyword_matching(self):
        """测试关键词匹配的大小写不敏感"""
        from grading.views import is_lab_report_file

        test_cases = [
            "PYTHON实验",
            "Python实验",
            "python实验",
            "COMPUTER LAB",
            "Computer Lab",
            "computer lab",
        ]

        for course_name in test_cases:
            with self.subTest(course_name=course_name):
                result = is_lab_report_file(course_name=course_name)
                self.assertTrue(result)

    def test_homework_type_overrides_course_type(self):
        """测试作业批次类型优先于课程类型"""
        # 实验课程但作业类型为普通作业（这种情况不应该发生，但要测试优先级）
        # 由于我们的测试数据中没有这种情况，创建一个新的
        from grading.models import Homework
        from grading.views import is_lab_report_file

        special_homework = Homework.objects.create(
            course=self.lab_course,  # 实验课
            title="特殊作业",
            homework_type="normal",  # 但作业类型为普通
            folder_name="特殊作业1",
        )

        result = is_lab_report_file(course_name="计算机网络实验", homework_folder="特殊作业1")

        # 应该返回False，因为作业类型优先
        self.assertFalse(result)

        # 清理
        special_homework.delete()

    def test_multiple_courses_same_name(self):
        """测试多个同名课程的情况"""
        from grading.models import Course, Homework
        from grading.views import is_lab_report_file

        # 创建另一个同名但不同类型的课程
        another_theory_course = Course.objects.create(
            semester=self.semester,
            teacher=self.user,
            name="数据结构",  # 与已有课程同名
            course_type="lab",  # 但类型不同
            location="实验室203",
        )

        # 为新课程创建作业
        another_homework = Homework.objects.create(
            course=another_theory_course,
            title="实验作业",
            homework_type="lab_report",
            folder_name="实验作业1",
        )

        # 查询时应该返回第一个匹配的结果
        result = is_lab_report_file(course_name="数据结构", homework_folder="实验作业1")

        self.assertTrue(result)

        # 清理
        another_homework.delete()
        another_theory_course.delete()

    def test_all_parameters_provided(self):
        """测试提供所有参数的情况"""
        import tempfile

        from grading.views import is_lab_report_file

        with tempfile.TemporaryDirectory() as temp_dir:
            file_path = os.path.join(temp_dir, "计算机网络实验", "实验1", "学生作业.docx")

            # 提供所有参数，应该优先使用course_name和homework_folder
            result = is_lab_report_file(
                course_name="计算机网络实验",
                homework_folder="实验1",
                file_path=file_path,
                base_dir=temp_dir,
            )

            self.assertTrue(result)

    def test_only_file_path_provided(self):
        """测试只提供文件路径的情况"""
        import tempfile

        from grading.views import is_lab_report_file

        with tempfile.TemporaryDirectory() as temp_dir:
            file_path = os.path.join(temp_dir, "计算机网络实验", "实验1", "学生作业.docx")

            # 只提供file_path，缺少base_dir，应该无法解析
            result = is_lab_report_file(file_path=file_path)

            self.assertFalse(result)

    def test_relative_path_parsing(self):
        """测试相对路径解析"""
        import tempfile

        from grading.views import is_lab_report_file

        with tempfile.TemporaryDirectory() as temp_dir:
            # 创建实际的目录结构
            course_dir = os.path.join(temp_dir, "计算机网络实验")
            homework_dir = os.path.join(course_dir, "实验1")
            os.makedirs(homework_dir, exist_ok=True)

            file_path = os.path.join(homework_dir, "学生作业.docx")

            result = is_lab_report_file(file_path=file_path, base_dir=temp_dir)

            self.assertTrue(result)


class WriteGradeToLabReportTest(TestCase):
    """测试 write_grade_to_lab_report 函数 - 实验报告评分写入功能"""

    def setUp(self):
        self.temp_dir = tempfile.mkdtemp()

    def tearDown(self):
        if os.path.exists(self.temp_dir):
            shutil.rmtree(self.temp_dir)

    def create_lab_report_with_signature_cell(self):
        """创建包含教师签字单元格的实验报告"""
        doc = Document()
        table = doc.add_table(rows=3, cols=2)

        # 添加"教师（签字）"单元格
        cell = table.rows[2].cells[0]
        cell.text = "教师（签字）：时间："

        return doc

    def create_lab_report_without_signature_cell(self):
        """创建不包含教师签字单元格的实验报告（格式错误）"""
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "学生姓名"
        table.rows[1].cells[0].text = "实验内容"

        return doc

    def test_write_grade_to_valid_lab_report(self):
        """测试向有效的实验报告写入评分"""
        from grading.views import write_grade_to_lab_report

        doc = self.create_lab_report_with_signature_cell()
        grade = "A"
        comment = "实验报告完成得非常出色"

        success, modified_grade, modified_comment = write_grade_to_lab_report(doc, grade, comment)

        self.assertTrue(success)
        self.assertEqual(modified_grade, "A")
        self.assertEqual(modified_comment, "实验报告完成得非常出色")

        # 验证内容已写入
        cell, _, _, _ = find_teacher_signature_cell(doc)
        self.assertIsNotNone(cell)
        self.assertIn(grade, cell.text)
        self.assertIn(comment, cell.text)

    def test_write_grade_without_comment(self):
        """测试写入评分但不提供评价（应自动生成）"""
        from grading.views import write_grade_to_lab_report

        doc = self.create_lab_report_with_signature_cell()
        grade = "B"

        success, modified_grade, modified_comment = write_grade_to_lab_report(
            doc, grade, comment=None
        )

        self.assertTrue(success)
        self.assertEqual(modified_grade, "B")
        self.assertIsNotNone(modified_comment)  # 应该自动生成评价

        # 验证内容已写入
        cell, _, _, _ = find_teacher_signature_cell(doc)
        self.assertIn(grade, cell.text)
        self.assertIn(modified_comment, cell.text)

    def test_format_error_detection(self):
        """测试格式错误检测：未找到"教师（签字）"单元格"""
        from grading.views import write_grade_to_lab_report

        doc = self.create_lab_report_without_signature_cell()
        grade = "A"
        comment = "很好"

        success, modified_grade, modified_comment = write_grade_to_lab_report(doc, grade, comment)

        # 应该返回失败
        self.assertFalse(success)

        # 应该自动设置D评分
        self.assertEqual(modified_grade, "D")

        # 应该设置锁定标记
        self.assertIn("【格式错误-已锁定】", modified_comment)
        self.assertIn("请按要求的格式写实验报告", modified_comment)
        self.assertIn("此评分不可修改", modified_comment)

    def test_format_error_locking_mechanism(self):
        """测试格式错误时的锁定机制"""
        from grading.views import write_grade_to_lab_report

        doc = self.create_lab_report_without_signature_cell()

        # 尝试写入A评分
        success, modified_grade, modified_comment = write_grade_to_lab_report(doc, "A", "优秀")

        self.assertFalse(success)
        self.assertEqual(modified_grade, "D")
        self.assertIn("【格式错误-已锁定】", modified_comment)

    def test_update_existing_grade(self):
        """测试更新已有评分"""
        from grading.views import write_grade_to_lab_report

        doc = self.create_lab_report_with_signature_cell()

        # 第一次写入
        cell, _, _, _ = find_teacher_signature_cell(doc)
        write_to_teacher_signature_cell(cell, "B", "还不错", "教师（签字）：张老师")

        # 更新评分
        success, modified_grade, modified_comment = write_grade_to_lab_report(
            doc, "A", "修改为优秀"
        )

        self.assertTrue(success)
        self.assertEqual(modified_grade, "A")
        self.assertEqual(modified_comment, "修改为优秀")

        # 验证更新后的内容
        cell, _, _, _ = find_teacher_signature_cell(doc)
        self.assertIn("A", cell.text)
        self.assertIn("修改为优秀", cell.text)
        self.assertIn("张老师", cell.text)  # 签字应该保留

    def test_preserve_existing_comment(self):
        """测试保留已有评价（当不提供新评价时）"""
        from grading.views import write_grade_to_lab_report

        doc = self.create_lab_report_with_signature_cell()

        # 第一次写入
        cell, _, _, _ = find_teacher_signature_cell(doc)
        write_to_teacher_signature_cell(cell, "B", "原始评价", "教师（签字）：测试")

        # 只更新评分，不提供新评价
        success, modified_grade, modified_comment = write_grade_to_lab_report(
            doc, "A", comment=None
        )

        self.assertTrue(success)
        self.assertEqual(modified_grade, "A")
        self.assertEqual(modified_comment, "原始评价")  # 应该保留原有评价

    def test_chinese_grade_format_error(self):
        """测试中文评分在格式错误时的处理"""
        from grading.views import write_grade_to_lab_report

        doc = self.create_lab_report_without_signature_cell()

        success, modified_grade, modified_comment = write_grade_to_lab_report(doc, "优秀", "很好")

        self.assertFalse(success)
        self.assertEqual(modified_grade, "D")
        self.assertIn("【格式错误-已锁定】", modified_comment)

    def test_exception_handling(self):
        """测试异常处理"""
        from unittest.mock import patch

        from grading.views import write_grade_to_lab_report

        doc = self.create_lab_report_with_signature_cell()

        # 模拟异常
        with patch("grading.views.find_teacher_signature_cell") as mock_find:
            mock_find.side_effect = Exception("Test exception")

            success, modified_grade, modified_comment = write_grade_to_lab_report(doc, "A", "测试")

            # 异常情况应该触发降级处理
            self.assertFalse(success)
            self.assertEqual(modified_grade, "D")
            self.assertIn("【格式错误-已锁定】", modified_comment)

    def test_write_with_file_save(self):
        """测试写入并保存文件"""
        from grading.views import write_grade_to_lab_report

        doc = self.create_lab_report_with_signature_cell()

        # 写入评分
        success, modified_grade, modified_comment = write_grade_to_lab_report(doc, "A", "优秀")

        self.assertTrue(success)

        # 保存文件
        file_path = os.path.join(self.temp_dir, "test_lab_report.docx")
        doc.save(file_path)

        # 重新加载并验证
        doc_reloaded = Document(file_path)
        cell, _, _, _ = find_teacher_signature_cell(doc_reloaded)

        self.assertIsNotNone(cell)
        self.assertIn("A", cell.text)
        self.assertIn("优秀", cell.text)

    def test_all_valid_grades(self):
        """测试所有有效的评分等级"""
        from grading.views import write_grade_to_lab_report

        valid_grades = ["A", "B", "C", "D", "E", "优秀", "良好", "中等", "及格", "不及格"]

        for grade in valid_grades:
            with self.subTest(grade=grade):
                doc = self.create_lab_report_with_signature_cell()

                success, modified_grade, modified_comment = write_grade_to_lab_report(
                    doc, grade, "测试评价"
                )

                self.assertTrue(success)
                self.assertEqual(modified_grade, grade)

                cell, _, _, _ = find_teacher_signature_cell(doc)
                self.assertIn(grade, cell.text)

    def test_signature_text_preservation(self):
        """测试签字文本的保留"""
        from grading.views import write_grade_to_lab_report

        doc = self.create_lab_report_with_signature_cell()

        # 设置初始签字文本
        cell, _, _, _ = find_teacher_signature_cell(doc)
        original_signature = "教师（签字）：张老师 时间：2024-01-01 备注：初次评分"
        write_to_teacher_signature_cell(cell, "B", "初次评价", original_signature)

        # 更新评分
        success, _, _ = write_grade_to_lab_report(doc, "A", "更新评价")

        self.assertTrue(success)

        # 验证签字文本被保留
        cell, _, _, _ = find_teacher_signature_cell(doc)
        self.assertIn("张老师", cell.text)
        self.assertIn("2024-01-01", cell.text)
        self.assertIn("初次评分", cell.text)

    def test_empty_document(self):
        """测试空文档（无表格）"""
        from grading.views import write_grade_to_lab_report

        doc = Document()  # 空文档，无表格

        success, modified_grade, modified_comment = write_grade_to_lab_report(doc, "A", "测试")

        self.assertFalse(success)
        self.assertEqual(modified_grade, "D")
        self.assertIn("【格式错误-已锁定】", modified_comment)

    def test_multiple_tables_with_signature(self):
        """测试多个表格中只有一个包含签字单元格"""
        from grading.views import write_grade_to_lab_report

        doc = Document()

        # 第一个表格（无签字单元格）
        table1 = doc.add_table(rows=2, cols=2)
        table1.rows[0].cells[0].text = "学生信息"

        # 第二个表格（有签字单元格）
        table2 = doc.add_table(rows=2, cols=2)
        table2.rows[1].cells[0].text = "教师（签字）："

        success, modified_grade, modified_comment = write_grade_to_lab_report(doc, "A", "很好")

        self.assertTrue(success)
        self.assertEqual(modified_grade, "A")

        # 验证写入到正确的表格
        cell, table_idx, _, _ = find_teacher_signature_cell(doc)
        self.assertEqual(table_idx, 1)  # 第二个表格
        self.assertIn("A", cell.text)

    def test_long_comment_handling(self):
        """测试长评价的处理"""
        from grading.views import write_grade_to_lab_report

        doc = self.create_lab_report_with_signature_cell()

        long_comment = "这是一个非常详细的评价内容，" * 20  # 很长的评价

        success, modified_grade, modified_comment = write_grade_to_lab_report(
            doc, "A", long_comment
        )

        self.assertTrue(success)
        self.assertEqual(modified_comment, long_comment)

        cell, _, _, _ = find_teacher_signature_cell(doc)
        self.assertIn(long_comment, cell.text)

    def test_special_characters_in_comment(self):
        """测试评价中包含特殊字符"""
        from grading.views import write_grade_to_lab_report

        doc = self.create_lab_report_with_signature_cell()

        special_comment = "评价：优秀！包含特殊字符：@#$%^&*()_+-=[]{}|;':\",./<>?"

        success, modified_grade, modified_comment = write_grade_to_lab_report(
            doc, "A", special_comment
        )

        self.assertTrue(success)
        self.assertEqual(modified_comment, special_comment)

        cell, _, _, _ = find_teacher_signature_cell(doc)
        self.assertIn(special_comment, cell.text)

    def test_half_width_parentheses_signature(self):
        """测试半角括号的教师(签字)"""
        from grading.views import write_grade_to_lab_report

        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        table.rows[1].cells[0].text = "教师(签字)："  # 半角括号

        success, modified_grade, modified_comment = write_grade_to_lab_report(doc, "A", "测试")

        self.assertTrue(success)
        self.assertEqual(modified_grade, "A")

        cell, _, _, _ = find_teacher_signature_cell(doc)
        self.assertIn("A", cell.text)


class ManualGradingIntegrationTest(TestCase):
    """手动评分功能集成测试"""

    def setUp(self):
        """设置测试环境"""
        from datetime import date

        from django.contrib.auth.models import User

        from grading.models import Course, Homework, Repository, Semester

        self.temp_dir = tempfile.mkdtemp()

        # 创建测试用户
        self.user = User.objects.create_user(
            username="testteacher", password="testpass123", is_staff=True
        )

        # 创建测试学期
        self.semester = Semester.objects.create(
            name="2024年春季学期",
            start_date=date(2024, 3, 1),
            end_date=date(2024, 7, 15),
            is_active=True,
        )

        # 创建测试课程
        self.course = Course.objects.create(
            semester=self.semester,
            teacher=self.user,
            name="计算机网络实验",
            course_type="lab",
            location="实验室201",
        )

        # 创建作业批次
        self.homework = Homework.objects.create(
            course=self.course, title="实验一", homework_type="lab_report", folder_name="实验1"
        )

        # 创建测试仓库
        self.repository = Repository.objects.create(
            name="测试仓库", owner=self.user, repo_type="local", path=self.temp_dir, is_active=True
        )

    def tearDown(self):
        """清理测试环境"""
        if os.path.exists(self.temp_dir):
            shutil.rmtree(self.temp_dir)

    def create_test_lab_report(self, filename="test_report.docx"):
        """创建测试用的实验报告文件"""
        doc = Document()

        # 添加标题
        doc.add_heading("实验报告", 0)
        doc.add_paragraph("学生姓名：张三")
        doc.add_paragraph("学号：20240001")

        # 添加评分表格
        table = doc.add_table(rows=3, cols=2)
        table.rows[0].cells[0].text = "实验内容"
        table.rows[0].cells[1].text = "网络协议分析"
        table.rows[1].cells[0].text = "实验结果"
        table.rows[1].cells[1].text = "完成所有实验要求"
        table.rows[2].cells[0].text = "教师（签字）："

        # 保存文件
        file_path = os.path.join(self.temp_dir, filename)
        doc.save(file_path)

        return file_path

    def create_test_normal_homework(self, filename="test_homework.docx"):
        """创建测试用的普通作业文件"""
        doc = Document()

        doc.add_heading("作业提交", 0)
        doc.add_paragraph("学生姓名：李四")
        doc.add_paragraph("学号：20240002")
        doc.add_paragraph("作业内容：")
        doc.add_paragraph("这是作业的主要内容...")

        # 保存文件
        file_path = os.path.join(self.temp_dir, filename)
        doc.save(file_path)

        return file_path

    def test_new_grade_submission_lab_report(self):
        """测试新建评分流程 - 实验报告"""
        from grading.views import write_grade_and_comment_to_file

        # 创建测试文件
        file_path = self.create_test_lab_report()

        # 提交新评分
        grade = "A"
        comment = "实验报告完成得非常出色，数据准确，分析透彻。"

        # 执行评分
        write_grade_and_comment_to_file(
            full_path=file_path,
            grade=grade,
            comment=comment,
            base_dir=self.temp_dir,
            is_lab_report=True,
        )

        # 验证评分已写入
        doc = Document(file_path)
        cell, _, _, _ = find_teacher_signature_cell(doc)

        self.assertIsNotNone(cell, "应该找到教师签字单元格")
        self.assertIn(grade, cell.text, "单元格应包含评分")
        self.assertIn(comment, cell.text, "单元格应包含评价")

    def test_new_grade_submission_normal_homework(self):
        """测试新建评分流程 - 普通作业"""
        from grading.views import write_grade_and_comment_to_file

        # 创建测试文件
        file_path = self.create_test_normal_homework()

        # 提交新评分
        grade = "B"
        comment = "作业完成较好，但还有改进空间。"

        # 执行评分
        write_grade_and_comment_to_file(
            full_path=file_path,
            grade=grade,
            comment=comment,
            base_dir=self.temp_dir,
            is_lab_report=False,
        )

        # 验证评分已写入
        doc = Document(file_path)

        # 检查文档末尾是否有评分段落
        found_grade = False
        found_comment = False

        for paragraph in doc.paragraphs:
            if grade in paragraph.text:
                found_grade = True
            if comment in paragraph.text:
                found_comment = True

        self.assertTrue(found_grade, "文档应包含评分")
        self.assertTrue(found_comment, "文档应包含评价")

    def test_update_existing_grade_lab_report(self):
        """测试更新评分流程 - 实验报告"""
        from grading.views import write_grade_and_comment_to_file

        # 创建测试文件
        file_path = self.create_test_lab_report()

        # 第一次评分
        initial_grade = "B"
        initial_comment = "初次评价"

        write_grade_and_comment_to_file(
            full_path=file_path,
            grade=initial_grade,
            comment=initial_comment,
            base_dir=self.temp_dir,
            is_lab_report=True,
        )

        # 验证第一次评分
        doc = Document(file_path)
        cell, _, _, _ = find_teacher_signature_cell(doc)
        self.assertIn(initial_grade, cell.text)
        self.assertIn(initial_comment, cell.text)

        # 更新评分
        updated_grade = "A"
        updated_comment = "修改后的评价，学生进步明显"

        write_grade_and_comment_to_file(
            full_path=file_path,
            grade=updated_grade,
            comment=updated_comment,
            base_dir=self.temp_dir,
            is_lab_report=True,
        )

        # 验证更新后的评分
        doc = Document(file_path)
        cell, _, _, _ = find_teacher_signature_cell(doc)

        self.assertIn(updated_grade, cell.text, "应包含更新后的评分")
        self.assertIn(updated_comment, cell.text, "应包含更新后的评价")
        self.assertNotIn(initial_comment, cell.text, "不应包含旧评价")

    def test_update_existing_grade_normal_homework(self):
        """测试更新评分流程 - 普通作业"""
        from grading.views import write_grade_and_comment_to_file

        # 创建测试文件
        file_path = self.create_test_normal_homework()

        # 第一次评分
        initial_grade = "C"
        initial_comment = "需要改进"

        write_grade_and_comment_to_file(
            full_path=file_path,
            grade=initial_grade,
            comment=initial_comment,
            base_dir=self.temp_dir,
            is_lab_report=False,
        )

        # 更新评分
        updated_grade = "B"
        updated_comment = "有所改进"

        write_grade_and_comment_to_file(
            full_path=file_path,
            grade=updated_grade,
            comment=updated_comment,
            base_dir=self.temp_dir,
            is_lab_report=False,
        )

        # 验证更新后的评分
        doc = Document(file_path)

        # 检查是否包含更新后的评分
        found_updated_grade = False
        found_updated_comment = False

        for paragraph in doc.paragraphs:
            if updated_grade in paragraph.text:
                found_updated_grade = True
            if updated_comment in paragraph.text:
                found_updated_comment = True

        self.assertTrue(found_updated_grade, "应包含更新后的评分")
        self.assertTrue(found_updated_comment, "应包含更新后的评价")

    def test_grade_type_switching_letter_to_chinese(self):
        """测试评分方式切换 - 字母评分切换到中文评分"""
        from grading.views import write_grade_and_comment_to_file

        # 创建测试文件
        file_path = self.create_test_lab_report()

        # 使用字母评分
        letter_grade = "A"
        comment = "优秀的实验报告"

        write_grade_and_comment_to_file(
            full_path=file_path,
            grade=letter_grade,
            comment=comment,
            base_dir=self.temp_dir,
            is_lab_report=True,
        )

        # 验证字母评分
        doc = Document(file_path)
        cell, _, _, _ = find_teacher_signature_cell(doc)
        self.assertIn(letter_grade, cell.text)

        # 切换到中文评分
        chinese_grade = "优秀"
        updated_comment = "实验报告质量很高"

        write_grade_and_comment_to_file(
            full_path=file_path,
            grade=chinese_grade,
            comment=updated_comment,
            base_dir=self.temp_dir,
            is_lab_report=True,
        )

        # 验证中文评分
        doc = Document(file_path)
        cell, _, _, _ = find_teacher_signature_cell(doc)

        self.assertIn(chinese_grade, cell.text, "应包含中文评分")
        self.assertIn(updated_comment, cell.text, "应包含更新后的评价")

    def test_grade_type_switching_chinese_to_letter(self):
        """测试评分方式切换 - 中文评分切换到字母评分"""
        from grading.views import write_grade_and_comment_to_file

        # 创建测试文件
        file_path = self.create_test_lab_report()

        # 使用中文评分
        chinese_grade = "良好"
        comment = "完成得不错"

        write_grade_and_comment_to_file(
            full_path=file_path,
            grade=chinese_grade,
            comment=comment,
            base_dir=self.temp_dir,
            is_lab_report=True,
        )

        # 验证中文评分
        doc = Document(file_path)
        cell, _, _, _ = find_teacher_signature_cell(doc)
        self.assertIn(chinese_grade, cell.text)

        # 切换到字母评分
        letter_grade = "B"
        updated_comment = "作业质量良好"

        write_grade_and_comment_to_file(
            full_path=file_path,
            grade=letter_grade,
            comment=updated_comment,
            base_dir=self.temp_dir,
            is_lab_report=True,
        )

        # 验证字母评分
        doc = Document(file_path)
        cell, _, _, _ = find_teacher_signature_cell(doc)

        self.assertIn(letter_grade, cell.text, "应包含字母评分")
        self.assertIn(updated_comment, cell.text, "应包含更新后的评价")

    def test_all_letter_grades(self):
        """测试所有字母评分等级"""
        from grading.views import write_grade_and_comment_to_file

        letter_grades = ["A", "B", "C", "D", "E"]

        for grade in letter_grades:
            with self.subTest(grade=grade):
                # 为每个评分创建新文件
                file_path = self.create_test_lab_report(f"test_{grade}.docx")

                comment = f"评分为{grade}的测试"

                write_grade_and_comment_to_file(
                    full_path=file_path,
                    grade=grade,
                    comment=comment,
                    base_dir=self.temp_dir,
                    is_lab_report=True,
                )

                # 验证评分
                doc = Document(file_path)
                cell, _, _, _ = find_teacher_signature_cell(doc)

                self.assertIsNotNone(cell, f"应该找到教师签字单元格 (grade={grade})")
                self.assertIn(grade, cell.text, f"应包含评分{grade}")
                self.assertIn(comment, cell.text, f"应包含评价")

    def test_all_chinese_grades(self):
        """测试所有中文评分等级"""
        from grading.views import write_grade_and_comment_to_file

        chinese_grades = ["优秀", "良好", "中等", "及格", "不及格"]

        for grade in chinese_grades:
            with self.subTest(grade=grade):
                # 为每个评分创建新文件
                file_path = self.create_test_lab_report(f"test_{grade}.docx")

                comment = f"评分为{grade}的测试"

                write_grade_and_comment_to_file(
                    full_path=file_path,
                    grade=grade,
                    comment=comment,
                    base_dir=self.temp_dir,
                    is_lab_report=True,
                )

                # 验证评分
                doc = Document(file_path)
                cell, _, _, _ = find_teacher_signature_cell(doc)

                self.assertIsNotNone(cell, f"应该找到教师签字单元格 (grade={grade})")
                self.assertIn(grade, cell.text, f"应包含评分{grade}")
                self.assertIn(comment, cell.text, f"应包含评价")

    def test_grade_without_comment(self):
        """测试只提供评分不提供评价"""
        from grading.views import write_grade_and_comment_to_file

        # 创建测试文件
        file_path = self.create_test_lab_report()

        # 只提供评分
        grade = "A"

        write_grade_and_comment_to_file(
            full_path=file_path,
            grade=grade,
            comment=None,
            base_dir=self.temp_dir,
            is_lab_report=True,
        )

        # 验证评分已写入
        doc = Document(file_path)
        cell, _, _, _ = find_teacher_signature_cell(doc)

        self.assertIsNotNone(cell)
        self.assertIn(grade, cell.text)

    def test_multiple_updates_preserve_signature(self):
        """测试多次更新评分时保留签字信息"""
        from grading.views import write_grade_and_comment_to_file

        # 创建测试文件
        file_path = self.create_test_lab_report()

        # 手动设置初始签字
        doc = Document(file_path)
        cell, _, _, _ = find_teacher_signature_cell(doc)
        original_signature = "教师（签字）：张老师 时间：2024-01-01"
        write_to_teacher_signature_cell(cell, "B", "初次评价", original_signature)
        doc.save(file_path)

        # 第一次更新
        write_grade_and_comment_to_file(
            full_path=file_path,
            grade="A",
            comment="第一次更新",
            base_dir=self.temp_dir,
            is_lab_report=True,
        )

        # 第二次更新
        write_grade_and_comment_to_file(
            full_path=file_path,
            grade="A",
            comment="第二次更新",
            base_dir=self.temp_dir,
            is_lab_report=True,
        )

        # 验证签字信息被保留
        doc = Document(file_path)
        cell, _, _, _ = find_teacher_signature_cell(doc)

        self.assertIn("张老师", cell.text, "应保留教师签字")
        self.assertIn("2024-01-01", cell.text, "应保留签字时间")
        self.assertIn("第二次更新", cell.text, "应包含最新评价")

    def test_format_error_handling(self):
        """测试格式错误处理 - 实验报告缺少签字单元格"""
        from grading.views import write_grade_and_comment_to_file

        # 创建没有签字单元格的文档
        doc = Document()
        doc.add_heading("实验报告", 0)
        doc.add_paragraph("这是一个格式错误的实验报告")

        file_path = os.path.join(self.temp_dir, "format_error.docx")
        doc.save(file_path)

        # 尝试评分
        grade = "A"
        comment = "测试评价"

        write_grade_and_comment_to_file(
            full_path=file_path,
            grade=grade,
            comment=comment,
            base_dir=self.temp_dir,
            is_lab_report=True,
        )

        # 验证降级处理：应该以普通作业方式写入，评分改为D，评价包含锁定标记
        doc = Document(file_path)

        # 检查是否包含格式错误标记
        found_lock_marker = False
        found_d_grade = False

        for paragraph in doc.paragraphs:
            if "【格式错误-已锁定】" in paragraph.text:
                found_lock_marker = True
            if "D" in paragraph.text:
                found_d_grade = True

        self.assertTrue(found_lock_marker, "应包含格式错误锁定标记")
        self.assertTrue(found_d_grade, "评分应自动改为D")

    def test_concurrent_grade_updates(self):
        """测试并发评分更新的一致性"""
        from grading.views import write_grade_and_comment_to_file

        # 创建测试文件
        file_path = self.create_test_lab_report()

        # 模拟快速连续的评分更新
        grades = ["B", "A", "A"]
        comments = ["第一次", "第二次", "第三次"]

        for grade, comment in zip(grades, comments):
            write_grade_and_comment_to_file(
                full_path=file_path,
                grade=grade,
                comment=comment,
                base_dir=self.temp_dir,
                is_lab_report=True,
            )

        # 验证最后一次更新的结果
        doc = Document(file_path)
        cell, _, _, _ = find_teacher_signature_cell(doc)

        self.assertIn("A", cell.text, "应包含最后的评分")
        self.assertIn("第三次", cell.text, "应包含最后的评价")

    def test_special_characters_in_grades(self):
        """测试评价中包含特殊字符"""
        from grading.views import write_grade_and_comment_to_file

        # 创建测试文件
        file_path = self.create_test_lab_report()

        # 包含特殊字符的评价
        grade = "A"
        comment = "评价：优秀！包含特殊字符：@#$%^&*()_+-=[]{}|;':\",./<>?"

        write_grade_and_comment_to_file(
            full_path=file_path,
            grade=grade,
            comment=comment,
            base_dir=self.temp_dir,
            is_lab_report=True,
        )

        # 验证特殊字符被正确保存
        doc = Document(file_path)
        cell, _, _, _ = find_teacher_signature_cell(doc)

        self.assertIn(comment, cell.text, "应正确保存包含特殊字符的评价")

    def test_long_comment_handling(self):
        """测试长评价的处理"""
        from grading.views import write_grade_and_comment_to_file

        # 创建测试文件
        file_path = self.create_test_lab_report()

        # 创建很长的评价
        long_comment = "这是一个非常详细的评价内容，" * 50

        write_grade_and_comment_to_file(
            full_path=file_path,
            grade="A",
            comment=long_comment,
            base_dir=self.temp_dir,
            is_lab_report=True,
        )

        # 验证长评价被正确保存
        doc = Document(file_path)
        cell, _, _, _ = find_teacher_signature_cell(doc)

        self.assertIn(long_comment, cell.text, "应正确保存长评价")

    def test_empty_file_handling(self):
        """测试空文件的处理"""
        from grading.views import write_grade_and_comment_to_file

        # 创建空文档
        doc = Document()
        file_path = os.path.join(self.temp_dir, "empty.docx")
        doc.save(file_path)

        # 尝试评分（作为实验报告）
        write_grade_and_comment_to_file(
            full_path=file_path,
            grade="A",
            comment="测试",
            base_dir=self.temp_dir,
            is_lab_report=True,
        )

        # 验证降级处理
        doc = Document(file_path)

        # 应该以普通作业方式写入
        found_content = False
        for paragraph in doc.paragraphs:
            if "D" in paragraph.text or "【格式错误-已锁定】" in paragraph.text:
                found_content = True
                break

        self.assertTrue(found_content, "应该有降级处理的内容")


class ClearLabReportGradeAndCommentTest(TestCase):
    """测试 clear_lab_report_grade_and_comment 函数 - 撤销实验报告评分功能"""

    def setUp(self):
        self.temp_dir = tempfile.mkdtemp()

    def tearDown(self):
        if os.path.exists(self.temp_dir):
            shutil.rmtree(self.temp_dir)

    def create_lab_report_with_grade(self):
        """创建包含评分和评价的实验报告"""
        doc = Document()
        table = doc.add_table(rows=2, cols=2)

        # 添加"教师（签字）"单元格，包含评分和评价
        cell = table.rows[1].cells[0]
        cell.text = "A\n作业完成得非常出色\n教师（签字）：张老师 时间：2024-01-01"

        return doc, cell

    def test_clear_grade_and_comment_success(self):
        """测试成功清除评分和评价"""
        from grading.views import clear_lab_report_grade_and_comment

        doc, original_cell = self.create_lab_report_with_grade()

        # 清除评分和评价
        result = clear_lab_report_grade_and_comment(doc)

        self.assertTrue(result)

        # 验证单元格内容
        cell, _, _, _ = find_teacher_signature_cell(doc)
        self.assertIsNotNone(cell)

        cell_text = cell.text
        # 评分和评价应该被清除
        self.assertNotIn("A", cell_text)
        self.assertNotIn("作业完成得非常出色", cell_text)
        # 签字文本应该保留
        self.assertIn("教师（签字）", cell_text)
        self.assertIn("张老师", cell_text)

    def test_clear_grade_only_signature_remains(self):
        """测试清除后只保留签字文本"""
        from grading.views import clear_lab_report_grade_and_comment

        doc, _ = self.create_lab_report_with_grade()

        clear_lab_report_grade_and_comment(doc)

        cell, _, _, _ = find_teacher_signature_cell(doc)
        cell_text = cell.text.strip()

        # 应该只包含签字文本
        self.assertTrue(cell_text.startswith("教师（签字）"))

    def test_clear_no_signature_cell(self):
        """测试没有教师签字单元格的情况"""
        from grading.views import clear_lab_report_grade_and_comment

        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "学生姓名"

        result = clear_lab_report_grade_and_comment(doc)

        self.assertFalse(result)

    def test_clear_empty_signature_cell(self):
        """测试签字单元格为空的情况"""
        from grading.views import clear_lab_report_grade_and_comment

        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        cell = table.rows[1].cells[0]
        cell.text = "B\n还不错\n教师（签字）："

        result = clear_lab_report_grade_and_comment(doc)

        self.assertTrue(result)

        # 验证清除后的内容
        cell, _, _, _ = find_teacher_signature_cell(doc)
        cell_text = cell.text.strip()
        self.assertEqual(cell_text, "教师（签字）：")

    def test_clear_with_file_save(self):
        """测试清除后保存文件"""
        from grading.views import clear_lab_report_grade_and_comment

        doc, _ = self.create_lab_report_with_grade()

        # 保存原始文档
        original_path = os.path.join(self.temp_dir, "test_with_grade.docx")
        doc.save(original_path)

        # 重新加载并清除
        doc = Document(original_path)
        result = clear_lab_report_grade_and_comment(doc)
        self.assertTrue(result)

        # 保存清除后的文档
        cleared_path = os.path.join(self.temp_dir, "test_cleared.docx")
        doc.save(cleared_path)

        # 重新加载并验证
        doc_reloaded = Document(cleared_path)
        cell, _, _, _ = find_teacher_signature_cell(doc_reloaded)

        self.assertIsNotNone(cell)
        cell_text = cell.text
        self.assertNotIn("A", cell_text)
        self.assertNotIn("作业完成得非常出色", cell_text)
        self.assertIn("教师（签字）", cell_text)

    def test_clear_chinese_grade(self):
        """测试清除中文评分"""
        from grading.views import clear_lab_report_grade_and_comment

        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        cell = table.rows[1].cells[0]
        cell.text = "优秀\n实验报告格式规范\n教师（签字）：李老师"

        result = clear_lab_report_grade_and_comment(doc)

        self.assertTrue(result)

        cell, _, _, _ = find_teacher_signature_cell(doc)
        cell_text = cell.text
        self.assertNotIn("优秀", cell_text)
        self.assertNotIn("实验报告格式规范", cell_text)
        self.assertIn("李老师", cell_text)

    def test_clear_multiple_times(self):
        """测试多次清除"""
        from grading.views import clear_lab_report_grade_and_comment

        doc, _ = self.create_lab_report_with_grade()

        # 第一次清除
        result1 = clear_lab_report_grade_and_comment(doc)
        self.assertTrue(result1)

        # 第二次清除（应该仍然成功，但内容已经是空的）
        result2 = clear_lab_report_grade_and_comment(doc)
        self.assertTrue(result2)

        # 验证签字文本仍然存在
        cell, _, _, _ = find_teacher_signature_cell(doc)
        self.assertIn("教师（签字）", cell.text)

    def test_clear_preserves_signature_formatting(self):
        """测试清除后保留签字文本的格式"""
        from grading.views import clear_lab_report_grade_and_comment

        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        cell = table.rows[1].cells[0]

        # 添加复杂的签字文本
        signature_text = "教师（签字）：王老师\n时间：2024-01-15\n备注：已审核"
        cell.text = f"B\n需要改进\n{signature_text}"

        result = clear_lab_report_grade_and_comment(doc)
        self.assertTrue(result)

        # 验证签字文本完整保留
        cell, _, _, _ = find_teacher_signature_cell(doc)
        cell_text = cell.text
        self.assertIn("王老师", cell_text)
        self.assertIn("2024-01-15", cell_text)
        self.assertIn("已审核", cell_text)

    def test_clear_with_half_width_parentheses(self):
        """测试清除半角括号的教师(签字)"""
        from grading.views import clear_lab_report_grade_and_comment

        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        cell = table.rows[1].cells[0]
        cell.text = "C\n一般\n教师(签字)：测试老师"

        result = clear_lab_report_grade_and_comment(doc)

        self.assertTrue(result)

        cell, _, _, _ = find_teacher_signature_cell(doc)
        cell_text = cell.text
        self.assertNotIn("C", cell_text)
        self.assertNotIn("一般", cell_text)
        self.assertIn("测试老师", cell_text)


class UndoGradingIntegrationTest(TestCase):
    """撤销评分功能集成测试"""

    def setUp(self):
        self.temp_dir = tempfile.mkdtemp()

    def tearDown(self):
        if os.path.exists(self.temp_dir):
            shutil.rmtree(self.temp_dir)

    def test_complete_undo_workflow_lab_report(self):
        """测试完整的撤销工作流：实验报告"""
        # 步骤1：创建带评分的实验报告
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        cell = table.rows[1].cells[0]

        # 使用统一函数写入评分
        signature_text = "教师（签字）：张老师 时间：2024-01-01"
        write_to_teacher_signature_cell(cell, "A", "作业完成得非常出色", signature_text)

        # 验证评分已写入
        grade, comment, sig = extract_grade_and_comment_from_cell(cell)
        self.assertEqual(grade, "A")
        self.assertEqual(comment, "作业完成得非常出色")

        # 步骤2：撤销评分
        from grading.views import clear_lab_report_grade_and_comment

        result = clear_lab_report_grade_and_comment(doc)
        self.assertTrue(result)

        # 步骤3：验证评分已清除
        grade_after, comment_after, sig_after = extract_grade_and_comment_from_cell(cell)
        self.assertIsNone(grade_after)
        self.assertIsNone(comment_after)
        # 签字文本应该保留
        self.assertIn("张老师", sig_after)

    def test_undo_then_regrade_lab_report(self):
        """测试撤销后重新评分"""
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        cell = table.rows[1].cells[0]

        # 第一次评分
        signature_text = "教师（签字）：李老师"
        write_to_teacher_signature_cell(cell, "B", "还不错", signature_text)

        # 撤销
        from grading.views import clear_lab_report_grade_and_comment

        clear_lab_report_grade_and_comment(doc)

        # 重新评分
        cell, _, _, _ = find_teacher_signature_cell(doc)
        _, _, sig = extract_grade_and_comment_from_cell(cell)
        write_to_teacher_signature_cell(cell, "A", "修改后的评价", sig)

        # 验证新评分
        grade, comment, _ = extract_grade_and_comment_from_cell(cell)
        self.assertEqual(grade, "A")
        self.assertEqual(comment, "修改后的评价")

    def test_undo_preserves_document_structure(self):
        """测试撤销不影响文档其他部分"""
        doc = Document()

        # 添加标题
        doc.add_heading("实验报告", 0)

        # 添加内容段落
        doc.add_paragraph("这是实验内容...")

        # 添加评分表格
        table = doc.add_table(rows=2, cols=2)
        cell = table.rows[1].cells[0]
        signature_text = "教师（签字）：王老师"
        write_to_teacher_signature_cell(cell, "A", "很好", signature_text)

        # 添加更多内容
        doc.add_paragraph("实验结论...")

        # 撤销评分
        from grading.views import clear_lab_report_grade_and_comment

        clear_lab_report_grade_and_comment(doc)

        # 验证文档结构完整
        self.assertEqual(len(doc.paragraphs), 3)  # 标题 + 2个段落
        self.assertEqual(len(doc.tables), 1)  # 表格仍然存在
        self.assertEqual(doc.paragraphs[0].text, "实验报告")
        self.assertEqual(doc.paragraphs[1].text, "这是实验内容...")
        self.assertEqual(doc.paragraphs[2].text, "实验结论...")
