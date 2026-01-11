import logging
from typing import Optional, Tuple

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

logger = logging.getLogger(__name__)


def _iter_tables(container):
    for table in container.tables:
        yield table
        for row in table.rows:
            for cell in row.cells:
                yield from _iter_tables(cell)


def _normalize_signature_text(text: str) -> str:
    if not text:
        return ""
    normalized = text.replace("（", "(").replace("）", ")").replace("：", ":")
    normalized = normalized.replace("\u3000", "")
    normalized = "".join(normalized.split())
    return normalized


def _has_signature_marker(text: str) -> bool:
    normalized = _normalize_signature_text(text)
    return (
        "教师(签字)" in normalized
        or "教师签字" in normalized
        or ("教师" in normalized and "签字" in normalized)
        or "指导教师" in normalized
        or "指导老师" in normalized
    )


def find_teacher_signature_cell(doc):
    """
    查找实验报告中包含"教师（签字）"的单元格

    Returns:
        tuple: (cell, table_idx, row_idx, col_idx) 如果找到，否则返回 (None, None, None, None)
    """
    for table_idx, table in enumerate(_iter_tables(doc)):
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                if _has_signature_marker(cell_text):
                    logger.info(
                        "找到'教师（签字）'单元格: 表格%d, 行%d, 列%d",
                        table_idx + 1,
                        row_idx + 1,
                        col_idx + 1,
                    )
                    return cell, table_idx, row_idx, col_idx
    return None, None, None, None


def extract_grade_and_comment_from_cell(cell) -> Tuple[Optional[str], Optional[str], str]:
    """
    从"教师（签字）"单元格中提取评分、评价和签字文本

    单元格格式：
    第一行：评分（如"A"）
    第二行：评价（如"作业完成得非常出色..."）
    第三行及之后：教师（签字）：时间：...

    Returns:
        tuple: (grade, comment, signature_text)
    """
    logger.info("=== 开始提取单元格内容 ===")

    cell_text = cell.text.strip()
    lines = cell_text.split("\n")

    logger.info("单元格总行数: %d", len(lines))
    for i, line in enumerate(lines):
        logger.info("  第%d行: %s...", i + 1, line[:50])

    grade = None
    comment = None
    signature_text = ""

    signature_line_idx = -1
    for i, line in enumerate(lines):
        if _has_signature_marker(line):
            signature_line_idx = i
            signature_text = "\n".join(lines[i:])
            logger.info("✓ 找到'教师（签字）'在第%d行", i + 1)
            break

    if signature_line_idx == -1:
        logger.warning("✗ 单元格中未找到'教师（签字）'文本")
        return None, None, ""

    before_signature = lines[:signature_line_idx]
    logger.info("'教师（签字）'之前有%d行内容", len(before_signature))

    if len(before_signature) >= 1:
        potential_grade = before_signature[0].strip()
        if potential_grade in ["A", "B", "C", "D", "E", "优秀", "良好", "中等", "及格", "不及格"]:
            grade = potential_grade
            logger.info("[OK] 提取到评分（第一行）: %s", grade)
        else:
            logger.warning("[WARN] 第一行不是有效评分: %s", potential_grade)

    if len(before_signature) >= 2:
        comment = before_signature[1].strip()
        if comment:
            logger.info("[OK] 提取到评价（第二行）: %s...", comment[:50])
        else:
            logger.info("[WARN] 第二行为空，无评价")
    else:
        logger.info("[WARN] 没有第二行，无评价")

    logger.info("[OK] 提取到签字文本: %s...", signature_text[:50])
    logger.info("=== 单元格内容提取完成 ===")

    return grade, comment, signature_text


def build_teacher_signature_text(teacher_name, sign_time):
    if not teacher_name:
        teacher_name = ""
    date_str = sign_time.strftime("%Y年%m月%d日") if sign_time else ""
    return f"教师（签字）：{teacher_name}\n时间：{date_str}"


def write_to_teacher_signature_cell(
    cell, grade, comment, signature_text, teacher_name=None, sign_time=None
):
    """
    向"教师（签字）"单元格写入评分和评价

    写入格式：
    第一行：评分（如"A"）
    第二行：评价（如"作业完成得非常出色..."）
    第三行及之后：教师（签字）：时间：...（保留原有内容）
    """
    logger.info("=== 开始写入教师签字单元格 ===")
    logger.info("评分: %s", grade)
    logger.info("评价: %s", comment)
    logger.info("签字文本: %s...", signature_text[:50] if signature_text else "无")

    for paragraph in cell.paragraphs:
        paragraph.clear()

    while len(cell.paragraphs) > 1:
        p = cell.paragraphs[-1]._element
        p.getparent().remove(p)

    logger.info("已清空单元格内容")

    p1 = cell.paragraphs[0]
    run1 = p1.add_run(grade)
    run1.font.size = Pt(14)
    run1.bold = True
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logger.info("✓ 已写入评分（第一行）: %s", grade)

    if comment:
        p2 = cell.add_paragraph()
        run2 = p2.add_run(comment)
        run2.font.size = Pt(11)
        p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        logger.info("✓ 已写入评价（第二行）: %s...", comment[:50])
    else:
        logger.info("✗ 未提供评价，跳过第二行")

    signature_output = signature_text
    if teacher_name or sign_time:
        signature_output = build_teacher_signature_text(teacher_name, sign_time)

    if signature_output:
        p3 = cell.add_paragraph()
        run3 = p3.add_run(signature_output)
        run3.font.size = Pt(10)
        p3.alignment = WD_ALIGN_PARAGRAPH.LEFT
        logger.info("✓ 已写入签字文本（第三行及之后）: %s...", signature_output[:50])
    else:
        logger.warning("✗ 未找到签字文本，可能导致格式不完整")

    logger.info("=== 教师签字单元格写入完成 ===")


def extract_grade_from_homework_doc(doc) -> Optional[str]:
    """
    从普通作业的段落中提取成绩。

    读取规则与写入逻辑对齐，仅识别固定前缀行：
    - 老师评分：
    - 评定分数：
    - 教师评分：
    """
    prefixes = ("老师评分：", "评定分数：", "教师评分：")
    for paragraph in reversed(doc.paragraphs):
        text = paragraph.text.strip()
        if not text:
            continue
        for prefix in prefixes:
            if text.startswith(prefix):
                grade = text[len(prefix):].strip()
                return grade or None
    return None
