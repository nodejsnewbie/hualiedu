"""
作业成绩写入成绩登分册工具类

包含三个核心工具类：
- GradeFileProcessor: 处理作业成绩文件
- RegistryManager: 管理Excel登分册
- NameMatcher: 学生姓名匹配
"""

import os
import re
import logging
import shutil
import fcntl
import errno
from datetime import datetime
from typing import Optional, Tuple, List, Dict

from docx import Document
from openpyxl import load_workbook
from openpyxl.workbook import Workbook
from openpyxl.worksheet.worksheet import Worksheet


logger = logging.getLogger(__name__)

CHINESE_NUMERAL_CHARS = "零〇一二两三四五六七八九十百千"
CHINESE_NUMERAL_DIGIT_MAP = {
    "零": 0,
    "〇": 0,
    "一": 1,
    "二": 2,
    "两": 2,
    "三": 3,
    "四": 4,
    "五": 5,
    "六": 6,
    "七": 7,
    "八": 8,
    "九": 9,
}
CHINESE_NUMERAL_UNIT_MAP = {
    "十": 10,
    "百": 100,
    "千": 1000,
}
CHINESE_NUMERAL_ALLOWED_SET = set(CHINESE_NUMERAL_CHARS)


def _convert_chinese_numeral_to_int(token: str) -> Optional[int]:
    """将中文数字（十以内/十几/几十/百以内）转换为整数。"""
    if not token:
        return None

    token = token.strip()
    if not token:
        return None

    total = 0
    current = 0
    has_value = False

    for ch in token:
        if ch in CHINESE_NUMERAL_DIGIT_MAP:
            current = CHINESE_NUMERAL_DIGIT_MAP[ch]
            has_value = True
        elif ch in CHINESE_NUMERAL_UNIT_MAP:
            multiplier = CHINESE_NUMERAL_UNIT_MAP[ch]
            if current == 0:
                # 处理“十、百、千”这种前面省略“1”的情况
                current = 1
            total += current * multiplier
            current = 0
            has_value = True
        else:
            # 出现无法识别的字符，放弃解析
            return None

    value = total + current
    if has_value and value > 0:
        return value
    return None


def _parse_homework_number_token(token: str) -> Optional[int]:
    """解析捕获到的数字（阿拉伯数字或中文数字）"""
    if not token:
        return None

    digit_match = re.search(r"\d+", token)
    if digit_match:
        return int(digit_match.group(0))

    chinese_part = "".join(ch for ch in token if ch in CHINESE_NUMERAL_ALLOWED_SET)
    if chinese_part:
        return _convert_chinese_numeral_to_int(chinese_part)

    return None


class GradeFileProcessor:
    """作业成绩文件处理器"""

    # 支持的成绩等级
    VALID_GRADES = ["A", "B", "C", "D", "E", "优秀", "良好", "中等", "及格", "不及格"]
    
    # 文件数量限制
    MAX_FILES_WARNING_THRESHOLD = 500

    @staticmethod
    def extract_student_name(file_path: str) -> Optional[str]:
        """
        从文件名提取学生姓名
        
        支持的格式：
        - 姓名_作业X.docx
        - 作业X_姓名.docx
        - 姓名-作业X.docx
        - 作业X-姓名.docx
        
        Args:
            file_path: 文件路径
            
        Returns:
            学生姓名，如果无法提取则返回None
        """
        try:
            logger.debug("开始提取学生姓名: %s", file_path)

            # 获取文件名（不含扩展名）
            filename = os.path.splitext(os.path.basename(file_path))[0]
            
            # 尝试多种分隔符
            for separator in ["_", "-", "—"]:
                if separator in filename:
                    parts = filename.split(separator)
                    
                    # 检查是否包含作业标识
                    for i, part in enumerate(parts):
                        if re.search(r"作业\d+|homework\d+|hw\d+", part, re.IGNORECASE):
                            # 作业标识在前，姓名在后
                            if i == 0 and len(parts) > 1:
                                student_name = parts[1].strip()
                                logger.debug("提取到学生姓名: %s (从文件名)", student_name)
                                return student_name
                            # 姓名在前，作业标识在后
                            elif i > 0:
                                student_name = parts[0].strip()
                                logger.debug("提取到学生姓名: %s (从文件名)", student_name)
                                return student_name
            
            # 如果文件名本身不包含“作业”“homework”等关键字，直接使用文件名
            if not re.search(r"(作业|homework|hw)\s*\d+", filename, re.IGNORECASE) and "作业" not in filename:
                logger.debug("提取到学生姓名: %s (从文件名整体)", filename.strip())
                return filename.strip()

            # 如果没有分隔符，尝试从路径中提取
            # 例如：/path/to/张三/作业1.docx
            parent_dir = os.path.basename(os.path.dirname(file_path))
            if parent_dir and "作业" not in parent_dir and not re.search(r"(homework|hw)\s*\d+", parent_dir, re.IGNORECASE):
                logger.debug("提取到学生姓名: %s (从目录名)", parent_dir.strip())
                return parent_dir.strip()
            
            logger.warning("无法从文件名提取学生姓名: %s", file_path)
            return None
            
        except (OSError, ValueError, AttributeError) as e:
            logger.error("提取学生姓名时出错: %s, 错误: %s", file_path, str(e), exc_info=True)
            return None

    @staticmethod
    def extract_homework_number_from_path(file_path: str) -> Optional[int]:
        """
        从目录路径提取作业次数（作业评分系统场景）
        
        支持的格式：
        - 第X次作业
        - 作业X
        - homework_X
        - hw_X
        
        Args:
            file_path: 文件路径或目录路径
            
        Returns:
            作业次数，如果无法识别则返回None
        """
        try:
            logger.debug("开始从路径提取作业次数: %s", file_path)

            # 从完整路径中查找作业次数
            # 匹配模式：第X次作业、作业X、homework_X、hw_X
            chinese_digits = CHINESE_NUMERAL_CHARS
            patterns = [
                rf"第([\d{chinese_digits}]+)次作业",
                rf"作业([\d{chinese_digits}]+)",
                rf"homework[_\s]?([\d{chinese_digits}]+)",
                rf"hw[_\s]?([\d{chinese_digits}]+)",
            ]
            
            for pattern in patterns:
                match = re.search(pattern, file_path, re.IGNORECASE)
                if match:
                    homework_number = _parse_homework_number_token(match.group(1))
                    if homework_number is not None:
                        logger.debug("提取到作业次数: %d (模式: %s)", homework_number, pattern)
                        return homework_number
            
            logger.warning("无法从路径提取作业次数: %s", file_path)
            return None
            
        except (ValueError, AttributeError) as e:
            logger.error("提取作业次数时出错: %s, 错误: %s", file_path, str(e), exc_info=True)
            return None

    @staticmethod
    def extract_homework_number_from_filename(file_path: str) -> Optional[int]:
        """
        从文件名提取作业次数（工具箱模块场景）
        
        支持的格式：
        - 第X次作业成绩.xlsx
        - 作业X成绩.xlsx
        - homework_X.xlsx
        - hw_X.xlsx
        
        Args:
            file_path: 文件路径
            
        Returns:
            作业次数，如果无法识别则返回None
        """
        try:
            logger.debug("开始从文件名提取作业次数: %s", file_path)

            # 获取文件名（不含扩展名）
            filename = os.path.splitext(os.path.basename(file_path))[0]
            
            # 匹配模式：第X次作业、作业X、homework_X、hw_X
            chinese_digits = CHINESE_NUMERAL_CHARS
            patterns = [
                rf"第([\d{chinese_digits}]+)次作业",
                rf"作业([\d{chinese_digits}]+)",
                rf"homework[_\s]?([\d{chinese_digits}]+)",
                rf"hw[_\s]?([\d{chinese_digits}]+)",
            ]
            
            for pattern in patterns:
                match = re.search(pattern, filename, re.IGNORECASE)
                if match:
                    homework_number = _parse_homework_number_token(match.group(1))
                    if homework_number is not None:
                        logger.debug("提取到作业次数: %d (模式: %s)", homework_number, pattern)
                        return homework_number
            
            logger.warning("无法从文件名提取作业次数: %s", file_path)
            return None
            
        except (OSError, ValueError, AttributeError) as e:
            logger.error("提取作业次数时出错: %s, 错误: %s", file_path, str(e), exc_info=True)
            return None

    @staticmethod
    def is_lab_report(file_path: str) -> bool:
        """
        判断是否为实验报告
        
        Args:
            file_path: 文件路径
            
        Returns:
            True表示是实验报告，False表示是普通作业
        """
        try:
            doc = Document(file_path)
            
            # 检查文档中是否包含实验报告的特征标记
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if "实验报告" in text or "实验名称" in text:
                    return True
            
            # 检查表格中是否有"教师（签字）："标记
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if "教师（签字）" in cell.text or "教师签字" in cell.text:
                            return True
            
            return False
            
        except (OSError, ValueError) as e:
            logger.error("判断是否为实验报告时出错: %s, 错误: %s", file_path, str(e))
            return False

    @staticmethod
    def extract_grade_from_word(file_path: str) -> Optional[str]:
        """
        从Word文档提取单个学生成绩（作业评分系统场景）
        
        Args:
            file_path: 文件路径
            
        Returns:
            成绩等级，如果无法提取则返回None
        """
        try:
            logger.debug("开始从Word文档提取成绩: %s", file_path)

            doc = Document(file_path)
            is_lab = GradeFileProcessor.is_lab_report(file_path)
            
            if is_lab:
                logger.debug("识别为实验报告，使用实验报告提取方法")
                # 实验报告：在"教师（签字）："单元格中查找评分
                grade = GradeFileProcessor._extract_grade_from_lab_report(doc)
            else:
                logger.debug("识别为普通作业，使用普通作业提取方法")
                # 普通作业：在文档末尾查找"老师评分："标记
                grade = GradeFileProcessor._extract_grade_from_homework(doc)

            if grade:
                logger.debug("成功提取成绩: %s", grade)
            else:
                logger.warning("未能提取到成绩: %s", file_path)

            return grade
                
        except (OSError, ValueError) as e:
            logger.error("提取成绩时出错: %s, 错误: %s", file_path, str(e), exc_info=True)
            return None

    @staticmethod
    def extract_grades_from_excel(file_path: str) -> List[Dict[str, str]]:
        """
        从Excel文件提取所有学生成绩（工具箱模块场景）

        Excel文件格式示例：
        | 姓名   | 成绩 |
        |--------|------|
        | 张三   | A    |
        | 李四   | B    |

        Args:
            file_path: Excel文件路径

        Returns:
            学生成绩列表，格式: [{"name": "张三", "grade": "A"}, ...]
            如果提取失败返回空列表
        """
        try:
            logger.debug("开始从Excel文件提取学生成绩: %s", file_path)

            # 性能优化：使用read_only模式加载Excel文件，data_only=True提高读取速度
            workbook = load_workbook(file_path, read_only=True, data_only=True)
            worksheet = workbook.active

            # 查找"姓名"和"成绩"列（支持表头出现在前几行）
            name_col_idx = None
            grade_col_idx = None
            header_row = None
            max_scan_rows = min(worksheet.max_row, 20)

            for row_idx in range(1, max_scan_rows + 1):
                header_cells = list(
                    worksheet.iter_rows(min_row=row_idx, max_row=row_idx, values_only=True)
                )
                if not header_cells:
                    continue

                header_values = header_cells[0]
                local_name_idx = None
                local_grade_idx = None

                for col_idx, cell_value in enumerate(header_values, start=1):
                    if not cell_value:
                        continue
                    cell_text = str(cell_value).strip()
                    if local_name_idx is None and "姓名" in cell_text:
                        local_name_idx = col_idx
                    if local_grade_idx is None and (
                        "成绩" in cell_text or "分数" in cell_text or "等级" in cell_text
                    ):
                        local_grade_idx = col_idx

                if local_name_idx is not None:
                    name_col_idx = local_name_idx
                    grade_col_idx = local_grade_idx
                    header_row = row_idx
                    logger.debug("检测到表头行: %d (姓名列: %d, 成绩列: %s)", header_row, name_col_idx, grade_col_idx)
                    break

            if name_col_idx is None or header_row is None:
                logger.error("Excel文件缺少姓名列: %s", file_path)
                workbook.close()
                return []

            if grade_col_idx is None:
                logger.error("Excel文件缺少成绩列: %s", file_path)
                workbook.close()
                return []

            # 性能优化：批量读取所有数据行，避免逐行访问
            grades = []
            for row in worksheet.iter_rows(min_row=header_row + 1, values_only=True):
                if len(row) >= max(name_col_idx, grade_col_idx):
                    name = row[name_col_idx - 1]
                    grade = row[grade_col_idx - 1]

                    if name and grade:
                        name = str(name).strip()
                        grade = str(grade).strip()

                        # 验证成绩格式
                        if grade in GradeFileProcessor.VALID_GRADES:
                            grades.append({"name": name, "grade": grade})
                            logger.debug("提取学生成绩: %s - %s", name, grade)
                        else:
                            logger.warning(
                                "无效的成绩格式: %s, 学生: %s, 文件: %s",
                                grade, name, file_path
                            )

            workbook.close()
            logger.info("从Excel文件提取了%d个学生成绩: %s", len(grades), file_path)
            return grades

        except (OSError, ValueError, KeyError) as e:
            logger.error("从Excel提取成绩时出错: %s, 错误: %s", file_path, str(e), exc_info=True)
            return []

    @staticmethod
    def _extract_grade_from_lab_report(doc: Document) -> Optional[str]:
        """从实验报告提取成绩"""
        # 在表格中查找"教师（签字）："单元格
        for table in doc.tables:
            for row in table.rows:
                for i, cell in enumerate(row.cells):
                    cell_text = cell.text.strip()
                    if "教师（签字）" in cell_text or "教师签字" in cell_text:
                        # 检查同一单元格或下一个单元格中的成绩
                        grade = GradeFileProcessor._find_grade_in_text(cell_text)
                        if grade:
                            return grade
                        
                        # 检查下一个单元格
                        if i + 1 < len(row.cells):
                            next_cell_text = row.cells[i + 1].text.strip()
                            grade = GradeFileProcessor._find_grade_in_text(next_cell_text)
                            if grade:
                                return grade
        
        return None

    @staticmethod
    def _extract_grade_from_homework(doc: Document) -> Optional[str]:
        """从普通作业提取成绩"""
        # 从文档末尾向前查找"老师评分："标记
        paragraphs = doc.paragraphs
        for paragraph in reversed(paragraphs):
            text = paragraph.text.strip()
            if "老师评分" in text or "教师评分" in text:
                grade = GradeFileProcessor._find_grade_in_text(text)
                if grade:
                    return grade
        
        return None

    @staticmethod
    def _find_grade_in_text(text: str) -> Optional[str]:
        """在文本中查找成绩等级"""
        for grade in GradeFileProcessor.VALID_GRADES:
            if grade in text:
                return grade
        return None



class RegistryManager:
    """成绩登分册管理器"""

    def __init__(self, registry_path: str):
        """
        初始化登分册管理器
        
        Args:
            registry_path: Excel登分册文件路径
        """
        self.registry_path = registry_path
        self.workbook: Optional[Workbook] = None
        self.worksheet: Optional[Worksheet] = None
        self.backup_path: Optional[str] = None
        self.name_column_index: Optional[int] = None
        self.student_names: Dict[str, int] = {}  # 姓名 -> 行号映射
        self.lock_file_handle = None  # 文件锁句柄
        self.lock_file_path = None  # 锁文件路径
        self.header_row_index: int = 1

    def _acquire_file_lock(self) -> bool:
        """
        获取文件锁（并发控制）
        
        Returns:
            True表示成功获取锁，False表示文件被占用
        """
        try:
            # 创建锁文件路径
            self.lock_file_path = f"{self.registry_path}.lock"
            
            # 尝试打开锁文件
            self.lock_file_handle = open(self.lock_file_path, 'w')
            
            # 尝试获取独占锁（非阻塞）
            try:
                fcntl.flock(self.lock_file_handle.fileno(), fcntl.LOCK_EX | fcntl.LOCK_NB)
                logger.info("成功获取文件锁: %s", self.lock_file_path)
                return True
            except IOError as e:
                if e.errno in (errno.EACCES, errno.EAGAIN):
                    logger.error("文件被其他进程占用: %s", self.registry_path)
                    self.lock_file_handle.close()
                    self.lock_file_handle = None
                    return False
                else:
                    raise
                    
        except Exception as e:
            logger.error("获取文件锁失败: %s", str(e), exc_info=True)
            if self.lock_file_handle:
                self.lock_file_handle.close()
                self.lock_file_handle = None
            return False

    def _release_file_lock(self):
        """释放文件锁（并发控制）"""
        try:
            if self.lock_file_handle:
                # 释放锁
                fcntl.flock(self.lock_file_handle.fileno(), fcntl.LOCK_UN)
                self.lock_file_handle.close()
                self.lock_file_handle = None
                
                # 删除锁文件
                if self.lock_file_path and os.path.exists(self.lock_file_path):
                    os.remove(self.lock_file_path)
                    logger.info("成功释放文件锁: %s", self.lock_file_path)
                    
        except Exception as e:
            logger.error("释放文件锁失败: %s", str(e), exc_info=True)

    def _check_file_in_use(self) -> Tuple[bool, Optional[str]]:
        """
        检测Excel文件是否被占用（并发控制）
        
        Returns:
            (是否被占用, 错误消息)
        """
        try:
            # 方法1：尝试以独占模式打开文件
            try:
                with open(self.registry_path, 'r+b') as f:
                    # 尝试获取文件锁
                    try:
                        fcntl.flock(f.fileno(), fcntl.LOCK_EX | fcntl.LOCK_NB)
                        fcntl.flock(f.fileno(), fcntl.LOCK_UN)
                        return False, None
                    except IOError as e:
                        if e.errno in (errno.EACCES, errno.EAGAIN):
                            logger.error("文件被占用: %s", self.registry_path)
                            return True, "成绩登分册文件被占用，请关闭后重试"
                        else:
                            raise
            except PermissionError:
                logger.error("无权限访问文件: %s", self.registry_path)
                return True, "无权限访问文件"
                
        except Exception as e:
            logger.error("检测文件占用状态失败: %s", str(e), exc_info=True)
            return False, None

    def load(self) -> bool:
        """
        加载Excel文件（带并发控制）

        Returns:
            True表示加载成功，False表示失败
        """
        try:
            logger.debug("开始加载登分册: %s", self.registry_path)

            if not os.path.exists(self.registry_path):
                logger.error("登分册文件不存在: %s", self.registry_path)
                return False

            # 并发控制：检测文件是否被占用
            is_in_use, error_msg = self._check_file_in_use()
            if is_in_use:
                logger.error("文件被占用: %s - %s", self.registry_path, error_msg)
                return False

            # 并发控制：获取文件锁
            if not self._acquire_file_lock():
                logger.error("无法获取文件锁: %s", self.registry_path)
                return False

            # 性能优化：不使用read_only模式（因为需要写入），但使用data_only=False保留公式
            # keep_vba=False 可以减少内存占用
            self.workbook = load_workbook(
                self.registry_path,
                data_only=False,
                keep_vba=False
            )
            self.worksheet = self.workbook.active

            logger.info(
                "成功加载登分册: %s (工作表: %s)",
                self.registry_path,
                self.worksheet.title
            )
            return True

        except (OSError, ValueError) as e:
            logger.error("加载登分册失败: %s, 错误: %s", self.registry_path, str(e), exc_info=True)
            # 加载失败时释放锁
            self._release_file_lock()
            return False

    def validate_format(self) -> Tuple[bool, Optional[str]]:
        """
        验证登分册格式
        
        Returns:
            (是否有效, 错误消息)
        """
        try:
            logger.debug("开始验证登分册格式")

            if not self.worksheet:
                return False, "工作表未加载"
            
            # 查找"姓名"列（允许表头位于前几行）
            self.name_column_index = None
            self.header_row_index = None
            max_scan_rows = min(self.worksheet.max_row, 20)

            for row_idx in range(1, max_scan_rows + 1):
                for col_idx in range(1, self.worksheet.max_column + 1):
                    cell_value = self.worksheet.cell(row_idx, col_idx).value
                    if cell_value and "姓名" in str(cell_value):
                        self.name_column_index = col_idx
                        self.header_row_index = row_idx
                        logger.debug("找到姓名列: 行%d 列%d", row_idx, col_idx)
                        break
                if self.name_column_index is not None:
                    break

            if self.name_column_index is None or self.header_row_index is None:
                logger.error("成绩登分册格式错误：缺少姓名列")
                return False, "成绩登分册格式错误：缺少姓名列"
            
            # 构建学生姓名映射
            self._build_student_name_map()
            
            logger.info("登分册格式验证通过，找到%d个学生", len(self.student_names))
            logger.debug("学生列表: %s", list(self.student_names.keys())[:10])  # 只记录前10个
            return True, None
            
        except (ValueError, AttributeError, KeyError) as e:
            error_msg = "验证登分册格式时出错: {}".format(str(e))
            logger.error(error_msg, exc_info=True)
            return False, error_msg

    def _build_student_name_map(self):
        """构建学生姓名到行号的映射（性能优化：缓存学生列表）"""
        logger.debug("开始构建学生姓名映射")
        self.student_names = {}

        # 性能优化：批量读取所有学生姓名，避免重复查询
        # 使用iter_rows批量读取，比逐个访问单元格快得多
        header_row = self.header_row_index or 1
        for row_idx, row in enumerate(
            self.worksheet.iter_rows(
                min_row=header_row + 1,
                min_col=self.name_column_index,
                max_col=self.name_column_index,
                values_only=True
            ),
            start=header_row + 1
        ):
            name = row[0]
            if name:
                name = str(name).strip()
                if name:
                    self.student_names[name] = row_idx
                    logger.debug("添加学生: %s (行%d)", name, row_idx)

        logger.info("学生姓名映射构建完成，共%d个学生（已缓存）", len(self.student_names))

    def find_student_row(self, student_name: str) -> Optional[int]:
        """
        查找学生行
        
        Args:
            student_name: 学生姓名
            
        Returns:
            行号，如果未找到则返回None
        """
        return self.student_names.get(student_name)

    def find_or_create_homework_column(self, homework_number: int) -> int:
        """
        查找或创建作业列
        
        Args:
            homework_number: 作业次数
            
        Returns:
            列索引
        """
        try:
            logger.debug("查找或创建作业列: 第%d次作业", homework_number)

            header_row = self.header_row_index or 1
            if not hasattr(self, "name_column_index") or not self.name_column_index:
                raise ValueError("姓名列未初始化，无法定位作业列")

            target_col = self.name_column_index + homework_number
            self.worksheet.cell(header_row, target_col)  # 确保单元格存在，但不写入标题
            return target_col
            
        except (ValueError, AttributeError) as e:
            logger.error("查找或创建作业列时出错: %s", str(e), exc_info=True)
            raise

    def write_grade(self, row: int, col: int, grade: str) -> Tuple[bool, Optional[str]]:
        """
        写入成绩到指定单元格
        
        Args:
            row: 行号
            col: 列号
            grade: 成绩
            
        Returns:
            (是否成功, 旧成绩值)
        """
        try:
            logger.debug("准备写入成绩: 行%d, 列%d, 成绩%s", row, col, grade)

            cell = self.worksheet.cell(row, col)
            old_grade = cell.value
            
            # 如果旧成绩和新成绩相同，跳过
            if old_grade == grade:
                logger.debug("成绩相同，跳过写入: 行%d, 列%d, 成绩%s", row, col, grade)
                return False, old_grade
            
            # 写入新成绩
            cell.value = grade
            
            if old_grade:
                logger.info("覆盖成绩: 行%d, 列%d, 旧成绩%s -> 新成绩%s", row, col, old_grade, grade)
            else:
                logger.info("写入成绩: 行%d, 列%d, 成绩%s", row, col, grade)
            
            return True, old_grade
            
        except (ValueError, AttributeError) as e:
            logger.error("写入成绩时出错: 行%d, 列%d, 错误: %s", row, col, str(e), exc_info=True)
            raise

    def save(self) -> bool:
        """
        保存Excel文件（带并发控制）
        
        Returns:
            True表示保存成功，False表示失败
        """
        try:
            logger.debug("开始保存登分册: %s", self.registry_path)

            if not self.workbook:
                logger.error("工作簿未加载，无法保存")
                return False
            
            # 并发控制：保存前再次检查文件锁
            if not self.lock_file_handle:
                logger.error("文件锁已丢失，无法保存: %s", self.registry_path)
                return False
            
            self.workbook.save(self.registry_path)
            logger.info("成功保存登分册: %s", self.registry_path)
            
            # 并发控制：保存成功后释放文件锁
            self._release_file_lock()
            return True
            
        except (OSError, ValueError) as e:
            logger.error("保存登分册失败: %s, 错误: %s", self.registry_path, str(e), exc_info=True)
            # 保存失败时也要释放锁
            self._release_file_lock()
            return False

    def create_backup(self) -> bool:
        """
        创建备份
        
        Returns:
            True表示备份成功，False表示失败
        """
        try:
            logger.debug("开始创建登分册备份")

            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            backup_filename = f"{os.path.splitext(self.registry_path)[0]}_backup_{timestamp}.xlsx"
            self.backup_path = backup_filename
            
            shutil.copy2(self.registry_path, self.backup_path)
            logger.info("创建备份成功: %s", self.backup_path)
            return True
            
        except (OSError, IOError) as e:
            logger.error("创建备份失败: %s", str(e), exc_info=True)
            return False

    def restore_from_backup(self) -> bool:
        """
        从备份恢复（带并发控制）
        
        Returns:
            True表示恢复成功，False表示失败
        """
        try:
            logger.warning("开始从备份恢复登分册")

            if not self.backup_path or not os.path.exists(self.backup_path):
                logger.error("备份文件不存在，无法恢复: %s", self.backup_path)
                return False
            
            shutil.copy2(self.backup_path, self.registry_path)
            logger.info("从备份恢复成功: %s -> %s", self.backup_path, self.registry_path)
            
            # 并发控制：恢复后释放文件锁
            self._release_file_lock()
            return True
            
        except (OSError, IOError) as e:
            logger.error("从备份恢复失败: %s", str(e), exc_info=True)
            # 恢复失败时也要释放锁
            self._release_file_lock()
            return False

    def delete_backup(self) -> bool:
        """
        删除备份文件
        
        Returns:
            True表示删除成功，False表示失败
        """
        try:
            logger.debug("开始删除备份文件")

            if self.backup_path and os.path.exists(self.backup_path):
                os.remove(self.backup_path)
                logger.info("删除备份成功: %s", self.backup_path)
                return True
            
            logger.debug("备份文件不存在，无需删除")
            return False
            
        except OSError as e:
            logger.error("删除备份失败: %s", str(e), exc_info=True)
            return False
    
    def __del__(self):
        """析构函数：确保释放文件锁"""
        self._release_file_lock()


class NameMatcher:
    """学生姓名匹配器"""

    @staticmethod
    def exact_match(name: str, name_list: List[str]) -> Optional[str]:
        """
        精确匹配
        
        Args:
            name: 要匹配的姓名
            name_list: 姓名列表
            
        Returns:
            匹配的姓名，如果未找到则返回None
        """
        if name in name_list:
            return name
        return None

    @staticmethod
    def fuzzy_match(name: str, name_list: List[str]) -> Tuple[Optional[str], List[str]]:
        """
        模糊匹配（去除空格和特殊字符）
        
        Args:
            name: 要匹配的姓名
            name_list: 姓名列表
            
        Returns:
            (匹配的姓名, 所有匹配结果列表)
            如果找到唯一匹配返回该姓名，如果有多个匹配或无匹配返回None
        """
        normalized_name = NameMatcher.normalize_name(name)
        matches = []
        
        for candidate in name_list:
            normalized_candidate = NameMatcher.normalize_name(candidate)
            if normalized_name == normalized_candidate:
                matches.append(candidate)
        
        if len(matches) == 1:
            return matches[0], matches
        elif len(matches) > 1:
            logger.warning("姓名'%s'有多个匹配: %s", name, matches)
            return None, matches
        else:
            return None, []

    @staticmethod
    def normalize_name(name: str) -> str:
        """
        规范化姓名（去除空格和特殊字符）
        
        Args:
            name: 原始姓名
            
        Returns:
            规范化后的姓名
        """
        if not name:
            return ""
        
        # 去除所有空格
        normalized = name.replace(" ", "").replace("\t", "").replace("\n", "")
        
        # 去除常见特殊字符
        special_chars = ["·", "•", ".", "。", "-", "_", "—"]
        for char in special_chars:
            normalized = normalized.replace(char, "")
        
        return normalized.strip()

    @staticmethod
    def match(name: str, name_list: List[str]) -> Tuple[Optional[str], str]:
        """
        匹配姓名（先精确匹配，再模糊匹配）
        
        Args:
            name: 要匹配的姓名
            name_list: 姓名列表
            
        Returns:
            (匹配的姓名, 匹配类型: 'exact'/'fuzzy'/'none'/'multiple')
        """
        logger.debug("开始匹配姓名: %s", name)

        # 先尝试精确匹配
        exact_result = NameMatcher.exact_match(name, name_list)
        if exact_result:
            logger.debug("精确匹配成功: %s", exact_result)
            return exact_result, "exact"
        
        # 再尝试模糊匹配
        fuzzy_result, all_matches = NameMatcher.fuzzy_match(name, name_list)
        if fuzzy_result:
            logger.debug("模糊匹配成功: %s -> %s", name, fuzzy_result)
            return fuzzy_result, "fuzzy"
        elif len(all_matches) > 1:
            logger.warning("姓名匹配到多个结果: %s -> %s", name, all_matches)
            return None, "multiple"
        else:
            logger.warning("姓名匹配失败: %s", name)
            return None, "none"
