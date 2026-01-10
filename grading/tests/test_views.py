"""
Comprehensive pytest tests for grading/views.py

This module tests the core view functions and utility functions in the grading views,
including file operations, validation, directory handling, and API endpoints.
"""

import json
import os
import tempfile
import shutil
from unittest.mock import Mock, patch, MagicMock
from io import BytesIO

import pytest
from django.test import TestCase, RequestFactory, Client
from django.contrib.auth import get_user_model
from django.http import JsonResponse, HttpResponse
from django.core.cache import cache
from docx import Document

from grading.models import (
    Repository, Course, Homework, Semester, GlobalConfig, 
    FileGradeStatus, Tenant, UserProfile
)
from grading.views import (
    get_base_directory,
    validate_file_path,
    get_teacher_display_name,
    validate_file_write_permission,
    validate_user_permissions,
    create_error_response,
    create_success_response,
    read_file_content,
    get_file_extension,
    is_lab_course_by_name,
    auto_detect_course_type,
    get_course_type_from_name,
    is_lab_report_file,
    get_directory_file_count_cached,
    clear_directory_file_count_cache,
    get_file_grade_info,
    _extract_homework_folder,
    _clean_relative_homework_path,
    _split_rel_path_parts,
    _is_homework_folder_rel_path,
    _is_homework_file_rel_path,
)

User = get_user_model()


class BaseViewTestCase(TestCase):
    """Base test case with common setup for view tests."""
    
    def setUp(self):
        """Set up test data."""
        self.factory = RequestFactory()
        self.client = Client()
        
        # Create test user
        self.user = User.objects.create_user(
            username='testuser',
            email='test@example.com',
            password='testpass123',
            is_staff=True
        )
        
        # Create tenant and user profile
        self.tenant = Tenant.objects.create(
            name='Test Tenant',
            domain='test.example.com'
        )
        
        self.user_profile = UserProfile.objects.create(
            user=self.user,
            tenant=self.tenant,
            repo_base_dir='~/test_jobs'
        )
        
        # Create test semester
        self.semester = Semester.objects.create(
            name='Test Semester',
            start_date='2024-01-01',
            end_date='2024-06-30',
            is_active=True
        )
        
        # Create test course
        self.course = Course.objects.create(
            name='Test Course',
            course_type='theory',
            semester=self.semester,
            teacher=self.user,
            location='Room 101'
        )
        
        # Create test repository
        self.repository = Repository.objects.create(
            name='test-repo',
            url='https://github.com/test/repo.git',
            owner=self.user,
            tenant=self.tenant,
            repo_type='local',
            local_path='/tmp/test-repo'
        )
        
        # Create temporary directory for testing
        self.temp_dir = tempfile.mkdtemp()
        self.addCleanup(shutil.rmtree, self.temp_dir)
        
        # Mock request with user and tenant
        self.request = self.factory.get('/')
        self.request.user = self.user
        self.request.user_profile = self.user_profile
        self.request.tenant = self.tenant


class TestUtilityFunctions(BaseViewTestCase):
    """Test utility functions in views.py."""
    
    def test_get_base_directory_with_request(self):
        """Test get_base_directory with request object."""
        result = get_base_directory(self.request)
        expected = os.path.expanduser('~/test_jobs')
        self.assertEqual(result, expected)
    
    def test_get_base_directory_without_request(self):
        """Test get_base_directory without request object."""
        GlobalConfig.objects.create(
            key='default_repo_base_dir',
            value='~/default_jobs'
        )
        
        result = get_base_directory()
        expected = os.path.expanduser('~/default_jobs')
        self.assertEqual(result, expected)
    
    def test_get_base_directory_fallback(self):
        """Test get_base_directory with fallback value."""
        result = get_base_directory()
        expected = os.path.expanduser('~/jobs')
        self.assertEqual(result, expected)
    
    def test_validate_file_path_success(self):
        """Test successful file path validation."""
        # Create test file
        test_file = os.path.join(self.temp_dir, 'test.txt')
        with open(test_file, 'w') as f:
            f.write('test content')
        
        is_valid, full_path, error_msg = validate_file_path(
            'test.txt', 
            base_dir=self.temp_dir
        )
        
        self.assertTrue(is_valid)
        self.assertEqual(full_path, test_file)
        self.assertIsNone(error_msg)
    
    def test_validate_file_path_missing_file(self):
        """Test file path validation with missing file."""
        is_valid, full_path, error_msg = validate_file_path(
            'nonexistent.txt',
            base_dir=self.temp_dir
        )
        
        self.assertFalse(is_valid)
        self.assertIsNone(full_path)
        self.assertEqual(error_msg, '文件不存在')
    
    def test_validate_file_path_directory_traversal(self):
        """Test file path validation prevents directory traversal."""
        is_valid, full_path, error_msg = validate_file_path(
            '../../../etc/passwd',
            base_dir=self.temp_dir
        )
        
        self.assertFalse(is_valid)
        self.assertIsNone(full_path)
        self.assertEqual(error_msg, '无权访问该文件')
    
    def test_validate_file_path_empty_path(self):
        """Test file path validation with empty path."""
        is_valid, full_path, error_msg = validate_file_path('')
        
        self.assertFalse(is_valid)
        self.assertIsNone(full_path)
        self.assertEqual(error_msg, '未提供文件路径')
    
    def test_get_teacher_display_name_with_full_name(self):
        """Test get_teacher_display_name with full name."""
        user = User.objects.create_user(
            username='teacher',
            first_name='John',
            last_name='Doe'
        )
        
        result = get_teacher_display_name(user)
        self.assertEqual(result, 'John Doe')
    
    def test_get_teacher_display_name_with_username_only(self):
        """Test get_teacher_display_name with username only."""
        user = User.objects.create_user(username='teacher')
        
        result = get_teacher_display_name(user)
        self.assertEqual(result, 'teacher')
    
    def test_get_teacher_display_name_with_none(self):
        """Test get_teacher_display_name with None user."""
        result = get_teacher_display_name(None)
        self.assertEqual(result, '')
    
    def test_validate_file_write_permission_success(self):
        """Test successful file write permission validation."""
        test_file = os.path.join(self.temp_dir, 'writable.txt')
        with open(test_file, 'w') as f:
            f.write('test')
        
        is_valid, error_msg = validate_file_write_permission(test_file)
        
        self.assertTrue(is_valid)
        self.assertIsNone(error_msg)
    
    def test_validate_user_permissions_authenticated_staff(self):
        """Test user permissions validation for authenticated staff user."""
        is_valid, error_msg = validate_user_permissions(self.request)
        
        self.assertTrue(is_valid)
        self.assertIsNone(error_msg)
    
    def test_validate_user_permissions_not_authenticated(self):
        """Test user permissions validation for unauthenticated user."""
        request = self.factory.get('/')
        request.user = Mock()
        request.user.is_authenticated = False
        
        is_valid, error_msg = validate_user_permissions(request)
        
        self.assertFalse(is_valid)
        self.assertEqual(error_msg, '请先登录')
    
    def test_validate_user_permissions_not_staff(self):
        """Test user permissions validation for non-staff user."""
        request = self.factory.get('/')
        request.user = Mock()
        request.user.is_authenticated = True
        request.user.is_staff = False
        
        is_valid, error_msg = validate_user_permissions(request)
        
        self.assertFalse(is_valid)
        self.assertEqual(error_msg, '无权限访问')


class TestResponseHelpers(BaseViewTestCase):
    """Test response helper functions."""
    
    def test_create_error_response_default(self):
        """Test create_error_response with default parameters."""
        response = create_error_response('Test error')
        
        self.assertIsInstance(response, JsonResponse)
        self.assertEqual(response.status_code, 400)
        
        data = json.loads(response.content)
        self.assertEqual(data['status'], 'error')
        self.assertEqual(data['message'], 'Test error')
    
    def test_create_error_response_custom_status(self):
        """Test create_error_response with custom status code."""
        response = create_error_response('Not found', status_code=404)
        
        self.assertEqual(response.status_code, 404)
    
    def test_create_error_response_success_format(self):
        """Test create_error_response with success format."""
        response = create_error_response('Test error', response_format='success')
        
        data = json.loads(response.content)
        self.assertEqual(data['success'], False)
        self.assertEqual(data['message'], 'Test error')
    
    def test_create_success_response_default(self):
        """Test create_success_response with default parameters."""
        response = create_success_response()
        
        self.assertIsInstance(response, JsonResponse)
        self.assertEqual(response.status_code, 200)
        
        data = json.loads(response.content)
        self.assertEqual(data['status'], 'success')
        self.assertEqual(data['message'], '操作成功')
    
    def test_create_success_response_with_data(self):
        """Test create_success_response with custom data."""
        test_data = {'key': 'value'}
        response = create_success_response(data=test_data, message='Custom message')
        
        data = json.loads(response.content)
        self.assertEqual(data['status'], 'success')
        self.assertEqual(data['message'], 'Custom message')
        self.assertEqual(data['key'], 'value')
    
    def test_create_success_response_success_format(self):
        """Test create_success_response with success format."""
        response = create_success_response(response_format='success')
        
        data = json.loads(response.content)
        self.assertEqual(data['success'], True)
        self.assertEqual(data['message'], '操作成功')


class TestFileOperations(BaseViewTestCase):
    """Test file operation functions."""
    
    def test_read_file_content_text_file(self):
        """Test reading text file content."""
        test_file = os.path.join(self.temp_dir, 'test.txt')
        test_content = 'Hello, World!\nThis is a test file.'
        
        with open(test_file, 'w', encoding='utf-8') as f:
            f.write(test_content)
        
        result = read_file_content(test_file)
        self.assertEqual(result, test_content)
    
    def test_read_file_content_docx_file(self):
        """Test reading DOCX file content."""
        test_file = os.path.join(self.temp_dir, 'test.docx')
        
        # Create a simple DOCX file
        doc = Document()
        doc.add_paragraph('First paragraph')
        doc.add_paragraph('Second paragraph')
        doc.save(test_file)
        
        result = read_file_content(test_file)
        self.assertIn('First paragraph', result)
        self.assertIn('Second paragraph', result)
    
    def test_read_file_content_nonexistent_file(self):
        """Test reading nonexistent file returns empty string."""
        result = read_file_content('/nonexistent/file.txt')
        self.assertEqual(result, '')
    
    def test_get_file_extension_with_extension(self):
        """Test get_file_extension with file that has extension."""
        result = get_file_extension('/path/to/file.docx')
        self.assertEqual(result, 'docx')
    
    def test_get_file_extension_without_extension(self):
        """Test get_file_extension with file that has no extension."""
        result = get_file_extension('/path/to/file')
        self.assertEqual(result, 'unknown')
    
    def test_get_file_extension_multiple_dots(self):
        """Test get_file_extension with file that has multiple dots."""
        result = get_file_extension('/path/to/file.backup.txt')
        self.assertEqual(result, 'txt')


class TestCourseTypeDetection(BaseViewTestCase):
    """Test course type detection functions."""
    
    def test_is_lab_course_by_name_with_lab_keyword(self):
        """Test is_lab_course_by_name with lab keyword."""
        test_cases = [
            '计算机实验',
            'Computer Lab',
            'Physics Experiment',
            '数据结构实训',
            'Software Practice'
        ]
        
        for course_name in test_cases:
            with self.subTest(course_name=course_name):
                result = is_lab_course_by_name(course_name)
                self.assertTrue(result)
    
    def test_is_lab_course_by_name_without_lab_keyword(self):
        """Test is_lab_course_by_name without lab keyword."""
        test_cases = [
            '高等数学',
            'English Literature',
            '计算机理论',
            'History'
        ]
        
        for course_name in test_cases:
            with self.subTest(course_name=course_name):
                result = is_lab_course_by_name(course_name)
                self.assertFalse(result)
    
    def test_is_lab_course_by_name_with_database_course(self):
        """Test is_lab_course_by_name with database course."""
        course = Course.objects.create(
            name='Database Lab',
            course_type='lab',
            semester=self.semester,
            teacher=self.user,
            location='Lab 1'
        )
        
        result = is_lab_course_by_name('Database Lab')
        self.assertTrue(result)
    
    def test_auto_detect_course_type_mixed(self):
        """Test auto_detect_course_type for mixed course."""
        test_cases = [
            '理论与实验',
            '理论+实验',
            'Theory and Lab Mixed'
        ]
        
        for course_name in test_cases:
            with self.subTest(course_name=course_name):
                result = auto_detect_course_type(course_name)
                self.assertEqual(result, 'mixed')
    
    def test_auto_detect_course_type_lab(self):
        """Test auto_detect_course_type for lab course."""
        test_cases = [
            '计算机实验',
            'Computer Lab',
            'Physics Experiment'
        ]
        
        for course_name in test_cases:
            with self.subTest(course_name=course_name):
                result = auto_detect_course_type(course_name)
                self.assertEqual(result, 'lab')
    
    def test_auto_detect_course_type_practice(self):
        """Test auto_detect_course_type for practice course."""
        test_cases = [
            '软件实训',
            'Software Practice',
            '工程实践'
        ]
        
        for course_name in test_cases:
            with self.subTest(course_name=course_name):
                result = auto_detect_course_type(course_name)
                self.assertEqual(result, 'practice')
    
    def test_auto_detect_course_type_theory(self):
        """Test auto_detect_course_type for theory course."""
        test_cases = [
            '高等数学',
            'English Literature',
            '计算机理论'
        ]
        
        for course_name in test_cases:
            with self.subTest(course_name=course_name):
                result = auto_detect_course_type(course_name)
                self.assertEqual(result, 'theory')
    
    def test_get_course_type_from_name_database_exists(self):
        """Test get_course_type_from_name when course exists in database."""
        result = get_course_type_from_name('Test Course')
        self.assertEqual(result, 'theory')
    
    def test_get_course_type_from_name_database_not_exists(self):
        """Test get_course_type_from_name when course doesn't exist in database."""
        result = get_course_type_from_name('Nonexistent Lab Course')
        self.assertEqual(result, 'lab')


class TestLabReportDetection(BaseViewTestCase):
    """Test lab report detection functions."""
    
    def test_extract_homework_folder_success(self):
        """Test _extract_homework_folder with valid path structure."""
        file_path = os.path.join(self.temp_dir, 'Test Course', 'Class1', 'Assignment1', 'file.docx')
        
        result = _extract_homework_folder(file_path, self.temp_dir, 'Test Course')
        self.assertEqual(result, 'Assignment1')
    
    def test_extract_homework_folder_invalid_structure(self):
        """Test _extract_homework_folder with invalid path structure."""
        file_path = os.path.join(self.temp_dir, 'file.docx')
        
        result = _extract_homework_folder(file_path, self.temp_dir, 'Test Course')
        self.assertIsNone(result)
    
    def test_is_lab_report_file_with_homework_in_database(self):
        """Test is_lab_report_file when homework exists in database."""
        # Create lab homework
        homework = Homework.objects.create(
            course=self.course,
            title='Lab Assignment 1',
            homework_type='lab_report',
            folder_name='lab1',
            tenant=self.tenant
        )
        
        result = is_lab_report_file(
            course_name='Test Course',
            homework_folder='lab1'
        )
        self.assertTrue(result)
    
    def test_is_lab_report_file_normal_homework(self):
        """Test is_lab_report_file with normal homework."""
        # Create normal homework
        homework = Homework.objects.create(
            course=self.course,
            title='Regular Assignment 1',
            homework_type='normal',
            folder_name='assignment1',
            tenant=self.tenant
        )
        
        result = is_lab_report_file(
            course_name='Test Course',
            homework_folder='assignment1'
        )
        self.assertFalse(result)
    
    def test_is_lab_report_file_course_type_fallback(self):
        """Test is_lab_report_file falls back to course type."""
        # Update course to lab type
        self.course.course_type = 'lab'
        self.course.save()
        
        result = is_lab_report_file(
            course_name='Test Course',
            homework_folder='unknown_homework'
        )
        self.assertTrue(result)


class TestDirectoryOperations(BaseViewTestCase):
    """Test directory operation functions."""
    
    @patch('grading.views.get_cache_manager')
    def test_get_directory_file_count_cached_with_cache_hit(self, mock_get_cache_manager):
        """Test get_directory_file_count_cached with cache hit."""
        mock_cache_manager = Mock()
        mock_cache_manager.get_file_count.return_value = 5
        mock_get_cache_manager.return_value = mock_cache_manager
        
        result = get_directory_file_count_cached('test_dir', self.temp_dir, self.request)
        
        self.assertEqual(result, 5)
        mock_cache_manager.get_file_count.assert_called_once_with('test_dir')
    
    @patch('grading.views.get_cache_manager')
    def test_get_directory_file_count_cached_with_cache_miss(self, mock_get_cache_manager):
        """Test get_directory_file_count_cached with cache miss."""
        mock_cache_manager = Mock()
        mock_cache_manager.get_file_count.return_value = None
        mock_cache_manager.check_file_count_threshold.return_value = {'warning': False}
        mock_get_cache_manager.return_value = mock_cache_manager
        
        # Create test directory with docx files
        test_dir = os.path.join(self.temp_dir, 'test_dir')
        os.makedirs(test_dir)
        
        # Create some .docx files
        for i in range(3):
            with open(os.path.join(test_dir, f'file{i}.docx'), 'w') as f:
                f.write('test')
        
        # Create some non-docx files (should be ignored)
        with open(os.path.join(test_dir, 'file.txt'), 'w') as f:
            f.write('test')
        
        result = get_directory_file_count_cached('test_dir', self.temp_dir, self.request)
        
        self.assertEqual(result, 3)
        mock_cache_manager.set_file_count.assert_called_once_with('test_dir', 3)
    
    @patch('grading.views.get_cache_manager')
    def test_clear_directory_file_count_cache(self, mock_get_cache_manager):
        """Test clear_directory_file_count_cache."""
        mock_cache_manager = Mock()
        mock_get_cache_manager.return_value = mock_cache_manager
        
        clear_directory_file_count_cache(self.request)
        
        mock_cache_manager.clear_file_count.assert_called_once()


class TestPathUtilities(BaseViewTestCase):
    """Test path utility functions."""
    
    def test_clean_relative_homework_path_valid(self):
        """Test _clean_relative_homework_path with valid path."""
        test_cases = [
            ('course/class/homework', 'course/class/homework'),
            ('course\\class\\homework', 'course/class/homework'),
            ('/course/class/homework/', 'course/class/homework'),
            ('  course/class/homework  ', 'course/class/homework'),
        ]
        
        for input_path, expected in test_cases:
            with self.subTest(input_path=input_path):
                result = _clean_relative_homework_path(input_path)
                self.assertEqual(result, expected)
    
    def test_clean_relative_homework_path_invalid(self):
        """Test _clean_relative_homework_path with invalid paths."""
        test_cases = [
            '',
            None,
            '/',
            '../../../etc/passwd',
            'course/../../../etc/passwd',
            '/absolute/path',
        ]
        
        for input_path in test_cases:
            with self.subTest(input_path=input_path):
                result = _clean_relative_homework_path(input_path)
                self.assertIsNone(result)
    
    def test_split_rel_path_parts(self):
        """Test _split_rel_path_parts function."""
        test_cases = [
            ('course/class/homework', ['course', 'class', 'homework']),
            ('course\\class\\homework', ['course', 'class', 'homework']),
            ('', []),
            (None, []),
            ('single', ['single']),
        ]
        
        for input_path, expected in test_cases:
            with self.subTest(input_path=input_path):
                result = _split_rel_path_parts(input_path)
                self.assertEqual(result, expected)
    
    def test_is_homework_folder_rel_path(self):
        """Test _is_homework_folder_rel_path function."""
        test_cases = [
            ('course/homework', 'course', True),
            ('course/class/homework', 'course', False),
            ('homework', 'course', False),
            ('other/homework', 'course', False),
        ]
        
        for rel_path, course_name, expected in test_cases:
            with self.subTest(rel_path=rel_path, course_name=course_name):
                result = _is_homework_folder_rel_path(rel_path, course_name)
                self.assertEqual(result, expected)
    
    def test_is_homework_file_rel_path(self):
        """Test _is_homework_file_rel_path function."""
        test_cases = [
            ('course/class/homework/file.docx', 'course', True),
            ('course/homework/file.docx', 'course', True),
            ('course/homework', 'course', False),
            ('homework/file.docx', 'course', False),
        ]
        
        for rel_path, course_name, expected in test_cases:
            with self.subTest(rel_path=rel_path, course_name=course_name):
                result = _is_homework_file_rel_path(rel_path, course_name)
                self.assertEqual(result, expected)


class TestFileGradeInfo(BaseViewTestCase):
    """Test file grade information functions."""
    
    def test_get_file_grade_info_text_file_with_grade(self):
        """Test get_file_grade_info with text file containing grade."""
        test_file = os.path.join(self.temp_dir, 'test.txt')
        content = """
        This is a test file.
        老师评分：A
        Some other content.
        """
        
        with open(test_file, 'w', encoding='utf-8') as f:
            f.write(content)
        
        result = get_file_grade_info(test_file)
        
        self.assertTrue(result['has_grade'])
        self.assertEqual(result['grade'], 'A')
        self.assertEqual(result['grade_type'], 'letter')
        self.assertFalse(result['in_table'])
    
    def test_get_file_grade_info_text_file_without_grade(self):
        """Test get_file_grade_info with text file without grade."""
        test_file = os.path.join(self.temp_dir, 'test.txt')
        content = "This is a test file without any grade."
        
        with open(test_file, 'w', encoding='utf-8') as f:
            f.write(content)
        
        result = get_file_grade_info(test_file)
        
        self.assertFalse(result['has_grade'])
        self.assertIsNone(result['grade'])
        self.assertIsNone(result['grade_type'])
    
    def test_get_file_grade_info_docx_file_with_grade(self):
        """Test get_file_grade_info with DOCX file containing grade."""
        test_file = os.path.join(self.temp_dir, 'test.docx')
        
        # Create DOCX with grade
        doc = Document()
        doc.add_paragraph('This is a test document.')
        doc.add_paragraph('老师评分：B')
        doc.save(test_file)
        
        result = get_file_grade_info(test_file)
        
        self.assertTrue(result['has_grade'])
        self.assertEqual(result['grade'], 'B')
        self.assertEqual(result['grade_type'], 'letter')
    
    def test_get_file_grade_info_percentage_grade(self):
        """Test get_file_grade_info with percentage grade."""
        test_file = os.path.join(self.temp_dir, 'test.txt')
        content = "老师评分：85"
        
        with open(test_file, 'w', encoding='utf-8') as f:
            f.write(content)
        
        result = get_file_grade_info(test_file)
        
        self.assertTrue(result['has_grade'])
        self.assertEqual(result['grade'], '85')
        self.assertEqual(result['grade_type'], 'percentage')
    
    def test_get_file_grade_info_text_grade(self):
        """Test get_file_grade_info with text grade."""
        test_file = os.path.join(self.temp_dir, 'test.txt')
        content = "老师评分：优秀"
        
        with open(test_file, 'w', encoding='utf-8') as f:
            f.write(content)
        
        result = get_file_grade_info(test_file)
        
        self.assertTrue(result['has_grade'])
        self.assertEqual(result['grade'], '优秀')
        self.assertEqual(result['grade_type'], 'text')


@pytest.mark.django_db
class TestViewIntegration:
    """Integration tests for view functions."""
    
    def test_index_view_authenticated_user(self, client, django_user_model):
        """Test index view with authenticated user."""
        user = django_user_model.objects.create_user(
            username='testuser',
            password='testpass123'
        )
        client.force_login(user)
        
        response = client.get('/')
        
        assert response.status_code == 200
        assert 'current_semester' in response.context
        assert 'user_courses' in response.context
        assert 'repository_stats' in response.context
    
    def test_index_view_anonymous_user(self, client):
        """Test index view with anonymous user."""
        response = client.get('/')
        
        assert response.status_code == 200
        assert response.context['user_courses'] == []
        assert response.context['repository_stats'] == {}


class TestErrorHandling(BaseViewTestCase):
    """Test error handling in view functions."""
    
    def test_read_file_content_permission_error(self):
        """Test read_file_content handles permission errors gracefully."""
        # Create a file and remove read permissions
        test_file = os.path.join(self.temp_dir, 'no_permission.txt')
        with open(test_file, 'w') as f:
            f.write('test')
        
        # Remove read permissions (on Unix systems)
        if hasattr(os, 'chmod'):
            os.chmod(test_file, 0o000)
            
            result = read_file_content(test_file)
            self.assertEqual(result, '')
            
            # Restore permissions for cleanup
            os.chmod(test_file, 0o644)
    
    def test_get_file_extension_edge_cases(self):
        """Test get_file_extension with edge cases."""
        test_cases = [
            ('', 'unknown'),
            ('.', 'unknown'),
            ('..', 'unknown'),
            ('.hidden', 'unknown'),
            ('file.', 'unknown'),
        ]
        
        for input_path, expected in test_cases:
            with self.subTest(input_path=input_path):
                result = get_file_extension(input_path)
                self.assertEqual(result, expected)
    
    def test_auto_detect_course_type_empty_name(self):
        """Test auto_detect_course_type with empty course name."""
        result = auto_detect_course_type('')
        self.assertEqual(result, 'theory')
        
        result = auto_detect_course_type(None)
        self.assertEqual(result, 'theory')


class TestCacheIntegration(BaseViewTestCase):
    """Test cache integration in view functions."""
    
    def setUp(self):
        super().setUp()
        # Clear cache before each test
        cache.clear()
    
    def tearDown(self):
        # Clear cache after each test
        cache.clear()
        super().tearDown()
    
    @patch('grading.views.get_cache_manager')
    def test_directory_file_count_caching(self, mock_get_cache_manager):
        """Test that directory file count is properly cached."""
        mock_cache_manager = Mock()
        mock_cache_manager.get_file_count.return_value = None
        mock_cache_manager.check_file_count_threshold.return_value = {'warning': False}
        mock_get_cache_manager.return_value = mock_cache_manager
        
        # Create test directory
        test_dir = os.path.join(self.temp_dir, 'cache_test')
        os.makedirs(test_dir)
        
        # First call should compute and cache
        result1 = get_directory_file_count_cached('cache_test', self.temp_dir, self.request)
        
        # Verify cache was set
        mock_cache_manager.set_file_count.assert_called_once()
        
        # Second call should use cache
        mock_cache_manager.get_file_count.return_value = result1
        result2 = get_directory_file_count_cached('cache_test', self.temp_dir, self.request)
        
        self.assertEqual(result1, result2)


if __name__ == '__main__':
    pytest.main([__file__])