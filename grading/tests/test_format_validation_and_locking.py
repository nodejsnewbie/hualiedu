"""
测试格式验证和锁定机制
Tests for format validation and file locking mechanisms

Property 12: 实验报告格式检测
Property 13: 文件锁定机制

Validates: Requirements 11.2, 11.6
"""

import os
import shutil
import tempfile

from django.test import TestCase
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from hypothesis import assume, given, settings
from hypothesis import strategies as st
from hypothesis.extra.django import TestCase as HypothesisTestCase

from grading.views import (
    find_teacher_signature_cell,
    write_grade_and_comment_to_file,
    write_grade_to_lab_report,
)


class FormatValidationTest(TestCase):
    """测试实验报告格式检测 - Property 12"""

    def setUp(self):
        self.temp_dir = tempfile.mkdtemp()

    def tearDown(self):
        if os.path.exists(self.temp_dir):
            shutil.rmtree(self.temp_dir)

    def create_valid_lab_report(self):
        """创建格式正确的实验报告（包含教师签字单元格）"""
        doc = Document()
        table = doc.add_table(rows=3, cols=2)

        # 添加"教师（签字）"单元格
        cell = table.rows[2].cells[0]
        cell.text = "教师（签字）：时间："

        return doc

    def create_invalid_lab_report(self):
        """创建格式错误的实验报告（缺少教师签字单元格）"""
        doc = Document()
        # 添加一些内容但不包含"教师（签字）"单元格
        doc.add_paragraph("实验报告内容")
        table = doc.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "学生姓名"
        table.rows[1].cells[0].text = "实验日期"

        return doc

    def test_valid_format_detection(self):
        """测试检测格式正确的实验报告"""
        doc = self.create_valid_lab_report()

        cell, table_idx, row_idx, col_idx = find_teacher_signature_cell(doc)

        self.assertIsNotNone(cell, "应该找到教师签字单元格")
        self.assertIsNotNone(table_idx)
        self.assertIsNotNone(row_idx)
        self.assertIsNotNone(col_idx)

    def test_invalid_format_detection(self):
        """测试检测格式错误的实验报告"""
        doc = self.create_invalid_lab_report()

        cell, table_idx, row_idx, col_idx = find_teacher_signature_cell(doc)

        self.assertIsNone(cell, "不应该找到教师签字单元格")
        self.assertIsNone(table_idx)
        self.assertIsNone(row_idx)
        self.assertIsNone(col_idx)

    def test_format_error_triggers_downgrade(self):
        """测试格式错误触发降级处理"""
        doc = self.create_invalid_lab_report()

        # 尝试写入评分
        success, modified_grade, modified_comment = write_grade_to_lab_report(doc, "A", "很好")

        # 应该返回失败，并设置D评分和锁定标记
        self.assertFalse(success, "格式错误应该返回失败")
        self.assertEqual(modified_grade, "D", "格式错误应该自动设置D评分")
        self.assertIn("【格式错误-已锁定】", modified_comment, "应该包含锁定标记")
        self.assertIn("请按要求的格式写实验报告", modified_comment)
        self.assertIn("此评分不可修改", modified_comment)

    def test_valid_format_allows_grading(self):
        """测试格式正确允许正常评分"""
        doc = self.create_valid_lab_report()

        # 尝试写入评分
        success, modified_grade, modified_comment = write_grade_to_lab_report(doc, "A", "很好")

        # 应该成功
        self.assertTrue(success, "格式正确应该允许评分")
        self.assertEqual(modified_grade, "A", "评分应该保持不变")
        self.assertEqual(modified_comment, "很好", "评价应该保持不变")

    def test_format_detection_with_half_width_parentheses(self):
        """测试检测半角括号的教师(签字)"""
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        cell = table.rows[1].cells[0]
        cell.text = "教师(签字)：时间："

        found_cell, _, _, _ = find_teacher_signature_cell(doc)

        self.assertIsNotNone(found_cell, "应该能识别半角括号")

    def test_format_detection_in_multiple_tables(self):
        """测试在多个表格中检测格式"""
        doc = Document()

        # 第一个表格（无签字单元格）
        table1 = doc.add_table(rows=2, cols=2)
        table1.rows[0].cells[0].text = "学生信息"

        # 第二个表格（有签字单元格）
        table2 = doc.add_table(rows=2, cols=2)
        table2.rows[1].cells[1].text = "教师（签字）："

        cell, table_idx, _, _ = find_teacher_signature_cell(doc)

        self.assertIsNotNone(cell)
        self.assertEqual(table_idx, 1, "应该在第二个表格中找到")

    def test_format_detection_empty_document(self):
        """测试空文档的格式检测"""
        doc = Document()

        cell, _, _, _ = find_teacher_signature_cell(doc)

        self.assertIsNone(cell, "空文档不应该找到签字单元格")

    def test_format_detection_document_without_tables(self):
        """测试没有表格的文档"""
        doc = Document()
        doc.add_paragraph("这是一个没有表格的文档")
        doc.add_paragraph("实验报告内容")

        cell, _, _, _ = find_teacher_signature_cell(doc)

        self.assertIsNone(cell, "没有表格的文档不应该找到签字单元格")


class FileLockingTest(TestCase):
    """测试文件锁定机制 - Property 13"""

    def setUp(self):
        self.temp_dir = tempfile.mkdtemp()

    def tearDown(self):
        if os.path.exists(self.temp_dir):
            shutil.rmtree(self.temp_dir)

    def create_locked_file(self):
        """创建已锁定的文件"""
        doc = Document()
        doc.add_paragraph("实验报告内容")
        doc.add_paragraph("老师评分：D")
        doc.add_paragraph("教师评价：【格式错误-已锁定】请按要求的格式写实验报告，此评分不可修改")

        file_path = os.path.join(self.temp_dir, "locked_report.docx")
        doc.save(file_path)
        return file_path

    def create_unlocked_file(self):
        """创建未锁定的文件"""
        doc = Document()
        doc.add_paragraph("实验报告内容")

        file_path = os.path.join(self.temp_dir, "unlocked_report.docx")
        doc.save(file_path)
        return file_path

    def test_locked_file_rejects_grading(self):
        """测试锁定文件拒绝评分操作"""
        locked_file = self.create_locked_file()

        # 尝试修改评分
        result = write_grade_and_comment_to_file(
            locked_file, grade="A", comment="新评价", is_lab_report=False
        )

        # 应该返回错误消息
        self.assertIsNotNone(result, "应该返回错误消息")
        self.assertIn("已被锁定", result)
        self.assertIn("不允许修改", result)

    def test_unlocked_file_allows_grading(self):
        """测试未锁定文件允许评分操作"""
        unlocked_file = self.create_unlocked_file()

        # 尝试添加评分
        result = write_grade_and_comment_to_file(
            unlocked_file, grade="A", comment="很好", is_lab_report=False
        )

        # 应该成功（返回None或无错误）
        self.assertIsNone(result, "未锁定文件应该允许评分")

        # 验证文件内容
        doc = Document(unlocked_file)
        text_content = "\n".join([p.text for p in doc.paragraphs])
        self.assertIn("A", text_content)
        self.assertIn("很好", text_content)

    def test_lock_marker_detection_full_width(self):
        """测试检测全角锁定标记"""
        doc = Document()
        doc.add_paragraph("【格式错误-已锁定】请按要求的格式写实验报告，此评分不可修改")

        file_path = os.path.join(self.temp_dir, "test.docx")
        doc.save(file_path)

        result = write_grade_and_comment_to_file(file_path, grade="A", is_lab_report=False)

        self.assertIsNotNone(result)
        self.assertIn("已被锁定", result)

    def test_lock_marker_detection_half_width(self):
        """测试检测半角锁定标记"""
        doc = Document()
        doc.add_paragraph("格式错误-已锁定 请按要求的格式写实验报告")

        file_path = os.path.join(self.temp_dir, "test.docx")
        doc.save(file_path)

        result = write_grade_and_comment_to_file(file_path, grade="A", is_lab_report=False)

        self.assertIsNotNone(result)
        self.assertIn("已被锁定", result)

    def test_lock_marker_in_middle_of_document(self):
        """测试锁定标记在文档中间"""
        doc = Document()
        doc.add_paragraph("实验报告内容")
        doc.add_paragraph("【格式错误-已锁定】此文件已锁定")
        doc.add_paragraph("更多内容")

        file_path = os.path.join(self.temp_dir, "test.docx")
        doc.save(file_path)

        result = write_grade_and_comment_to_file(file_path, grade="B", is_lab_report=False)

        self.assertIsNotNone(result)
        self.assertIn("已被锁定", result)

    def test_format_error_creates_locked_file(self):
        """测试格式错误创建锁定文件"""
        # 创建格式错误的实验报告
        doc = Document()
        doc.add_paragraph("实验报告内容")
        # 没有"教师（签字）"表格

        file_path = os.path.join(self.temp_dir, "format_error.docx")
        doc.save(file_path)

        # 尝试按实验报告格式写入
        result = write_grade_and_comment_to_file(
            file_path, grade="A", comment="很好", is_lab_report=True
        )

        # 应该返回格式警告
        self.assertIsNotNone(result)
        self.assertIn("格式不正确", result)
        self.assertIn("D评分", result)
        self.assertIn("锁定", result)

        # 验证文件被锁定
        doc_reloaded = Document(file_path)
        text_content = "\n".join([p.text for p in doc_reloaded.paragraphs])
        self.assertIn("【格式错误-已锁定】", text_content)
        self.assertIn("D", text_content)

    def test_locked_file_preserves_lock_marker(self):
        """测试锁定文件保留锁定标记"""
        locked_file = self.create_locked_file()

        # 尝试修改（应该被拒绝）
        write_grade_and_comment_to_file(locked_file, grade="A", is_lab_report=False)

        # 验证锁定标记仍然存在
        doc = Document(locked_file)
        text_content = "\n".join([p.text for p in doc.paragraphs])
        self.assertIn("【格式错误-已锁定】", text_content)


class PropertyBasedFormatValidationTest(HypothesisTestCase):
    """基于属性的格式验证测试"""

    def setUp(self):
        self.temp_dir = tempfile.mkdtemp()

    def tearDown(self):
        if os.path.exists(self.temp_dir):
            shutil.rmtree(self.temp_dir)

    @given(
        has_signature_cell=st.booleans(),
        grade=st.sampled_from(["A", "B", "C", "D", "E", "优秀", "良好", "中等", "及格", "不及格"]),
        comment=st.text(
            alphabet=st.characters(blacklist_categories=("Cc", "Cs")), min_size=1, max_size=100
        ),
    )
    @settings(max_examples=100, deadline=None)
    def test_property_format_detection_consistency(self, has_signature_cell, grade, comment):
        """
        Property 12: 实验报告格式检测
        For any Word document, the format validation function should correctly
        identify whether it contains the "教师（签字）" cell.

        Validates: Requirements 11.2
        """
        # 创建文档
        doc = Document()

        if has_signature_cell:
            # 创建包含签字单元格的文档
            table = doc.add_table(rows=2, cols=2)
            table.rows[1].cells[0].text = "教师（签字）：时间："
        else:
            # 创建不包含签字单元格的文档
            doc.add_paragraph("实验报告内容")

        # 检测格式
        cell, _, _, _ = find_teacher_signature_cell(doc)

        # 验证检测结果与实际情况一致
        if has_signature_cell:
            self.assertIsNotNone(cell, "包含签字单元格的文档应该被检测到")
        else:
            self.assertIsNone(cell, "不包含签字单元格的文档不应该被检测到")

        # 尝试写入评分
        success, modified_grade, modified_comment = write_grade_to_lab_report(doc, grade, comment)

        # 验证写入结果与格式检测一致
        if has_signature_cell:
            self.assertTrue(success, "格式正确应该允许写入")
            self.assertEqual(modified_grade, grade, "评分应该保持不变")
        else:
            self.assertFalse(success, "格式错误应该拒绝写入")
            self.assertEqual(modified_grade, "D", "格式错误应该设置D评分")
            self.assertIn("【格式错误-已锁定】", modified_comment, "应该包含锁定标记")

    @given(
        initial_grade=st.sampled_from(["A", "B", "C"]),
        new_grade=st.sampled_from(["A", "B", "C", "D", "E"]),
    )
    @settings(max_examples=50, deadline=None)
    def test_property_file_locking_prevents_modification(self, initial_grade, new_grade):
        """
        Property 13: 文件锁定机制
        For any file containing the "【格式错误-已锁定】" marker,
        all subsequent grading operations should be rejected.

        Validates: Requirements 11.6
        """
        # 创建锁定文件
        doc = Document()
        doc.add_paragraph("实验报告内容")
        doc.add_paragraph(f"老师评分：{initial_grade}")
        doc.add_paragraph("教师评价：【格式错误-已锁定】请按要求的格式写实验报告，此评分不可修改")

        file_path = os.path.join(self.temp_dir, f"locked_{initial_grade}_{new_grade}.docx")
        doc.save(file_path)

        # 尝试修改评分
        result = write_grade_and_comment_to_file(
            file_path, grade=new_grade, comment="新评价", is_lab_report=False
        )

        # 应该被拒绝
        self.assertIsNotNone(result, "锁定文件应该拒绝修改")
        self.assertIn("已被锁定", result)

        # 验证文件内容未被修改
        doc_reloaded = Document(file_path)
        text_content = "\n".join([p.text for p in doc_reloaded.paragraphs])

        # 原始评分应该保持不变
        self.assertIn(initial_grade, text_content, "原始评分应该保持不变")
        # 不应该包含新评分（除非新评分与原始评分相同）
        if new_grade != initial_grade:
            # 检查新评分是否被写入（不应该）
            # 由于锁定机制，文件不应该被修改
            self.assertIn("【格式错误-已锁定】", text_content, "锁定标记应该保留")


class IntegrationFormatAndLockingTest(TestCase):
    """格式验证和锁定机制的集成测试"""

    def setUp(self):
        self.temp_dir = tempfile.mkdtemp()

    def tearDown(self):
        if os.path.exists(self.temp_dir):
            shutil.rmtree(self.temp_dir)

    def test_complete_format_error_workflow(self):
        """测试完整的格式错误工作流"""
        # 步骤1：创建格式错误的实验报告
        doc = Document()
        doc.add_paragraph("实验报告内容")
        doc.add_paragraph("实验结果")

        file_path = os.path.join(self.temp_dir, "format_error_workflow.docx")
        doc.save(file_path)

        # 步骤2：尝试评分（应该触发格式错误）
        result = write_grade_and_comment_to_file(
            file_path, grade="A", comment="很好", is_lab_report=True
        )

        # 应该返回格式警告
        self.assertIsNotNone(result)
        self.assertIn("格式不正确", result)

        # 步骤3：验证文件被锁定
        doc_after_first = Document(file_path)
        text_after_first = "\n".join([p.text for p in doc_after_first.paragraphs])
        self.assertIn("【格式错误-已锁定】", text_after_first)
        self.assertIn("D", text_after_first)

        # 步骤4：尝试再次修改（应该被拒绝）
        result2 = write_grade_and_comment_to_file(
            file_path, grade="A", comment="修改评价", is_lab_report=False
        )

        self.assertIsNotNone(result2)
        self.assertIn("已被锁定", result2)

        # 步骤5：验证文件内容未被修改
        doc_after_second = Document(file_path)
        text_after_second = "\n".join([p.text for p in doc_after_second.paragraphs])
        self.assertIn("【格式错误-已锁定】", text_after_second)
        self.assertIn("D", text_after_second)
        self.assertNotIn("修改评价", text_after_second)

    def test_valid_format_no_locking(self):
        """测试格式正确不触发锁定"""
        # 创建格式正确的实验报告
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        table.rows[1].cells[0].text = "教师（签字）：时间："

        file_path = os.path.join(self.temp_dir, "valid_format.docx")
        doc.save(file_path)

        # 第一次评分
        result1 = write_grade_and_comment_to_file(
            file_path, grade="B", comment="良好", is_lab_report=True
        )

        # 应该成功，无警告
        self.assertIsNone(result1)

        # 验证文件未被锁定
        doc_after_first = Document(file_path)
        text_after_first = "\n".join([p.text for p in doc_after_first.paragraphs])
        self.assertNotIn("【格式错误-已锁定】", text_after_first)

        # 第二次修改评分（应该允许）
        result2 = write_grade_and_comment_to_file(
            file_path, grade="A", comment="优秀", is_lab_report=True
        )

        # 应该成功
        self.assertIsNone(result2)

        # 验证文件被更新
        doc_after_second = Document(file_path)
        # 检查表格中的内容
        cell = doc_after_second.tables[0].rows[1].cells[0]
        self.assertIn("A", cell.text)
        self.assertIn("优秀", cell.text)

    def test_format_error_downgrade_to_paragraph(self):
        """测试格式错误降级为段落写入"""
        # 创建格式错误的实验报告
        doc = Document()
        doc.add_paragraph("实验报告内容")

        file_path = os.path.join(self.temp_dir, "downgrade_test.docx")
        doc.save(file_path)

        # 尝试按实验报告格式写入
        result = write_grade_and_comment_to_file(
            file_path, grade="A", comment="很好", is_lab_report=True
        )

        # 应该返回格式警告
        self.assertIsNotNone(result)

        # 验证降级为段落写入
        doc_reloaded = Document(file_path)
        paragraphs = [p.text for p in doc_reloaded.paragraphs]

        # 应该包含评分和评价段落（而不是表格）
        has_grade_paragraph = any("老师评分：D" in p for p in paragraphs)
        has_comment_paragraph = any("【格式错误-已锁定】" in p for p in paragraphs)

        self.assertTrue(has_grade_paragraph, "应该包含评分段落")
        self.assertTrue(has_comment_paragraph, "应该包含评价段落")

        # 不应该有表格（或表格中没有评分）
        if doc_reloaded.tables:
            for table in doc_reloaded.tables:
                for row in table.rows:
                    for cell in row.cells:
                        # 表格中不应该有评分内容
                        self.assertNotIn("老师评分", cell.text)
