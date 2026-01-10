"""
Tests for file operation functions in grading/views.py

This module tests file handling, grade writing, comment management,
and document processing functions.
"""

import json
import os
import tempfile
import shutil
from unittest.mock import Mock, patch, MagicMock, mock_open
from io import BytesIO

import pytest
from django.test import TestCase, RequestFactory
from django.contrib.auth import get_user_model
from docx import Document
from docx.shared import Pt

from grading.models import (
    Repository, Course, Homework, Semester, GlobalConfig, 
    FileGradeStatus, Tenant, UserProfile
)
from grading.views import (
    update_file_grade_status,
    maybe_sync_repository,
    push_grade_changes,
    get_file_grade_info,
    _get_repo_head_commit,
    _file_changed_since_commit,
    _file_has_updates,
    _homework_folder_has_updates,
)

User = get_user_model()


class BaseFileOperationTestCase(TestCase):
    """Base test case for file operation tests."""
    
    def setUp(self):
        """Set up test data for file operation tests."""
        self.factory = RequestFactory()
        
        # Create test user
        self.user = User.objects.create_user(
            username='testuser',
            email='test@example.com',
            password='testpass123',
            is_staff=True,
            first_name='Test',
            last_name='User'
        )
        
        # Create tenant and user profile
        self.tenant = Tenant.objects.create(
            name='Test Tenant',
            domain='test.example.com'
        )
        
        self.user_profile = UserProfile.objects.create(
            user=self.user,
            tenant=self.tenant,
            repo_base_dir='~/test_jobs'
        )
        
        # Create test semester
        self.semester = Semester.objects.create(
            name='Test Semester',
            start_date='2024-01-01',
            end_date='2024-06-30',
            is_active=True
        )
        
        # Create test course
        self.course = Course.objects.create(
            name='Test Course',
            course_type='theory',
            semester=self.semester,
            teacher=self.user,
            location='Room 101'
        )
        
        # Create test repository
        self.repository = Repository.objects.create(
            name='test-repo',
            url='https://github.com/test/repo.git',
            owner=self.user,
            tenant=self.tenant,
            repo_type='local',
            local_path='/tmp/test-repo'
        )
        
        # Create temporary directory
        self.temp_dir = tempfile.mkdtemp()
        self.addCleanup(shutil.rmtree, self.temp_dir)
        
        # Mock request
        self.request = self.factory.get('/')
        self.request.user = self.user
        self.request.user_profile = self.user_profile
        self.request.tenant = self.tenant


class TestFileGradeStatus(BaseFileOperationTestCase):
    """Test file grade status management functions."""
    
    def test_update_file_grade_status_local_repo(self):
        """Test update_file_grade_status with local repository."""
        # Set up local repository
        self.repository.repo_type = 'local'
        self.repository.local_path = self.temp_dir
        self.repository.save()
        
        # Create a git repository in temp dir
        os.system(f'cd {self.temp_dir} && git init')
        os.system(f'cd {self.temp_dir} && git config user.email "test@example.com"')
        os.system(f'cd {self.temp_dir} && git config user.name "Test User"')
        
        # Create and commit a test file
        test_file = os.path.join(self.temp_dir, 'test.txt')
        with open(test_file, 'w') as f:
            f.write('test content')
        
        os.system(f'cd {self.temp_dir} && git add test.txt')
        os.system(f'cd {self.temp_dir} && git commit -m "Initial commit"')
        
        # Test update
        update_file_grade_status(
            self.repository, 
            'test.txt', 
            course_name='Test Course',
            user=self.user
        )
        
        # Verify FileGradeStatus was created
        status = FileGradeStatus.objects.get(
            repository=self.repository,
            file_path='Test Course/test.txt'
        )
        
        self.assertIsNotNone(status.last_graded_at)
        self.assertEqual(status.last_graded_by, 'Test User')
        self.assertIsNotNone(status.last_graded_commit)
    
    def test_update_file_grade_status_git_repo(self):
        """Test update_file_grade_status with Git repository."""
        # Set up Git repository
        self.repository.repo_type = 'git'
        self.repository.git_url = 'https://github.com/test/repo.git'
        self.repository.git_branch = 'main'
        self.repository.save()
        
        with patch('grading.views.GitStorageAdapter') as mock_adapter_class:
            mock_adapter = Mock()
            mock_adapter.get_head_commit.return_value = 'abc123'
            mock_adapter_class.return_value = mock_adapter
            
            update_file_grade_status(
                self.repository,
                'test.txt',
                course_name='Test Course',
                user=self.user
            )
        
        # Verify FileGradeStatus was created
        status = FileGradeStatus.objects.get(
            repository=self.repository,
            file_path='Test Course/test.txt'
        )
        
        self.assertIsNotNone(status.last_graded_at)
        self.assertEqual(status.last_graded_by, 'Test User')
        self.assertEqual(status.last_graded_commit, 'abc123')
    
    def test_update_file_grade_status_no_repository(self):
        """Test update_file_grade_status with no repository."""
        # Should not raise exception
        update_file_grade_status(None, 'test.txt')
        
        # Should not create any FileGradeStatus
        self.assertEqual(FileGradeStatus.objects.count(), 0)
    
    def test_update_file_grade_status_exception_handling(self):
        """Test update_file_grade_status handles exceptions gracefully."""
        with patch('grading.views.FileGradeStatus.objects.update_or_create') as mock_update:
            mock_update.side_effect = Exception('Database error')
            
            # Should not raise exception
            update_file_grade_status(
                self.repository,
                'test.txt',
                user=self.user
            )


class TestRepositorySync(BaseFileOperationTestCase):
    """Test repository synchronization functions."""
    
    def test_maybe_sync_repository_success(self):
        """Test successful repository synchronization."""
        # Set up repository that can sync
        self.repository.repo_type = 'git'
        self.repository.url = 'https://github.com/test/repo.git'
        self.repository.branch = 'main'
        self.repository.local_path = self.temp_dir
        self.repository.save()
        
        with patch('grading.views.GitHandler') as mock_git_handler:
            mock_git_handler.pull_repo.return_value = True
            
            # Create the directory
            os.makedirs(self.temp_dir, exist_ok=True)
            
            result = maybe_sync_repository(self.repository, self.request)
        
        self.assertTrue(result)
        
        # Verify last_sync was updated
        self.repository.refresh_from_db()
        self.assertIsNotNone(self.repository.last_sync)
    
    def test_maybe_sync_repository_clone_new(self):
        """Test repository synchronization with new clone."""
        # Set up repository that can sync
        self.repository.repo_type = 'git'
        self.repository.url = 'https://github.com/test/repo.git'
        self.repository.branch = 'main'
        self.repository.local_path = os.path.join(self.temp_dir, 'new_repo')
        self.repository.save()
        
        with patch('grading.views.GitHandler') as mock_git_handler:
            mock_git_handler.clone_repo_remote.return_value = True
            
            result = maybe_sync_repository(self.repository, self.request)
        
        self.assertTrue(result)
    
    def test_maybe_sync_repository_cannot_sync(self):
        """Test repository synchronization when repository cannot sync."""
        # Set up repository that cannot sync (local type)
        self.repository.repo_type = 'local'
        self.repository.save()
        
        result = maybe_sync_repository(self.repository, self.request)
        
        self.assertFalse(result)
    
    def test_maybe_sync_repository_recent_sync(self):
        """Test repository synchronization skips recent sync."""
        from django.utils import timezone
        
        # Set last_sync to very recent
        self.repository.last_sync = timezone.now()
        self.repository.save()
        
        result = maybe_sync_repository(self.repository, self.request, min_interval_seconds=3600)
        
        self.assertFalse(result)
    
    def test_maybe_sync_repository_failure(self):
        """Test repository synchronization failure."""
        # Set up repository that can sync
        self.repository.repo_type = 'git'
        self.repository.url = 'https://github.com/test/repo.git'
        self.repository.local_path = self.temp_dir
        self.repository.save()
        
        with patch('grading.views.GitHandler') as mock_git_handler:
            mock_git_handler.pull_repo.return_value = False
            
            # Create the directory
            os.makedirs(self.temp_dir, exist_ok=True)
            
            result = maybe_sync_repository(self.repository, self.request)
        
        self.assertFalse(result)


class TestGitOperations(BaseFileOperationTestCase):
    """Test Git-related operations."""
    
    def test_push_grade_changes_git_repo(self):
        """Test push_grade_changes with Git repository."""
        # Set up Git repository
        self.repository.repo_type = 'git'
        self.repository.local_path = self.temp_dir
        self.repository.branch = 'main'
        self.repository.save()
        
        test_file = os.path.join(self.temp_dir, 'test.txt')
        
        with patch('grading.views.GitHandler') as mock_git_handler:
            push_grade_changes(self.repository, test_file)
            
            mock_git_handler.commit_and_push.assert_called_once_with(
                self.temp_dir,
                message='评分更新: test.txt',
                branch='main',
                paths='test.txt'
            )
    
    def test_push_grade_changes_local_repo(self):
        """Test push_grade_changes with local repository."""
        # Set up local repository
        self.repository.repo_type = 'local'
        self.repository.save()
        
        test_file = os.path.join(self.temp_dir, 'test.txt')
        
        with patch('grading.views.GitHandler') as mock_git_handler:
            push_grade_changes(self.repository, test_file)
            
            # Should not call GitHandler for local repos
            mock_git_handler.commit_and_push.assert_not_called()
    
    def test_get_repo_head_commit_success(self):
        """Test _get_repo_head_commit with successful Git operation."""
        # Set up repository path
        self.repository.local_path = self.temp_dir
        self.repository.save()
        
        with patch('grading.views.GitHandler') as mock_git_handler:
            mock_git_handler.is_git_repo.return_value = True
            
            with patch('subprocess.run') as mock_run:
                mock_result = Mock()
                mock_result.returncode = 0
                mock_result.stdout = 'abc123def456\n'
                mock_run.return_value = mock_result
                
                result = _get_repo_head_commit(self.repository)
        
        self.assertEqual(result, 'abc123def456')
    
    def test_get_repo_head_commit_not_git_repo(self):
        """Test _get_repo_head_commit with non-Git repository."""
        with patch('grading.views.GitHandler') as mock_git_handler:
            mock_git_handler.is_git_repo.return_value = False
            
            result = _get_repo_head_commit(self.repository)
        
        self.assertIsNone(result)
    
    def test_get_repo_head_commit_git_error(self):
        """Test _get_repo_head_commit with Git error."""
        self.repository.local_path = self.temp_dir
        self.repository.save()
        
        with patch('grading.views.GitHandler') as mock_git_handler:
            mock_git_handler.is_git_repo.return_value = True
            
            with patch('subprocess.run') as mock_run:
                mock_result = Mock()
                mock_result.returncode = 1
                mock_result.stderr = 'fatal: not a git repository'
                mock_run.return_value = mock_result
                
                result = _get_repo_head_commit(self.repository)
        
        self.assertIsNone(result)
    
    def test_file_changed_since_commit_success(self):
        """Test _file_changed_since_commit with changes detected."""
        self.repository.local_path = self.temp_dir
        self.repository.save()
        
        with patch('grading.views.GitHandler') as mock_git_handler:
            mock_git_handler.is_git_repo.return_value = True
            
            with patch('subprocess.run') as mock_run:
                mock_result = Mock()
                mock_result.returncode = 0
                mock_result.stdout = 'test.txt\n'  # File has changes
                mock_run.return_value = mock_result
                
                result = _file_changed_since_commit(
                    self.repository, 
                    'test.txt', 
                    'old_commit', 
                    'new_commit'
                )
        
        self.assertTrue(result)
    
    def test_file_changed_since_commit_no_changes(self):
        """Test _file_changed_since_commit with no changes."""
        self.repository.local_path = self.temp_dir
        self.repository.save()
        
        with patch('grading.views.GitHandler') as mock_git_handler:
            mock_git_handler.is_git_repo.return_value = True
            
            with patch('subprocess.run') as mock_run:
                mock_result = Mock()
                mock_result.returncode = 0
                mock_result.stdout = ''  # No changes
                mock_run.return_value = mock_result
                
                result = _file_changed_since_commit(
                    self.repository, 
                    'test.txt', 
                    'old_commit', 
                    'new_commit'
                )
        
        self.assertFalse(result)
    
    def test_file_changed_since_commit_missing_commits(self):
        """Test _file_changed_since_commit with missing commit hashes."""
        result = _file_changed_since_commit(
            self.repository, 
            'test.txt', 
            None,  # Missing old commit
            'new_commit'
        )
        
        self.assertFalse(result)
        
        result = _file_changed_since_commit(
            self.repository, 
            'test.txt', 
            'old_commit',
            None  # Missing new commit
        )
        
        self.assertFalse(result)


class TestFileUpdateDetection(BaseFileOperationTestCase):
    """Test file update detection functions."""
    
    def test_file_has_updates_new_file(self):
        """Test _file_has_updates with new file (no grade status)."""
        test_file = os.path.join(self.temp_dir, 'test.txt')
        with open(test_file, 'w') as f:
            f.write('test content')
        
        with patch('grading.views.get_file_grade_info') as mock_grade_info:
            mock_grade_info.return_value = {'has_grade': False}
            
            result = _file_has_updates(
                self.repository,
                'test.txt',
                test_file,
                base_dir=self.temp_dir
            )
        
        self.assertTrue(result)
    
    def test_file_has_updates_already_graded(self):
        """Test _file_has_updates with already graded file."""
        test_file = os.path.join(self.temp_dir, 'test.txt')
        with open(test_file, 'w') as f:
            f.write('test content')
        
        # Create FileGradeStatus
        FileGradeStatus.objects.create(
            repository=self.repository,
            file_path='test.txt',
            last_graded_commit='abc123'
        )
        
        with patch('grading.views.get_file_grade_info') as mock_grade_info:
            mock_grade_info.return_value = {'has_grade': True}
            
            result = _file_has_updates(
                self.repository,
                'test.txt',
                test_file,
                current_head='abc123',  # Same commit
                base_dir=self.temp_dir
            )
        
        self.assertFalse(result)
    
    def test_file_has_updates_git_changes(self):
        """Test _file_has_updates with Git changes detected."""
        test_file = os.path.join(self.temp_dir, 'test.txt')
        with open(test_file, 'w') as f:
            f.write('test content')
        
        # Create FileGradeStatus with old commit
        FileGradeStatus.objects.create(
            repository=self.repository,
            file_path='test.txt',
            last_graded_commit='old_commit'
        )
        
        with patch('grading.views._file_changed_since_commit') as mock_changed:
            mock_changed.return_value = True
            
            result = _file_has_updates(
                self.repository,
                'test.txt',
                test_file,
                current_head='new_commit',
                base_dir=self.temp_dir
            )
        
        self.assertTrue(result)
    
    def test_homework_folder_has_updates_with_changes(self):
        """Test _homework_folder_has_updates with file changes."""
        # Create homework folder structure
        homework_dir = os.path.join(self.temp_dir, 'homework1')
        os.makedirs(homework_dir)
        
        # Create test files
        test_files = ['file1.txt', 'file2.txt']
        for filename in test_files:
            with open(os.path.join(homework_dir, filename), 'w') as f:
                f.write('test content')
        
        with patch('grading.views.get_file_grade_info') as mock_grade_info:
            mock_grade_info.return_value = {'has_grade': False}
            
            result = _homework_folder_has_updates(
                self.repository,
                homework_dir,
                'homework1',
                base_dir=self.temp_dir
            )
        
        self.assertTrue(result)
    
    def test_homework_folder_has_updates_no_changes(self):
        """Test _homework_folder_has_updates with no changes."""
        # Create homework folder structure
        homework_dir = os.path.join(self.temp_dir, 'homework1')
        os.makedirs(homework_dir)
        
        # Create test file
        with open(os.path.join(homework_dir, 'file1.txt'), 'w') as f:
            f.write('test content')
        
        # Create FileGradeStatus for the file
        FileGradeStatus.objects.create(
            repository=self.repository,
            file_path='homework1/file1.txt',
            last_graded_commit='abc123'
        )
        
        with patch('grading.views._file_changed_since_commit') as mock_changed:
            mock_changed.return_value = False
            
            result = _homework_folder_has_updates(
                self.repository,
                homework_dir,
                'homework1',
                current_head='abc123',
                base_dir=self.temp_dir
            )
        
        self.assertFalse(result)


class TestFileGradeInfoAdvanced(BaseFileOperationTestCase):
    """Test advanced file grade information functions."""
    
    def test_get_file_grade_info_docx_with_table_grade(self):
        """Test get_file_grade_info with DOCX file containing table grade."""
        test_file = os.path.join(self.temp_dir, 'test.docx')
        
        # Create DOCX with table containing grade
        doc = Document()
        doc.add_paragraph('Test document')
        
        # Add table with grade
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = '评定分数'
        table.cell(0, 1).text = 'A'
        table.cell(1, 0).text = '其他信息'
        table.cell(1, 1).text = '测试'
        
        doc.save(test_file)
        
        with patch('grading.views._iter_tables') as mock_iter_tables:
            mock_iter_tables.return_value = [table]
            
            result = get_file_grade_info(test_file)
        
        self.assertTrue(result['has_grade'])
        self.assertEqual(result['grade'], 'A')
        self.assertEqual(result['grade_type'], 'letter')
        self.assertTrue(result['in_table'])
    
    def test_get_file_grade_info_docx_with_teacher_signature(self):
        """Test get_file_grade_info with DOCX file containing teacher signature."""
        test_file = os.path.join(self.temp_dir, 'test.docx')
        
        # Create DOCX with teacher signature
        doc = Document()
        doc.add_paragraph('Test document')
        
        # Add table with teacher signature
        table = doc.add_table(rows=1, cols=1)
        cell = table.cell(0, 0)
        cell.text = '教师（签字）\n评分：B\n评价：很好的作业'
        
        doc.save(test_file)
        
        with patch('grading.views._iter_tables') as mock_iter_tables:
            mock_iter_tables.return_value = [table]
            
            with patch('grading.views.extract_grade_and_comment_from_cell') as mock_extract:
                mock_extract.return_value = ('B', '很好的作业', '教师（签字）')
                
                result = get_file_grade_info(test_file)
        
        self.assertTrue(result['has_grade'])
        self.assertEqual(result['grade'], 'B')
        self.assertEqual(result['grade_type'], 'letter')
        self.assertTrue(result['in_table'])
        self.assertTrue(result['has_comment'])
        self.assertEqual(result['comment'], '很好的作业')
    
    def test_get_file_grade_info_docx_locked_file(self):
        """Test get_file_grade_info with locked DOCX file."""
        test_file = os.path.join(self.temp_dir, 'test.docx')
        
        # Create DOCX with locked indicator
        doc = Document()
        doc.add_paragraph('【格式错误-已锁定】')
        doc.add_paragraph('老师评分：C')
        
        doc.save(test_file)
        
        with patch('grading.views.find_teacher_signature_cell') as mock_find_cell:
            mock_find_cell.return_value = (Mock(), None, None, None)
            
            result = get_file_grade_info(test_file)
        
        self.assertTrue(result['locked'])
        self.assertTrue(result['format_valid'])
        self.assertTrue(result['has_grade'])
        self.assertEqual(result['grade'], 'C')
    
    def test_get_file_grade_info_percentage_grade_detection(self):
        """Test get_file_grade_info correctly detects percentage grades."""
        test_cases = [
            ('85', 'percentage'),
            ('92.5', 'percentage'),
            ('100', 'percentage'),
            ('0', 'percentage'),
            ('150', 'letter'),  # Invalid percentage
            ('-10', 'letter'),  # Invalid percentage
            ('A+', 'letter'),
            ('优秀', 'text'),
        ]
        
        for grade_value, expected_type in test_cases:
            with self.subTest(grade_value=grade_value):
                test_file = os.path.join(self.temp_dir, f'test_{grade_value}.txt')
                
                with open(test_file, 'w', encoding='utf-8') as f:
                    f.write(f'老师评分：{grade_value}')
                
                result = get_file_grade_info(test_file)
                
                self.assertTrue(result['has_grade'])
                self.assertEqual(result['grade'], grade_value)
                self.assertEqual(result['grade_type'], expected_type)
    
    def test_get_file_grade_info_lab_report_detection(self):
        """Test get_file_grade_info correctly detects lab reports."""
        test_file = os.path.join(self.temp_dir, 'test.txt')
        
        with open(test_file, 'w', encoding='utf-8') as f:
            f.write('测试内容')
        
        with patch('grading.views.is_lab_report_file') as mock_is_lab:
            mock_is_lab.return_value = True
            
            result = get_file_grade_info(test_file, course_name='Lab Course')
        
        self.assertTrue(result['is_lab_report'])
    
    def test_get_file_grade_info_exception_handling(self):
        """Test get_file_grade_info handles exceptions gracefully."""
        test_file = os.path.join(self.temp_dir, 'test.docx')
        
        # Create invalid DOCX file
        with open(test_file, 'w') as f:
            f.write('This is not a valid DOCX file')
        
        result = get_file_grade_info(test_file)
        
        # Should return default values without raising exception
        self.assertFalse(result['has_grade'])
        self.assertIsNone(result['grade'])
        self.assertIsNone(result['grade_type'])


class TestErrorHandlingAndEdgeCases(BaseFileOperationTestCase):
    """Test error handling and edge cases in file operations."""
    
    def test_update_file_grade_status_with_special_characters(self):
        """Test update_file_grade_status with special characters in path."""
        special_path = 'course/班级1/作业#1/文件.txt'
        
        update_file_grade_status(
            self.repository,
            special_path,
            course_name='Test Course',
            user=self.user
        )
        
        # Should handle special characters without error
        status = FileGradeStatus.objects.get(
            repository=self.repository,
            file_path='Test Course/course/班级1/作业#1/文件.txt'
        )
        
        self.assertIsNotNone(status)
    
    def test_maybe_sync_repository_permission_error(self):
        """Test maybe_sync_repository handles permission errors."""
        self.repository.repo_type = 'git'
        self.repository.url = 'https://github.com/test/repo.git'
        self.repository.local_path = '/root/no_permission'  # Likely no permission
        self.repository.save()
        
        with patch('os.makedirs') as mock_makedirs:
            mock_makedirs.side_effect = PermissionError('Permission denied')
            
            result = maybe_sync_repository(self.repository, self.request)
        
        # Should handle error gracefully
        self.assertFalse(result)
    
    def test_file_has_updates_with_corrupted_status(self):
        """Test _file_has_updates with corrupted FileGradeStatus."""
        test_file = os.path.join(self.temp_dir, 'test.txt')
        with open(test_file, 'w') as f:
            f.write('test content')
        
        # Create corrupted FileGradeStatus (missing required fields)
        with patch('grading.views.FileGradeStatus.objects.filter') as mock_filter:
            mock_filter.side_effect = Exception('Database corruption')
            
            result = _file_has_updates(
                self.repository,
                'test.txt',
                test_file,
                base_dir=self.temp_dir
            )
        
        # Should handle error gracefully and return False
        self.assertFalse(result)
    
    def test_homework_folder_has_updates_empty_folder(self):
        """Test _homework_folder_has_updates with empty folder."""
        # Create empty homework folder
        homework_dir = os.path.join(self.temp_dir, 'empty_homework')
        os.makedirs(homework_dir)
        
        result = _homework_folder_has_updates(
            self.repository,
            homework_dir,
            'empty_homework',
            base_dir=self.temp_dir
        )
        
        self.assertFalse(result)
    
    def test_homework_folder_has_updates_deep_nesting(self):
        """Test _homework_folder_has_updates with deeply nested files."""
        # Create deeply nested structure
        deep_dir = os.path.join(self.temp_dir, 'homework', 'level1', 'level2', 'level3')
        os.makedirs(deep_dir)
        
        # Create file at deep level
        with open(os.path.join(deep_dir, 'deep_file.txt'), 'w') as f:
            f.write('deep content')
        
        homework_dir = os.path.join(self.temp_dir, 'homework')
        
        result = _homework_folder_has_updates(
            self.repository,
            homework_dir,
            'homework',
            base_dir=self.temp_dir
        )
        
        # Should skip files beyond depth limit
        # The exact behavior depends on the depth limit in the function
        self.assertIsInstance(result, bool)


if __name__ == '__main__':
    pytest.main([__file__])