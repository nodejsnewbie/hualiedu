"""
作业成绩写入成绩登分册工具类单元测试

测试三个核心工具类：
- GradeFileProcessor: 处理作业成绩文件
- RegistryManager: 管理Excel登分册
- NameMatcher: 学生姓名匹配
"""

import os
import tempfile
import shutil
from unittest.mock import Mock, patch, MagicMock

from docx import Document
from openpyxl import Workbook

from grading.grade_registry_writer import (
    GradeFileProcessor,
    RegistryManager,
    NameMatcher,
)
from .base import BaseTestCase


class GradeFileProcessorExtractStudentNameTest(BaseTestCase):
    """测试GradeFileProcessor.extract_student_name方法"""

    def test_extract_name_with_underscore_name_first(self):
        """测试使用下划线分隔，姓名在前"""
        result = GradeFileProcessor.extract_student_name("张三_作业1.docx")
        self.assertEqual(result, "张三")

    def test_extract_name_with_underscore_homework_first(self):
        """测试使用下划线分隔，作业标识在前"""
        result = GradeFileProcessor.extract_student_name("作业1_张三.docx")
        self.assertEqual(result, "张三")

    def test_extract_name_with_hyphen(self):
        """测试使用连字符分隔"""
        result = GradeFileProcessor.extract_student_name("张三-作业1.docx")
        self.assertEqual(result, "张三")

    def test_extract_name_with_em_dash(self):
        """测试使用长破折号分隔"""
        result = GradeFileProcessor.extract_student_name("张三—作业1.docx")
        self.assertEqual(result, "张三")
    
    def test_extract_name_no_separator(self):
        """文件名无分隔符时应直接返回文件名"""
        result = GradeFileProcessor.extract_student_name("李四.docx")
        self.assertEqual(result, "李四")


class GradeFileProcessorExtractHomeworkNumberTest(BaseTestCase):
    """测试作业批次提取支持中文数字"""

    def test_extract_from_path_with_chinese_characters(self):
        """目录名包含中文数字时也能识别批次"""
        path = "/tmp/Web前端开发/23计算机5班/第一次作业"
        self.assertEqual(GradeFileProcessor.extract_homework_number_from_path(path), 1)

    def test_extract_from_path_with_double_digit_chinese(self):
        """支持两位中文数字（第十二次）"""
        path = "/tmp/Web前端开发/23计算机5班/第十二次作业"
        self.assertEqual(GradeFileProcessor.extract_homework_number_from_path(path), 12)

    def test_extract_from_filename_with_chinese_characters(self):
        """Excel 文件名中的中文数字可识别"""
        filename = "/tmp/成绩/第十次作业成绩.xlsx"
        self.assertEqual(GradeFileProcessor.extract_homework_number_from_filename(filename), 10)


class GradeFileProcessorFindGradeInTextTest(BaseTestCase):
    """
    测试GradeFileProcessor._find_grade_in_text方法
    
    验证支持：
    - 字母等级：A/B/C/D/E
    - 文字等级：优秀/良好/中等/及格/不及格
    - 百分制：0-100的数字
    
    需求: 4.1, 4.3, 4.4, 7.1-7.7
    """

    def test_find_letter_grade(self):
        """测试识别字母等级"""
        self.assertEqual(GradeFileProcessor._find_grade_in_text("老师评分：A"), "A")
        self.assertEqual(GradeFileProcessor._find_grade_in_text("评分：B"), "B")
        self.assertEqual(GradeFileProcessor._find_grade_in_text("成绩：C"), "C")

    def test_find_text_grade(self):
        """测试识别文字等级"""
        self.assertEqual(GradeFileProcessor._find_grade_in_text("老师评分：优秀"), "优秀")
        self.assertEqual(GradeFileProcessor._find_grade_in_text("评分：良好"), "良好")
        self.assertEqual(GradeFileProcessor._find_grade_in_text("成绩：及格"), "及格")

    def test_find_percentage_grade_integer(self):
        """测试识别百分制成绩（整数）"""
        self.assertEqual(GradeFileProcessor._find_grade_in_text("老师评分：85"), "85")
        self.assertEqual(GradeFileProcessor._find_grade_in_text("评分：90分"), "90")
        self.assertEqual(GradeFileProcessor._find_grade_in_text("成绩：100"), "100")
        self.assertEqual(GradeFileProcessor._find_grade_in_text("得分：0"), "0")

    def test_find_percentage_grade_decimal(self):
        """测试识别百分制成绩（小数）"""
        self.assertEqual(GradeFileProcessor._find_grade_in_text("老师评分：85.5"), "85.5")
        self.assertEqual(GradeFileProcessor._find_grade_in_text("评分：90.5分"), "90.5")
        self.assertEqual(GradeFileProcessor._find_grade_in_text("成绩：78.8"), "78.8")

    def test_find_percentage_grade_out_of_range(self):
        """测试超出范围的百分制成绩应被忽略"""
        # 超过100的分数应该被忽略，尝试匹配其他等级
        result = GradeFileProcessor._find_grade_in_text("老师评分：150")
        self.assertIsNone(result)
        
        # 负数应该被忽略
        result = GradeFileProcessor._find_grade_in_text("老师评分：-10")
        self.assertIsNone(result)

    def test_find_grade_priority(self):
        """测试百分制优先于字母等级"""
        # 如果文本中同时包含数字和字母，应该优先识别数字
        self.assertEqual(GradeFileProcessor._find_grade_in_text("老师评分：85 A"), "85")

    def test_find_no_grade(self):
        """测试没有成绩的情况"""
        self.assertIsNone(GradeFileProcessor._find_grade_in_text("这是一段没有成绩的文本"))
        self.assertIsNone(GradeFileProcessor._find_grade_in_text(""))


class GradeFileProcessorValidateLabReportCommentTest(BaseTestCase):
    """
    测试GradeFileProcessor.validate_lab_report_comment方法
    
    验证实验报告评价验证功能
    
    需求: 4.5, 5.2, 7.1-7.7
    """

    def setUp(self):
        """设置测试环境"""
        super().setUp()
        self.test_dir = tempfile.mkdtemp()

    def tearDown(self):
        """清理测试环境"""
        shutil.rmtree(self.test_dir, ignore_errors=True)
        super().tearDown()

    def test_validate_non_lab_report(self):
        """测试非实验报告不需要验证评价"""
        # 创建一个普通作业文件
        doc_path = os.path.join(self.test_dir, "张三_作业1.docx")
        doc = Document()
        doc.add_paragraph("老师评分：A")
        doc.save(doc_path)
        
        # 非实验报告应该验证通过
        is_valid, error_msg = GradeFileProcessor.validate_lab_report_comment(doc_path)
        self.assertTrue(is_valid)
        self.assertIsNone(error_msg)

    @patch('grading.grade_registry_writer.GradeFileProcessor.is_lab_report')
    def test_validate_lab_report_with_comment(self, mock_is_lab):
        """测试实验报告有评价时验证通过"""
        mock_is_lab.return_value = True
        
        # 创建一个实验报告文件，包含评价
        doc_path = os.path.join(self.test_dir, "张三_实验报告.docx")
        doc = Document()
        
        # 添加表格，模拟实验报告格式
        table = doc.add_table(rows=2, cols=2)
        cell = table.rows[0].cells[0]
        cell.text = "85\n这是一个很好的实验报告\n教师（签字）："
        
        doc.save(doc_path)
        
        # 有评价应该验证通过
        is_valid, error_msg = GradeFileProcessor.validate_lab_report_comment(doc_path)
        self.assertTrue(is_valid)
        self.assertIsNone(error_msg)

    @patch('grading.grade_registry_writer.GradeFileProcessor.is_lab_report')
    def test_validate_lab_report_without_comment(self, mock_is_lab):
        """测试实验报告没有评价时验证失败"""
        mock_is_lab.return_value = True
        
        # 创建一个实验报告文件，只有评分没有评价
        doc_path = os.path.join(self.test_dir, "张三_实验报告.docx")
        doc = Document()
        
        # 添加表格，模拟实验报告格式（只有评分）
        table = doc.add_table(rows=2, cols=2)
        cell = table.rows[0].cells[0]
        cell.text = "85\n教师（签字）："
        
        doc.save(doc_path)
        
        # 没有评价应该验证失败
        is_valid, error_msg = GradeFileProcessor.validate_lab_report_comment(doc_path)
        self.assertFalse(is_valid)
        self.assertIn("实验报告必须添加评价", error_msg)


class RegistryManagerConcurrencyTest(BaseTestCase):
    """测试RegistryManager的并发控制功能"""

    def setUp(self):
        """设置测试环境"""
        super().setUp()
        self.test_dir = tempfile.mkdtemp()
        self.registry_path = os.path.join(self.test_dir, "成绩登分册.xlsx")
        self._create_test_registry()

    def tearDown(self):
        """清理测试环境"""
        if os.path.exists(self.test_dir):
            shutil.rmtree(self.test_dir)
        super().tearDown()

    def _create_test_registry(self):
        """创建测试用的登分册文件"""
        workbook = Workbook()
        worksheet = workbook.active
        worksheet.title = "成绩表"

        # 添加表头
        worksheet.cell(1, 1).value = "姓名"
        worksheet.cell(1, 2).value = "学号"

        # 添加学生数据
        students = [
            ("张三", "20210001"),
            ("李四", "20210002"),
            ("王五", "20210003"),
        ]
        for i, (name, student_id) in enumerate(students, start=2):
            worksheet.cell(i, 1).value = name
            worksheet.cell(i, 2).value = student_id

        workbook.save(self.registry_path)
        workbook.close()

    def test_acquire_file_lock_success(self):
        """测试成功获取文件锁"""
        manager = RegistryManager(self.registry_path)

        # 加载文件应该成功获取锁
        self.assertTrue(manager.load())
        self.assertIsNotNone(manager.lock_file_handle)
        self.assertTrue(os.path.exists(manager.lock_file_path))

        # 清理
        manager._release_file_lock()

    def test_file_lock_prevents_concurrent_access(self):
        """测试文件锁阻止并发访问"""
        manager1 = RegistryManager(self.registry_path)
        manager2 = RegistryManager(self.registry_path)

        # 第一个管理器加载文件并获取锁
        self.assertTrue(manager1.load())

        # 第二个管理器尝试加载应该失败（文件被锁定）
        self.assertFalse(manager2.load())

        # 释放第一个管理器的锁
        manager1._release_file_lock()

        # 现在第二个管理器应该能够加载
        self.assertTrue(manager2.load())
        manager2._release_file_lock()

    def test_release_file_lock(self):
        """测试释放文件锁"""
        manager = RegistryManager(self.registry_path)

        # 获取锁
        manager.load()
        lock_file_path = manager.lock_file_path

        # 释放锁
        manager._release_file_lock()

        # 验证锁文件被删除
        self.assertFalse(os.path.exists(lock_file_path))
        self.assertIsNone(manager.lock_file_handle)

    def test_check_file_in_use(self):
        """测试检测文件是否被占用"""
        manager1 = RegistryManager(self.registry_path)
        manager2 = RegistryManager(self.registry_path)

        # 第一个管理器加载文件
        manager1.load()

        # 第二个管理器尝试加载应该失败（因为文件被锁定）
        # _check_file_in_use 在 load() 中被调用
        self.assertFalse(manager2.load())

        # 释放锁后应该能够加载
        manager1._release_file_lock()
        self.assertTrue(manager2.load())
        manager2._release_file_lock()

    def test_save_releases_lock(self):
        """测试保存文件后释放锁"""
        manager = RegistryManager(self.registry_path)

        # 加载并修改
        manager.load()
        manager.validate_format()
        lock_file_path = manager.lock_file_path

        # 保存应该释放锁
        self.assertTrue(manager.save())
        self.assertFalse(os.path.exists(lock_file_path))
        self.assertIsNone(manager.lock_file_handle)

    def test_restore_from_backup_releases_lock(self):
        """测试从备份恢复后释放锁"""
        manager = RegistryManager(self.registry_path)

        # 加载并创建备份
        manager.load()
        manager.validate_format()
        manager.create_backup()
        lock_file_path = manager.lock_file_path

        # 恢复应该释放锁
        self.assertTrue(manager.restore_from_backup())
        self.assertFalse(os.path.exists(lock_file_path))
        self.assertIsNone(manager.lock_file_handle)

    def test_destructor_releases_lock(self):
        """测试析构函数释放锁"""
        manager = RegistryManager(self.registry_path)
        manager.load()
        lock_file_path = manager.lock_file_path

        # 删除对象应该释放锁
        del manager

        # 验证锁文件被删除
        self.assertFalse(os.path.exists(lock_file_path))

    def test_concurrent_write_conflict_handling(self):
        """测试并发写入冲突处理"""
        manager1 = RegistryManager(self.registry_path)
        manager2 = RegistryManager(self.registry_path)

        # 第一个管理器加载并修改
        self.assertTrue(manager1.load())
        self.assertTrue(manager1.validate_format()[0])
        col = manager1.find_or_create_homework_column(1)
        row = manager1.find_student_row("张三")
        manager1.write_grade(row, col, "A")

        # 第二个管理器尝试加载应该失败
        self.assertFalse(manager2.load())

        # 第一个管理器保存
        self.assertTrue(manager1.save())

        # 现在第二个管理器可以加载
        self.assertTrue(manager2.load())
        self.assertTrue(manager2.validate_format()[0])

        # 验证第一个管理器的修改已保存
        row = manager2.find_student_row("张三")
        grade = manager2.worksheet.cell(row, col).value
        self.assertEqual(grade, "A")

        manager2._release_file_lock()

    def test_load_failure_releases_lock(self):
        """测试加载失败时释放锁"""
        # 创建一个不存在的文件路径
        invalid_path = os.path.join(self.test_dir, "不存在的文件.xlsx")
        manager = RegistryManager(invalid_path)

        # 加载应该失败
        self.assertFalse(manager.load())

        # 验证没有锁文件残留
        lock_file_path = f"{invalid_path}.lock"
        self.assertFalse(os.path.exists(lock_file_path))
        self.assertIsNone(manager.lock_file_handle)


class RegistryManagerHeaderDetectionTest(BaseTestCase):
    """测试登分册表头自动检测"""

    def setUp(self):
        super().setUp()
        self.test_dir = tempfile.mkdtemp()
        self.registry_path = os.path.join(self.test_dir, "registry.xlsx")
        workbook = Workbook()
        worksheet = workbook.active
        worksheet.append(["学校成绩表"])
        worksheet.append(["", "", ""])
        worksheet.append(["序号", "学号", "姓名", "第1次作业"])
        worksheet.append([1, "1001", "张三", "A"])
        workbook.save(self.registry_path)
        workbook.close()

    def tearDown(self):
        if os.path.exists(self.test_dir):
            shutil.rmtree(self.test_dir)
        super().tearDown()

    def test_validate_format_detects_header_not_in_first_row(self):
        manager = RegistryManager(self.registry_path)
        self.assertTrue(manager.load())
        is_valid, error = manager.validate_format()
        self.assertTrue(is_valid, msg=error)
        self.assertEqual(manager.header_row_index, 3)
        self.assertEqual(manager.name_column_index, 3)
        self.assertEqual(manager.find_student_row("张三"), 4)
        manager._release_file_lock()
