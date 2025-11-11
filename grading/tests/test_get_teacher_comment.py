"""
测试获取教师评价功能
Tests for get_teacher_comment view function
"""

import os
import tempfile
import shutil
from docx import Document

from django.test import TestCase, RequestFactory
from django.contrib.auth.models import User
from django.http import JsonResponse

from grading.models import Repository, Semester, Course, Homework
from grading.views import get_teacher_comment, write_to_teacher_signature_cell


class GetTeacherCommentTest(TestCase):
    """测试 get_teacher_comment 视图函数"""

    def setUp(self):
        """设置测试环境"""
        self.temp_dir = tempfile.mkdtemp()
        self.factory = RequestFactory()
        
        # 创建测试用户
        self.user = User.objects.create_user(username='testuser', password='testpass')
        
        # 创建测试学期
        from datetime import date
        self.semester = Semester.objects.create(
            name='2024年春季学期',
            start_date=date(2024, 3, 1),
            end_date=date(2024, 7, 15),
            is_active=True
        )
        
        # 创建测试课程
        self.lab_course = Course.objects.create(
            semester=self.semester,
            teacher=self.user,
            name='计算机网络实验',
            course_type='lab',
            location='实验室201'
        )
        
        self.theory_course = Course.objects.create(
            semester=self.semester,
            teacher=self.user,
            name='数据结构',
            course_type='theory',
            location='教学楼101'
        )
        
        # 创建测试作业批次
        self.lab_homework = Homework.objects.create(
            course=self.lab_course,
            title='实验一',
            homework_type='lab_report',
            folder_name='实验1'
        )
        
        self.normal_homework = Homework.objects.create(
            course=self.theory_course,
            title='第一次作业',
            homework_type='normal',
            folder_name='作业1'
        )
        
        # 创建测试仓库
        self.repo = Repository.objects.create(
            name='测试仓库',
            owner=self.user,
            repo_type='local',
            local_path=self.temp_dir,
            is_active=True
        )

    def tearDown(self):
        """清理测试环境"""
        if os.path.exists(self.temp_dir):
            shutil.rmtree(self.temp_dir)

    def create_lab_report_with_comment(self, comment_text):
        """创建包含评价的实验报告"""
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        
        # 添加"教师（签字）"单元格并写入评价
        cell = table.rows[1].cells[0]
        write_to_teacher_signature_cell(
            cell,
            grade="A",
            comment=comment_text,
            signature_text="教师（签字）：张老师 时间：2024-01-01"
        )
        
        return doc

    def create_normal_homework_with_comment(self, comment_text):
        """创建包含评价的普通作业"""
        doc = Document()
        doc.add_paragraph("这是作业内容")
        doc.add_paragraph(f"教师评价：{comment_text}")
        
        return doc

    def test_get_comment_from_lab_report(self):
        """测试从实验报告中获取评价"""
        # 创建测试文件
        comment_text = "实验报告完成得很好，数据准确。"
        doc = self.create_lab_report_with_comment(comment_text)
        
        # 保存文件
        course_dir = os.path.join(self.temp_dir, '计算机网络实验')
        homework_dir = os.path.join(course_dir, '实验1')
        os.makedirs(homework_dir, exist_ok=True)
        
        file_path = os.path.join(homework_dir, '学生作业.docx')
        doc.save(file_path)
        
        # 创建请求
        relative_path = os.path.relpath(file_path, self.temp_dir)
        request = self.factory.get(
            '/get_teacher_comment/',
            {
                'file_path': relative_path,
                'repo_id': str(self.repo.id),
                'course': '计算机网络实验',
                'homework_folder': '实验1'
            }
        )
        request.user = self.user
        
        # 调用视图函数
        response = get_teacher_comment(request)
        
        # 验证响应
        self.assertIsInstance(response, JsonResponse)
        data = response.json()
        self.assertTrue(data['success'])
        self.assertEqual(data['comment'], comment_text)

    def test_get_comment_from_normal_homework(self):
        """测试从普通作业中获取评价"""
        # 创建测试文件
        comment_text = "作业完成得不错，继续努力。"
        doc = self.create_normal_homework_with_comment(comment_text)
        
        # 保存文件
        course_dir = os.path.join(self.temp_dir, '数据结构')
        homework_dir = os.path.join(course_dir, '作业1')
        os.makedirs(homework_dir, exist_ok=True)
        
        file_path = os.path.join(homework_dir, '学生作业.docx')
        doc.save(file_path)
        
        # 创建请求
        relative_path = os.path.relpath(file_path, self.temp_dir)
        request = self.factory.get(
            '/get_teacher_comment/',
            {
                'file_path': relative_path,
                'repo_id': str(self.repo.id),
                'course': '数据结构',
                'homework_folder': '作业1'
            }
        )
        request.user = self.user
        
        # 调用视图函数
        response = get_teacher_comment(request)
        
        # 验证响应
        self.assertIsInstance(response, JsonResponse)
        data = response.json()
        self.assertTrue(data['success'])
        self.assertEqual(data['comment'], comment_text)

    def test_get_comment_no_comment_found(self):
        """测试文件中没有评价的情况"""
        # 创建没有评价的文件
        doc = Document()
        doc.add_paragraph("这是作业内容")
        
        # 保存文件
        course_dir = os.path.join(self.temp_dir, '数据结构')
        homework_dir = os.path.join(course_dir, '作业1')
        os.makedirs(homework_dir, exist_ok=True)
        
        file_path = os.path.join(homework_dir, '学生作业.docx')
        doc.save(file_path)
        
        # 创建请求
        relative_path = os.path.relpath(file_path, self.temp_dir)
        request = self.factory.get(
            '/get_teacher_comment/',
            {
                'file_path': relative_path,
                'repo_id': str(self.repo.id),
                'course': '数据结构',
                'homework_folder': '作业1'
            }
        )
        request.user = self.user
        
        # 调用视图函数
        response = get_teacher_comment(request)
        
        # 验证响应
        self.assertIsInstance(response, JsonResponse)
        data = response.json()
        self.assertTrue(data['success'])
        self.assertEqual(data['comment'], '暂无评价')

    def test_get_comment_file_not_found(self):
        """测试文件不存在的情况"""
        # 创建请求（文件不存在）
        request = self.factory.get(
            '/get_teacher_comment/',
            {
                'file_path': '不存在的文件.docx',
                'repo_id': str(self.repo.id),
                'course': '数据结构',
                'homework_folder': '作业1'
            }
        )
        request.user = self.user
        
        # 调用视图函数
        response = get_teacher_comment(request)
        
        # 验证响应
        self.assertIsInstance(response, JsonResponse)
        data = response.json()
        self.assertFalse(data['success'])

    def test_get_comment_no_file_path(self):
        """测试未提供文件路径的情况"""
        # 创建请求（缺少file_path参数）
        request = self.factory.get('/get_teacher_comment/', {})
        request.user = self.user
        
        # 调用视图函数
        response = get_teacher_comment(request)
        
        # 验证响应
        self.assertIsInstance(response, JsonResponse)
        data = response.json()
        self.assertFalse(data['success'])
        self.assertEqual(data['message'], '未提供文件路径')

    def test_get_comment_with_ai_evaluation(self):
        """测试获取AI评价"""
        # 创建包含AI评价的文件
        comment_text = "AI评价：代码结构清晰，逻辑正确。"
        doc = Document()
        doc.add_paragraph("这是作业内容")
        doc.add_paragraph(comment_text)
        
        # 保存文件
        course_dir = os.path.join(self.temp_dir, '数据结构')
        homework_dir = os.path.join(course_dir, '作业1')
        os.makedirs(homework_dir, exist_ok=True)
        
        file_path = os.path.join(homework_dir, '学生作业.docx')
        doc.save(file_path)
        
        # 创建请求
        relative_path = os.path.relpath(file_path, self.temp_dir)
        request = self.factory.get(
            '/get_teacher_comment/',
            {
                'file_path': relative_path,
                'repo_id': str(self.repo.id),
                'course': '数据结构',
                'homework_folder': '作业1'
            }
        )
        request.user = self.user
        
        # 调用视图函数
        response = get_teacher_comment(request)
        
        # 验证响应
        self.assertIsInstance(response, JsonResponse)
        data = response.json()
        self.assertTrue(data['success'])
        # 应该提取冒号后的内容
        self.assertEqual(data['comment'], "代码结构清晰，逻辑正确。")
