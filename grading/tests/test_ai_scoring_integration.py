"""
集成测试：AI评分功能
Integration tests for AI scoring functionality

测试需求: 6.1-6.10 (AI评分), 8.1-8.7 (批量AI评分)
"""

import os
import tempfile
import shutil
from unittest.mock import patch, MagicMock

from django.test import TestCase, Client
from django.contrib.auth.models import User
from django.urls import reverse
from docx import Document

from grading.models import Repository, Course, Semester, Homework


class AIScoreViewTest(TestCase):
    """测试AI评分视图函数"""

    def setUp(self):
        """设置测试环境"""
        # 创建测试用户
        self.user = User.objects.create_user(
            username='testuser',
            password='testpass123',
            is_staff=True
        )
        self.client = Client()
        self.client.login(username='testuser', password='testpass123')
        
        # 创建临时目录
        self.temp_dir = tempfile.mkdtemp()
        
        # 创建测试学期
        from datetime import date
        self.semester = Semester.objects.create(
            name='2024年春季学期',
            start_date=date(2024, 3, 1),
            end_date=date(2024, 7, 15),
            is_active=True
        )
        
        # 创建测试仓库
        self.repo = Repository.objects.create(
            name='测试仓库',
            owner=self.user,
            repo_type='local',
            path=self.temp_dir,
            is_active=True
        )
        
        # 创建测试文档
        self.test_file_path = os.path.join(self.temp_dir, 'test_homework.docx')
        doc = Document()
        doc.add_paragraph('这是一份测试作业，内容包括Python编程基础知识。')
        doc.add_paragraph('学生完成了所有练习题，代码运行正常。')
        doc.save(self.test_file_path)

    def tearDown(self):
        """清理测试环境"""
        if os.path.exists(self.temp_dir):
            shutil.rmtree(self.temp_dir)

    @patch('grading.views.volcengine_score_homework')
    def test_ai_score_success(self):
        """测试AI评分成功流程 - 需求 6.1, 6.2"""
        # Mock AI服务返回
        mock_score = 85
        mock_comment = "作业完成得很好，代码规范，逻辑清晰。"
        
        with patch('grading.views.volcengine_score_homework', return_value=(mock_score, mock_comment)):
            response = self.client.post(reverse('grading:ai_score'), {
                'file_path': 'test_homework.docx',
                'repo_id': self.repo.id
            })
        
        self.assertEqual(response.status_code, 200)
        data = response.json()
        self.assertEqual(data['status'], 'success')
        self.assertEqual(data['ai_score'], mock_score)
        self.assertEqual(data['ai_grade'], 'B')  # 85分对应B等级
        self.assertEqual(data['ai_comment'], mock_comment)

    @patch('grading.views.volcengine_score_homework')
    def test_ai_score_confirm_write(self):
        """测试确认AI评分并写入文件 - 需求 6.3, 6.4, 6.5, 6.6, 6.7"""
        # Mock AI服务返回
        mock_score = 90
        mock_comment = "优秀的作业"
        
        with patch('grading.views.volcengine_score_homework', return_value=(mock_score, mock_comment)):
            # 第一步：获取AI评分建议
            response1 = self.client.post(reverse('grading:ai_score'), {
                'file_path': 'test_homework.docx',
                'repo_id': self.repo.id
            })
            
            self.assertEqual(response1.status_code, 200)
            data1 = response1.json()
            ai_grade = data1['ai_grade']
            ai_comment = data1['ai_comment']
            
            # 第二步：确认并写入
            response2 = self.client.post(reverse('grading:ai_score'), {
                'file_path': 'test_homework.docx',
                'repo_id': self.repo.id,
                'confirm': 'true',
                'ai_grade': ai_grade,
                'ai_comment': ai_comment
            })
            
            self.assertEqual(response2.status_code, 200)
            data2 = response2.json()
            self.assertEqual(data2['status'], 'success')
            self.assertIn('已保存', data2['message'])
            
            # 验证文件已被修改
            doc = Document(self.test_file_path)
            doc_text = '\n'.join([p.text for p in doc.paragraphs])
            self.assertIn(ai_grade, doc_text)
            self.assertIn('AI评价', doc_text)

    def test_ai_score_locked_file(self):
        """测试已锁定文件的AI评分禁用 - 需求 6.8"""
        # 创建已锁定的文件
        locked_file_path = os.path.join(self.temp_dir, 'locked_homework.docx')
        doc = Document()
        doc.add_paragraph('测试内容')
        doc.add_paragraph('【格式错误-已锁定】请按要求的格式写实验报告，此评分不可修改')
        doc.save(locked_file_path)
        
        response = self.client.post(reverse('grading:ai_score'), {
            'file_path': 'locked_homework.docx',
            'repo_id': self.repo.id
        })
        
        self.assertEqual(response.status_code, 400)
        data = response.json()
        self.assertEqual(data['status'], 'error')
        self.assertIn('锁定', data['message'])

    @patch('grading.views.volcengine_score_homework')
    def test_ai_score_api_failure(self):
        """测试AI服务失败处理 - 需求 6.9"""
        # Mock AI服务返回失败
        with patch('grading.views.volcengine_score_homework', return_value=(None, "API调用失败")):
            response = self.client.post(reverse('grading:ai_score'), {
                'file_path': 'test_homework.docx',
                'repo_id': self.repo.id
            })
        
        self.assertEqual(response.status_code, 400)
        data = response.json()
        self.assertEqual(data['status'], 'error')
        self.assertIn('失败', data['message'])

    def test_ai_score_missing_file_path(self):
        """测试缺少文件路径参数"""
        response = self.client.post(reverse('grading:ai_score'), {
            'repo_id': self.repo.id
        })
        
        self.assertEqual(response.status_code, 400)
        data = response.json()
        self.assertEqual(data['status'], 'error')
        self.assertIn('文件路径', data['message'])

    def test_ai_score_invalid_file_path(self):
        """测试无效的文件路径"""
        response = self.client.post(reverse('grading:ai_score'), {
            'file_path': 'nonexistent.docx',
            'repo_id': self.repo.id
        })
        
        self.assertEqual(response.status_code, 400)
        data = response.json()
        self.assertEqual(data['status'], 'error')

    def test_ai_score_unauthenticated(self):
        """测试未认证用户访问"""
        self.client.logout()
        response = self.client.post(reverse('grading:ai_score'), {
            'file_path': 'test_homework.docx',
            'repo_id': self.repo.id
        })
        
        # 应该重定向到登录页面
        self.assertEqual(response.status_code, 302)

    @patch('grading.views.volcengine_score_homework')
    def test_ai_score_grade_conversion(self):
        """测试分数到等级的转换"""
        test_cases = [
            (95, 'A'),
            (85, 'B'),
            (75, 'C'),
            (65, 'D'),
            (55, 'E'),
        ]
        
        for score, expected_grade in test_cases:
            with self.subTest(score=score):
                with patch('grading.views.volcengine_score_homework', return_value=(score, "测试评价")):
                    response = self.client.post(reverse('grading:ai_score'), {
                        'file_path': 'test_homework.docx',
                        'repo_id': self.repo.id
                    })
                    
                    self.assertEqual(response.status_code, 200)
                    data = response.json()
                    self.assertEqual(data['ai_grade'], expected_grade)


class BatchAIScoreViewTest(TestCase):
    """测试批量AI评分视图函数"""

    def setUp(self):
        """设置测试环境"""
        # 创建测试用户
        self.user = User.objects.create_user(
            username='testuser',
            password='testpass123',
            is_staff=True
        )
        self.client = Client()
        self.client.login(username='testuser', password='testpass123')
        
        # 创建临时目录
        self.temp_dir = tempfile.mkdtemp()
        self.test_dir = os.path.join(self.temp_dir, 'homework_batch')
        os.makedirs(self.test_dir)
        
        # 创建测试学期
        from datetime import date
        self.semester = Semester.objects.create(
            name='2024年春季学期',
            start_date=date(2024, 3, 1),
            end_date=date(2024, 7, 15),
            is_active=True
        )
        
        # 创建测试仓库
        self.repo = Repository.objects.create(
            name='测试仓库',
            owner=self.user,
            repo_type='local',
            path=self.temp_dir,
            is_active=True
        )
        
        # 创建多个测试文档
        for i in range(3):
            file_path = os.path.join(self.test_dir, f'homework_{i+1}.docx')
            doc = Document()
            doc.add_paragraph(f'这是第{i+1}份作业的内容。')
            doc.save(file_path)

    def tearDown(self):
        """清理测试环境"""
        if os.path.exists(self.temp_dir):
            shutil.rmtree(self.temp_dir)

    @patch('grading.views.volcengine_score_homework')
    def test_batch_ai_score_success(self):
        """测试批量AI评分成功流程 - 需求 8.1, 8.2, 8.4"""
        # Mock AI服务返回
        def mock_ai_score(content):
            return (85, "作业完成得很好")
        
        with patch('grading.views.volcengine_score_homework', side_effect=mock_ai_score):
            response = self.client.post(reverse('grading:batch_ai_score'), {
                'dir_path': 'homework_batch',
                'repo_id': self.repo.id
            })
        
        self.assertEqual(response.status_code, 200)
        data = response.json()
        self.assertEqual(data['status'], 'success')
        self.assertEqual(data['total'], 3)
        self.assertEqual(data['success'], 3)
        self.assertEqual(data['failed'], 0)
        self.assertEqual(data['skipped'], 0)

    @patch('grading.views.volcengine_score_homework')
    @patch('grading.views.rate_limit_api_request')
    def test_batch_ai_score_rate_limit(self):
        """测试批量AI评分的速率限制 - 需求 8.3"""
        # Mock AI服务返回
        def mock_ai_score(content):
            return (85, "作业完成得很好")
        
        mock_rate_limit = MagicMock()
        
        with patch('grading.views.volcengine_score_homework', side_effect=mock_ai_score):
            with patch('grading.views.rate_limit_api_request', mock_rate_limit):
                response = self.client.post(reverse('grading:batch_ai_score'), {
                    'dir_path': 'homework_batch',
                    'repo_id': self.repo.id
                })
        
        # 验证速率限制函数被调用了3次（每个文件一次）
        self.assertEqual(mock_rate_limit.call_count, 3)

    @patch('grading.views.volcengine_score_homework')
    def test_batch_ai_score_partial_failure(self):
        """测试批量AI评分部分失败 - 需求 8.6"""
        # Mock AI服务：第二个文件失败
        call_count = [0]
        def mock_ai_score(content):
            call_count[0] += 1
            if call_count[0] == 2:
                return (None, "API调用失败")
            return (85, "作业完成得很好")
        
        with patch('grading.views.volcengine_score_homework', side_effect=mock_ai_score):
            response = self.client.post(reverse('grading:batch_ai_score'), {
                'dir_path': 'homework_batch',
                'repo_id': self.repo.id
            })
        
        self.assertEqual(response.status_code, 200)
        data = response.json()
        self.assertEqual(data['status'], 'success')
        self.assertEqual(data['total'], 3)
        self.assertEqual(data['success'], 2)
        self.assertEqual(data['failed'], 1)

    def test_batch_ai_score_with_locked_files(self):
        """测试批量AI评分跳过已锁定文件"""
        # 创建一个已锁定的文件
        locked_file_path = os.path.join(self.test_dir, 'locked_homework.docx')
        doc = Document()
        doc.add_paragraph('测试内容')
        doc.add_paragraph('【格式错误-已锁定】请按要求的格式写实验报告，此评分不可修改')
        doc.save(locked_file_path)
        
        with patch('grading.views.volcengine_score_homework', return_value=(85, "很好")):
            response = self.client.post(reverse('grading:batch_ai_score'), {
                'dir_path': 'homework_batch',
                'repo_id': self.repo.id
            })
        
        self.assertEqual(response.status_code, 200)
        data = response.json()
        self.assertEqual(data['total'], 4)  # 3个正常 + 1个锁定
        self.assertEqual(data['skipped'], 1)  # 锁定的文件被跳过

    def test_batch_ai_score_empty_directory(self):
        """测试批量AI评分空目录"""
        empty_dir = os.path.join(self.temp_dir, 'empty_dir')
        os.makedirs(empty_dir)
        
        response = self.client.post(reverse('grading:batch_ai_score'), {
            'dir_path': 'empty_dir',
            'repo_id': self.repo.id
        })
        
        self.assertEqual(response.status_code, 400)
        data = response.json()
        self.assertEqual(data['status'], 'error')
        self.assertIn('没有', data['message'])

    def test_batch_ai_score_missing_dir_path(self):
        """测试缺少目录路径参数"""
        response = self.client.post(reverse('grading:batch_ai_score'), {
            'repo_id': self.repo.id
        })
        
        self.assertEqual(response.status_code, 400)
        data = response.json()
        self.assertEqual(data['status'], 'error')
        self.assertIn('目录路径', data['message'])

    def test_batch_ai_score_invalid_directory(self):
        """测试无效的目录路径"""
        response = self.client.post(reverse('grading:batch_ai_score'), {
            'dir_path': 'nonexistent_dir',
            'repo_id': self.repo.id
        })
        
        self.assertEqual(response.status_code, 400)
        data = response.json()
        self.assertEqual(data['status'], 'error')

    @patch('grading.views.volcengine_score_homework')
    def test_batch_ai_score_result_summary(self):
        """测试批量AI评分结果摘要 - 需求 8.7"""
        with patch('grading.views.volcengine_score_homework', return_value=(85, "很好")):
            response = self.client.post(reverse('grading:batch_ai_score'), {
                'dir_path': 'homework_batch',
                'repo_id': self.repo.id
            })
        
        self.assertEqual(response.status_code, 200)
        data = response.json()
        
        # 验证结果摘要包含所有必要信息
        self.assertIn('total', data)
        self.assertIn('success', data)
        self.assertIn('failed', data)
        self.assertIn('skipped', data)
        self.assertIn('details', data)
        self.assertIsInstance(data['details'], list)
        
        # 验证每个文件的详细信息
        for detail in data['details']:
            self.assertIn('file', detail)
            self.assertIn('status', detail)
            if detail['status'] == 'success':
                self.assertIn('grade', detail)
                self.assertIn('score', detail)
                self.assertIn('comment', detail)


class AIScoreErrorHandlingTest(TestCase):
    """测试AI评分错误处理"""

    def setUp(self):
        """设置测试环境"""
        self.user = User.objects.create_user(
            username='testuser',
            password='testpass123',
            is_staff=True
        )
        self.client = Client()
        self.client.login(username='testuser', password='testpass123')
        
        self.temp_dir = tempfile.mkdtemp()
        
        self.repo = Repository.objects.create(
            name='测试仓库',
            owner=self.user,
            repo_type='local',
            path=self.temp_dir,
            is_active=True
        )

    def tearDown(self):
        """清理测试环境"""
        if os.path.exists(self.temp_dir):
            shutil.rmtree(self.temp_dir)

    def test_ai_score_empty_file(self):
        """测试空文件的AI评分"""
        empty_file_path = os.path.join(self.temp_dir, 'empty.docx')
        doc = Document()
        doc.save(empty_file_path)
        
        with patch('grading.views.volcengine_score_homework', return_value=(None, "内容为空")):
            response = self.client.post(reverse('grading:ai_score'), {
                'file_path': 'empty.docx',
                'repo_id': self.repo.id
            })
        
        self.assertEqual(response.status_code, 400)

    @patch('grading.views.volcengine_score_homework')
    def test_ai_score_exception_handling(self):
        """测试AI评分异常处理"""
        # Mock AI服务抛出异常
        with patch('grading.views.volcengine_score_homework', side_effect=Exception("网络错误")):
            # 创建测试文件
            test_file_path = os.path.join(self.temp_dir, 'test.docx')
            doc = Document()
            doc.add_paragraph('测试内容')
            doc.save(test_file_path)
            
            response = self.client.post(reverse('grading:ai_score'), {
                'file_path': 'test.docx',
                'repo_id': self.repo.id
            })
        
        self.assertEqual(response.status_code, 400)
        data = response.json()
        self.assertEqual(data['status'], 'error')

    def test_ai_score_permission_denied(self):
        """测试无权限访问文件"""
        # 创建另一个用户的仓库
        other_user = User.objects.create_user(
            username='otheruser',
            password='testpass123',
            is_staff=True
        )
        other_repo = Repository.objects.create(
            name='其他仓库',
            owner=other_user,
            repo_type='local',
            path=self.temp_dir,
            is_active=True
        )
        
        response = self.client.post(reverse('grading:ai_score'), {
            'file_path': 'test.docx',
            'repo_id': other_repo.id
        })
        
        self.assertEqual(response.status_code, 400)
        data = response.json()
        self.assertEqual(data['status'], 'error')
