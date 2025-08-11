"""
Django tests for AI scoring functionality.
"""

import json
import os
import tempfile
from unittest.mock import MagicMock, patch

from django.contrib.auth.models import User
from django.test import TestCase
from django.urls import reverse

from grading.models import GlobalConfig


class AIScoringTest(TestCase):
    """Test cases for AI scoring functionality."""

    def setUp(self):
        """Set up test data."""
        # Create test user
        self.user = User.objects.create_user(
            username="testuser", password="testpass123", is_staff=True
        )

        # Create temporary directory for test files
        self.temp_dir = tempfile.mkdtemp()
        self.test_file_path = os.path.join(self.temp_dir, "test_document.docx")

        # Create GlobalConfig
        self.config = GlobalConfig.objects.create(repo_base_dir=self.temp_dir)

        # Create a simple test file
        from docx import Document

        doc = Document()
        doc.add_paragraph("Test document content")
        doc.save(self.test_file_path)

    def tearDown(self):
        """Clean up test data."""
        import shutil

        shutil.rmtree(self.temp_dir, ignore_errors=True)

    @patch("grading.views.volcengine_score_homework")
    def test_ai_score_view_success(self, mock_ai_score):
        """Test successful AI scoring."""
        # Mock AI scoring response
        mock_ai_score.return_value = (85, "Excellent work! Very detailed and well-structured.")

        self.client.login(username="testuser", password="testpass123")

        data = {"path": "test_document.docx"}
        response = self.client.post(reverse("grading:ai_score"), data)

        self.assertEqual(response.status_code, 200)
        response_data = json.loads(response.content)
        self.assertEqual(response_data["status"], "success")
        self.assertEqual(response_data["score"], 85)
        self.assertEqual(response_data["grade"], "B")
        self.assertIn("Excellent work", response_data["comment"])

    @patch("grading.views.volcengine_score_homework")
    def test_ai_score_view_failure(self, mock_ai_score):
        """Test AI scoring failure."""
        # Mock AI scoring failure
        mock_ai_score.side_effect = Exception("API error")

        self.client.login(username="testuser", password="testpass123")

        data = {"path": "test_document.docx"}
        response = self.client.post(reverse("grading:ai_score"), data)

        self.assertEqual(response.status_code, 500)
        response_data = json.loads(response.content)
        self.assertEqual(response_data["status"], "error")

    def test_ai_score_view_missing_data(self):
        """Test AI scoring with missing data."""
        self.client.login(username="testuser", password="testpass123")

        data = {}  # Missing path
        response = self.client.post(reverse("grading:ai_score"), data)

        self.assertEqual(response.status_code, 400)

    def test_ai_score_view_unauthenticated(self):
        """Test AI scoring without authentication."""
        data = {"path": "test_document.docx"}
        response = self.client.post(reverse("grading:ai_score"), data)

        self.assertEqual(response.status_code, 302)  # Redirect to login

    def test_ai_score_view_already_graded(self):
        """Test AI scoring when file already has a grade."""
        # Create a document with existing grade
        from docx import Document

        doc = Document()
        doc.add_paragraph("Test document content")
        # Add a table with grade
        table = doc.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "评定分数"
        table.cell(0, 1).text = "A"
        doc.save(self.test_file_path)

        self.client.login(username="testuser", password="testpass123")

        data = {"path": "test_document.docx"}
        response = self.client.post(reverse("grading:ai_score"), data)

        self.assertEqual(response.status_code, 400)
        response_data = json.loads(response.content)
        self.assertEqual(response_data["status"], "error")
        self.assertIn("已有评分", response_data["message"])
        self.assertIn("A", response_data["message"])

    @patch("grading.views.volcengine_score_homework")
    def test_batch_ai_score_view_success(self, mock_ai_score):
        """Test successful batch AI scoring."""
        # Mock AI scoring response
        mock_ai_score.return_value = (90, "Great work!")

        self.client.login(username="testuser", password="testpass123")

        data = {"path": "test_directory"}
        response = self.client.post(reverse("grading:batch_ai_score"), data)

        self.assertEqual(response.status_code, 200)
        response_data = json.loads(response.content)
        self.assertEqual(response_data["status"], "success")

    def test_batch_ai_score_view_missing_data(self):
        """Test batch AI scoring with missing data."""
        self.client.login(username="testuser", password="testpass123")

        data = {}  # Missing path
        response = self.client.post(reverse("grading:batch_ai_score"), data)

        self.assertEqual(response.status_code, 400)

    def test_batch_ai_score_view_unauthenticated(self):
        """Test batch AI scoring without authentication."""
        data = {"path": "test_directory"}
        response = self.client.post(reverse("grading:batch_ai_score"), data)

        self.assertEqual(response.status_code, 302)  # Redirect to login

    def test_batch_ai_score_view_with_graded_files(self):
        """Test batch AI scoring when some files already have grades."""
        # Create test directory
        test_dir = os.path.join(self.temp_dir, "test_directory")
        os.makedirs(test_dir, exist_ok=True)

        # Create a file with existing grade
        graded_file_path = os.path.join(test_dir, "graded_document.docx")
        from docx import Document

        doc = Document()
        doc.add_paragraph("Test document content")
        # Add a table with grade
        table = doc.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "评定分数"
        table.cell(0, 1).text = "B"
        doc.save(graded_file_path)

        # Create a file without grade
        ungraded_file_path = os.path.join(test_dir, "ungraded_document.docx")
        doc = Document()
        doc.add_paragraph("Test document content")
        doc.save(ungraded_file_path)

        self.client.login(username="testuser", password="testpass123")

        data = {"path": "test_directory"}
        response = self.client.post(reverse("grading:batch_ai_score"), data)

        self.assertEqual(response.status_code, 200)
        response_data = json.loads(response.content)
        self.assertEqual(response_data["status"], "success")

        # Check that we have results for both files
        self.assertEqual(len(response_data["results"]), 2)

        # Find the graded file result
        graded_result = next(
            (r for r in response_data["results"] if r["file"] == "graded_document.docx"), None
        )
        self.assertIsNotNone(graded_result)
        self.assertFalse(graded_result["success"])
        self.assertIn("已有评分", graded_result["error"])
        self.assertIn("B", graded_result["error"])


class AIScoringFunctionTest(TestCase):
    """Test cases for AI scoring functions."""

    @patch("grading.views.Ark")
    def test_volcengine_score_homework_success(self, mock_ark):
        """Test successful volcengine AI scoring."""
        # Mock Ark client
        mock_client = MagicMock()
        mock_ark.return_value = mock_client

        # Mock API response
        mock_response = MagicMock()
        mock_response.choices = [MagicMock()]
        mock_response.choices[0].message.content = "分数：85分\n评价：作业完成得很好，思路清晰"
        mock_client.chat.completions.create.return_value = mock_response

        # Mock environment variable
        with patch.dict("os.environ", {"ARK_API_KEY": "test_key"}):
            from grading.views import volcengine_score_homework

            score, comment = volcengine_score_homework("Test homework content")

            self.assertEqual(score, 85)
            self.assertIn("作业完成得很好", comment)

    @patch("grading.views.Ark")
    def test_volcengine_score_homework_no_api_key(self, mock_ark):
        """Test AI scoring without API key."""
        # Mock environment variable (no API key)
        with patch.dict("os.environ", {}, clear=True):
            from grading.views import volcengine_score_homework

            score, comment = volcengine_score_homework("Test homework content")

            self.assertIsNone(score)
            self.assertEqual(comment, "API密钥未配置")

    @patch("grading.views.Ark")
    def test_volcengine_score_homework_api_error(self, mock_ark):
        """Test AI scoring with API error."""
        # Mock Ark client
        mock_client = MagicMock()
        mock_ark.return_value = mock_client

        # Mock API error
        mock_client.chat.completions.create.side_effect = Exception("API error")

        # Mock environment variable
        with patch.dict("os.environ", {"ARK_API_KEY": "test_key"}):
            from grading.views import volcengine_score_homework

            score, comment = volcengine_score_homework("Test homework content")

            self.assertIsNone(score)
            self.assertEqual(comment, "")

    @patch("grading.views.Ark")
    def test_volcengine_score_homework_no_score_in_response(self, mock_ark):
        """Test AI scoring with no score in response."""
        # Mock Ark client
        mock_client = MagicMock()
        mock_ark.return_value = mock_client

        # Mock API response without score
        mock_response = MagicMock()
        mock_response.choices = [MagicMock()]
        mock_response.choices[0].message.content = "作业完成得很好，思路清晰"
        mock_client.chat.completions.create.return_value = mock_response

        # Mock environment variable
        with patch.dict("os.environ", {"ARK_API_KEY": "test_key"}):
            from grading.views import volcengine_score_homework

            score, comment = volcengine_score_homework("Test homework content")

            self.assertIsNone(score)
            self.assertEqual(comment, "作业完成得很好，思路清晰")

    def test_convert_score_to_grade(self):
        """Test score to grade conversion."""
        from grading.views import convert_score_to_grade

        # Test grade conversions
        self.assertEqual(convert_score_to_grade(95), "A")
        self.assertEqual(convert_score_to_grade(85), "B")
        self.assertEqual(convert_score_to_grade(75), "C")
        self.assertEqual(convert_score_to_grade(65), "D")
        self.assertEqual(convert_score_to_grade(55), "E")
        self.assertEqual(convert_score_to_grade(None), "N/A")


class AIScoringIntegrationTest(TestCase):
    """Integration tests for AI scoring."""

    def setUp(self):
        """Set up test data."""
        self.user = User.objects.create_user(
            username="testuser", password="testpass123", is_staff=True
        )
        self.client.login(username="testuser", password="testpass123")
        self.temp_dir = tempfile.mkdtemp()
        self.config = GlobalConfig.objects.create(repo_base_dir=self.temp_dir)
        # Create a test file in the temp directory
        from docx import Document

        doc = Document()
        doc.add_paragraph("test content")
        doc.save(os.path.join(self.temp_dir, "test_document.docx"))
        os.makedirs(os.path.join(self.temp_dir, "test_directory"))
        doc.save(os.path.join(self.temp_dir, "test_directory", "test1.docx"))
        doc.save(os.path.join(self.temp_dir, "test_directory", "test2.docx"))

    @patch("grading.views.volcengine_score_homework")
    @patch("grading.views.write_grade_and_comment_to_file")
    def test_ai_scoring_integration(self, mock_write_file, mock_ai_score):
        """Test AI scoring integration with file operations."""
        # Mock AI scoring response
        mock_ai_score.return_value = (90, "Excellent work!")

        # Mock file operations
        mock_write_file.return_value = None

        data = {"path": "test_document.docx"}
        response = self.client.post(reverse("grading:ai_score"), data)

        self.assertEqual(response.status_code, 200)
        response_data = json.loads(response.content)
        self.assertEqual(response_data["status"], "success")

        # Verify unified file operation was called
        mock_write_file.assert_called_once()

    @patch("grading.views.volcengine_score_homework")
    def test_ai_scoring_with_different_content_types(self, mock_ai_score):
        """Test AI scoring with different content types."""
        # Mock AI scoring response
        mock_ai_score.return_value = (88, "Good work!")

        content_types = [
            "This is a simple text homework.",
            "这是一个中文作业内容。",
            "Homework with numbers: 1, 2, 3, 4, 5",
            "Very long homework content " * 100,  # Long content
        ]

        for content in content_types:
            score, comment = mock_ai_score(content)
            self.assertEqual(score, 88)
            self.assertEqual(comment, "Good work!")

    def test_ai_scoring_error_handling(self):
        """Test AI scoring error handling."""
        # Test with invalid file path
        data = {"path": "nonexistent_file.docx"}
        response = self.client.post(reverse("grading:ai_score"), data)

        self.assertEqual(response.status_code, 200)
        response_data = json.loads(response.content)
        self.assertEqual(response_data["status"], "error")

    @patch("grading.views.volcengine_score_homework")
    def test_batch_ai_scoring_with_multiple_files(self, mock_ai_score):
        """Test batch AI scoring with multiple files."""
        # Mock AI scoring response
        mock_ai_score.return_value = (85, "Good work!")

        data = {"path": "test_directory"}
        response = self.client.post(reverse("grading:batch_ai_score"), data)

        self.assertEqual(response.status_code, 200)
        response_data = json.loads(response.content)
        self.assertEqual(response_data["status"], "success")
        self.assertIn("results", response_data)
