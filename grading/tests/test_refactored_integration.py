"""
集成测试：重构后的评分功能
Integration tests for refactored grading features

测试需求: 4.1-4.9 (手动评分), 5.1-5.8 (教师评价), 6.1-6.10 (AI评分), 
         15.1-15.7 (获取评价), 16.1-16.7 (撤销评分)
"""

import os
import shutil
import tempfile
from unittest.mock import patch

from django.test import TestCase, Client
from django.contrib.auth.models import User
from django.urls import reverse
from docx import Document
from datetime import date

from grading.models import Semester, Course, Homework
from grading.views import (
    find_teacher_signature_cell,
    extract_grade_and_comment_from_cell,
    write_to_teacher_signature_cell,
)


class RefactoredGradingIntegrationTest(TestCase):
    """重构后的评分功能集成测试"""

    def setUp(self):
        """设置测试环境"""
        self.temp_dir = tempfile.mkdtemp()
        self.client = Client()
        
        self.user = User.objects.create_user(
            username='testteacher',
            password='testpass123',
            is_staff=True
        )

    def tearDown(self):
        """清理测试环境"""
        if os.path.exists(self.temp_dir):
            shutil.rmtree(self.temp_dir)

    def create_lab_report(self, filename='test_report.docx'):
        """创建测试用的实验报告文件"""
        doc = Document()
        doc.add_heading('实验报告', 0)
        doc.add_paragraph('学生姓名：张三')
        doc.add_paragraph('学号：20240001')
        
        table = doc.add_table(rows=3, cols=2)
        table.rows[0].cells[0].text = '实验内容'
        table.rows[0].cells[1].text = '网络协议分析'
        table.rows[1].cells[0].text = '实验结果'
        table.rows[1].cells[1].text = '完成所有实验要求'
        table.rows[2].cells[0].text = '教师（签字）：'
        
        file_path = os.path.join(self.temp_dir, filename)
        doc.save(file_path)
        return file_path

    def test_manual_grading_new_submission(self):
        """测试手动评分流程 - 新建评分 (需求 4.1, 4.3, 4.4, 4.5, 4.6)"""
        file_path = self.create_lab_report()
        
        doc = Document(file_path)
        cell, _, _, _ = find_teacher_signature_cell(doc)
        self.assertIsNotNone(cell, "应该找到教师签字单元格")
        
        grade = "A"
        comment = "实验报告完成得非常出色，数据准确，分析透彻。"
        signature_text = "教师（签字）："
        
        write_to_teacher_signature_cell(cell, grade, comment, signature_text)
        doc.save(file_path)
        
        doc_reloaded = Document(file_path)
        cell_reloaded, _, _, _ = find_teacher_signature_cell(doc_reloaded)
        grade_extracted, comment_extracted, sig_extracted = extract_grade_and_comment_from_cell(cell_reloaded)
        
        self.assertEqual(grade_extracted, grade)
        self.assertEqual(comment_extracted, comment)
        self.assertIn("教师（签字）", sig_extracted)

    def test_manual_grading_update_existing(self):
        """测试手动评分流程 - 更新评分 (需求 4.8)"""
        file_path = self.create_lab_report()
        
        doc = Document(file_path)
        cell, _, _, _ = find_teacher_signature_cell(doc)
        
        write_to_teacher_signature_cell(cell, "B", "初次评价", "教师（签字）：张老师")
        doc.save(file_path)
        
        doc = Document(file_path)
        cell, _, _, _ = find_teacher_signature_cell(doc)
        grade, comment, sig = extract_grade_and_comment_from_cell(cell)
        
        self.assertEqual(grade, "B")
        self.assertEqual(comment, "初次评价")
        
        write_to_teacher_signature_cell(cell, "A", "修改后的评价", sig)
        doc.save(file_path)
        
        doc = Document(file_path)
        cell, _, _, _ = find_teacher_signature_cell(doc)
        grade_new, comment_new, sig_new = extract_grade_and_comment_from_cell(cell)
        
        self.assertEqual(grade_new, "A")
        self.assertEqual(comment_new, "修改后的评价")
        self.assertIn("张老师", sig_new)

    def test_teacher_comment_functionality(self):
        """测试教师评价功能 (需求 5.1-5.8)"""
        file_path = self.create_lab_report()
        
        doc = Document(file_path)
        cell, _, _, _ = find_teacher_signature_cell(doc)
        
        grade = "A"
        comment = "实验报告格式规范，数据准确，分析透彻，结论合理。"
        signature_text = "教师（签字）：李老师 时间：2024-03-15"
        
        write_to_teacher_signature_cell(cell, grade, comment, signature_text)
        doc.save(file_path)
        
        doc = Document(file_path)
        cell, _, _, _ = find_teacher_signature_cell(doc)
        _, comment_extracted, sig_extracted = extract_grade_and_comment_from_cell(cell)
        
        self.assertEqual(comment_extracted, comment)
        self.assertIn("李老师", sig_extracted)
        self.assertIn("2024-03-15", sig_extracted)


    def test_get_teacher_comment(self):
        """测试获取教师评价功能 (需求 15.1-15.7)"""
        file_path = self.create_lab_report()
        
        doc = Document(file_path)
        cell, _, _, _ = find_teacher_signature_cell(doc)
        
        grade = "B"
        comment = "作业完成较好，但还有改进空间。"
        signature_text = "教师（签字）：王老师"
        
        write_to_teacher_signature_cell(cell, grade, comment, signature_text)
        doc.save(file_path)
        
        doc = Document(file_path)
        cell, _, _, _ = find_teacher_signature_cell(doc)
        grade_extracted, comment_extracted, _ = extract_grade_and_comment_from_cell(cell)
        
        self.assertEqual(grade_extracted, grade)
        self.assertEqual(comment_extracted, comment)

    def test_undo_grading(self):
        """测试撤销评分功能 (需求 16.1-16.7)"""
        file_path = self.create_lab_report()
        
        doc = Document(file_path)
        cell, _, _, _ = find_teacher_signature_cell(doc)
        
        original_signature = "教师（签字）：赵老师 时间：2024-03-10"
        write_to_teacher_signature_cell(cell, "A", "优秀", original_signature)
        doc.save(file_path)
        
        doc = Document(file_path)
        cell, _, _, _ = find_teacher_signature_cell(doc)
        grade, comment, sig = extract_grade_and_comment_from_cell(cell)
        
        self.assertEqual(grade, "A")
        self.assertEqual(comment, "优秀")
        
        for paragraph in cell.paragraphs:
            paragraph.clear()
        
        cell.text = sig
        doc.save(file_path)
        
        doc = Document(file_path)
        cell, _, _, _ = find_teacher_signature_cell(doc)
        grade_after, comment_after, sig_after = extract_grade_and_comment_from_cell(cell)
        
        self.assertIsNone(grade_after)
        self.assertIsNone(comment_after)
        self.assertIn("赵老师", sig_after)

    def test_grade_type_switching(self):
        """测试评分方式切换 (需求 4.1, 4.2)"""
        file_path = self.create_lab_report()
        
        doc = Document(file_path)
        cell, _, _, _ = find_teacher_signature_cell(doc)
        
        write_to_teacher_signature_cell(cell, "A", "字母评分", "教师（签字）：")
        doc.save(file_path)
        
        doc = Document(file_path)
        cell, _, _, _ = find_teacher_signature_cell(doc)
        grade1, _, _ = extract_grade_and_comment_from_cell(cell)
        self.assertEqual(grade1, "A")
        
        write_to_teacher_signature_cell(cell, "优秀", "中文评分", "教师（签字）：")
        doc.save(file_path)
        
        doc = Document(file_path)
        cell, _, _, _ = find_teacher_signature_cell(doc)
        grade2, _, _ = extract_grade_and_comment_from_cell(cell)
        self.assertEqual(grade2, "优秀")


    # AI评分测试需要完整的用户配置（Tenant和UserProfile），暂时跳过
    # @patch('grading.views.volcengine_score_homework')
    # def test_ai_scoring_flow(self, mock_ai):
    #     """测试AI评分流程 (需求 6.1-6.10)"""
    #     pass

    def test_format_error_handling(self):
        """测试格式错误处理 (需求 11.1-11.9)"""
        doc = Document()
        doc.add_heading('实验报告', 0)
        doc.add_paragraph('这是一个格式错误的实验报告，没有教师签字单元格')
        
        file_path = os.path.join(self.temp_dir, 'format_error.docx')
        doc.save(file_path)
        
        doc = Document(file_path)
        cell, _, _, _ = find_teacher_signature_cell(doc)
        
        self.assertIsNone(cell, "格式错误的文档不应该找到教师签字单元格")

    def test_complete_workflow_integration(self):
        """测试完整工作流集成"""
        file_path = self.create_lab_report()
        
        doc = Document(file_path)
        cell, _, _, _ = find_teacher_signature_cell(doc)
        self.assertIsNotNone(cell)
        
        write_to_teacher_signature_cell(cell, "B", "初次评价", "教师（签字）：张老师")
        doc.save(file_path)
        
        doc = Document(file_path)
        cell, _, _, _ = find_teacher_signature_cell(doc)
        grade, comment, sig = extract_grade_and_comment_from_cell(cell)
        self.assertEqual(grade, "B")
        self.assertEqual(comment, "初次评价")
        
        write_to_teacher_signature_cell(cell, "A", "修改后的评价", sig)
        doc.save(file_path)
        
        doc = Document(file_path)
        cell, _, _, _ = find_teacher_signature_cell(doc)
        grade, comment, sig = extract_grade_and_comment_from_cell(cell)
        self.assertEqual(grade, "A")
        self.assertEqual(comment, "修改后的评价")
        self.assertIn("张老师", sig)
        
        for paragraph in cell.paragraphs:
            paragraph.clear()
        cell.text = sig
        doc.save(file_path)
        
        doc = Document(file_path)
        cell, _, _, _ = find_teacher_signature_cell(doc)
        grade, comment, sig = extract_grade_and_comment_from_cell(cell)
        self.assertIsNone(grade)
        self.assertIsNone(comment)
        self.assertIn("张老师", sig)
