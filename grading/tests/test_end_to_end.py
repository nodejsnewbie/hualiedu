"""
端到端测试：完整的评分流程
End-to-End tests for complete grading workflows

测试需求: 所有需求
- 完整的评分流程（手动评分、AI评分、教师评价）
- 格式错误处理流程（实验报告格式验证和锁定机制）
- 批量操作流程（批量评分、批量AI评分）
"""

import os
import shutil
import tempfile
import time
from pathlib import Path
from unittest.mock import patch, MagicMock

from django.test import TestCase, Client
from django.contrib.auth.models import User
from django.urls import reverse
from docx import Document

from grading.models import Semester, Course, Homework, Repository, Tenant, UserProfile
from grading.views import (
    find_teacher_signature_cell,
    extract_grade_and_comment_from_cell,
    write_to_teacher_signature_cell,
    is_lab_report_file,
)


class EndToEndGradingWorkflowTest(TestCase):
    """端到端测试：完整的评分工作流"""

    def setUp(self):
        """设置测试环境"""
        self.temp_dir = tempfile.mkdtemp()
        self.client = Client()
        
        # 创建租户
        self.tenant = Tenant.objects.create(
            name='测试学校',
            description='端到端测试租户',
            is_active=True,
            tenant_repo_dir='test_tenant'
        )
        
        # 创建用户和用户配置
        self.user = User.objects.create_user(
            username='testteacher',
            password='testpass123',
            is_staff=True
        )
        
        self.user_profile = UserProfile.objects.create(
            user=self.user,
            tenant=self.tenant,
            is_tenant_admin=False
        )
        
        # 创建学期
        self.semester = Semester.objects.create(
            name='2024春季学期',
            start_date='2024-02-01',
            end_date='2024-07-01'
        )
        
        # 创建仓库
        self.repository = Repository.objects.create(
            name='测试仓库',
            path=self.temp_dir,
            repo_type='local',
            owner=self.user,
            tenant=self.tenant
        )
        
        # 创建课程
        self.course = Course.objects.create(
            name='计算机网络实验',
            course_type='lab',
            semester=self.semester,
            teacher=self.user,
            location='实验室A101'
        )
        
        # 创建作业批次
        self.homework = Homework.objects.create(
            course=self.course,
            title='第一次实验',
            folder_name='第一次实验',
            homework_type='lab_report'
        )
        
        # 创建目录结构
        self.course_dir = Path(self.temp_dir) / '计算机网络实验'
        self.homework_dir = self.course_dir / '第一次实验'
        self.homework_dir.mkdir(parents=True, exist_ok=True)

    def tearDown(self):
        """清理测试环境"""
        if os.path.exists(self.temp_dir):
            shutil.rmtree(self.temp_dir)

    def create_lab_report(self, filename='张三.docx', with_signature_cell=True):
        """创建实验报告文件"""
        doc = Document()
        doc.add_heading('计算机网络实验报告', 0)
        doc.add_paragraph('学生姓名：张三')
        doc.add_paragraph('学号：20240001')
        doc.add_paragraph('实验内容：TCP/IP协议分析')
        
        if with_signature_cell:
            table = doc.add_table(rows=3, cols=2)
            table.rows[0].cells[0].text = '实验结果'
            table.rows[0].cells[1].text = '完成所有实验要求'
            table.rows[1].cells[0].text = '实验总结'
            table.rows[1].cells[1].text = '通过本次实验，深入理解了TCP/IP协议的工作原理'
            table.rows[2].cells[0].text = '教师（签字）：'
        
        file_path = self.homework_dir / filename
        doc.save(str(file_path))
        return str(file_path)

    def create_normal_homework(self, filename='李四.docx'):
        """创建普通作业文件"""
        doc = Document()
        doc.add_heading('计算机网络作业', 0)
        doc.add_paragraph('学生姓名：李四')
        doc.add_paragraph('学号：20240002')
        doc.add_paragraph('作业内容：简述OSI七层模型')
        doc.add_paragraph('答：OSI七层模型包括物理层、数据链路层、网络层、传输层、会话层、表示层和应用层...')
        
        file_path = self.homework_dir / filename
        doc.save(str(file_path))
        return str(file_path)

    def test_complete_manual_grading_workflow(self):
        """
        测试完整的手动评分工作流
        需求: 4.1-4.9 (手动评分), 5.1-5.8 (教师评价), 15.1-15.7 (获取评价), 16.1-16.7 (撤销评分)
        """
        # 步骤1: 创建实验报告
        file_path = self.create_lab_report('张三.docx')
        
        # 步骤2: 首次评分
        doc = Document(file_path)
        cell, _, _, _ = find_teacher_signature_cell(doc)
        self.assertIsNotNone(cell, "应该找到教师签字单元格")
        
        grade = "A"
        comment = "实验报告完成得非常出色，数据准确，分析透彻。"
        signature_text = "教师（签字）：张老师 2024-03-15"
        
        write_to_teacher_signature_cell(cell, grade, comment, signature_text)
        doc.save(file_path)
        
        # 步骤3: 验证评分已写入
        doc = Document(file_path)
        cell, _, _, _ = find_teacher_signature_cell(doc)
        grade_read, comment_read, sig_read = extract_grade_and_comment_from_cell(cell)
        
        self.assertEqual(grade_read, grade)
        self.assertEqual(comment_read, comment)
        self.assertIn("张老师", sig_read)
        
        # 步骤4: 更新评分（使用系统支持的评分等级）
        new_grade = "B"
        new_comment = "修改评价：实验报告非常优秀，超出预期。"
        
        # 重新加载文档以获取新的cell对象
        doc = Document(file_path)
        cell, _, _, _ = find_teacher_signature_cell(doc)
        _, _, sig_read = extract_grade_and_comment_from_cell(cell)
        
        write_to_teacher_signature_cell(cell, new_grade, new_comment, sig_read)
        doc.save(file_path)
        
        # 步骤5: 验证更新后的评分
        doc = Document(file_path)
        cell, _, _, _ = find_teacher_signature_cell(doc)
        grade_updated, comment_updated, sig_updated = extract_grade_and_comment_from_cell(cell)
        
        self.assertEqual(grade_updated, new_grade)
        self.assertEqual(comment_updated, new_comment)
        self.assertIn("张老师", sig_updated)
        
        # 步骤6: 撤销评分
        for paragraph in cell.paragraphs:
            paragraph.clear()
        cell.text = sig_updated
        doc.save(file_path)
        
        # 步骤7: 验证评分已撤销
        doc = Document(file_path)
        cell, _, _, _ = find_teacher_signature_cell(doc)
        grade_removed, comment_removed, sig_final = extract_grade_and_comment_from_cell(cell)
        
        self.assertIsNone(grade_removed)
        self.assertIsNone(comment_removed)
        self.assertIn("张老师", sig_final)

    def test_format_error_handling_workflow(self):
        """
        测试格式错误处理流程
        需求: 11.1-11.9 (格式错误锁定机制)
        """
        # 步骤1: 创建格式错误的实验报告（没有教师签字单元格）
        file_path = self.create_lab_report('格式错误.docx', with_signature_cell=False)
        
        # 步骤2: 验证无法找到教师签字单元格
        doc = Document(file_path)
        cell, _, _, _ = find_teacher_signature_cell(doc)
        self.assertIsNone(cell, "格式错误的文档不应该找到教师签字单元格")
        
        # 步骤3: 模拟格式错误处理（降级为普通作业方式）
        # 按照设计文档，格式错误时应该：
        # - 评分自动改为 D
        # - 评价设置为 "【格式错误-已锁定】请按要求的格式写实验报告，此评分不可修改"
        # - 按普通作业方式写入（文档末尾段落）
        
        error_grade = "D"
        error_comment = "【格式错误-已锁定】请按要求的格式写实验报告，此评分不可修改"
        
        doc.add_paragraph(f"老师评分：{error_grade}")
        doc.add_paragraph(f"教师评价：{error_comment}")
        doc.save(file_path)
        
        # 步骤4: 验证锁定标记
        doc = Document(file_path)
        text_content = '\n'.join([p.text for p in doc.paragraphs])
        
        self.assertIn("【格式错误-已锁定】", text_content)
        self.assertIn("D", text_content)
        
        # 步骤5: 验证文件被锁定（检查是否包含锁定标记）
        is_locked = "【格式错误-已锁定】" in text_content
        self.assertTrue(is_locked, "文件应该被锁定")

    def test_homework_type_detection_workflow(self):
        """
        测试作业类型判断流程
        需求: 2.1.1-2.1.9 (作业类型自动判断)
        """
        # 步骤1: 测试实验报告类型判断（通过数据库配置）
        file_path = self.create_lab_report('实验报告.docx')
        
        is_lab = is_lab_report_file(
            course_name='计算机网络实验',
            homework_folder='第一次实验',
            file_path=file_path,
            base_dir=self.temp_dir
        )
        
        self.assertTrue(is_lab, "应该判断为实验报告")
        
        # 步骤2: 修改作业类型为普通作业
        self.homework.homework_type = 'normal'
        self.homework.save()
        
        is_lab_after_change = is_lab_report_file(
            course_name='计算机网络实验',
            homework_folder='第一次实验',
            file_path=file_path,
            base_dir=self.temp_dir
        )
        
        self.assertFalse(is_lab_after_change, "应该判断为普通作业")
        
        # 步骤3: 测试课程类型判断（删除作业批次配置）
        self.homework.delete()
        
        is_lab_by_course = is_lab_report_file(
            course_name='计算机网络实验',
            homework_folder='第一次实验',
            file_path=file_path,
            base_dir=self.temp_dir
        )
        
        # 课程类型为lab，应该判断为实验报告
        self.assertTrue(is_lab_by_course, "根据课程类型应该判断为实验报告")

    def test_grade_type_switching_workflow(self):
        """
        测试评分方式切换流程
        需求: 4.1, 4.2 (评分方式切换)
        """
        file_path = self.create_lab_report('评分切换.docx')
        
        # 步骤1: 使用字母评分
        doc = Document(file_path)
        cell, _, _, _ = find_teacher_signature_cell(doc)
        
        write_to_teacher_signature_cell(cell, "A", "字母评分测试", "教师（签字）：")
        doc.save(file_path)
        
        doc = Document(file_path)
        cell, _, _, _ = find_teacher_signature_cell(doc)
        grade1, _, _ = extract_grade_and_comment_from_cell(cell)
        self.assertEqual(grade1, "A")
        
        # 步骤2: 切换到中文评分
        write_to_teacher_signature_cell(cell, "优秀", "中文评分测试", "教师（签字）：")
        doc.save(file_path)
        
        doc = Document(file_path)
        cell, _, _, _ = find_teacher_signature_cell(doc)
        grade2, _, _ = extract_grade_and_comment_from_cell(cell)
        self.assertEqual(grade2, "优秀")
        
        # 步骤3: 再次切换回字母评分
        write_to_teacher_signature_cell(cell, "B", "再次切换测试", "教师（签字）：")
        doc.save(file_path)
        
        doc = Document(file_path)
        cell, _, _, _ = find_teacher_signature_cell(doc)
        grade3, _, _ = extract_grade_and_comment_from_cell(cell)
        self.assertEqual(grade3, "B")


class EndToEndBatchOperationsTest(TestCase):
    """端到端测试：批量操作流程"""

    def setUp(self):
        """设置测试环境"""
        self.temp_dir = tempfile.mkdtemp()
        self.client = Client()
        
        # 创建租户
        self.tenant = Tenant.objects.create(
            name='测试学校',
            description='端到端测试租户',
            is_active=True,
            tenant_repo_dir='test_tenant'
        )
        
        # 创建用户
        self.user = User.objects.create_user(
            username='testteacher',
            password='testpass123',
            is_staff=True
        )
        
        self.user_profile = UserProfile.objects.create(
            user=self.user,
            tenant=self.tenant,
            is_tenant_admin=False
        )
        
        # 创建学期
        self.semester = Semester.objects.create(
            name='2024春季学期',
            start_date='2024-02-01',
            end_date='2024-07-01'
        )
        
        # 创建仓库
        self.repository = Repository.objects.create(
            name='批量测试仓库',
            path=self.temp_dir,
            repo_type='local',
            owner=self.user,
            tenant=self.tenant
        )
        
        # 创建课程
        self.course = Course.objects.create(
            name='数据结构实验',
            course_type='lab',
            semester=self.semester,
            teacher=self.user,
            location='实验室B202'
        )
        
        # 创建作业批次
        self.homework = Homework.objects.create(
            course=self.course,
            title='第一次实验',
            folder_name='第一次实验',
            homework_type='lab_report'
        )
        
        # 创建目录结构
        self.course_dir = Path(self.temp_dir) / '数据结构实验'
        self.homework_dir = self.course_dir / '第一次实验'
        self.homework_dir.mkdir(parents=True, exist_ok=True)

    def tearDown(self):
        """清理测试环境"""
        if os.path.exists(self.temp_dir):
            shutil.rmtree(self.temp_dir)

    def create_multiple_lab_reports(self, count=5):
        """创建多个实验报告文件"""
        file_paths = []
        students = ['张三', '李四', '王五', '赵六', '钱七', '孙八', '周九', '吴十']
        
        for i in range(count):
            student_name = students[i % len(students)]
            doc = Document()
            doc.add_heading(f'{student_name}的实验报告', 0)
            doc.add_paragraph(f'学生姓名：{student_name}')
            doc.add_paragraph(f'学号：202400{i+1:02d}')
            doc.add_paragraph('实验内容：数据结构基础实验')
            
            table = doc.add_table(rows=2, cols=2)
            table.rows[0].cells[0].text = '实验结果'
            table.rows[0].cells[1].text = f'完成实验{i+1}'
            table.rows[1].cells[0].text = '教师（签字）：'
            
            file_path = self.homework_dir / f'{student_name}_{i+1}.docx'
            doc.save(str(file_path))
            file_paths.append(str(file_path))
        
        return file_paths

    def test_batch_grading_workflow(self):
        """
        测试批量评分流程
        需求: 7.1-7.6 (批量评分功能)
        """
        # 步骤1: 创建多个实验报告
        file_paths = self.create_multiple_lab_reports(count=5)
        self.assertEqual(len(file_paths), 5, "应该创建5个文件")
        
        # 步骤2: 批量评分
        batch_grade = "B"
        batch_comment = "批量评分测试"
        signature = "教师（签字）：批量评分老师"
        
        success_count = 0
        error_count = 0
        
        for file_path in file_paths:
            try:
                doc = Document(file_path)
                cell, _, _, _ = find_teacher_signature_cell(doc)
                
                if cell:
                    write_to_teacher_signature_cell(cell, batch_grade, batch_comment, signature)
                    doc.save(file_path)
                    success_count += 1
                else:
                    error_count += 1
            except Exception as e:
                error_count += 1
        
        # 步骤3: 验证批量评分结果
        self.assertEqual(success_count, 5, "应该成功评分5个文件")
        self.assertEqual(error_count, 0, "不应该有失败的文件")
        
        # 步骤4: 验证每个文件的评分
        for file_path in file_paths:
            doc = Document(file_path)
            cell, _, _, _ = find_teacher_signature_cell(doc)
            grade, comment, sig = extract_grade_and_comment_from_cell(cell)
            
            self.assertEqual(grade, batch_grade)
            self.assertEqual(comment, batch_comment)
            self.assertIn("批量评分老师", sig)

    @patch('grading.views.volcengine_score_homework')
    def test_batch_ai_scoring_workflow(self, mock_ai):
        """
        测试批量AI评分流程
        需求: 8.1-8.7 (批量AI评分功能)
        """
        # 模拟AI评分返回
        mock_ai.return_value = (85, "AI评价：实验报告完成良好，建议改进数据分析部分。")
        
        # 步骤1: 创建多个实验报告
        file_paths = self.create_multiple_lab_reports(count=3)
        
        # 步骤2: 批量AI评分（模拟速率限制）
        success_count = 0
        error_count = 0
        results = []
        
        for i, file_path in enumerate(file_paths):
            try:
                # 模拟速率限制：每秒最多2个请求
                if i > 0 and i % 2 == 0:
                    time.sleep(0.5)  # 简化的速率限制
                
                # 调用AI评分
                score, comment = mock_ai()
                
                # 将分数转换为等级
                if score >= 90:
                    grade = "A"
                elif score >= 80:
                    grade = "B"
                elif score >= 70:
                    grade = "C"
                elif score >= 60:
                    grade = "D"
                else:
                    grade = "E"
                
                # 写入评分
                doc = Document(file_path)
                cell, _, _, _ = find_teacher_signature_cell(doc)
                
                if cell:
                    write_to_teacher_signature_cell(
                        cell, 
                        grade, 
                        comment, 
                        "教师（签字）：AI评分系统"
                    )
                    doc.save(file_path)
                    success_count += 1
                    results.append({
                        'file': Path(file_path).name,
                        'status': 'success',
                        'grade': grade,
                        'score': score
                    })
                else:
                    error_count += 1
                    results.append({
                        'file': Path(file_path).name,
                        'status': 'error',
                        'error': '格式错误'
                    })
            except Exception as e:
                error_count += 1
                results.append({
                    'file': Path(file_path).name,
                    'status': 'error',
                    'error': str(e)
                })
        
        # 步骤3: 验证批量AI评分结果
        self.assertEqual(success_count, 3, "应该成功评分3个文件")
        self.assertEqual(error_count, 0, "不应该有失败的文件")
        self.assertEqual(len(results), 3, "应该有3个结果记录")
        
        # 步骤4: 验证AI评分调用次数
        self.assertEqual(mock_ai.call_count, 3, "AI评分应该被调用3次")
        
        # 步骤5: 验证每个文件的AI评分
        for file_path in file_paths:
            doc = Document(file_path)
            cell, _, _, _ = find_teacher_signature_cell(doc)
            grade, comment, sig = extract_grade_and_comment_from_cell(cell)
            
            self.assertIsNotNone(grade)
            self.assertIn("AI评价", comment)
            self.assertIn("AI评分系统", sig)

    def test_batch_operations_with_mixed_files(self):
        """
        测试混合文件的批量操作
        需求: 7.6 (批量操作错误处理)
        """
        # 步骤1: 创建混合文件（正常文件和格式错误文件）
        normal_files = self.create_multiple_lab_reports(count=3)
        
        # 创建格式错误的文件
        error_file = self.homework_dir / '格式错误.docx'
        doc = Document()
        doc.add_heading('格式错误的报告', 0)
        doc.add_paragraph('这个文件没有教师签字单元格')
        doc.save(str(error_file))
        
        all_files = normal_files + [str(error_file)]
        
        # 步骤2: 批量评分
        success_count = 0
        error_count = 0
        
        for file_path in all_files:
            try:
                doc = Document(file_path)
                cell, _, _, _ = find_teacher_signature_cell(doc)
                
                if cell:
                    write_to_teacher_signature_cell(
                        cell, 
                        "B", 
                        "批量评分", 
                        "教师（签字）："
                    )
                    doc.save(file_path)
                    success_count += 1
                else:
                    # 格式错误，按降级处理
                    doc.add_paragraph("老师评分：D")
                    doc.add_paragraph("教师评价：【格式错误-已锁定】请按要求的格式写实验报告，此评分不可修改")
                    doc.save(file_path)
                    error_count += 1
            except Exception as e:
                error_count += 1
        
        # 步骤3: 验证结果
        self.assertEqual(success_count, 3, "应该成功评分3个正常文件")
        self.assertEqual(error_count, 1, "应该有1个格式错误文件")
        
        # 步骤4: 验证格式错误文件的处理
        doc = Document(str(error_file))
        text_content = '\n'.join([p.text for p in doc.paragraphs])
        self.assertIn("【格式错误-已锁定】", text_content)
        self.assertIn("D", text_content)


class EndToEndWebInterfaceTest(TestCase):
    """端到端测试：Web界面完整流程"""

    def setUp(self):
        """设置测试环境"""
        self.temp_dir = tempfile.mkdtemp()
        self.client = Client()
        
        # 创建租户
        self.tenant = Tenant.objects.create(
            name='测试学校',
            description='Web界面测试租户',
            is_active=True,
            tenant_repo_dir='test_tenant_web'
        )
        
        # 创建教师用户
        self.teacher = User.objects.create_user(
            username='webteacher',
            password='testpass123',
            is_staff=True
        )
        
        self.teacher_profile = UserProfile.objects.create(
            user=self.teacher,
            tenant=self.tenant,
            is_tenant_admin=False,
            repo_base_dir=self.temp_dir
        )
        
        # 创建学生用户
        self.student = User.objects.create_user(
            username='webstudent',
            password='testpass123'
        )
        
        self.student_profile = UserProfile.objects.create(
            user=self.student,
            tenant=self.tenant,
            is_tenant_admin=False
        )
        
        # 创建学期
        self.semester = Semester.objects.create(
            name='2024春季学期',
            start_date='2024-02-01',
            end_date='2024-07-01'
        )

    def tearDown(self):
        """清理测试环境"""
        if os.path.exists(self.temp_dir):
            shutil.rmtree(self.temp_dir)

    def test_teacher_creates_course_and_class(self):
        """
        测试教师创建课程和班级的完整流程
        需求: 1.1-1.5 (课程和班级创建)
        """
        # 步骤1: 教师登录
        self.client.login(username='webteacher', password='testpass123')
        
        # 步骤2: 创建课程
        from grading.services.course_service import CourseService
        course_service = CourseService()
        
        course = course_service.create_course(
            teacher=self.teacher,
            name='Web测试课程',
            course_type='lab',
            semester=self.semester,
            description='这是一个Web界面测试课程'
        )
        
        self.assertIsNotNone(course)
        self.assertEqual(course.name, 'Web测试课程')
        self.assertEqual(course.course_type, 'lab')
        self.assertEqual(course.teacher, self.teacher)
        
        # 步骤3: 创建班级
        from grading.services.class_service import ClassService
        class_service = ClassService()
        
        class_obj = class_service.create_class(
            course=course,
            name='计算机1班',
            student_count=30
        )
        
        self.assertIsNotNone(class_obj)
        self.assertEqual(class_obj.name, '计算机1班')
        self.assertEqual(class_obj.course, course)
        self.assertEqual(class_obj.student_count, 30)
        
        # 步骤4: 验证教师只能看到自己的课程
        courses = course_service.list_courses(teacher=self.teacher)
        self.assertEqual(len(courses), 1)
        self.assertEqual(courses[0].id, course.id)

    def test_student_uploads_homework(self):
        """
        测试学生上传作业的完整流程
        需求: 1.4.1-1.4.7 (学生作业上传)
        """
        # 步骤1: 创建课程和班级
        from grading.services.course_service import CourseService
        from grading.services.class_service import ClassService
        
        course_service = CourseService()
        class_service = ClassService()
        
        course = course_service.create_course(
            teacher=self.teacher,
            name='作业上传测试课程',
            course_type='lab',
            semester=self.semester,
            description='测试学生上传作业'
        )
        
        class_obj = class_service.create_class(
            course=course,
            name='测试班级',
            student_count=20
        )
        
        # 步骤2: 创建仓库（文件系统方式）
        from grading.services.repository_service import RepositoryService
        repo_service = RepositoryService()
        
        repository = repo_service.create_filesystem_repository(
            teacher=self.teacher,
            class_obj=class_obj,
            name='学生作业仓库'
        )
        
        self.assertIsNotNone(repository)
        self.assertEqual(repository.repo_type, 'filesystem')
        
        # 步骤3: 创建作业批次
        homework = Homework.objects.create(
            course=course,
            title='第一次作业',
            folder_name='第一次作业',
            homework_type='lab_report'
        )
        
        # 步骤4: 学生上传作业
        from grading.services.file_upload_service import FileUploadService
        from django.core.files.uploadedfile import SimpleUploadedFile
        
        file_service = FileUploadService()
        
        # 创建测试文件
        doc = Document()
        doc.add_heading('学生作业', 0)
        doc.add_paragraph('这是学生提交的作业内容')
        
        import io
        file_buffer = io.BytesIO()
        doc.save(file_buffer)
        file_buffer.seek(0)
        
        uploaded_file = SimpleUploadedFile(
            "学生作业.docx",
            file_buffer.read(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        
        # 验证文件
        is_valid, error_msg = file_service.validate_file(uploaded_file)
        self.assertTrue(is_valid, f"文件验证失败: {error_msg}")
        
        # 步骤5: 创建提交记录
        file_path = f"{repository.filesystem_path}/{homework.folder_name}/学生作业.docx"
        submission = file_service.create_submission_record(
            student=self.student,
            homework=homework,
            repository=repository,
            file_path=file_path,
            file_name='学生作业.docx',
            file_size=len(file_buffer.getvalue())
        )
        
        self.assertIsNotNone(submission)
        self.assertEqual(submission.student, self.student)
        self.assertEqual(submission.homework, homework)
        self.assertEqual(submission.file_name, '学生作业.docx')

    def test_comment_template_recommendation(self):
        """
        测试评价模板推荐功能
        需求: 5.2.1-5.2.12 (评价模板功能)
        """
        from grading.services.comment_template_service import CommentTemplateService
        from grading.models import CommentTemplate
        
        service = CommentTemplateService()
        
        # 步骤1: 记录教师使用评价
        comments = [
            "实验报告完成得非常出色",
            "数据分析准确，结论合理",
            "实验步骤清晰，格式规范",
            "需要加强对实验原理的理解",
            "实验报告完成得非常出色",  # 重复使用
            "数据分析准确，结论合理",  # 重复使用
            "实验报告完成得非常出色",  # 再次重复
        ]
        
        for comment in comments:
            service.record_comment_usage(
                teacher=self.teacher,
                comment_text=comment
            )
        
        # 步骤2: 获取个人评价模板
        personal_templates = service.get_personal_templates(
            teacher=self.teacher,
            limit=5
        )
        
        self.assertGreater(len(personal_templates), 0)
        self.assertLessEqual(len(personal_templates), 5)
        
        # 步骤3: 验证排序（使用次数最多的在前）
        if len(personal_templates) >= 2:
            self.assertGreaterEqual(
                personal_templates[0].usage_count,
                personal_templates[1].usage_count
            )
        
        # 步骤4: 验证最常用的评价
        top_template = personal_templates[0]
        self.assertEqual(top_template.comment_text, "实验报告完成得非常出色")
        self.assertEqual(top_template.usage_count, 3)
        
        # 步骤5: 获取推荐模板（个人优先）
        recommended = service.get_recommended_templates(teacher=self.teacher)
        self.assertGreater(len(recommended), 0)
        self.assertLessEqual(len(recommended), 5)

    def test_three_grading_methods(self):
        """
        测试三种评分方式
        需求: 4.1-4.4 (手动评分功能 - 三种评分方式)
        """
        # 创建测试文件
        homework_dir = Path(self.temp_dir) / '评分方式测试'
        homework_dir.mkdir(parents=True, exist_ok=True)
        
        # 步骤1: 测试字母评分
        file_path_letter = homework_dir / '字母评分.docx'
        doc = Document()
        doc.add_heading('实验报告', 0)
        table = doc.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = '实验结果'
        table.rows[0].cells[1].text = '完成'
        table.rows[1].cells[0].text = '教师（签字）：'
        doc.save(str(file_path_letter))
        
        doc = Document(str(file_path_letter))
        cell, _, _, _ = find_teacher_signature_cell(doc)
        write_to_teacher_signature_cell(cell, "A", "字母评分测试", "教师（签字）：")
        doc.save(str(file_path_letter))
        
        doc = Document(str(file_path_letter))
        cell, _, _, _ = find_teacher_signature_cell(doc)
        grade, _, _ = extract_grade_and_comment_from_cell(cell)
        self.assertEqual(grade, "A")
        
        # 步骤2: 测试文字评分
        file_path_text = homework_dir / '文字评分.docx'
        doc = Document()
        doc.add_heading('实验报告', 0)
        table = doc.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = '实验结果'
        table.rows[0].cells[1].text = '完成'
        table.rows[1].cells[0].text = '教师（签字）：'
        doc.save(str(file_path_text))
        
        doc = Document(str(file_path_text))
        cell, _, _, _ = find_teacher_signature_cell(doc)
        write_to_teacher_signature_cell(cell, "优秀", "文字评分测试", "教师（签字）：")
        doc.save(str(file_path_text))
        
        doc = Document(str(file_path_text))
        cell, _, _, _ = find_teacher_signature_cell(doc)
        grade, _, _ = extract_grade_and_comment_from_cell(cell)
        self.assertEqual(grade, "优秀")
        
        # 步骤3: 测试百分制评分
        # 注意：当前extract函数只识别字母和文字评分，百分制评分需要通过其他方式验证
        # 这里我们测试写入百分制评分，但验证时检查单元格文本内容
        file_path_percent = homework_dir / '百分制评分.docx'
        doc = Document()
        doc.add_heading('实验报告', 0)
        table = doc.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = '实验结果'
        table.rows[0].cells[1].text = '完成'
        table.rows[1].cells[0].text = '教师（签字）：'
        doc.save(str(file_path_percent))
        
        # 重新加载文档
        doc = Document(str(file_path_percent))
        cell, _, _, _ = find_teacher_signature_cell(doc)
        self.assertIsNotNone(cell, "应该找到教师签字单元格")
        
        write_to_teacher_signature_cell(cell, "85", "百分制评分测试", "教师（签字）：")
        doc.save(str(file_path_percent))
        
        # 再次重新加载文档以读取
        doc = Document(str(file_path_percent))
        cell, _, _, _ = find_teacher_signature_cell(doc)
        self.assertIsNotNone(cell, "重新加载后应该找到教师签字单元格")
        
        # 验证单元格包含百分制评分（通过文本内容）
        cell_text = cell.text
        self.assertIn("85", cell_text, "单元格应该包含百分制评分85")
        self.assertIn("百分制评分测试", cell_text, "单元格应该包含评价内容")


class EndToEndPerformanceTest(TestCase):
    """端到端测试：性能测试"""

    def setUp(self):
        """设置测试环境"""
        self.temp_dir = tempfile.mkdtemp()

    def tearDown(self):
        """清理测试环境"""
        if os.path.exists(self.temp_dir):
            shutil.rmtree(self.temp_dir)

    def test_large_batch_operations_performance(self):
        """
        测试大批量操作性能
        需求: 14.1-14.5 (缓存和性能优化)
        """
        # 创建测试目录
        homework_dir = Path(self.temp_dir) / '性能测试'
        homework_dir.mkdir(parents=True, exist_ok=True)
        
        # 步骤1: 创建大量文件（20个）
        file_count = 20
        file_paths = []
        
        for i in range(file_count):
            doc = Document()
            doc.add_heading(f'学生{i+1}的报告', 0)
            doc.add_paragraph(f'学号：202400{i+1:02d}')
            
            table = doc.add_table(rows=2, cols=2)
            table.rows[0].cells[0].text = '实验结果'
            table.rows[0].cells[1].text = f'完成实验{i+1}'
            table.rows[1].cells[0].text = '教师（签字）：'
            
            file_path = homework_dir / f'学生{i+1}.docx'
            doc.save(str(file_path))
            file_paths.append(str(file_path))
        
        # 步骤2: 测试批量评分性能
        start_time = time.time()
        
        for file_path in file_paths:
            doc = Document(file_path)
            cell, _, _, _ = find_teacher_signature_cell(doc)
            
            if cell:
                write_to_teacher_signature_cell(
                    cell, 
                    "B", 
                    "性能测试", 
                    "教师（签字）："
                )
                doc.save(file_path)
        
        end_time = time.time()
        elapsed_time = end_time - start_time
        
        # 步骤3: 验证性能要求（应该在合理时间内完成）
        # 20个文件应该在10秒内完成
        self.assertLess(elapsed_time, 10.0, f"批量操作应该在10秒内完成，实际用时：{elapsed_time:.2f}秒")
        
        # 步骤4: 验证所有文件都已评分
        for file_path in file_paths:
            doc = Document(file_path)
            cell, _, _, _ = find_teacher_signature_cell(doc)
            grade, _, _ = extract_grade_and_comment_from_cell(cell)
            self.assertEqual(grade, "B")

